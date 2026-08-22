import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.domain.canonical_jultp_template import (
    CANONICAL_STRUCTURE,
    subsection_alias_match,
)
from app.services.jutlp_validator import (
    check_discussion_subsections,
    check_dot_points,
    check_method_subsections,
    check_required_sections,
    check_section_order,
    validate,
)

PACK = "tests/jutlp_sample_docx_test_pack"


def rules(report):
    return {r["rule_id"]: r["status"] for r in report["results"]}


def results_by_rule(results):
    return {r["rule_id"]: r for r in results}


def _make_numbered(paragraph):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)


def test_dot_points_allowed_only_in_practitioner_notes(tmp_path):
    path = tmp_path / "input.docx"
    doc = Document()
    doc.add_paragraph("Practitioner Notes")
    _make_numbered(doc.add_paragraph("Allowed practitioner note."))
    doc.add_paragraph("Keywords")
    _make_numbered(doc.add_paragraph("Body list item."))
    doc.save(path)

    fails = [r for r in check_dot_points(str(path)) if r["status"] == "fail"]

    assert len(fails) == 1
    assert fails[0]["para_idx"] == 3


def test_required_literature_and_method_messages_include_expected_location():
    parsed = {
        "main_sections": [
            "Introduction",
            "Results",
            "Discussion",
            "Conclusion",
            "Acknowledgements",
            "References",
        ]
    }

    rd = results_by_rule(check_required_sections(parsed))

    assert rd["SEC002"]["status"] == "fail"
    assert rd["SEC002"]["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Literature"]
    assert rd["SEC003"]["status"] == "fail"
    assert rd["SEC003"]["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Method"]


def test_intro_literature_method_order_reports_specific_pairs():
    parsed = {
        "main_sections": [
            "Method",
            "Literature",
            "Introduction",
            "Results",
            "Discussion",
            "Conclusion",
            "Acknowledgements",
            "References",
        ]
    }

    result = check_section_order(parsed)[0]

    assert result["rule_id"] == "SEC009"
    assert result["status"] == "fail"
    assert "'Literature' should appear after 'Introduction'" in result["message"]
    assert "'Method' should appear after 'Literature'" in result["message"]


def test_combined_findings_discussion_reports_long_policy_comment():
    parsed = {
        "main_sections": [
            "Introduction",
            "Literature",
            "Method",
            "Findings and Discussion",
            "Conclusion",
            "Acknowledgements",
            "References",
        ]
    }

    from app.services.jutlp_validator import check_combined_results_discussion

    result = check_combined_results_discussion(parsed)[0]

    assert result["rule_id"] == "SPE001"
    assert result["status"] == "warn"
    assert CANONICAL_STRUCTURE["combined_results_discussion_query"] in result["message"]


class TestValidIdentified:
    @pytest.fixture(scope="class")
    def report(self):
        return validate(f"{PACK}/01_valid_identified.docx")

    def test_no_failures(self, report):
        assert report["fail"] == 0

    def test_all_required_sections_pass(self, report):
        rd = rules(report)
        for rule_id in ["SEC001", "SEC002", "SEC003", "SEC004", "SEC005", "SEC006", "SEC007", "SEC008"]:
            assert rd[rule_id] == "pass", f"{rule_id} should pass on valid doc"

    def test_section_order_pass(self, report):
        assert rules(report)["SEC009"] == "pass"

    def test_all_method_subsections_pass(self, report):
        rd = rules(report)
        for rule_id in ["MET001", "MET002", "MET003", "MET004", "MET005"]:
            assert rd[rule_id] == "pass", f"{rule_id} should pass on valid doc"

    def test_all_discussion_subsections_pass(self, report):
        rd = rules(report)
        for rule_id in ["DIS001", "DIS002", "DIS003"]:
            assert rd[rule_id] == "pass", f"{rule_id} should pass on valid doc"

    def test_front_page_passes(self, report):
        rd = rules(report)
        for rule_id in ["FP001", "FP002", "FP003", "FP004", "FP005", "FP006"]:
            assert rd[rule_id] == "pass", f"{rule_id} should pass on valid doc"

    def test_abstract_passes(self, report):
        rd = rules(report)
        assert rd["FP007"] == "pass"
        assert rd["FP008"] == "pass"

    def test_practitioner_notes_pass(self, report):
        assert rules(report)["FP009"] == "pass"

    def test_keywords_pass(self, report):
        assert rules(report)["FP010"] == "pass"

    def test_no_forbidden_styles(self, report):
        assert rules(report)["STY001"] == "pass"


class TestMissingMethodSubsection:
    @pytest.fixture(scope="class")
    def report(self):
        return validate(f"{PACK}/02_missing_method_subsection.docx")

    def test_exactly_one_failure(self, report):
        assert report["fail"] == 1

    def test_participants_missing(self, report):
        assert rules(report)["MET002"] == "fail"

    def test_other_method_subsections_still_pass(self, report):
        rd = rules(report)
        for rule_id in ["MET001", "MET003", "MET004", "MET005"]:
            assert rd[rule_id] == "pass", f"{rule_id} should still pass"

    def test_sections_unaffected(self, report):
        rd = rules(report)
        for rule_id in ["SEC001", "SEC002", "SEC003", "SEC004", "SEC005", "SEC006", "SEC007", "SEC008"]:
            assert rd[rule_id] == "pass"


class TestFrontPageIssues:
    @pytest.fixture(scope="class")
    def report(self):
        return validate(f"{PACK}/03_front_page_issues.docx")

    def test_exactly_two_failures(self, report):
        assert report["fail"] == 2

    def test_practitioner_notes_count_wrong(self, report):
        assert rules(report)["FP009"] == "fail"

    def test_too_many_keywords(self, report):
        assert rules(report)["FP010"] == "fail"

    def test_sections_unaffected(self, report):
        rd = rules(report)
        for rule_id in ["SEC001", "SEC002", "SEC003", "SEC004", "SEC005"]:
            assert rd[rule_id] == "pass"


class TestStructureAndEndmatterIssues:
    @pytest.fixture(scope="class")
    def report(self):
        return validate(f"{PACK}/04_structure_and_endmatter_issues.docx")

    def test_discussion_section_missing(self, report):
        # Fixture has "Results and Discussion" as a single heading. Under
        # this validator (post-merge), a combined section still fails
        # SEC005 — SPE001 also fires, but check_required_sections doesn't
        # downgrade the combined branch to a warn. If that policy changes
        # in future, update both branches of check_required_sections.
        assert rules(report)["SEC005"] == "fail"

    def test_guidance_notes_present(self, report):
        assert rules(report)["STY001"] == "fail"

    def test_discussion_subsections_all_fail(self, report):
        # check_discussion_subsections fails DIS001/002/003 when Discussion
        # is absent OR only present combined (e.g. "Results and Discussion").
        # The fixture has the combined form, so all three subsection rules
        # report fail.
        rd = rules(report)
        for rule_id in ["DIS001", "DIS002", "DIS003"]:
            assert rd[rule_id] == "fail", f"{rule_id} should fail — Discussion section combined"

    def test_other_sections_unaffected(self, report):
        rd = rules(report)
        for rule_id in ["SEC001", "SEC002", "SEC003", "SEC004", "SEC007", "SEC008"]:
            assert rd[rule_id] == "pass"


class TestValidDeidentified:
    @pytest.fixture(scope="class")
    def report(self):
        return validate(f"{PACK}/05_valid_deidentified.docx")

    def test_no_failures(self, report):
        assert report["fail"] == 0

    def test_all_required_sections_pass(self, report):
        rd = rules(report)
        for rule_id in ["SEC001", "SEC002", "SEC003", "SEC004", "SEC005", "SEC006", "SEC007", "SEC008"]:
            assert rd[rule_id] == "pass"

    def test_no_forbidden_styles(self, report):
        assert rules(report)["STY001"] == "pass"


class TestDeidentifiedWithAuthorLeak:
    @pytest.fixture(scope="class")
    def report(self):
        return validate(f"{PACK}/06_deidentified_with_author_leak.docx")

    def test_structural_checks_pass(self, report):
        rd = rules(report)
        for rule_id in ["SEC001", "SEC002", "SEC003", "SEC004", "SEC005", "SEC006", "SEC007", "SEC008"]:
            assert rd[rule_id] == "pass"

    def test_no_structural_failures(self, report):
        assert report["fail"] == 0


class TestSubsectionAliasMatch:
    """Direct tests of the alias matcher — no docx fixture needed."""

    def test_canonical_label_matches(self):
        assert subsection_alias_match("Research Design", ["Research Design"]) is True

    def test_alias_matches(self):
        assert subsection_alias_match(
            "Research Design", ["Setting of the Research"]
        ) is True

    def test_alias_match_is_case_insensitive(self):
        assert subsection_alias_match(
            "Measures", ["research instrument development and validation"]
        ) is True

    def test_alias_match_strips_trailing_punctuation(self):
        assert subsection_alias_match(
            "Research Design", ["Setting of the Research:"]
        ) is True

    def test_unrelated_heading_does_not_match(self):
        assert subsection_alias_match(
            "Participants", ["Demographic Profile"]
        ) is False

    def test_empty_found_subs_does_not_match(self):
        assert subsection_alias_match("Analysis", []) is False


class TestMethodSubsectionsWithAliases:
    def test_setting_of_the_research_satisfies_research_design(self):
        parsed = {
            "main_sections": ["Method"],
            "subsections": {
                "Method": [
                    "Setting of the Research",
                    "Participants",
                    "Research Instrument Development and Validation",
                    "Procedure",
                    "Analysis",
                ],
            },
        }
        rules_out = {r["rule_id"]: r["status"] for r in check_method_subsections(parsed)}
        assert rules_out["MET001"] == "pass"  # Research Design via alias
        assert rules_out["MET002"] == "pass"  # Participants canonical
        assert rules_out["MET003"] == "pass"  # Measures via alias
        assert rules_out["MET004"] == "pass"
        assert rules_out["MET005"] == "pass"

    def test_real_missing_subsection_still_fails(self):
        parsed = {
            "main_sections": ["Method"],
            "subsections": {
                "Method": [
                    "Setting of the Research",
                    "Participants",
                    "Measures",
                    # No Procedure or alias for it
                    "Analysis",
                ],
            },
        }
        rules_out = {r["rule_id"]: r["status"] for r in check_method_subsections(parsed)}
        assert rules_out["MET004"] == "fail"  # Procedure missing
        assert rules_out["MET001"] == "pass"
        assert rules_out["MET005"] == "pass"


class TestNonTemplateSubheadings:
    """Method/Discussion subheadings that aren't in the template's expected set
    are flagged with an advisory (MET900/DIS900) warn."""

    def test_method_extra_subheadings_flagged(self):
        parsed = {
            "main_sections": ["Method"],
            "subsections": {
                "Method": [
                    "Setting of the Research",   # alias → Research Design
                    "Research Approach",         # NOT canonical
                    "Participants",
                    "Research Instrument Development and Validation",  # alias → Measures
                    "Interviews",                # NOT canonical
                ],
            },
        }
        results = check_method_subsections(parsed)
        by_id = {r["rule_id"]: r for r in results}
        assert "MET900" in by_id
        assert by_id["MET900"]["status"] == "warn"
        # "Interviews" is genuinely non-canonical → flagged as extra.
        assert "Interviews" in by_id["MET900"]["message"]
        # Aliased headings are NOT reported as extras ("Setting of the
        # Research" and "Research Approach" both alias Research Design).
        assert "Setting of the Research" not in by_id["MET900"]["message"]
        assert "Research Approach" not in by_id["MET900"]["message"]

    def test_method_all_canonical_no_extra_result(self):
        parsed = {
            "main_sections": ["Method"],
            "subsections": {
                "Method": [
                    "Research Design", "Participants", "Measures",
                    "Procedure", "Analysis",
                ],
            },
        }
        ids = {r["rule_id"] for r in check_method_subsections(parsed)}
        assert "MET900" not in ids

    def test_discussion_extra_subheadings_flagged(self):
        parsed = {
            "main_sections": ["Discussion"],
            "subsections": {
                "Discussion": [
                    "Emerging but Developing Ethical Thinking",   # NOT canonical
                    "Practical Implications",
                ],
            },
        }
        by_id = {r["rule_id"]: r for r in check_discussion_subsections(parsed)}
        assert "DIS900" in by_id
        assert by_id["DIS900"]["status"] == "warn"
        assert "Emerging but Developing Ethical Thinking" in by_id["DIS900"]["message"]

    def test_discussion_all_canonical_no_extra_result(self):
        parsed = {
            "main_sections": ["Discussion"],
            "subsections": {
                "Discussion": [
                    "Practical Implications",
                    "Theoretical Implications",
                    "Limitations and Future Research",
                ],
            },
        }
        ids = {r["rule_id"] for r in check_discussion_subsections(parsed)}
        assert "DIS900" not in ids


class TestDiscussionSubsectionsWithAliases:
    def test_limitations_alias_passes_dis003(self):
        parsed = {
            "main_sections": ["Discussion"],
            "subsections": {
                "Discussion": [
                    "Limitations",  # alias for "Limitations and Future Research"
                ],
            },
        }
        rules_out = {r["rule_id"]: r["status"] for r in check_discussion_subsections(parsed)}
        assert rules_out["DIS001"] == "fail"  # Practical Implications still missing
        assert rules_out["DIS002"] == "fail"  # Theoretical Implications still missing
        assert rules_out["DIS003"] == "pass"  # satisfied by "Limitations" alias
