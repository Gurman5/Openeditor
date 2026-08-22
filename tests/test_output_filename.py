from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.services.output_filename import (
    _first_author_last_name,
    build_output_filename,
    build_output_filename_from_author_line,
)


def test_first_author_last_name_removes_affiliation_marker():
    name = _first_author_last_name(
        "Alex Chena, Priya Nairb, and Sam Perrya",
        ["a", "b"],
    )

    assert name == "Chen"


def test_output_filename_uses_first_author_and_increments(tmp_path):
    docx_path = "tests/jutlp_sample_docx_test_pack/01_valid_identified.docx"
    year = datetime.now().year

    first_name = build_output_filename(docx_path, str(tmp_path))
    first_path = tmp_path / first_name
    first_path.write_text("existing")

    second_name = build_output_filename(docx_path, str(tmp_path))

    assert first_name == f"Chen_JUTLP_{year}_CopyEdit1.docx"
    assert second_name == f"Chen_JUTLP_{year}_CopyEdit2.docx"


def test_output_filename_handles_plain_affiliation_markers(tmp_path):
    docx_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Article title")
    doc.add_paragraph("Aina Francha, Sam Writerb")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("Citation")
    doc.add_paragraph("Example citation (2025).")
    doc.save(docx_path)

    filename = build_output_filename(str(docx_path), str(tmp_path))

    assert filename == "Franch_JUTLP_2025_CopyEdit1.docx"


def test_output_filename_uses_author_line_before_affiliation_break(tmp_path):
    docx_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Article title")
    doc.add_paragraph("Professor Michael Sankey \nCharles Darwin University, Australia")
    doc.add_paragraph("Citation")
    doc.add_paragraph("Sankey, M. (2024). Example.")
    doc.save(docx_path)

    filename = build_output_filename(str(docx_path), str(tmp_path))

    assert filename == "Sankey_JUTLP_2024_CopyEdit1.docx"


def test_output_filename_ignores_parenthetical_degrees():
    assert _first_author_last_name("Imre Fekete (PhD), college associate professor", []) == "Fekete"


def test_output_filename_can_use_corrected_author_line(tmp_path):
    docx_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Alignment title")
    doc.add_paragraph("More title text")
    doc.add_paragraph("Citation")
    doc.add_paragraph("Example citation (2025).")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.save(docx_path)

    filename = build_output_filename_from_author_line(
        str(docx_path),
        str(tmp_path),
        "Alex Chen^a, Priya Nair^b, and Sam Perry^a",
    )

    assert filename == "Chen_JUTLP_2025_CopyEdit1.docx"


def test_corrected_author_line_uses_doc_affiliation_markers(tmp_path):
    docx_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Article title")
    doc.add_paragraph("Placeholder authors")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("Citation")
    doc.add_paragraph("Example citation (2026).")
    doc.save(docx_path)

    filename = build_output_filename_from_author_line(
        str(docx_path),
        str(tmp_path),
        "Patsie Pollya, Sam Writerb",
    )

    assert filename == "Polly_JUTLP_2026_CopyEdit1.docx"


def test_corrected_author_line_uses_doc_without_superscript_markers(tmp_path):
    docx_path = tmp_path / "input.docx"

    doc = Document()
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Article title")
    author = doc.add_paragraph()
    author.style = "Authors"
    author.add_run("Prof. Patsie Polly")
    marker = author.add_run("a")
    marker.font.superscript = True
    doc.add_paragraph("Citation")
    doc.add_paragraph("Example citation (2026).")
    doc.save(docx_path)

    filename = build_output_filename_from_author_line(
        str(docx_path),
        str(tmp_path),
        "Prof. Patsie Pollya",
    )

    assert filename == "Polly_JUTLP_2026_CopyEdit1.docx"
