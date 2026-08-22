import re
import zipfile
from unittest.mock import patch

from docx import Document
from lxml import etree

from app.pipelines.feedback_gen_pipeline import _renumber_final_author_queries
from app.services.body_llm_edits import _llm_edits_for_paragraph
from app.services.editorial_review_comments import _group_similar_comments
from app.services.language_corrections import apply_au_spelling_corrections
from app.services.output_generation_samfix import (
    _apply_intra_paragraph_tracked_replace,
    _make_comment_element,
    _renumber_author_queries_by_anchor_order,
    _write_single_comment_docx,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"
TIMESTAMP_RE = re.compile(
    r"\b\d{4}-\d{2}-\d{2}(?:[ T]\d{2}:\d{2}(?::\d{2})?)?(?:Z|\s*UTC(?:\+?0)?)?\b",
    re.IGNORECASE,
)


def _docx_xml(path, part):
    with zipfile.ZipFile(path, "r") as z:
        return etree.fromstring(z.read(part))


def _comment_text(comment):
    return " ".join("".join(t.text or "" for t in comment.iter(f"{WQ}t")).split())


def _comment_texts(path):
    root = _docx_xml(path, "word/comments.xml")
    return [_comment_text(comment) for comment in root.findall(f"{WQ}comment")]


def _tracked_deletions(path):
    root = _docx_xml(path, "word/document.xml")
    return [el.text or "" for el in root.findall(f".//{WQ}delText")]


def _tracked_insertions(path):
    root = _docx_xml(path, "word/document.xml")
    return [
        "".join(t.text or "" for t in ins.findall(f".//{WQ}t"))
        for ins in root.findall(f".//{WQ}ins")
    ]


def _paragraph_with_runs(run_texts):
    paragraph = etree.Element(f"{WQ}p", nsmap={"w": W})
    for run_text in run_texts:
        run = etree.SubElement(paragraph, f"{WQ}r")
        text = etree.SubElement(run, f"{WQ}t")
        text.text = run_text
    return paragraph


def _accepted_visible_text(paragraph):
    parts = []
    for text in paragraph.iter(f"{WQ}t"):
        parent = text.getparent()
        while parent is not None:
            if parent.tag == f"{WQ}del":
                break
            parent = parent.getparent()
        else:
            parts.append(text.text or "")
    return "".join(parts)


def _assert_no_visible_timestamp(text):
    assert TIMESTAMP_RE.search(text) is None
    assert "UTC" not in text.upper()
    assert "UTC+0" not in text.upper()


def test_tracked_replace_edits_standalone_text_not_substring_inside_word():
    paragraph = _paragraph_with_runs([
        "We reutilize tools before we utilize shared resources."
    ])

    changed = _apply_intra_paragraph_tracked_replace(
        paragraph,
        "utilize",
        "utilise",
        42,
    )

    assert changed is True
    assert [el.text for el in paragraph.findall(f".//{WQ}delText")] == ["utilize"]
    assert [
        "".join(t.text or "" for t in ins.findall(f".//{WQ}t"))
        for ins in paragraph.findall(f".//{WQ}ins")
    ] == ["utilise"]
    assert _accepted_visible_text(paragraph) == (
        "We reutilize tools before we utilise shared resources."
    )


@patch("app.services.body_llm_edits.call_llm_json")
def test_copyedit_rejects_vague_reason_without_change_type_or_exact_text(mock_call):
    mock_call.return_value = {
        "content": {
            "edits": [
                {
                    "find": "utilize",
                    "replace": "use",
                    "reason": "Improves clarity.",
                }
            ]
        }
    }

    accepted = _llm_edits_for_paragraph(
        "Students utilize shared tools during the workshop."
    )

    assert accepted == []


@patch("app.services.body_llm_edits.call_llm_json")
def test_copyedit_accepts_specific_reason_with_type_and_exact_change(mock_call):
    reason = 'Style: replaced "utilize" with "use" for concise wording.'
    mock_call.return_value = {
        "content": {
            "edits": [
                {
                    "find": "utilize",
                    "replace": "use",
                    "reason": reason,
                }
            ]
        }
    }

    accepted = _llm_edits_for_paragraph(
        "Students utilize shared tools during the workshop."
    )

    assert accepted == [{"find": "utilize", "replace": "use", "reason": reason}]


def test_visible_copyeditor_comment_text_excludes_timestamp_metadata(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("This paragraph needs an editor comment.")
    doc.save(source)

    _write_single_comment_docx(
        str(source),
        str(output),
        0,
        'Formatting: changed paragraph style from "Body Text" to "Normal".',
    )

    comments_root = _docx_xml(output, "word/comments.xml")
    comment = comments_root.find(f"{WQ}comment")
    assert comment is not None
    assert comment.get(f"{WQ}date")

    visible_text = _comment_text(comment)
    assert "Formatting:" in visible_text
    assert comment.get(f"{WQ}date") not in visible_text
    _assert_no_visible_timestamp(visible_text)


def test_au_spelling_comment_does_not_leak_timestamp_or_utc_text(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("We organize the color-coded center activities.")
    doc.save(source)

    apply_au_spelling_corrections(str(source), str(output))

    comments_root = _docx_xml(output, "word/comments.xml")
    comments = comments_root.findall(f"{WQ}comment")
    assert comments
    for comment in comments:
        assert comment.get(f"{WQ}date")
        visible_text = _comment_text(comment)
        assert comment.get(f"{WQ}date") not in visible_text
        _assert_no_visible_timestamp(visible_text)


def test_repeated_au_spelling_fixes_create_one_summary_comment_with_word_pairs(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph(
        "We organize seminars, realize benefits, and organize follow-up activities."
    )
    doc.save(source)

    _, corrections = apply_au_spelling_corrections(str(source), str(output))

    assert corrections == [
        {"original": "organize", "replacement": "organise"},
        {"original": "realize", "replacement": "realise"},
        {"original": "organize", "replacement": "organise"},
    ]
    assert _tracked_deletions(output) == ["organize", "realize", "organize"]
    assert _tracked_insertions(output) == ["organise", "realise", "organise"]

    comments_root = _docx_xml(output, "word/comments.xml")
    comments = comments_root.findall(f"{WQ}comment")
    document_root = _docx_xml(output, "word/document.xml")
    assert len(comments) == 1
    assert len(document_root.findall(f".//{WQ}commentReference")) == 1

    summary_text = _comment_text(comments[0])
    assert "Australian English" in summary_text
    assert "organize \u2192 organise" in summary_text
    assert "realize \u2192 realise" in summary_text


def test_au_spelling_summary_keeps_other_issue_type_comment_separate(tmp_path):
    source = tmp_path / "source.docx"
    with_existing_comment = tmp_path / "with_existing_comment.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("We organize seminars and realize shared benefits.")
    doc.save(source)
    _write_single_comment_docx(
        str(source),
        str(with_existing_comment),
        0,
        'Grammar: changed "was" to "were" for subject-verb agreement.',
    )

    apply_au_spelling_corrections(str(with_existing_comment), str(output))

    comments = _comment_texts(output)
    grammar_comments = [text for text in comments if "Grammar:" in text]
    spelling_comments = [text for text in comments if "Australian English" in text]
    assert len(comments) == 2
    assert len(grammar_comments) == 1
    assert len(spelling_comments) == 1
    assert "subject-verb agreement" in grammar_comments[0]
    assert "organize" not in grammar_comments[0]
    assert "realize" not in grammar_comments[0]
    assert "organize \u2192 organise" in spelling_comments[0]
    assert "realize \u2192 realise" in spelling_comments[0]
    assert "Grammar:" not in spelling_comments[0]


def test_final_author_query_renumbering_includes_late_spelling_comments(tmp_path):
    source = tmp_path / "source.docx"
    with_later_comment = tmp_path / "with_later_comment.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("We organize the abstract analysis.")
    doc.add_paragraph("Introduction")
    doc.add_paragraph("Later body paragraph.")
    doc.save(source)

    _write_single_comment_docx(
        str(source),
        str(with_later_comment),
        3,
        "Later body comment.",
    )
    apply_au_spelling_corrections(str(with_later_comment), str(output))

    assert _comment_texts(output) == [
        "Author Query 1. Later body comment.",
        (
            'Author Query 2. US spelling "organize" was changed to '
            'Australian English "organise". No later repeats of this '
            "spelling issue were found."
        ),
    ]

    changed = _renumber_final_author_queries(str(output))

    assert changed is True
    assert _comment_texts(output) == [
        (
            'Author Query 1. US spelling "organize" was changed to '
            'Australian English "organise". No later repeats of this '
            "spelling issue were found."
        ),
        "Author Query 2. Later body comment.",
    ]


def test_editorial_review_groups_similar_acronym_comments():
    comments = _group_similar_comments([
        {
            "anchor_pos": 5,
            "message": "[High] Abstract\nPlease check acronym use for accessibility.",
        },
        {
            "anchor_pos": 2,
            "message": "[High] Introduction\nPlease check acronym use for accessibility.",
        },
        {
            "anchor_pos": 8,
            "message": "[Medium] Discussion\nClarify the implication.",
        },
    ])

    assert comments == [
        {
            "anchor_pos": 2,
            "message": (
                "Please check acronym use across the manuscript. 2 similar acronym "
                "notes were grouped here to keep the comment pane manageable. "
                "Expand acronyms on first use and use the acronym form thereafter."
            ),
        },
        {
            "anchor_pos": 8,
            "message": "[Medium] Discussion\nClarify the implication.",
        },
    ]


def test_author_query_renumbering_is_dense_with_phantom_and_orphan_comments(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(source)

    _write_single_comment_docx(str(source), str(output), 1, "Second anchor.")

    with zipfile.ZipFile(output, "r") as zin:
        doc_root = etree.fromstring(zin.read("word/document.xml"))
        comments_root = etree.fromstring(zin.read("word/comments.xml"))
        existing = comments_root.find(f"{WQ}comment")
        existing.set(f"{WQ}id", "20")
        for marker in doc_root.iter():
            if marker.get(f"{WQ}id") == "1":
                marker.set(f"{WQ}id", "20")

        body = doc_root.find(f"{WQ}body")
        first_para = body.find(f"{WQ}p")
        phantom = etree.Element(f"{WQ}commentRangeStart", nsmap={"w": W})
        phantom.set(f"{WQ}id", "99")
        first_para.insert(0, phantom)

        comments_root.append(_make_comment_element(50, "Orphan comment."))
        new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        new_comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        tmp = tmp_path / "mutated.docx"
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_doc_xml)
                elif item.filename == "word/comments.xml":
                    zout.writestr(item, new_comments_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))
    tmp.replace(output)

    assert _renumber_author_queries_by_anchor_order(str(output)) is True

    assert _comment_texts(output) == [
        "Author Query 1. Second anchor.",
        "Author Query 2. Orphan comment.",
    ]
