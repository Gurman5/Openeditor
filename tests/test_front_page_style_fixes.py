import zipfile
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

from app.domain.canonical_jultp_template import CANONICAL_STRUCTURE
from app.services.document_analysis_services import load_paragraphs
from app.services.language_corrections import WQ
from app.services.output_generation_samfix import (
    AFFILIATIONS_REQUIRED_STYLE_ID,
    AUTHORS_REQUIRED_STYLE_ID,
    CITATION_FOOTER_STYLE_ID,
    FRONT_PAGE_BANNER_COMMENT,
    FRONT_PAGE_TEXT_BOX_COMMENT,
    KEYWORDS_REQUIRED_HEADING_STYLE_ID,
    KEYWORDS_REQUIRED_TEXT_STYLE_ID,
    _apply_abstract_plan,
    _apply_author_plan,
    _apply_citation_plan,
    _apply_document_body_plan,
    _apply_front_page_asset_plan,
    _apply_heading_1_text_normalization,
    _apply_heading_2_title_case,
    _apply_heading_center_alignment,
    _apply_heading_keep_next,
    _apply_intro_page_break,
    _apply_keywords_plan,
    _apply_missing_keywords_stub,
    _apply_missing_practitioner_stub,
    _apply_normal_style_fix,
    _apply_title_plan,
    _apply_tracked_style_fixes,
    _author_naming_pattern_valid,
    _authors_with_multiple_affiliation_markers,
    _format_author_query_text,
    _is_deidentified_manuscript,
    _remove_line_numbers,
    _renumber_author_queries_by_anchor_order,
    _to_title_case_title,
    _write_single_comment_docx,
    abstractFormatCheck,
    abstractFound,
    authorFound,
    build_author_check_plan,
    build_edited_document,
    build_front_page_asset_check_plan,
    citationFormatCheck,
    citationFound,
    documentBodyFormatCheck,
    documentBodyFound,
    keywordsFormatCheck,
    keywordsFound,
    titleFound,
)


def test_title_case_fix_handles_colon_and_all_caps_titles():
    assert _to_title_case_title("THE IMPACT OF AI: AN APA STUDY IN STEM") == (
        "The Impact of AI: An APA Study in STEM"
    )
    assert _to_title_case_title("USING GENAI IN FEEDBACK") == "Using GenAI in Feedback"


def test_title_case_preserves_acronyms_in_hyphenated_compounds():
    """Regression test for the "AI-Mediated" → "Ai-mediated" bug.

    Hyphenated compounds must be split per segment so the embedded acronym
    keeps its canonical case AND the subsequent segment stays capitalised.
    """
    assert _to_title_case_title(
        "Investigating Students' Perceptions of Ethical Principles in "
        "Translation Teaching in the AI-Mediated Era"
    ) == (
        "Investigating Students' Perceptions of Ethical Principles in "
        "Translation Teaching in the AI-Mediated Era"
    )


def test_title_case_restores_lowercase_known_acronyms():
    """Lowercase known-acronym tokens are restored to canonical casing —
    so an author who typed the title in sentence case still gets
    ``AI`` / ``STEM`` instead of ``Ai`` / ``Stem``."""
    assert _to_title_case_title("investigating ai use in stem") == (
        "Investigating AI Use in STEM"
    )
    assert _to_title_case_title("ai-mediated era") == "AI-Mediated Era"


def test_title_case_keeps_small_words_lowercase_inside_compounds():
    """Compounds like ``state-of-the-art`` must retain the small-word rule
    for their interior segments — the second-through-last segments aren't
    bumped to capitalised just because they sit inside a hyphenated word."""
    assert _to_title_case_title("A State-of-the-Art Solution") == (
        "A State-of-the-Art Solution"
    )
    assert _to_title_case_title("a state-of-the-art solution") == (
        "A State-of-the-Art Solution"
    )


def test_title_case_fix_writes_final_title_text(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    old_title = "using genai in feedback"
    new_title = "Using GenAI in Feedback"
    doc = Document()
    title = doc.add_paragraph()
    for text in ["using ", "g", "enai in ", "f", "eedback"]:
        title.add_run(text)
    doc.save(input_path)

    plan = {
        "action": "replace_title_and_comment",
        "anchor_pos": 0,
        "message": "",
        "needs_title_case_fix": True,
        "style_ok": True,
        "title_found": True,
        "title_text": old_title,
        "corrected_title": new_title,
    }
    assert _apply_title_plan(str(input_path), str(output_path), plan) is True

    assert Document(output_path).paragraphs[0].text == new_title
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    assert list(root.iter(f"{WQ}ins")) == []
    assert list(root.iter(f"{WQ}del")) == []


def _paragraph_texts_by_style(docx_path, style_name):
    doc = Document(docx_path)
    return [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style is not None and paragraph.style.name == style_name
    ]


def test_samfix_author_query_comment_text_is_numbered():
    assert _format_author_query_text(2, "Check this manuscript detail.") == (
        "Author Query 2. Check this manuscript detail."
    )


def test_author_naming_pattern_rejects_trailing_title():
    assert _author_naming_pattern_valid(
        "Samuel Perry Mr^a, and John Smith^b"
    ) is False
    assert _author_naming_pattern_valid(
        "Samuel Perry Mr.^a, and John Smith^b"
    ) is False


@patch("app.services.output_generation_samfix.call_llm_json")
def test_author_plan_flags_trailing_title_as_incorrect(mock_call, tmp_path):
    mock_call.side_effect = [
        {
            "content": {
                "is_author_naming_correct": False,
                "corrected_authors_line": "Mr. Samuel Perry^a, and John Smith^b",
                "reason": "Title appears after the first author's name.",
            },
        },
        Exception("skip affiliation check"),
    ]
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title")
    doc.add_paragraph("")
    doc.add_paragraph("Samuel Perry Mr^a, and John Smith^b", style="Authors")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "replace_in_place"
    assert plan["corrected_authors_line"] == "Mr. Samuel Perry^a, and John Smith^b"


@patch("app.services.output_generation_samfix.call_llm_json")
def test_author_plan_comments_on_trailing_title_fix(mock_call, tmp_path):
    mock_call.side_effect = [
        {
            "content": {
                "is_author_naming_correct": False,
                "corrected_authors_line": "Mr. Samuel Perry^a, and John Smith^b",
                "reason": "Title appears after the first author's name.",
            },
        },
        Exception("skip affiliation check"),
    ]
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title")
    doc.add_paragraph("")
    doc.add_paragraph("Samuel Perry Mr^a, and John Smith^b", style="Authors")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))
    changed = _apply_author_plan(str(input_path), str(output_path), plan)

    assert changed is True
    comments = _comment_texts(output_path)
    assert len(comments) == 1
    assert "Authors paragraph was changed because:" in comments[0]
    assert "Title appears after the first author's name." in comments[0]


def _comment_texts(docx_path):
    with zipfile.ZipFile(docx_path, "r") as z:
        comments_xml = z.read("word/comments.xml")
    root = etree.fromstring(comments_xml)
    texts = []
    for comment in root.findall(f"{WQ}comment"):
        texts.append("".join(t.text or "" for t in comment.iter(f"{WQ}t")))
    return texts


def _doc_has_textbox(docx_path):
    with zipfile.ZipFile(docx_path, "r") as z:
        doc_xml = z.read("word/document.xml")
    root = etree.fromstring(doc_xml)
    for element in root.iter():
        if etree.QName(element).localname in {"txbxContent", "textbox"}:
            return True
    return False


def _doc_has_header_image(docx_path):
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if not (name.startswith("word/header") and name.endswith(".xml")):
                continue
            root = etree.fromstring(z.read(name))
            if any(etree.QName(el).localname in {"blip", "imagedata"} for el in root.iter()):
                return True
    return False


def _first_section_header_has_image(docx_path):
    rel_q = "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
    with zipfile.ZipFile(docx_path, "r") as z:
        doc_root = etree.fromstring(z.read("word/document.xml"))
        rels_root = etree.fromstring(z.read("word/_rels/document.xml.rels"))
        sect_pr = next(doc_root.iter(f"{WQ}sectPr"), None)
        header_type = "first" if sect_pr.find(f"{WQ}titlePg") is not None else "default"
        for header_ref in sect_pr.findall(f"{WQ}headerReference"):
            if header_ref.get(f"{WQ}type") != header_type:
                continue
            rel_id = header_ref.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            target = next((r.get("Target", "") for r in rels_root.findall(rel_q) if r.get("Id") == rel_id), "")
            header_path = target if target.startswith("word/") else "word/" + target
            header_root = etree.fromstring(z.read(header_path))
            return any(etree.QName(el).localname in {"blip", "imagedata"} for el in header_root.iter())
    return False


def test_remove_line_numbers_strips_section_setting(tmp_path):
    input_path = tmp_path / "input.docx"
    tmp_docx = tmp_path / "tmp.docx"

    doc = Document()
    doc.add_paragraph("Body text")
    doc.save(input_path)

    with zipfile.ZipFile(input_path, "r") as zin:
        root = etree.fromstring(zin.read("word/document.xml"))
        sect_pr = root.find(f".//{WQ}sectPr")
        etree.SubElement(sect_pr, f"{WQ}lnNumType").set(f"{WQ}countBy", "1")
        new_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, new_xml if item.filename == "word/document.xml" else zin.read(item.filename))
    tmp_docx.replace(input_path)

    assert _remove_line_numbers(str(input_path)) is True
    with zipfile.ZipFile(input_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    assert root.find(f".//{WQ}lnNumType") is None


def test_missing_authors_inserts_default_placeholders_without_blank_line(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))
    assert plan["action"] == "insert_default_authors"
    assert _apply_author_plan(str(input_path), str(output_path), plan) is True

    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    paragraphs = root.findall(f".//{WQ}body/{WQ}p")
    texts = ["".join(t.text or "" for t in p.iter(f"{WQ}t")) for p in paragraphs]
    styles = [p.find(f"{WQ}pPr/{WQ}pStyle") for p in paragraphs]

    assert texts[:4] == [
        "Example article title",
        "First Authorᵃ, Second Authorᵃ, and Third Authorᵇ",
        "a Affiliation 1; b Affiliation 2",
        "Abstract",
    ]
    assert styles[1].get(f"{WQ}val") == AUTHORS_REQUIRED_STYLE_ID
    assert styles[2].get(f"{WQ}val") == AFFILIATIONS_REQUIRED_STYLE_ID
    assert paragraphs[1].find(f".//{WQ}commentReference") is not None
    assert paragraphs[2].find(f".//{WQ}commentReference") is not None
    assert load_paragraphs(str(output_path))[1].text == texts[1]
    assert abstractFound(str(output_path))["affiliations_found"] is True
    comments = _comment_texts(output_path)
    assert any("Replace these placeholder author names" in text for text in comments)
    assert any("Replace these placeholder affiliations" in text for text in comments)


def test_missing_authors_before_unstyled_introduction_get_default_placeholders(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("Introduction")
    doc.add_paragraph(
        "This body paragraph has citations (Smith, 2020; Jones, 2021), but it is not an authors line."
    )
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "insert_default_authors"
    assert plan["default_affiliations"] == ("a Affiliation 1; b Affiliation 2",)


def test_proposal_title_and_author_labels_are_detected(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Special Issue Article Submission for the call")
    doc.add_paragraph("Proposed article title:")
    doc.add_paragraph(
        "Between promise and practice: Bridging ethical AI literacy gaps across students, educators, and policy"
    )
    doc.add_paragraph("Author:")
    doc.add_paragraph("Imre Fekete (PhD), college associate professor of Budapest University")
    doc.save(input_path)

    title_state = titleFound(str(input_path))
    author_state = authorFound(str(input_path))
    abstract_plan = abstractFormatCheck(str(input_path), abstractFound(str(input_path)))

    assert title_state["anchor_pos"] == 2
    assert author_state["anchor_pos"] == 4
    assert abstract_plan["insert_default_affiliations"] is True
    assert abstract_plan["default_affiliations_insert_pos"] == 5


def test_missing_authors_removes_blank_spacer_before_abstract(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))
    assert _apply_author_plan(str(input_path), str(output_path), plan) is True

    assert [p.text for p in load_paragraphs(str(output_path))][:4] == [
        "Example article title",
        "First Authorᵃ, Second Authorᵃ, and Third Authorᵇ",
        "a Affiliation 1; b Affiliation 2",
        "Abstract",
    ]


def test_existing_default_authors_are_not_moved_on_second_pass(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Author Affiliations", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title", style="Article Title")
    doc.add_paragraph("First Authorᵃ, Second Authorᵃ, and Third Authorᵇ", style="Authors")
    doc.add_paragraph("a Affiliation 1; b Affiliation 2", style="Author Affiliations")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "none"


def test_existing_valid_authors_directly_under_title_are_not_moved(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Author Affiliations", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title", style="Article Title")
    doc.add_paragraph("Samuel Perry^a, and John Smith^b", style="Authors")
    doc.add_paragraph("a RMIT University; b RMIT University", style="Author Affiliations")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "none"


def test_missing_abstract_plan_still_schedules_practitioner_stub(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Author Affiliations", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Heading Front Page", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title", style="Article Title")
    doc.add_paragraph("First Authorᵃ, Second Authorᵃ, and Third Authorᵇ", style="Authors")
    doc.add_paragraph("a Affiliation 1; b Affiliation 2", style="Author Affiliations")
    doc.add_paragraph("Keywords", style="Heading Front Page")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(input_path)

    plan = abstractFormatCheck(str(input_path), abstractFound(str(input_path)))

    assert plan["reason"] == "Abstract heading not found"
    assert plan["insert_practitioner_stub"] is True
    assert plan["ensure_practitioner_section"] is True
    assert plan["practitioner_insert_after_pos"] == 2


def test_missing_abstract_practitioner_stub_inserts_before_existing_keywords(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Author Affiliations", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Heading Front Page", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title", style="Article Title")
    doc.add_paragraph("First Authorᵃ, Second Authorᵃ, and Third Authorᵇ", style="Authors")
    doc.add_paragraph("a Affiliation 1; b Affiliation 2", style="Author Affiliations")
    doc.add_paragraph("Keywords", style="Heading Front Page")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(input_path)

    plan = abstractFormatCheck(str(input_path), abstractFound(str(input_path)))
    changed = _apply_missing_practitioner_stub(str(input_path), str(output_path), plan)

    assert changed is True
    texts = [p.text.strip() for p in load_paragraphs(str(output_path))[:6]]
    assert texts == [
        "Example article title",
        "First Authorᵃ, Second Authorᵃ, and Third Authorᵇ",
        "a Affiliation 1; b Affiliation 2",
        "Practitioner Notes",
        "",
        "Keywords",
    ]
    assert any("Practitioner Notes section not found" in text for text in _comment_texts(output_path))


def test_missing_abstract_practitioner_stub_uses_inserted_affiliations_anchor(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Heading Front Page", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title", style="Article Title")
    doc.add_paragraph("First Authorᵃ, Second Authorᵃ, and Third Authorᵇ", style="Authors")
    doc.add_paragraph("Keywords", style="Heading Front Page")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(input_path)

    plan = abstractFormatCheck(str(input_path), abstractFound(str(input_path)))
    assert plan["insert_default_affiliations"] is True
    assert _apply_abstract_plan(str(input_path), str(output_path), plan) is True
    assert _apply_missing_practitioner_stub(str(output_path), str(output_path), plan) is True

    texts = [p.text.strip() for p in load_paragraphs(str(output_path))[:6]]
    assert texts == [
        "Example article title",
        "First Authorᵃ, Second Authorᵃ, and Third Authorᵇ",
        "a Affiliation 1; b Affiliation 2",
        "Practitioner Notes",
        "",
        "Keywords",
    ]


def test_missing_authors_reuses_existing_affiliations(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("b University X, Australia")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))
    assert plan["default_affiliations"] == ()
    assert _apply_author_plan(str(input_path), str(output_path), plan) is True

    texts = [p.text for p in load_paragraphs(str(output_path))]
    assert texts[:4] == [
        "Example article title",
        "First Authorᵃ, Second Authorᵃ, and Third Authorᵇ",
        "a RMIT University, Australia",
        "b University X, Australia",
    ]
    assert "a Affiliation 1" not in texts
    assert "b Affiliation 2" not in texts


def test_missing_authors_adds_second_default_affiliation_when_only_one_exists(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    affiliation = doc.add_paragraph("a RMIT University, Australia")
    affiliation.add_run().add_break()
    affiliation.add_run("*corresponding author")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))
    assert plan["default_affiliations"] == ()
    assert plan["append_affiliation_text"] == "b Affiliation 2"
    assert _apply_author_plan(str(input_path), str(output_path), plan) is True

    assert [p.text for p in load_paragraphs(str(output_path))][:4] == [
        "Example article title",
        "First Authorᵃ, Second Authorᵃ, and Third Authorᵇ",
        "a RMIT University, Australia; b Affiliation 2\n*corresponding author",
        "Abstract",
    ]
    assert any("Replace these placeholder affiliations" in text for text in _comment_texts(output_path))


def test_existing_authors_missing_affiliations_get_default_affiliation_paragraph(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("First Author^a, Second Author^a, and Third Author^b")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    state = abstractFound(str(input_path))
    plan = abstractFormatCheck(str(input_path), state)

    assert plan["insert_default_affiliations"] is True
    assert _apply_abstract_plan(str(input_path), str(output_path), plan) is True

    paragraphs = load_paragraphs(str(output_path))
    assert paragraphs[2].text == "a Affiliation 1; b Affiliation 2"
    assert paragraphs[2].style == "Author Affiliations"
    assert any("Replace these placeholder affiliations" in text for text in _comment_texts(output_path))

    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    output_paragraphs = root.findall(f".//{WQ}body/{WQ}p")
    p_style = output_paragraphs[2].find(f"{WQ}pPr/{WQ}pStyle")
    assert p_style.get(f"{WQ}val") == AFFILIATIONS_REQUIRED_STYLE_ID
    assert output_paragraphs[2].find(f".//{WQ}commentReference") is not None


def test_author_naming_rejects_multiple_affiliation_markers():
    assert _author_naming_pattern_valid(
        "First Author^a, Second Author^a, and Third Author^b"
    ) is True
    assert _author_naming_pattern_valid(
        "First Author^a^b, Second Author^a, and Third Author^b"
    ) is False
    assert _authors_with_multiple_affiliation_markers(
        "First Authorᵃᵇ, Second Authorᵃ, and Third Authorᵇ"
    ) == ["First Authorᵃᵇ"]


def test_author_with_multiple_affiliations_gets_query_not_rewrite(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("First Author^a^b, Second Author^a, and Third Author^b")
    doc.add_paragraph("a RMIT University, Australia; b University X, Australia")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "add_authors_naming_comment"
    assert "one primary affiliation" in plan["message"]
    assert plan["corrected_authors_line"] == (
        "First Author^a^b, Second Author^a, and Third Author^b"
    )
    assert _apply_author_plan(str(input_path), str(output_path), plan) is True
    assert any("one primary affiliation" in text for text in _comment_texts(output_path))

    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    assert root.findall(f".//{WQ}body/{WQ}p")[1].find(f".//{WQ}commentReference") is not None


def test_author_with_formatted_superscript_multiple_affiliations_gets_query(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    author = doc.add_paragraph()
    author.add_run("First Author")
    marker = author.add_run("ab")
    marker.font.superscript = True
    author.add_run(", Second Author")
    marker = author.add_run("a")
    marker.font.superscript = True
    author.add_run(", and Third Author")
    marker = author.add_run("b")
    marker.font.superscript = True
    doc.add_paragraph("a RMIT University, Australia; b University X, Australia")
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "add_authors_naming_comment"
    assert "First Author^a^b" in plan["message"]


@patch("app.services.output_generation_samfix.call_llm_json")
def test_author_naming_failure_comment_is_actionable_not_internal(mock_call, tmp_path):
    mock_call.side_effect = [
        {
            "content": {
                "is_author_naming_correct": False,
                "corrected_authors_line": "Between Promise and Practice: Bridging Ethical AI Literacy",
                "reason": "still not valid",
            },
        },
        Exception("skip affiliation check"),
    ]
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.styles.add_style("Authors", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Example article title")
    doc.add_paragraph(
        "Between Promise and Practice: Bridging Ethical AI Literacy",
        style="Authors",
    )
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    plan = build_author_check_plan(str(input_path))

    assert plan["action"] == "add_authors_naming_comment"
    assert "Check the author line formatting" in plan["message"]
    assert "First Author^a, Second Author^a, and Third Author^b" in plan["message"]
    assert "LLM" not in plan["message"]
    assert "still incorrect after LLM rewrite" not in plan["message"]


def test_author_queries_are_renumbered_by_document_anchor_order(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    doc.save(input_path)

    _write_single_comment_docx(str(input_path), str(output_path), 2, "Later comment.")
    _write_single_comment_docx(str(output_path), str(output_path), 0, "Earlier comment.")

    assert _comment_texts(output_path) == [
        "Author Query 1. Later comment.",
        "Author Query 2. Earlier comment.",
    ]

    changed = _renumber_author_queries_by_anchor_order(str(output_path))

    assert changed is True
    assert _comment_texts(output_path) == [
        "Author Query 1. Earlier comment.",
        "Author Query 2. Later comment.",
    ]


def test_front_page_asset_plan_comments_when_banner_and_text_box_are_missing(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Abstract body.")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text.")
    doc.save(input_path)

    plan = build_front_page_asset_check_plan(str(input_path))

    assert plan["action"] == "add_front_page_asset_comments"
    assert plan["comments"] == [
        {"anchor_pos": 0, "message": FRONT_PAGE_BANNER_COMMENT},
        {"anchor_pos": 0, "message": FRONT_PAGE_TEXT_BOX_COMMENT},
    ]


def test_front_page_asset_plan_inserts_banner_and_textbox(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Abstract body.")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text.")
    doc.save(input_path)

    plan = build_front_page_asset_check_plan(str(input_path))
    applied = _apply_front_page_asset_plan(str(input_path), str(output_path), plan)

    assert applied is True
    with zipfile.ZipFile(output_path, "r") as z:
        assert "word/comments.xml" not in z.namelist()
    assert _doc_has_header_image(output_path) is True
    assert _doc_has_textbox(output_path) is True
    assert build_front_page_asset_check_plan(str(output_path))["action"] == "none"


def test_front_page_asset_plan_inserts_banner_in_first_section_header(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Abstract body.")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text.")
    doc.save(input_path)

    plan = build_front_page_asset_check_plan(str(input_path))
    applied = _apply_front_page_asset_plan(str(input_path), str(output_path), plan)

    assert applied is True
    assert _first_section_header_has_image(output_path) is True
    assert build_front_page_asset_check_plan(str(output_path))["action"] == "none"


def test_heading_1_text_normalization_changes_known_aliases_without_comments(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature Review").style = "Heading 1"
    doc.add_paragraph("Review of Literature").style = "Heading 1"
    doc.add_paragraph("Methodology").style = "Heading 1"
    doc.add_paragraph("Findings").style = "Heading 1"
    doc.add_paragraph("Conclusions").style = "Heading 1"
    doc.add_paragraph("Concluding Remarks").style = "Heading 1"
    doc.save(input_path)

    changed = _apply_heading_1_text_normalization(str(input_path), str(output_path))

    template_heading_names = CANONICAL_STRUCTURE["main_sections"]
    assert "Literature" in template_heading_names
    assert "Method" in template_heading_names
    assert changed is True
    assert _paragraph_texts_by_style(output_path, "Heading 1") == [
        "Introduction",
        "Literature",
        "Literature",
        "Method",
        "Results",
        "Conclusion",
        "Conclusion",
    ]

    with zipfile.ZipFile(output_path, "r") as z:
        assert "word/comments.xml" not in z.namelist()


def test_combined_results_discussion_heading_is_not_silently_normalized(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Results and Discussion").style = "Heading 1"
    doc.save(input_path)

    changed = _apply_heading_1_text_normalization(str(input_path), str(output_path))

    assert changed is False


def test_heading_1_text_normalization_returns_false_when_no_aliases(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.save(input_path)

    changed = _apply_heading_1_text_normalization(str(input_path), str(output_path))

    assert changed is False
    assert output_path.exists() is False


def test_heading_2_title_case_normalization_changes_subheadings(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("research design and participants", style="Heading 2")
    doc.add_paragraph("state-of-the-art approach", style="Heading 2")
    doc.add_paragraph("Normal body text.")
    doc.save(input_path)

    changed = _apply_heading_2_title_case(str(input_path), str(output_path))

    assert changed is True
    assert _paragraph_texts_by_style(output_path, "Heading 2") == [
        "Research Design and Participants",
        "State-of-the-Art Approach",
    ]
    assert Document(output_path).paragraphs[2].text == "Normal body text."


@patch("app.services.output_generation_samfix.call_llm_json")
def test_abstract_plan_includes_affiliation_style_fix(mock_call, tmp_path):
    mock_call.side_effect = Exception("skip llm")
    input_path = tmp_path / "front.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("Sam Person^a and Lee Writer^b")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This abstract paragraph gives enough text for a simple test.")
    doc.add_paragraph("")
    doc.add_paragraph("Practioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_paragraph("Keywords")
    doc.add_paragraph("copy editing; testing")
    doc.save(input_path)

    state = abstractFound(str(input_path))
    plan = abstractFormatCheck(str(input_path), state)

    expected_fix = {
        "paragraph_index": 2,
        "new_style": AFFILIATIONS_REQUIRED_STYLE_ID,
    }
    assert expected_fix in plan["style_fixes"]


@patch("app.services.output_generation_samfix.call_llm_json")
def test_split_abstract_word_limit_counts_all_paragraphs(mock_call, tmp_path):
    mock_call.side_effect = Exception("skip llm")
    input_path = tmp_path / "front.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("Sam Person^a")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("Abstract")
    for _ in range(5):
        doc.add_paragraph(" ".join(["word"] * 60))
    doc.add_paragraph("Keywords")
    doc.add_paragraph("copy editing; testing")
    doc.save(input_path)

    state = abstractFound(str(input_path))
    plan = abstractFormatCheck(str(input_path), state)
    messages = "\n\n".join(comment["message"] for comment in plan.get("comments", []))

    assert state["abstract_word_count"] == 300
    assert "Abstract is too long. Current: 300 word" in messages


def test_affiliation_style_fix_removes_direct_space_before(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example title")
    doc.add_paragraph("Sam Person^a")
    affiliation = doc.add_paragraph("a RMIT University, Australia")
    affiliation.paragraph_format.space_before = Pt(12)
    doc.add_paragraph("Abstract")
    doc.save(input_path)

    style_fixes = [
        {
            "paragraph_index": 2,
            "new_style": AFFILIATIONS_REQUIRED_STYLE_ID,
        }
    ]
    changed = _apply_tracked_style_fixes(str(input_path), str(output_path), style_fixes, 7000)
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")

    root = etree.fromstring(xml)
    body = root.find(f"{WQ}body")
    paragraphs = [el for el in body if el.tag == f"{WQ}p"]
    affiliation_paragraph = paragraphs[2]

    p_style = affiliation_paragraph.find(f"{WQ}pPr/{WQ}pStyle")
    spacing = affiliation_paragraph.find(f"{WQ}pPr/{WQ}spacing")
    tracked_change = affiliation_paragraph.find(f"{WQ}pPr/{WQ}pPrChange")

    assert p_style.get(f"{WQ}val") == AFFILIATIONS_REQUIRED_STYLE_ID
    if spacing is not None:
        assert spacing.get(f"{WQ}before") is None
    assert tracked_change is not None
    old_spacing = tracked_change.find(f"{WQ}pPr/{WQ}spacing")
    assert old_spacing is not None
    assert old_spacing.get(f"{WQ}before") == "240"


@patch("app.services.output_generation_samfix.call_llm_json")
def test_multiple_affiliation_lines_do_not_keep_style_spacing_between_them(mock_call, tmp_path):
    mock_call.side_effect = Exception("skip llm")
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example title")
    doc.add_paragraph("Sam Person^a and Lee Writer^b")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("b RMIT University, Australia")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This abstract paragraph gives enough text for a simple test.")
    doc.add_paragraph("Practioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_paragraph("Keywords")
    doc.add_paragraph("copy editing; testing")
    doc.save(input_path)

    state = abstractFound(str(input_path))
    plan = abstractFormatCheck(str(input_path), state)

    first_affiliation_fix = None
    for fix in plan["style_fixes"]:
        if fix.get("paragraph_index") == 2:
            first_affiliation_fix = fix

    assert first_affiliation_fix is not None
    assert first_affiliation_fix.get("clear_after_spacing") is True

    changed = _apply_tracked_style_fixes(str(input_path), str(output_path), plan["style_fixes"], 7000)
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")

    root = etree.fromstring(xml)
    body = root.find(f"{WQ}body")
    paragraphs = [el for el in body if el.tag == f"{WQ}p"]
    first_affiliation = paragraphs[2]
    second_affiliation = paragraphs[3]

    first_spacing = first_affiliation.find(f"{WQ}pPr/{WQ}spacing")
    second_spacing = second_affiliation.find(f"{WQ}pPr/{WQ}spacing")

    assert first_spacing is not None
    assert first_spacing.get(f"{WQ}after") == "0"
    if second_spacing is not None:
        assert second_spacing.get(f"{WQ}after") is None


def test_inline_keywords_line_is_split_into_heading_and_text(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Practioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_paragraph("Keywords: feedback; learning")
    doc.save(input_path)

    state = keywordsFound(str(input_path))
    plan = keywordsFormatCheck(str(input_path), state)

    assert plan["split_inline_keywords"] is True
    assert "Add one keywords line directly under the Keywords heading." not in plan["message"]
    plan["keywords_heading_pos"] = 0
    plan["style_fixes"][0]["paragraph_index"] = 0

    changed = _apply_keywords_plan(str(input_path), str(output_path), plan)
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)

    styles_by_text = {}
    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t")).strip()
        if text != "":
            p_style = para.find(f"{WQ}pPr/{WQ}pStyle")
            if p_style is not None:
                styles_by_text[text] = p_style.get(f"{WQ}val")

    assert styles_by_text["Keywords"] == KEYWORDS_REQUIRED_HEADING_STYLE_ID
    assert styles_by_text["feedback; learning"] == KEYWORDS_REQUIRED_TEXT_STYLE_ID
    assert "Keywords: feedback; learning" not in styles_by_text


def test_body_citation_moves_to_footer_and_intro_gets_page_break(tmp_path):
    input_path = tmp_path / "input.docx"
    moved_path = tmp_path / "moved.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Keywords")
    doc.add_paragraph("feedback; learning")
    doc.add_paragraph("Citation")
    doc.add_paragraph("Author, A. (2025). Article title. Journal, 1(1). https://doi.org/10.0000/example")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text.")
    doc.save(input_path)

    state = citationFound(str(input_path))
    plan = citationFormatCheck(str(input_path), state)

    assert plan["action"] == "move_citation_to_footer"

    changed = _apply_citation_plan(str(input_path), str(moved_path), plan)
    assert changed is True

    page_break_changed = _apply_intro_page_break(str(moved_path), str(output_path))
    assert page_break_changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        document_xml = z.read("word/document.xml")
        footer_names = [name for name in z.namelist() if name.startswith("word/footer")]
        footer_xml = z.read(footer_names[0])

    document_root = etree.fromstring(document_xml)
    body_texts = []
    intro_page_break = None
    for para in document_root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t")).strip()
        if text != "":
            body_texts.append(text)
        if text == "Introduction":
            intro_page_break = para.find(f"{WQ}pPr/{WQ}pageBreakBefore")

    assert "Citation" not in body_texts
    assert intro_page_break is not None

    footer_root = etree.fromstring(footer_xml)
    footer_text = " ".join(t.text or "" for t in footer_root.iter(f"{WQ}t"))
    footer_styles = []
    for para in footer_root.iter(f"{WQ}p"):
        p_style = para.find(f"{WQ}pPr/{WQ}pStyle")
        if p_style is not None:
            footer_styles.append(p_style.get(f"{WQ}val"))

    assert "Author, A. (2025)." in footer_text
    assert CITATION_FOOTER_STYLE_ID in footer_styles


def test_body_indent_issue_is_tracked_formatting_not_comment(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    paragraph = doc.add_paragraph("This body paragraph has a wrong indent.")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body text.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body text.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)

    assert plan["action"] == "apply_document_body_formatting"
    assert len(plan["format_fixes"]) == 1
    assert plan["format_fixes"][0]["indent"] is True

    changed = _apply_document_body_plan(str(input_path), str(output_path), plan)
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        names = z.namelist()
        xml = z.read("word/document.xml")

    assert "word/comments.xml" not in names

    root = etree.fromstring(xml)
    body = root.find(f"{WQ}body")
    paragraphs = [el for el in body if el.tag == f"{WQ}p"]
    body_paragraph = paragraphs[1]
    indent = body_paragraph.find(f"{WQ}pPr/{WQ}ind")
    tracked_change = body_paragraph.find(f"{WQ}pPr/{WQ}pPrChange")

    assert indent is not None
    assert indent.get(f"{WQ}left") == "0"
    assert tracked_change is not None


def test_normal_style_is_replaced_with_template(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("Body text.")
    doc.save(input_path)

    changed = _apply_normal_style_fix(str(input_path), str(output_path))
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/styles.xml"))

    normal_style = root.find(f".//{WQ}style[@{WQ}styleId='Normal']")
    assert normal_style.find(f"{WQ}pPr/{WQ}jc").get(f"{WQ}val") == "both"
    spacing = normal_style.find(f"{WQ}pPr/{WQ}spacing")
    assert spacing.get(f"{WQ}after") == "120"
    assert spacing.get(f"{WQ}line") == "276"
    assert normal_style.find(f"{WQ}rPr/{WQ}rFonts").get(f"{WQ}ascii") == "Arial"


def test_body_formatting_fix_sets_justified_spacing_and_font(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("This body paragraph has wrong direct formatting.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body text.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body text.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)

    assert plan["action"] == "apply_document_body_formatting"
    fix = plan["format_fixes"][0]
    assert fix["alignment"] is True
    assert fix["line_spacing"] is True
    assert fix["spacing"] is True
    assert fix["font_name"] is True
    assert fix["font_size"] is True

    changed = _apply_document_body_plan(str(input_path), str(output_path), plan)
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    para = next(
        p for p in root.iter(f"{WQ}p")
        if "".join(t.text or "" for t in p.iter(f"{WQ}t")) == "This body paragraph has wrong direct formatting."
    )
    assert para.find(f"{WQ}pPr/{WQ}jc").get(f"{WQ}val") == "both"
    spacing = para.find(f"{WQ}pPr/{WQ}spacing")
    assert spacing.get(f"{WQ}line") == "276"
    assert spacing.get(f"{WQ}after") == "120"
    r_pr = para.find(f"{WQ}r/{WQ}rPr")
    assert r_pr.find(f"{WQ}rFonts").get(f"{WQ}ascii") == "Arial"
    assert r_pr.find(f"{WQ}sz").get(f"{WQ}val") == "22"


def test_body_required_section_comments_include_expected_locations(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Introduction body text.")
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body text.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("Conclusion").style = "Heading 1"
    doc.add_paragraph("Acknowledgements").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    comments_text = "\n\n".join(comment["message"] for comment in plan["comments"])

    assert CANONICAL_STRUCTURE["missing_section_queries"]["Literature"] in comments_text
    assert CANONICAL_STRUCTURE["missing_section_queries"]["Method"] in comments_text


def test_missing_method_and_results_remain_separate_comments(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Introduction body text.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body text.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    messages = [comment["message"] for comment in plan["comments"]]

    assert any(message == CANONICAL_STRUCTURE["missing_section_queries"]["Method"] for message in messages)
    assert any(message == CANONICAL_STRUCTURE["missing_section_queries"]["Results"] for message in messages)
    assert not any(CANONICAL_STRUCTURE["missing_section_queries"]["Method"] + "\n\n" + CANONICAL_STRUCTURE["missing_section_queries"]["Results"] in message for message in messages)


def test_missing_results_anchors_at_heading_in_literature_discussion_gap(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Introduction body text.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body text.")
    doc.add_paragraph("Implementation of the EF Fellowship Program").style = "Heading 1"
    doc.add_paragraph("Implementation body.")
    doc.add_paragraph("Project Outcomes").style = "Heading 1"
    doc.add_paragraph("Outcomes body.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    results_comment = next(
        comment for comment in plan["comments"]
        if comment["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Results"]
    )

    assert results_comment["anchor_pos"] > state["literature_index"]
    assert results_comment["anchor_pos"] < state["discussion_index"]


def test_missing_literature_anchors_at_first_heading_after_introduction(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Intro body.")
    doc.add_paragraph("Background").style = "Heading 1"
    doc.add_paragraph("Background body.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    literature_comment = next(
        comment for comment in plan["comments"]
        if comment["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Literature"]
    )

    assert literature_comment["anchor_pos"] == state["intro_index"] + 2


def test_missing_method_anchors_at_second_heading_after_introduction(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Intro body.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body.")
    doc.add_paragraph("Project Design").style = "Heading 1"
    doc.add_paragraph("Project design body.")
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    method_comment = next(
        comment for comment in plan["comments"]
        if comment["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Method"]
    )

    assert method_comment["anchor_pos"] == state["literature_index"] + 2


def test_missing_results_anchors_at_first_heading_after_method(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Method body.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    results_comment = next(
        comment for comment in plan["comments"]
        if comment["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Results"]
    )

    assert results_comment["anchor_pos"] == state["discussion_index"]


def test_combined_results_discussion_heading_gets_long_policy_comment(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Results and Discussion").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    comments_text = "\n\n".join(comment["message"] for comment in plan["comments"])

    assert CANONICAL_STRUCTURE["combined_results_discussion_query"] in comments_text


def test_body_required_section_comments_report_intro_literature_method_order(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Introduction body text.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body text.")
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body text.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("Conclusion").style = "Heading 1"
    doc.add_paragraph("Acknowledgements").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    comments_text = "\n\n".join(comment["message"] for comment in plan["comments"])

    assert "Method should appear after Literature." in comments_text


def test_missing_method_comment_anchors_at_literature_not_document_top(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Title page text.")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Introduction body text.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body text.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("Conclusion").style = "Heading 1"
    doc.add_paragraph("Acknowledgements").style = "Heading 1"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    comments = plan["comments"]

    literature_anchor = state["literature_index"]
    method_comment = next(
        comment for comment in comments
        if comment["message"] == CANONICAL_STRUCTURE["missing_section_queries"]["Method"]
    )

    assert literature_anchor != state["first_heading1_index"]
    assert method_comment["anchor_pos"] == literature_anchor


def test_method_missing_required_heading2_subsection_gets_comment(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Intro body.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    # No Participants-equivalent heading at all — the synonym list accepts
    # "Sampling", "Sample", "Respondents", etc. as Participants aliases, so
    # we use a non-equivalent placeholder to keep the missing-subsection
    # behaviour exercised.
    doc.add_paragraph("Demographic Profile").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)

    assert plan["action"] == "add_document_body_comment"
    assert "Method is missing required Heading 2 subheading(s): Participants." in plan["message"]


def test_method_body_text_does_not_satisfy_required_heading2_subheadings(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants")
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure")
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)

    method_comments = [
        comment["message"]
        for comment in plan.get("comments", [])
        if comment["message"].startswith("Method is missing required Heading 2 subheading(s):")
    ]
    assert method_comments == [
        "Method is missing required Heading 2 subheading(s): Participants, Procedure."
    ]


def test_missing_required_subheadings_are_bunched_into_one_comment(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)

    missing_subheading_comments = [
        comment
        for comment in plan.get("comments", [])
        if "missing required Heading 2 subheading(s)" in comment["message"]
    ]
    assert missing_subheading_comments == [
        {
            "anchor_pos": state["method_index"],
            "message": (
                "Method is missing required Heading 2 subheading(s): Participants, Procedure."
                "\n\n"
                "Discussion is missing required Heading 2 subheading(s): "
                "Theoretical Implications, Limitations and Future Research."
            ),
        }
    ]


def test_method_section_allows_single_sentence_paragraphs(tmp_path):
    input_path = tmp_path / "input.docx"
    long_sentence = (
        "This paragraph deliberately contains more than thirty five words because it describes the sampling site, "
        "recruitment procedure, measurement context, data collection timing, participant briefing, consent process, "
        "analytic preparation, survey administration, interview scheduling, field note handling, and data storage procedures in one extended sentence for testing purposes."
    )

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Intro body.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph(long_sentence)
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    messages = "\n\n".join(comment["message"] for comment in plan.get("comments", []))

    assert "A single sentence does not constitute a paragraph" not in messages


def test_discussion_section_still_flags_single_sentence_paragraphs(tmp_path):
    input_path = tmp_path / "input.docx"
    long_sentence = (
        "This paragraph deliberately contains more than thirty five words because it describes the sampling site, "
        "recruitment procedure, measurement context, data collection timing, participant briefing, consent process, "
        "analytic preparation, survey administration, interview scheduling, field note handling, and data storage procedures in one extended sentence for testing purposes."
    )

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Intro body.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Literature body.")
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Results body.")
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph(long_sentence)
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    messages = "\n\n".join(comment["message"] for comment in plan.get("comments", []))

    assert "A single sentence does not constitute a paragraph" in messages


def test_body_numbered_author_placeholder_gets_query(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Author 1 conducted the interviews and Author 2 reviewed the transcripts.")
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    messages = "\n\n".join(comment["message"] for comment in plan.get("comments", []))

    assert "Use author surnames rather than numbered placeholders" in messages


def test_credit_numbered_author_placeholder_gets_query(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("Acknowledgements").style = "Heading 1"
    doc.add_paragraph(
        "Conflict of interest: none. Funding: none. Ethics approval was granted. "
        "AI use: no generative AI was used. CRediT author contributions: "
        "Author 1: Conceptualization; Author 2: Writing."
    )
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    messages = "\n\n".join(comment["message"] for comment in plan.get("comments", []))

    assert "Use author surnames rather than numbered placeholders" in messages


def test_front_page_numbered_author_placeholder_not_body_query(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Example article title")
    doc.add_paragraph("[Author 1] [Author 2]")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Literature").style = "Heading 1"
    doc.add_paragraph("Method").style = "Heading 1"
    doc.add_paragraph("Research Design").style = "Heading 2"
    doc.add_paragraph("Participants").style = "Heading 2"
    doc.add_paragraph("Measures").style = "Heading 2"
    doc.add_paragraph("Procedure").style = "Heading 2"
    doc.add_paragraph("Analysis").style = "Heading 2"
    doc.add_paragraph("Results").style = "Heading 1"
    doc.add_paragraph("Discussion").style = "Heading 1"
    doc.add_paragraph("Practical Implications").style = "Heading 2"
    doc.add_paragraph("Theoretical Implications").style = "Heading 2"
    doc.add_paragraph("Limitations and Future Research").style = "Heading 2"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    state = documentBodyFound(str(input_path))
    plan = documentBodyFormatCheck(str(input_path), state)
    messages = "\n\n".join(comment["message"] for comment in plan.get("comments", []))

    assert "Use author surnames rather than numbered placeholders" not in messages


def test_main_headings_are_centered_but_front_page_headings_are_left(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    title = doc.add_paragraph("Example title")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style = OxmlElement("w:pStyle")
    title_style.set(qn("w:val"), "ArticleTitle")
    title._p.get_or_add_pPr().insert(0, title_style)
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Practioner Notes")
    doc.add_paragraph("Keywords")
    doc.add_paragraph("Introduction")
    doc.add_paragraph("Literature")
    doc.add_paragraph("Method")
    doc.add_paragraph("Conclusion")
    doc.save(input_path)

    changed = _apply_heading_center_alignment(str(input_path), str(output_path))
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")

    root = etree.fromstring(xml)
    alignment_by_text = {}
    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t")).strip()
        if text == "":
            continue
        jc = para.find(f"{WQ}pPr/{WQ}jc")
        if jc is not None:
            alignment_by_text[text] = jc.get(f"{WQ}val")
        else:
            alignment_by_text[text] = ""

    assert alignment_by_text["Example title"] == "left"
    assert alignment_by_text["Abstract"] == ""
    assert alignment_by_text["Practioner Notes"] == "left"
    assert alignment_by_text["Keywords"] == "left"
    assert alignment_by_text["Introduction"] == "center"
    assert alignment_by_text["Literature"] == "center"
    assert alignment_by_text["Method"] == "center"
    assert alignment_by_text["Conclusion"] == "center"


def _make_numbered(paragraph):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)


def test_body_plan_flags_word_numbered_normal_paragraphs_but_allows_practitioner_notes(tmp_path):
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Practitioner Notes")
    _make_numbered(doc.add_paragraph("Allowed practitioner note."))
    doc.add_paragraph("Keywords")
    doc.add_paragraph("Introduction", style="Heading 1")
    _make_numbered(doc.add_paragraph("Body list item styled Normal."))
    doc.add_paragraph("References", style="Heading 1")
    doc.save(input_path)

    plan = documentBodyFormatCheck(str(input_path), documentBodyFound(str(input_path)))
    dot_comments = [
        comment for comment in plan["comments"]
        if "dot-point" in comment["message"].lower()
    ]

    # The dot-point message was reworded in commit 995bdc2 ("Consolidate
    # dot-point/acronym comments, …") to be more explicit about the
    # Practitioner-Notes exception. The intent of the test — that body
    # numbered paragraphs get flagged at anchor_pos 5 (one bullet, not
    # multiple, not on Practitioner Notes) — is unchanged.
    assert dot_comments == [
        {
            "anchor_pos": 5,
            "message": (
                "Dot-point lists are not permitted in academic journal "
                "manuscripts outside the Practitioner Notes section. "
                "Please convert this dot-point list to continuous prose."
            ),
        }
    ]


def test_missing_keywords_stub_removes_blank_spacer_after_practitioner_notes(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Abstract body")
    doc.add_paragraph("Practitioner Notes")
    note = doc.add_paragraph("Placeholder note.")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), "PractitionerNotes")
    note._p.get_or_add_pPr().append(p_style)
    doc.add_paragraph("")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(input_path)

    plan = {"ensure_keywords_section": True, "anchor_pos": 0, "missing_keywords_comment": ""}
    assert _apply_missing_keywords_stub(str(input_path), str(output_path), plan) is True

    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    paragraphs = root.findall(f".//{WQ}body/{WQ}p")
    texts = ["".join(t.text or "" for t in p.iter(f"{WQ}t")) for p in paragraphs[:6]]
    styles = [
        p.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val")
        if p.find(f"{WQ}pPr/{WQ}pStyle") is not None
        else ""
        for p in paragraphs[:6]
    ]

    assert texts[:6] == [
        "Abstract",
        "Abstract body",
        "Practitioner Notes",
        "Placeholder note.",
        "Keywords",
        "Introduction",
    ]
    assert styles[4] == KEYWORDS_REQUIRED_HEADING_STYLE_ID


def test_headings_keep_with_following_text_even_with_blank_line(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("")
    doc.add_paragraph("Body text.")
    heading_2 = doc.add_paragraph("Subheading", style="Heading 2")
    heading_2.paragraph_format.keep_with_next = False
    heading_2.paragraph_format.keep_together = False
    doc.add_paragraph("Subheading body.")
    doc.add_paragraph("Nested subheading", style="Heading 3")
    doc.add_paragraph("")
    doc.add_paragraph("Nested body.")
    doc.save(input_path)

    changed = _apply_heading_keep_next(str(input_path), str(output_path))
    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        paras = etree.fromstring(z.read("word/document.xml")).findall(f".//{WQ}body/{WQ}p")

    assert paras[0].find(f"{WQ}pPr/{WQ}keepNext") is not None
    assert paras[0].find(f"{WQ}pPr/{WQ}keepLines") is not None
    assert paras[1].find(f"{WQ}pPr/{WQ}keepNext") is not None
    assert paras[2].find(f"{WQ}pPr/{WQ}keepNext") is None
    assert paras[3].find(f"{WQ}pPr/{WQ}keepNext") is not None
    assert paras[3].find(f"{WQ}pPr/{WQ}keepLines") is not None
    assert paras[3].find(f"{WQ}pPr/{WQ}keepNext").get(f"{WQ}val") is None
    assert paras[3].find(f"{WQ}pPr/{WQ}keepLines").get(f"{WQ}val") is None
    assert paras[4].find(f"{WQ}pPr/{WQ}keepNext") is None
    assert paras[5].find(f"{WQ}pPr/{WQ}keepNext") is not None
    assert paras[5].find(f"{WQ}pPr/{WQ}keepLines") is not None
    assert paras[6].find(f"{WQ}pPr/{WQ}keepNext") is not None
    assert paras[7].find(f"{WQ}pPr/{WQ}keepNext") is None


def test_build_rechecks_keywords_after_abstract_merge(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Example title")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("First abstract paragraph.")
    doc.add_paragraph("Second abstract paragraph.")
    doc.add_paragraph("")
    doc.add_paragraph("Practioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_paragraph("Keywords: feedback; learning")
    doc.add_paragraph("Citation")
    doc.add_paragraph("Example citation.")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text.")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example.")
    doc.save(input_path)

    with patch("app.services.output_generation_samfix.call_llm_json") as mock_llm:
        mock_llm.side_effect = Exception("skip llm")
        with patch("app.services.output_generation_samfix.build_body_edit_plan") as mock_body:
            mock_body.return_value = {"action": "none", "edits": []}
            with patch("app.services.output_generation_samfix.build_editorial_review_comment_plan") as mock_review:
                mock_review.return_value = {"action": "none", "comments": []}
                build_edited_document(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)

    styles_by_text = {}
    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t")).strip()
        if text != "":
            p_style = para.find(f"{WQ}pPr/{WQ}pStyle")
            if p_style is not None:
                styles_by_text[text] = p_style.get(f"{WQ}val")

    assert styles_by_text["Keywords"] == KEYWORDS_REQUIRED_HEADING_STYLE_ID
    assert styles_by_text["feedback; learning"] == KEYWORDS_REQUIRED_TEXT_STYLE_ID


# ── Deidentified-manuscript detection ─────────────────────────────────────────
# Covers Round 2 fix B — the filename-only check failed because abstractFormatCheck
# receives the post-normalisation temp filename. The content-based check has to
# pick up the deidentification signal from the document body even when the
# filename is generic.


def _make_deid_state():
    """Minimal abstract_state stub — the deidentified check only reads docxpath."""
    return {"affiliations_found": False, "abstract_heading_pos": 0, "anchor_pos": 0}


def test_deidentified_detected_by_filename(tmp_path):
    p = tmp_path / "Deidentified_manuscript.docx"
    Document().save(p)
    assert _is_deidentified_manuscript(str(p), _make_deid_state()) is True


def test_deidentified_detected_by_placeholder_text(tmp_path):
    """Plain `[Authors removed for blind review]` placeholder anywhere in the
    front-page paragraphs is sufficient — no template style required."""
    p = tmp_path / "manuscript.docx"  # neutral filename
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("[Authors removed for blind review]")
    doc.add_paragraph("Abstract")
    doc.save(p)
    assert _is_deidentified_manuscript(str(p), _make_deid_state()) is True


def test_deidentified_detected_by_numbered_author_placeholder(tmp_path):
    p = tmp_path / "manuscript.docx"
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("[Author 1] [Author 2]")
    doc.add_paragraph("Abstract")
    doc.save(p)
    assert _is_deidentified_manuscript(str(p), _make_deid_state()) is True


def test_deidentified_detected_by_affiliations_removed_placeholder(tmp_path):
    """The placeholder regex covers the 'Affiliations removed' phrasing too."""
    p = tmp_path / "manuscript.docx"
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Jane Doe and John Smith")
    doc.add_paragraph("[Affiliations anonymised for review]")
    doc.add_paragraph("Abstract")
    doc.save(p)
    assert _is_deidentified_manuscript(str(p), _make_deid_state()) is True


def test_non_deidentified_manuscript_returns_false(tmp_path):
    """A manuscript that genuinely lacks affiliations (but has authors and
    no placeholders) should still get the affiliation comment — the
    content-based detector must NOT match here."""
    p = tmp_path / "manuscript.docx"
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Jane Doe and John Smith")
    doc.add_paragraph("Abstract")
    doc.save(p)
    assert _is_deidentified_manuscript(str(p), _make_deid_state()) is False


# ── LLM-generated Keywords (missing-section path) ───────────────────────────


def _build_doc_missing_keywords(input_path):
    """Build a minimal manuscript whose front matter has Practitioner Notes
    but NO Keywords section."""
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Abstract")
    doc.add_paragraph(
        "This article investigates collaborative learning across "
        "a twelve-week semester block in higher education."
    )
    doc.add_paragraph("Practitioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Opening body paragraph for the introduction section.")
    doc.save(input_path)


def _collect_inserted_paragraphs(output_path):
    """Return ordered (style_id, text) for every paragraph wrapped in <w:ins>."""
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)

    out = []
    for para in root.iter(f"{WQ}p"):
        ins = para.find(f"{WQ}ins")
        if ins is None:
            continue
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t")).strip()
        style_el = para.find(f"{WQ}pPr/{WQ}pStyle")
        style_id = style_el.get(f"{WQ}val") if style_el is not None else None
        out.append((style_id, text))
    return out


def test_missing_keywords_section_inserts_heading_and_ai_body(tmp_path):
    """When the LLM returns five keywords, the apply pass writes BOTH a
    tracked-insertion heading and a tracked-insertion comma-separated
    body paragraph immediately after Practitioner Notes."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _build_doc_missing_keywords(input_path)

    with patch(
        "app.services.output_generation_samfix._maybe_generate_keywords_for_missing_section",
        return_value=["alpha", "beta", "gamma", "delta", "epsilon"],
    ):
        state = keywordsFound(str(input_path))
        plan = keywordsFormatCheck(str(input_path), state)
        applied = _apply_missing_keywords_stub(str(input_path), str(output_path), plan)

    assert applied is True
    inserted = _collect_inserted_paragraphs(output_path)
    inserted_texts = [t for _, t in inserted]
    assert "Keywords" in inserted_texts
    assert "alpha, beta, gamma, delta, epsilon" in inserted_texts
    # Plan dict carries the AI-suggested-keywords comment text rather than
    # the original "please add" placeholder.
    assert "AI-suggested" in plan["missing_keywords_comment"]


def test_missing_keywords_fallback_when_generator_returns_empty(tmp_path):
    """When the LLM is unavailable or every candidate is filtered out the
    pass falls back to the original behaviour: heading-only insertion
    plus the "please add 5 keywords" comment."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _build_doc_missing_keywords(input_path)

    with patch(
        "app.services.output_generation_samfix._maybe_generate_keywords_for_missing_section",
        return_value=[],
    ):
        state = keywordsFound(str(input_path))
        plan = keywordsFormatCheck(str(input_path), state)
        applied = _apply_missing_keywords_stub(str(input_path), str(output_path), plan)

    assert applied is True
    inserted = _collect_inserted_paragraphs(output_path)
    inserted_texts = [t for _, t in inserted]
    assert "Keywords" in inserted_texts
    # No body paragraph with comma-separated keywords.
    assert all("," not in t for t in inserted_texts if t != "Keywords")
    # Comment falls back to the original "please add" wording.
    assert "please add" in plan["missing_keywords_comment"]
    assert "AI-suggested" not in plan["missing_keywords_comment"]


def test_keywords_present_paper_untouched_by_generation(tmp_path):
    """A paper that already has a Keywords section should NOT trigger the
    AI-generation path — regression coverage."""
    input_path = tmp_path / "input.docx"

    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Practitioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_paragraph("Keywords")
    doc.add_paragraph("feedback, engagement")
    doc.save(input_path)

    called = {"count": 0}

    def _spy(*args, **kwargs):
        called["count"] += 1
        return ["should", "never", "appear"]

    with patch(
        "app.services.output_generation_samfix._maybe_generate_keywords_for_missing_section",
        side_effect=_spy,
    ):
        state = keywordsFound(str(input_path))
        plan = keywordsFormatCheck(str(input_path), state)

    # `ensure_keywords_section` is only set on the missing branch; for a
    # paper that already has Keywords the plan must NOT carry it and the
    # generator must NOT have been consulted.
    assert plan.get("ensure_keywords_section") is not True
    assert called["count"] == 0


def test_keywords_heading_with_no_body_does_not_eat_introduction_paragraph(tmp_path):
    """Regression: when the Keywords heading has no real body and the doc
    lacks a `Heading 1` Introduction (so `get_front_page` returns the
    whole document), the scan must NOT promote the Introduction's first
    body paragraph to ``keywords_paragraph`` — otherwise that paragraph
    gets restyled as keywords."""
    input_path = tmp_path / "input.docx"
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Practitioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("- Second note")
    doc.add_paragraph("- Third note")
    doc.add_paragraph("- Fourth note")
    doc.add_paragraph("- Fifth note")
    doc.add_paragraph("Keywords")
    # NO real keywords body. The Introduction H1 is also intentionally
    # not styled — this exercises the `get_front_page` fall-through to
    # "whole document" path that the previous loop didn't defend against.
    doc.add_paragraph(
        "Recent advances in higher education research have surfaced a number "
        "of methodological challenges that previously went unexamined in the "
        "literature."
    )
    doc.save(input_path)

    state = keywordsFound(str(input_path))
    # The long prose sentence must NOT be picked up as the keywords line.
    assert state["keywords_paragraph_found"] is False


def test_keywords_heading_with_proper_body_still_detected(tmp_path):
    """Sanity check: a manuscript with a real keywords list still detects
    it correctly. Guards against the predicate over-rejecting legitimate
    cases."""
    input_path = tmp_path / "input.docx"
    doc = Document()
    doc.add_paragraph("An Article Title")
    doc.add_paragraph("Practitioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("Keywords")
    doc.add_paragraph("feedback, engagement, peer review")
    doc.save(input_path)

    state = keywordsFound(str(input_path))
    assert state["keywords_paragraph_found"] is True
    assert state["keywords_count"] == 3


def test_keywords_scan_stops_at_any_heading_not_just_introduction(tmp_path):
    """Previously the scan only broke on `Heading 1 + "Introduction"`,
    so any other heading between Keywords and Introduction (or a
    Heading 2 immediately after) would be consumed as body text."""
    input_path = tmp_path / "input.docx"
    doc = Document()
    doc.add_paragraph("Article Title")
    doc.add_paragraph("Practitioner Notes")
    doc.add_paragraph("- First note")
    doc.add_paragraph("Keywords")
    # No keyword body line — next paragraph is a Heading 1 with text OTHER
    # than "Introduction". The scan must break here, not promote it.
    doc.add_heading("Background", level=1)
    doc.add_paragraph("Body prose follows.")
    doc.save(input_path)

    state = keywordsFound(str(input_path))
    assert state["keywords_paragraph_found"] is False
