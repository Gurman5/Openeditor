from unittest.mock import patch

from docx import Document

from app.domain.editorial_feedback import EditorialNote, EditorialReviewResult
from app.services.editorial_review_comments import build_editorial_review_comment_plan


def _plan(docx_path, notes):
    review = EditorialReviewResult(notes=notes)
    with patch("app.services.editorial_review_comments.validate", return_value={}):
        with patch("app.services.editorial_review_comments.run_editorial_review", return_value=review):
            return build_editorial_review_comment_plan(str(docx_path))


def _length_note(section, quote, suggestion="Split this paragraph."):
    return EditorialNote(
        "general",
        "low",
        section,
        "Paragraphs can't be longer than 3 sentences.",
        suggestion,
        quote,
    )


def test_paragraph_length_comments_anchor_only_to_body_prose(tmp_path):
    docx_path = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph("Front page text. It has sentences. But it is not body.")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(
        "This paragraph has four sentences. It continues. It keeps going. It ends.",
        style="Normal",
    )
    doc.add_paragraph("Findings Overview", style="Heading 2")
    doc.add_paragraph("Table 1. One. Two. Three. Four.", style="Normal")
    pseudo = doc.add_paragraph()
    pseudo.add_run("Pseudo Heading").bold = True
    doc.add_paragraph("Literature", style="Heading 1")
    doc.save(docx_path)

    notes = [
        _length_note("Introduction", "This paragraph has four sentences."),
        _length_note("Introduction", "Findings Overview", "Do not treat headings as body paragraphs."),
        _length_note("Introduction", "Table 1. One. Two.", "Do not treat table captions as body paragraphs."),
        _length_note("", "Front page text.", "Do not treat front-page text as body paragraphs."),
        _length_note("Introduction", "Pseudo Heading", "Do not treat pseudo-headings as body paragraphs."),
        _length_note("Introduction", "Missing quote.", "Skip paragraph notes with no body quote match."),
    ]

    plan = _plan(docx_path, notes)

    assert plan["action"] == "add_editorial_review_comments"
    assert [c["anchor_pos"] for c in plan["comments"]] == [2]


def test_paragraph_length_note_without_section_searches_body_only(tmp_path):
    docx_path = tmp_path / "body_only.docx"
    doc = Document()
    doc.add_paragraph("Repeated quote. It appears before the body.")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Repeated quote. It appears in the body.", style="Normal")
    doc.save(docx_path)

    plan = _plan(docx_path, [_length_note("", "Repeated quote.")])

    assert [c["anchor_pos"] for c in plan["comments"]] == [2]


def test_paragraph_length_quote_respects_note_section(tmp_path):
    docx_path = tmp_path / "sections.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Shared quote. Intro body text.", style="Normal")
    doc.add_paragraph("Method", style="Heading 1")
    doc.add_paragraph("Shared quote. Method body text.", style="Normal")
    doc.save(docx_path)

    plan = _plan(docx_path, [_length_note("Method", "Shared quote.")])

    assert [c["anchor_pos"] for c in plan["comments"]] == [3]


def test_non_paragraph_notes_still_fall_back_to_section_heading(tmp_path):
    docx_path = tmp_path / "section_note.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Intro body text.", style="Normal")
    doc.save(docx_path)

    plan = _plan(
        docx_path,
        [
            EditorialNote(
                "introduction_quality",
                "medium",
                "Introduction",
                "The introduction could state the contribution more clearly.",
                "Add a focused contribution sentence.",
                "",
            )
        ]
    )

    assert plan["action"] == "add_editorial_review_comments"
    assert [c["anchor_pos"] for c in plan["comments"]] == [0]
