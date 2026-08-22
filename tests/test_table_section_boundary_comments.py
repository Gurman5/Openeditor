"""Tests for the table section-boundary (open/close with a table) check."""

import zipfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

from app.services.table_section_boundary_comments import (
    apply_table_section_boundary_comments,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _comment_texts(docx_path: str) -> list[str]:
    with zipfile.ZipFile(docx_path, "r") as z:
        if "word/comments.xml" not in z.namelist():
            return []
        root = etree.fromstring(z.read("word/comments.xml"))
    return ["".join(t.text or "" for t in c.iter(f"{WQ}t")) for c in root.findall(f"{WQ}comment")]


def _new_doc():
    doc = Document()
    for name in ("Table Number", "Table Title", "Table Note"):
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except Exception:
            pass
    return doc


def _add_table(doc, rows=2):
    t = doc.add_table(rows=rows, cols=2)
    for r in range(rows):
        for c in range(2):
            t.rows[r].cells[c].text = f"r{r}c{c}"
    return t


def test_section_opening_with_table_flagged(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Table 1", style="Table Number")
    doc.add_paragraph("Some statistics", style="Table Title")
    _add_table(doc)
    doc.add_paragraph("This paragraph interprets the table.")  # closing prose
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_section_boundary_comments(str(path), str(out))

    assert len(actions) == 1
    assert "opens directly with a table" in _comment_texts(str(out))[0]


def test_section_closing_with_table_flagged(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Lead-in prose introducing the table below.")
    doc.add_paragraph("Table 1", style="Table Number")
    _add_table(doc)
    doc.add_paragraph("Note. Adapted from X.", style="Table Note")  # part of block
    doc.add_paragraph("Discussion", style="Heading 1")  # next section
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_section_boundary_comments(str(path), str(out))

    assert len(actions) == 1
    assert "ends with a table" in _comment_texts(str(out))[0]


def test_section_that_is_only_a_table_gets_single_combined_comment(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Table 1", style="Table Number")
    doc.add_paragraph("Title", style="Table Title")
    _add_table(doc)
    doc.add_paragraph("Note. X", style="Table Note")
    doc.add_paragraph("Discussion", style="Heading 1")
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_section_boundary_comments(str(path), str(out))

    assert len(actions) == 1
    assert "consists of a table" in _comment_texts(str(out))[0]


def test_table_with_prose_on_both_sides_not_flagged(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Intro prose before the table.")
    doc.add_paragraph("Table 1", style="Table Number")
    _add_table(doc)
    doc.add_paragraph("Closing prose after the table.")
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_section_boundary_comments(str(path), str(out))

    assert actions == []


def test_opening_and_closing_different_tables_both_flagged(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Table 1", style="Table Number")
    _add_table(doc)                              # opens
    doc.add_paragraph("Middle prose.")
    doc.add_paragraph("Table 2", style="Table Number")
    _add_table(doc)                              # closes
    doc.add_paragraph("Discussion", style="Heading 1")
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_section_boundary_comments(str(path), str(out))

    texts = _comment_texts(str(out))
    assert len(actions) == 2
    assert any("opens directly" in t for t in texts)
    assert any("ends with a table" in t for t in texts)


def test_references_section_table_not_flagged(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Table 1", style="Table Number")
    _add_table(doc)
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_section_boundary_comments(str(path), str(out))

    assert actions == []


def test_no_tables_is_noop(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Just prose, no tables.")
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    returned_id, actions = apply_table_section_boundary_comments(str(path), str(out), 7)

    assert actions == []
    assert returned_id == 7  # comment-only: change id unchanged


def test_output_is_valid_docx(tmp_path):
    doc = _new_doc()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Table 1", style="Table Number")
    _add_table(doc)
    path = tmp_path / "in.docx"
    doc.save(str(path))
    out = tmp_path / "out.docx"
    apply_table_section_boundary_comments(str(path), str(out))
    assert Document(str(out))
