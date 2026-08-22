"""Tests for the capital-N table-header notation comment."""

import zipfile

from docx import Document
from lxml import etree

from app.services.table_n_notation_comments import apply_table_n_notation_comments

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _comment_texts(docx_path: str) -> list[str]:
    with zipfile.ZipFile(docx_path, "r") as z:
        if "word/comments.xml" not in z.namelist():
            return []
        root = etree.fromstring(z.read("word/comments.xml"))
    texts = []
    for c in root.findall(f"{WQ}comment"):
        texts.append("".join(t.text or "" for t in c.iter(f"{WQ}t")))
    return texts


def _table_doc(tmp_path, headers, *, body_before="Introduction"):
    """Build a docx with one table whose first row holds ``headers``."""
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph(body_before, style="Heading 1")
    table = doc.add_table(rows=2, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[1].cells[i].text = str(i + 1)
    doc.save(str(path))
    return str(path)


def test_capital_n_header_gets_comment(tmp_path):
    src = _table_doc(tmp_path, ["Items", "N", "Mean", "Std. Deviation"])
    out = tmp_path / "out.docx"

    _, actions = apply_table_n_notation_comments(src, str(out))

    assert len(actions) == 1
    texts = _comment_texts(str(out))
    assert len(texts) == 1
    assert "capital" in texts[0].lower()
    assert "subsample" in texts[0].lower()


def test_lowercase_n_header_not_flagged(tmp_path):
    src = _table_doc(tmp_path, ["Items", "n", "Mean"])
    out = tmp_path / "out.docx"

    _, actions = apply_table_n_notation_comments(src, str(out))

    assert actions == []
    assert _comment_texts(str(out)) == []


def test_numeric_data_cells_not_flagged(tmp_path):
    """A data cell holding a number, and a header that merely contains N as
    part of a word, must not trigger the comment."""
    src = _table_doc(tmp_path, ["Name", "Count", "Notes"])
    out = tmp_path / "out.docx"

    _, actions = apply_table_n_notation_comments(src, str(out))

    assert actions == []


def test_standalone_n_in_body_prose_not_flagged(tmp_path):
    """A lone "N" outside any table must not be flagged."""
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("N")  # not in a table
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_n_notation_comments(str(path), str(out))

    assert actions == []


def test_multiple_tables_each_flagged(tmp_path):
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    for _ in range(2):
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Items"
        table.rows[0].cells[1].text = "N"
        table.rows[1].cells[0].text = "a"
        table.rows[1].cells[1].text = "1"
    doc.save(str(path))
    out = tmp_path / "out.docx"

    _, actions = apply_table_n_notation_comments(str(path), str(out))

    assert len(actions) == 2


def test_next_change_id_threaded_unchanged(tmp_path):
    """Comment-only pass must return the change id it was given."""
    src = _table_doc(tmp_path, ["Items", "N"])
    out = tmp_path / "out.docx"

    returned_id, _ = apply_table_n_notation_comments(src, str(out), 42)

    assert returned_id == 42
