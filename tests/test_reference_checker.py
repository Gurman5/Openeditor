import pytest  # noqa: F401
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.services import reference_checker as rc
from app.services.document_analysis_services import load_paragraphs
from app.services.reference_checker import (
    _CITE_NARRATIVE_RE,
    _extract_ref_author_part,
    _is_appendix_heading,
    _iter_paren_citations,
    _normalise_surname,
    _normalise_text,
    _ref_surname_keys,
    _references_window,
    check_and_report,
    check_entry_author,
    check_entry_year,
    check_text_dois,
    extract_references,
)

PACK = "tests/jutlp_sample_docx_test_pack"
REFERENCE_STYLE = "APA 7 Reference List Entry"


def _build_docx_with_sections(tmp_path, blocks):
    """Build a docx from a list of (kind, text) blocks.

    kind: "h1" → Heading 1, "ref" → reference-styled paragraph,
    "body" → plain paragraph.
    """
    doc = Document()
    doc.styles.add_style(REFERENCE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    for kind, text in blocks:
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "ref":
            doc.add_paragraph(text, style=REFERENCE_STYLE)
        else:
            doc.add_paragraph(text)
    path = tmp_path / "doc.docx"
    doc.save(path)
    return str(path)


class TestReferencesBoundary:
    """References checking must stop at the References section — appendix
    content after it (which JUTLP doesn't allow but some papers include)
    must never be treated as reference entries."""

    def test_appendix_reference_styled_paragraph_is_excluded(self, tmp_path):
        path = _build_docx_with_sections(tmp_path, [
            ("h1", "Introduction"),
            ("body", "Body prose here with enough length to look real."),
            ("h1", "References"),
            ("ref", "Smith, J. (2020). A real study. Journal, 1(1), 1."),
            ("ref", "Jones, A. (2021). Another study. Journal, 2(2), 2."),
            ("h1", "Appendix A"),
            # Appendix item ALSO styled as a reference entry — the bug.
            ("ref", "What makes a research study good or credible?"),
        ])
        refs = extract_references(path)
        assert len(refs) == 2
        assert all("Appendix" not in r for r in refs)
        assert not any("research study good or credible" in r for r in refs)

    def test_appendix_heading_not_styled_as_h1_still_closes_window(self, tmp_path):
        """The real-world bug: the appendix heading carries the reference
        style (NOT Heading 1), so the window must close on the heading TEXT.
        Otherwise every appendix prompt is checked as a reference entry."""
        path = _build_docx_with_sections(tmp_path, [
            ("h1", "References"),
            ("ref", "Smith, J. (2020). A real study. Journal, 1(1), 1."),
            ("ref", "Jones, A. (2021). Another study. Journal, 2(2), 2."),
            # Appendix heading styled as a reference entry, not Heading 1.
            ("ref", "Appendix A. Research-Skills Self-Assessment Prompts"),
            ("ref", "What makes a research study good or credible?"),
            ("ref", "How do you deal with uncertainty or conflicting data?"),
        ])
        refs = extract_references(path)
        assert len(refs) == 2
        assert not any("appendix" in r.lower() for r in refs)
        assert not any("research study good or credible" in r.lower() for r in refs)

    def test_window_bounds_at_next_heading(self, tmp_path):
        path = _build_docx_with_sections(tmp_path, [
            ("h1", "References"),
            ("ref", "Smith, J. (2020). A study."),
            ("h1", "Appendix A"),
            ("ref", "Appendix question one?"),
        ])
        paragraphs = load_paragraphs(path)
        start, end = _references_window(paragraphs)
        assert start is not None
        # The Appendix heading closes the window before the appendix entry.
        appendix_h1 = next(
            p.index for p in paragraphs
            if p.style == "Heading 1" and p.text.strip() == "Appendix A"
        )
        assert end == appendix_h1

    def test_no_references_heading_falls_back_to_docwide(self, tmp_path):
        """A paper with reference-styled paragraphs but NO References heading
        still has its entries collected (no regression)."""
        path = _build_docx_with_sections(tmp_path, [
            ("body", "Some body text."),
            ("ref", "Smith, J. (2020). A study with no References heading."),
            ("ref", "Jones, A. (2021). A second one."),
        ])
        paragraphs = load_paragraphs(path)
        start, _ = _references_window(paragraphs)
        assert start is None  # no heading
        refs = extract_references(path)
        assert len(refs) == 2

    def test_references_at_end_with_no_appendix_unchanged(self, tmp_path):
        path = _build_docx_with_sections(tmp_path, [
            ("h1", "Introduction"),
            ("body", "Body prose."),
            ("h1", "References"),
            ("ref", "Smith, J. (2020). A study."),
            ("ref", "Jones, A. (2021). Another."),
        ])
        refs = extract_references(path)
        assert len(refs) == 2


class TestIsAppendixHeading:
    @pytest.mark.parametrize("text", [
        "Appendix",
        "Appendix A",
        "Appendix A. Research-Skills Self-Assessment Prompts",
        "Appendix 1",
        "appendix b",
        "  Appendix C ",
        "Appendices",
    ])
    def test_matches_appendix_headings(self, text):
        assert _is_appendix_heading(text)

    @pytest.mark.parametrize("text", [
        "",
        "Smith, J. (2020). A study.",
        "The appendix contains extra data.",   # 'appendix' not at the start
        "Appended materials follow.",          # 'append' but not 'appendix'
        "References",
    ])
    def test_rejects_non_appendix_text(self, text):
        assert not _is_appendix_heading(text)


def rules(report):
    return {r["rule_id"]: r["status"] for r in report["results"]}


class TestExtractReferences:
    def test_finds_four_references(self):
        refs = extract_references(f"{PACK}/01_valid_identified.docx")
        assert len(refs) == 4

    def test_references_are_strings(self):
        refs = extract_references(f"{PACK}/01_valid_identified.docx")
        assert all(isinstance(r, str) for r in refs)

    def test_no_empty_entries(self):
        refs = extract_references(f"{PACK}/01_valid_identified.docx")
        assert all(len(r) > 0 for r in refs)


class TestValidDoc:
    @pytest.fixture(scope="class")
    def report(self):
        return check_and_report(f"{PACK}/01_valid_identified.docx")

    def test_references_section_found(self, report):
        assert rules(report)["REF001"] == "pass"

    def test_entries_found(self, report):
        assert rules(report)["REF002"] == "pass"

    # ``test_alphabetical_order`` retired: REF003 (the alphabetical
    # check) was removed from the reference checker on develop.

    def test_no_failures(self, report):
        assert report["fail"] == 0

    def test_all_entries_have_years(self, report):
        year_rules = {k: v for k, v in rules(report).items() if k.endswith("_YEAR")}
        assert all(s == "pass" for s in year_rules.values())

    def test_all_entries_have_valid_author_format(self, report):
        auth_rules = {k: v for k, v in rules(report).items() if k.endswith("_AUTH")}
        assert all(s == "pass" for s in auth_rules.values())


class TestEntryYearCheck:
    def test_valid_year(self):
        result = check_entry_year(1, "Smith, J. (2023). Some title. Journal, 1(1), 1-10.")
        assert result["status"] == "pass"

    def test_year_with_letter_suffix(self):
        result = check_entry_year(1, "Smith, J. (2023a). Some title. Journal, 1(1), 1-10.")
        assert result["status"] == "pass"

    def test_missing_year(self):
        result = check_entry_year(1, "Smith, J. Some title without a year.")
        assert result["status"] == "fail"

    def test_rule_id_format(self):
        result = check_entry_year(3, "Smith, J. (2023). Title.")
        assert result["rule_id"] == "REFE003_YEAR"


class TestEntryAuthorCheck:
    def test_valid_single_author(self):
        result = check_entry_author(1, "Smith, J. (2023). Title.")
        assert result["status"] == "pass"

    def test_valid_author_with_middle_initial(self):
        result = check_entry_author(1, "Panadero, E. A. (2022). Title.")
        assert result["status"] == "pass"

    def test_invalid_author_format(self):
        result = check_entry_author(1, "John Smith (2023). Title.")
        assert result["status"] == "warn"

    def test_rule_id_format(self):
        result = check_entry_author(2, "Smith, J. (2023). Title.")
        assert result["rule_id"] == "REFE002_AUTH"


# Note: ``check_alphabetical_order`` was removed from the production
# code on develop; the tests that used to exercise it have been retired
# along with the function. Restore them if/when the check returns.


# ── Citation-consistency unit tests (CONS001/CONS002) ────────────────────────
# Regression coverage for the false positives reported on the JUTLP
# re-review manuscript:
#   - O'Hagan / O’Hagan with curly apostrophes captured as just "Hagan"
#   - Kiraly's (2014) possessive captured as "Kiralys"
#   - Drugan & Megone (2011) narrative form picking up "Megone"
#   - Organisational refs like "PACTE. (2003)..." or "European Master's
#     in Translation. (2022)..." normalising the entire reference text
#     as the "surname".


class TestNormaliseSurname:
    def test_strips_curly_apostrophe(self):
        assert _normalise_surname("O’Hagan") == "ohagan"

    def test_strips_straight_apostrophe(self):
        assert _normalise_surname("O'Hagan") == "ohagan"

    def test_strips_combining_accent_marks(self):
        # `ú` (U+00FA) decomposes to `u` + combining acute (U+0301);
        # the combining mark gets dropped so the bare letter survives.
        assert _normalise_surname("O’Mathúna") == "omathuna"
        assert _normalise_surname("Mathúna") == "mathuna"

    def test_strips_possessive_s_with_straight_apostrophe(self):
        assert _normalise_surname("Kiraly's") == "kiraly"

    def test_strips_possessive_s_with_curly_apostrophe(self):
        assert _normalise_surname("Kiraly’s") == "kiraly"

    def test_takes_first_author_only(self):
        assert _normalise_surname("Drugan, J., & Megone, C.") == "drugan"


class TestCitationRegexes:
    def test_curly_apostrophe_name_captured_after_text_normalisation(self):
        body = _normalise_text("As O’Hagan (2020) shows…")
        m = _CITE_NARRATIVE_RE.search(body)
        assert m is not None
        assert m.group(1) == "O'Hagan"

    def test_narrative_two_author_capture_first_with_ampersand(self):
        m = _CITE_NARRATIVE_RE.search("Drugan & Megone (2011) argue…")
        assert m is not None
        assert m.group(1) == "Drugan"

    def test_narrative_two_author_capture_first_with_and(self):
        m = _CITE_NARRATIVE_RE.search("Drugan and Megone (2011) argue…")
        assert m is not None
        assert m.group(1) == "Drugan"

    def test_narrative_et_al(self):
        m = _CITE_NARRATIVE_RE.search("Belhassen et al. (2025) reported…")
        assert m is not None
        assert m.group(1) == "Belhassen"

    def test_paren_with_curly_apostrophe_author(self):
        body = _normalise_text("(O’Hagan, 2020)")
        citations = list(_iter_paren_citations(body))
        assert len(citations) == 1
        surname, year = citations[0]
        assert surname.strip() == "O'Hagan"
        assert year == "2020"


class TestParenCitationIterator:
    """Coverage for the multi-citation parenthesis parser."""

    def test_single_citation(self):
        citations = list(_iter_paren_citations("(Smith, 2020)"))
        assert citations == [("Smith", "2020")]

    def test_two_citations_separated_by_semicolon(self):
        body = "Recent work (Smith, 2020; Jones, 2021) suggests…"
        surnames = [s for s, _ in _iter_paren_citations(body)]
        assert surnames == ["Smith", "Jones"]

    def test_four_citations_in_one_parenthesis_regression(self):
        """The reported bug: a four-citation paren block produced ZERO
        matches because the old regex required `)` directly after the
        first year."""
        body = (
            "Several authors agree "
            "(Alkhawaja, 2024; Belhassen et al., 2025; "
            "Bo, 2023; Hidayati & Nihayah, 2024) "
            "on this point."
        )
        surnames = [s.strip() for s, _ in _iter_paren_citations(body)]
        assert surnames == ["Alkhawaja", "Belhassen", "Bo", "Hidayati"]

    def test_citation_with_et_al(self):
        body = "(Belhassen et al., 2025)"
        citations = list(_iter_paren_citations(body))
        assert citations == [("Belhassen", "2025")]

    def test_citation_with_ampersand_coauthor(self):
        body = "(Hidayati & Nihayah, 2024)"
        citations = list(_iter_paren_citations(body))
        # Co-author consumed by the optional group; surname captured is first author.
        assert citations[0][0].strip() == "Hidayati"
        assert citations[0][1] == "2024"

    def test_citation_with_year_suffix(self):
        """APA disambiguators like ``2020a`` must survive."""
        body = "(Smith, 2020a; Smith, 2020b)"
        citations = list(_iter_paren_citations(body))
        assert [year for _, year in citations] == ["2020a", "2020b"]

    def test_prose_with_no_year_parens_ignored(self):
        body = "An ordinary (parenthetical aside) with no year."
        assert list(_iter_paren_citations(body)) == []

    def test_leading_lowercase_word_inside_parens_ignored(self):
        """A `(see ...)` aside without a Surname,Year pair must yield
        nothing — even though the paren contains a year, the inner
        regex requires a capitalised surname."""
        body = "Refer to figure 1 (see also figure 2, page 2020)."
        assert list(_iter_paren_citations(body)) == []


class TestRefSurnameKeys:
    def test_person_ref(self):
        ref = "Drugan, J., & Megone, C. (2011). Bringing ethics into translator training."
        assert _ref_surname_keys(ref) == {"drugan"}

    def test_org_ref_without_commas_keeps_only_org_name(self):
        ref = "PACTE. (2003). Building a translation competence model."
        keys = _ref_surname_keys(ref)
        # Previous bug: keys would contain the entire 80+ character ref string
        # normalised as one surname. Now: just the org name.
        assert "pacte" in keys
        assert all(len(k) <= 20 for k in keys)

    def test_org_ref_emits_acronym_alias(self):
        ref = (
            "European Master's in Translation. (2022). EMT competence framework. "
            "https://commission.europa.eu/system/files/2022-11/emt.pdf"
        )
        keys = _ref_surname_keys(ref)
        assert "emt" in keys
        assert "europeanmastersintranslation" in keys

    def test_person_ref_does_not_get_acronym_alias(self):
        """Person refs (with a comma) must NOT have an acronym alias
        generated — otherwise ``Drugan, J., & Megone, C.`` would yield
        ``DM`` etc., creating spurious matches."""
        ref = "Drugan, J., & Megone, C. (2011). Bringing ethics into translator training."
        keys = _ref_surname_keys(ref)
        assert keys == {"drugan"}


class TestExtractRefAuthorPart:
    def test_person_ref(self):
        # `_extract_ref_author_part` strips trailing ". " so the
        # initial-period after the final author initial is consumed.
        assert _extract_ref_author_part(
            "Smith, J., & Jones, A. (2024). Title."
        ) == "Smith, J., & Jones, A"

    def test_org_ref_with_only_period_before_year(self):
        assert _extract_ref_author_part(
            "PACTE. (2003). Building a translation competence model."
        ) == "PACTE"

    def test_org_ref_with_long_title(self):
        assert _extract_ref_author_part(
            "European Master's in Translation. (2022). EMT competence framework."
        ) == "European Master's in Translation"

    def test_no_year_returns_whole_string(self):
        # Defensive: if a reference is malformed and has no year, we
        # still return something stable rather than crashing.
        assert _extract_ref_author_part("Smith J Title").startswith("Smith")


# ── Plain-text DOI validation (DOIT***) ─────────────────────────────────────


def _stub_lookup(mapping):
    """Return a fake ``_lookup_doi`` that resolves DOIs from a dict.
    ``None`` value means "404 from CrossRef"."""
    def _call(doi):
        return mapping.get(doi)
    return _call


def _cr_item(year, surname):
    return {
        "issued": {"date-parts": [[year]]},
        "author": [{"family": surname}],
    }


class TestCheckTextDOIs:
    def test_passes_when_crossref_record_matches(self, monkeypatch):
        monkeypatch.setattr(
            rc, "_lookup_doi",
            _stub_lookup({"10.1234/abcd": _cr_item(2020, "Smith")}),
        )
        refs = [
            "Smith, J. (2020). A study. Journal, 1(1), 1. "
            "https://doi.org/10.1234/abcd"
        ]
        results = check_text_dois(refs, delay=0)
        assert len(results) == 1
        assert results[0]["rule_id"] == "DOIT001"
        assert results[0]["status"] == "pass"

    def test_fails_when_crossref_returns_404(self, monkeypatch):
        monkeypatch.setattr(
            rc, "_lookup_doi", _stub_lookup({"10.1234/missing": None})
        )
        refs = ["Smith, J. (2020). A study. https://doi.org/10.1234/missing"]
        results = check_text_dois(refs, delay=0)
        assert len(results) == 1
        assert results[0]["status"] == "fail"
        assert "10.1234/missing" in results[0]["message"]

    def test_fails_on_year_mismatch(self, monkeypatch):
        monkeypatch.setattr(
            rc, "_lookup_doi",
            _stub_lookup({"10.1234/yearbad": _cr_item(1999, "Smith")}),
        )
        refs = ["Smith, J. (2020). A study. doi: 10.1234/yearbad"]
        results = check_text_dois(refs, delay=0)
        assert results[0]["status"] == "fail"
        assert "year in reference" in results[0]["message"]

    def test_fails_on_author_mismatch(self, monkeypatch):
        monkeypatch.setattr(
            rc, "_lookup_doi",
            _stub_lookup({"10.1234/authbad": _cr_item(2020, "Jones")}),
        )
        refs = ["Smith, J. (2020). A study. 10.1234/authbad"]
        results = check_text_dois(refs, delay=0)
        assert results[0]["status"] == "fail"
        assert "author in reference" in results[0]["message"]

    def test_reference_without_doi_emits_no_row(self, monkeypatch):
        called = {"n": 0}

        def _spy(doi):
            called["n"] += 1
            return None

        monkeypatch.setattr(rc, "_lookup_doi", _spy)
        refs = ["Smith, J. (1985). An older study without a DOI."]
        results = check_text_dois(refs, delay=0)
        assert results == []
        assert called["n"] == 0

    def test_multiple_dois_in_one_reference_emit_one_row_each(self, monkeypatch):
        monkeypatch.setattr(
            rc, "_lookup_doi",
            _stub_lookup({
                "10.1234/first":  _cr_item(2020, "Smith"),
                "10.5678/second": _cr_item(2020, "Smith"),
            }),
        )
        refs = [
            "Smith, J. (2020). A study. "
            "https://doi.org/10.1234/first See also: 10.5678/second"
        ]
        results = check_text_dois(refs, delay=0)
        ids = sorted(r["rule_id"] for r in results)
        assert ids == ["DOIT001_1", "DOIT001_2"]
        assert all(r["status"] == "pass" for r in results)

    def test_duplicate_doi_in_one_reference_deduped(self, monkeypatch):
        """A reference that quotes the same DOI twice (bare + URL form)
        should only hit CrossRef once."""
        call_log: list[str] = []

        def _spy(doi):
            call_log.append(doi)
            return _cr_item(2020, "Smith")

        monkeypatch.setattr(rc, "_lookup_doi", _spy)
        refs = ["Smith, J. (2020). A study. 10.1234/same https://doi.org/10.1234/same"]
        results = check_text_dois(refs, delay=0)
        assert len(results) == 1
        assert len(call_log) == 1
