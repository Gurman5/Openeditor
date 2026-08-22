from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.document_analysis_services import (
    detect_heading_level_1,
    extract_heading_level_1_sections,
    extract_keywords,
    extract_main_sections,
    extract_subsections,
    load_paragraphs,
    parse_docx_structure,
)


def _add_inserted_paragraph(doc, text, style=None):
    """Append a paragraph whose text sits inside a tracked insertion (w:ins)."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "900")
    ins.set(qn("w:author"), "CopyEditor AI")
    ins.set(qn("w:date"), "2026-01-01T00:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    ins.append(r)
    p._p.append(ins)
    return p


def _add_paragraph_with_deletion(doc, before, deleted, after):
    """Append 'before <del>deleted</del> after' — deleted text is a w:delText."""
    p = doc.add_paragraph()
    p.add_run(before)
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "901")
    dele.set(qn("w:author"), "CopyEditor AI")
    dele.set(qn("w:date"), "2026-01-01T00:00:00Z")
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = deleted
    r.append(dt)
    dele.append(r)
    p._p.append(dele)
    p.add_run(after)
    return p


def _add_paragraph_with_textbox(doc, text, box_text):
    """Append a paragraph that hosts a floating text box (nested w:p)."""
    p = doc.add_paragraph(text)
    run = OxmlElement("w:r")
    txbx = OxmlElement("w:txbxContent")
    np = OxmlElement("w:p")
    nr = OxmlElement("w:r")
    nt = OxmlElement("w:t")
    nt.text = box_text
    nr.append(nt)
    np.append(nr)
    txbx.append(np)
    run.append(txbx)
    p._p.append(run)
    return p


def test_load_paragraphs_excludes_textbox_content(tmp_path):
    """A floating text box anchored on a paragraph (e.g. the JUTLP editorial
    'Editors [LEAVE BLANK]' box on the Abstract) must NOT merge into the host
    paragraph's text — otherwise the exact-text section match breaks."""
    docx_path = tmp_path / "tb.docx"
    doc = Document()
    _add_paragraph_with_textbox(doc, "Abstract", "Editors [LEAVE BLANK] Section: Senior Editor")
    doc.save(str(docx_path))

    texts = [p.text for p in load_paragraphs(str(docx_path))]
    assert "Abstract" in texts          # exact — not "Abstract Editors [LEAVE BLANK]…"
    assert all("LEAVE BLANK" not in t for t in texts)


def test_load_paragraphs_includes_tracked_insertions(tmp_path):
    """A tracked-inserted heading/line must be visible to load_paragraphs —
    python-docx's plain .text drops w:ins runs, which hid bot-inserted content
    (e.g. the Keywords section) from any pass that re-reads the output."""
    docx_path = tmp_path / "ins.docx"
    doc = Document()
    _add_inserted_paragraph(doc, "Keywords", style="Heading 1")
    _add_inserted_paragraph(doc, "alpha, beta, gamma")
    doc.save(str(docx_path))

    texts = [p.text for p in load_paragraphs(str(docx_path))]
    assert "Keywords" in texts
    assert "alpha, beta, gamma" in texts


def test_load_paragraphs_excludes_tracked_deletions(tmp_path):
    """Deleted text (w:delText) must NOT appear — the accepted-changes view."""
    docx_path = tmp_path / "del.docx"
    doc = Document()
    _add_paragraph_with_deletion(doc, "Hello ", "removed ", "world")
    doc.save(str(docx_path))

    texts = [p.text for p in load_paragraphs(str(docx_path))]
    assert "Hello world" in texts
    assert all("removed" not in t for t in texts)


def test_extract_keywords_sees_tracked_inserted_section(tmp_path):
    """Regression for the re-validation blind spot: an inserted Keywords
    heading + keyword line are picked up so re-validating the output doesn't
    falsely report keywords missing."""
    docx_path = tmp_path / "kw.docx"
    doc = Document()
    doc.add_paragraph("Practitioner Notes")
    doc.add_paragraph("A note.")
    _add_inserted_paragraph(doc, "Keywords")
    _add_inserted_paragraph(doc, "translation ethics, artificial intelligence")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(str(docx_path))

    keywords = extract_keywords(load_paragraphs(str(docx_path)))
    assert "translation ethics" in keywords
    assert "artificial intelligence" in keywords


def test_detect_heading_level_1_returns_text_and_positions_in_order(tmp_path):
    docx_path = tmp_path / "headings.docx"

    doc = Document()
    doc.add_paragraph("Title", style="Title")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Intro body.", style="Normal")
    doc.add_paragraph("Research Design", style="Heading 2")
    doc.add_paragraph("Method", style="Heading 1")
    doc.add_paragraph("", style="Heading 1")
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    assert detect_heading_level_1(str(docx_path)) == [
        {"text": "Introduction", "position": 1},
        {"text": "Method", "position": 4},
        {"text": "Results", "position": 6},
    ]


def test_parse_docx_structure_exposes_heading_level_1_sections(tmp_path):
    docx_path = tmp_path / "structure.docx"

    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body.", style="Normal")
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    parsed = parse_docx_structure(str(docx_path))

    assert parsed["main_sections"] == ["Introduction", "References"]
    assert parsed["heading_level_1_sections"] == [
        {"text": "Introduction", "position": 0},
        {"text": "References", "position": 2},
    ]


def test_extract_heading_level_1_sections_accepts_loaded_paragraphs(tmp_path):
    docx_path = tmp_path / "loaded.docx"

    doc = Document()
    doc.add_paragraph("Abstract", style="Normal")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))

    assert extract_heading_level_1_sections(paragraphs) == [
        {"text": "Introduction", "position": 1},
    ]


# ── Pseudo-H2 detection (Round 3) ─────────────────────────────────────────────
# extract_subsections must accept paragraphs that look like H2 subheadings even
# when they're styled Normal — common for manuscripts that mark subheadings
# only visually (bold + short) before Sam-fix's restyling pass runs.


def _make_bold_paragraph(doc, text):
    """Append a paragraph whose single run is explicitly bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def test_pseudo_heading2_bold_normal_detected_as_subsection(tmp_path):
    docx_path = tmp_path / "pseudo_h2.docx"
    doc = Document()
    doc.add_paragraph("Method", style="Heading 1")
    _make_bold_paragraph(doc, "Setting of the Research")
    doc.add_paragraph(
        "The research was conducted in 2025 at a university."
    )
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    subs = extract_subsections(paragraphs, "Method")
    assert subs == ["Setting of the Research"]


def test_pseudo_heading2_with_terminal_punctuation_is_excluded(tmp_path):
    """Bold + Normal but ending in a period — that's a sentence, not a heading."""
    docx_path = tmp_path / "pseudo_h2_period.docx"
    doc = Document()
    doc.add_paragraph("Method", style="Heading 1")
    _make_bold_paragraph(doc, "This is a bold emphasised sentence.")
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert extract_subsections(paragraphs, "Method") == []


def test_table_caption_not_treated_as_subsection(tmp_path):
    """`Table 1.` is bold and short but it's a caption, not a subheading."""
    docx_path = tmp_path / "pseudo_h2_caption.docx"
    doc = Document()
    doc.add_paragraph("Method", style="Heading 1")
    _make_bold_paragraph(doc, "Table 1.")
    _make_bold_paragraph(doc, "Figure 3:")
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert extract_subsections(paragraphs, "Method") == []


def test_real_heading_2_still_detected(tmp_path):
    """Regression: a real Heading 2 paragraph still comes back unchanged."""
    docx_path = tmp_path / "real_h2.docx"
    doc = Document()
    doc.add_paragraph("Method", style="Heading 1")
    doc.add_paragraph("Participants", style="Heading 2")
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert extract_subsections(paragraphs, "Method") == ["Participants"]


def test_long_bold_paragraph_not_treated_as_subsection(tmp_path):
    """Subheadings are short — a 200-char bold paragraph is just emphasis."""
    docx_path = tmp_path / "long_bold.docx"
    doc = Document()
    doc.add_paragraph("Method", style="Heading 1")
    _make_bold_paragraph(doc, "x" * 200)
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert extract_subsections(paragraphs, "Method") == []


def test_method_subsections_end_to_end_with_pseudo_h2(tmp_path):
    """End-to-end: validator's check_method_subsections passes MET001-MET003
    on a manuscript whose Method subheadings are styled bold-Normal."""
    from app.services.jutlp_validator import check_method_subsections

    docx_path = tmp_path / "method_pseudo.docx"
    doc = Document()
    doc.add_paragraph("Method", style="Heading 1")
    _make_bold_paragraph(doc, "Setting of the Research")  # alias for Research Design
    doc.add_paragraph("Body.")
    _make_bold_paragraph(doc, "Participants")  # canonical
    doc.add_paragraph("Body.")
    _make_bold_paragraph(doc, "Research Instrument Development and Validation")  # alias for Measures
    doc.add_paragraph("Body.")
    doc.add_paragraph("Results", style="Heading 1")
    doc.save(docx_path)

    parsed = parse_docx_structure(str(docx_path))
    rules_out = {r["rule_id"]: r["status"] for r in check_method_subsections(parsed)}
    assert rules_out["MET001"] == "pass"
    assert rules_out["MET002"] == "pass"
    assert rules_out["MET003"] == "pass"
    assert rules_out["MET004"] == "fail"  # Procedure genuinely missing
    assert rules_out["MET005"] == "fail"  # Analysis genuinely missing


# ── Unstyled main-section headings (Phase-ordering gap) ───────────────────────
# The validator runs in Phase 1; the Sam-fix pass that restyles a paragraph to
# Heading 1 by matching its text (e.g. "Discussion") runs in Phase 3. So a
# section the author left as Normal/bold text must still be recognised as a main
# section here, or it is wrongly reported missing for a section that is plainly
# present and will be restyled.


def test_unstyled_discussion_recognised_as_main_section(tmp_path):
    docx_path = tmp_path / "unstyled_discussion.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Intro body.", style="Normal")
    doc.add_paragraph("Discussion", style="Normal")  # author never applied H1
    doc.add_paragraph("Discussion body.", style="Normal")
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert "Discussion" in extract_main_sections(paragraphs)


def test_bold_unstyled_discussion_recognised(tmp_path):
    docx_path = tmp_path / "bold_discussion.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    _make_bold_paragraph(doc, "Discussion")  # visually a heading only
    doc.add_paragraph("Discussion body.", style="Normal")
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert "Discussion" in extract_main_sections(paragraphs)


def test_unstyled_alias_recognised(tmp_path):
    """A canonical alias ("Findings" for Results) left unstyled is recognised."""
    docx_path = tmp_path / "unstyled_alias.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Findings", style="Normal")  # alias of Results
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert "Findings" in extract_main_sections(paragraphs)


def test_body_sentence_starting_with_section_name_not_recognised(tmp_path):
    """Only a whole-paragraph match counts — a sentence is not a heading."""
    docx_path = tmp_path / "sentence.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(
        "Discussion of the findings is presented below.", style="Normal"
    )
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert extract_main_sections(paragraphs) == ["Introduction", "References"]


def test_discussion_required_section_passes_when_unstyled(tmp_path):
    """End-to-end: SEC rule for Discussion passes even when the heading is
    Normal-styled, matching what the Phase-3 restyle pass will produce."""
    from app.services.jutlp_validator import check_required_sections

    docx_path = tmp_path / "discussion_sec.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Literature", style="Heading 1")
    doc.add_paragraph("Method", style="Heading 1")
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Discussion", style="Normal")  # unstyled heading
    doc.add_paragraph("Discussion body.", style="Normal")
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    parsed = parse_docx_structure(str(docx_path))
    results = check_required_sections(parsed)
    discussion = next(r for r in results if r["message"].startswith("'Discussion'"))
    assert discussion["status"] == "pass"


def test_unstyled_discussion_subsections_detected(tmp_path):
    """get_section_bounds must locate an unstyled Discussion so its pseudo-H2
    subheadings are still scanned."""
    docx_path = tmp_path / "discussion_subs.docx"
    doc = Document()
    doc.add_paragraph("Results", style="Heading 1")
    doc.add_paragraph("Discussion", style="Normal")  # unstyled main heading
    _make_bold_paragraph(doc, "Practical Implications")
    doc.add_paragraph("Body.")
    doc.add_paragraph("References", style="Heading 1")
    doc.save(docx_path)

    paragraphs = load_paragraphs(str(docx_path))
    assert "Practical Implications" in extract_subsections(paragraphs, "Discussion")
