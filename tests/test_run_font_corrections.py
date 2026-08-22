"""Tests for run-level font normalisation (Arial; references at 11pt)."""

import zipfile

from docx import Document
from docx.shared import Pt
from lxml import etree

from app.services.run_font_corrections import apply_run_font_corrections

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _current_fonts_and_sizes(path):
    """Return per-run (font, sz) from the *current* rPr (ignoring rPrChange)."""
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    out = []
    for r in root.iter(f"{WQ}r"):
        if r.find(f"{WQ}t") is None:
            continue
        rpr = r.find(f"{WQ}rPr")
        rf = rpr.find(f"{WQ}rFonts") if rpr is not None else None
        sz = rpr.find(f"{WQ}sz") if rpr is not None else None
        out.append((
            rf.get(f"{WQ}ascii") if rf is not None else None,
            sz.get(f"{WQ}val") if sz is not None else None,
        ))
    return out


def _set_run_font(p, name, size_pt):
    for run in p.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)


def test_direct_times_run_becomes_arial(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Body text in the wrong font.")
    _set_run_font(p, "Times New Roman", 12)
    src = tmp_path / "in.docx"
    doc.save(str(src))
    out = tmp_path / "out.docx"

    _, actions = apply_run_font_corrections(str(src), str(out))
    assert len(actions) >= 1
    fonts = [f for f, _ in _current_fonts_and_sizes(str(out))]
    assert "Times New Roman" not in fonts
    assert all(f in (None, "Arial") for f in fonts)


def _find_or_make(parent, tag):
    el = parent.find(tag)
    if el is None:
        el = etree.SubElement(parent, tag)
    return el


def test_reference_runs_set_to_arial_11pt(tmp_path):
    doc = Document()
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Adams, B. (2020). A study. Journal of Things.")
    src0 = tmp_path / "s0.docx"
    doc.save(str(src0))
    # Force the reference paragraph's pStyle id + a wrong direct font/size.
    with zipfile.ZipFile(str(src0)) as z:
        root = etree.fromstring(z.read("word/document.xml"))
        members = {i.filename: z.read(i.filename) for i in z.infolist()}
    body = root.find(f"{WQ}body")
    ref_p = body.findall(f"{WQ}p")[-1]
    pPr = _find_or_make(ref_p, f"{WQ}pPr")
    _find_or_make(pPr, f"{WQ}pStyle").set(f"{WQ}val", "APA7ReferenceListEntry")
    for r in ref_p.iter(f"{WQ}r"):
        rpr = _find_or_make(r, f"{WQ}rPr")
        rf = _find_or_make(rpr, f"{WQ}rFonts")
        for a in ("ascii", "hAnsi", "cs"):
            rf.set(f"{WQ}{a}", "Times New Roman")
        _find_or_make(rpr, f"{WQ}sz").set(f"{WQ}val", "24")
    members["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    src = tmp_path / "in.docx"
    with zipfile.ZipFile(str(src), "w", zipfile.ZIP_DEFLATED) as zout:
        for n, d in members.items():
            zout.writestr(n, d)

    out = tmp_path / "out.docx"
    _, actions = apply_run_font_corrections(str(src), str(out))
    assert any(a["ref"] for a in actions)
    # the reference run is now Arial 22 (11pt)
    fonts_sizes = _current_fonts_and_sizes(str(out))
    ref_run = fonts_sizes[-1]
    assert ref_run == ("Arial", "22")


def test_inheriting_run_untouched(tmp_path):
    # A run with no direct font (inherits Normal) is left alone outside refs.
    doc = Document()
    doc.add_paragraph("Plain inheriting paragraph.")
    src = tmp_path / "in.docx"
    doc.save(str(src))
    out = tmp_path / "out.docx"
    _, actions = apply_run_font_corrections(str(src), str(out))
    assert actions == []


def test_reconstructed_reference_run_inside_ins_styled_directly(tmp_path):
    """A reconstructed reference run (inside <w:ins>, as the APA-7 format pass
    emits) is set to Arial 11pt DIRECTLY — no rPrChange nested in the insertion."""
    doc = Document()
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("placeholder")
    src0 = tmp_path / "s0.docx"
    doc.save(str(src0))
    with zipfile.ZipFile(str(src0)) as z:
        root = etree.fromstring(z.read("word/document.xml"))
        members = {i.filename: z.read(i.filename) for i in z.infolist()}
    body = root.find(f"{WQ}body")
    ref_p = body.findall(f"{WQ}p")[-1]
    # style it as a reference entry
    pPr = ref_p.find(f"{WQ}pPr") or etree.SubElement(ref_p, f"{WQ}pPr")
    etree.SubElement(pPr, f"{WQ}pStyle").set(f"{WQ}val", "APA7ReferenceListEntry")
    # remove the placeholder run, add an inserted (tracked) run in Times
    for r in ref_p.findall(f"{WQ}r"):
        ref_p.remove(r)
    ins = etree.SubElement(ref_p, f"{WQ}ins")
    ins.set(f"{WQ}id", "500")
    run = etree.SubElement(ins, f"{WQ}r")
    rpr = etree.SubElement(run, f"{WQ}rPr")
    rf = etree.SubElement(rpr, f"{WQ}rFonts")
    for a in ("ascii", "hAnsi", "cs"):
        rf.set(f"{WQ}{a}", "Times New Roman")
    etree.SubElement(run, f"{WQ}t").text = "Smith, J. (2020). A reconstructed reference. Journal."
    members["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    src = tmp_path / "in.docx"
    with zipfile.ZipFile(str(src), "w", zipfile.ZIP_DEFLATED) as zout:
        for n, d in members.items():
            zout.writestr(n, d)

    out = tmp_path / "out.docx"
    apply_run_font_corrections(str(src), str(out))

    with zipfile.ZipFile(str(out)) as z:
        root2 = etree.fromstring(z.read("word/document.xml"))
    ins_run = root2.find(f"{WQ}body").findall(f"{WQ}p")[-1].find(f"{WQ}ins").find(f"{WQ}r")
    rpr2 = ins_run.find(f"{WQ}rPr")
    assert rpr2.find(f"{WQ}rFonts").get(f"{WQ}ascii") == "Arial"
    assert rpr2.find(f"{WQ}sz").get(f"{WQ}val") == "22"
    # styled directly — no rPrChange nested inside the insertion
    assert rpr2.find(f"{WQ}rPrChange") is None
