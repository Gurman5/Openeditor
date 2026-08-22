"""Tests for heading corrections (number stripping + section renaming) and the
number-tolerance of the structure matchers."""

import zipfile
from unittest.mock import patch

import pytest
from docx import Document
from lxml import etree

from app.domain.canonical_jultp_template import (
    CANONICAL_STRUCTURE,
    SECTION_RENAME_MAP,
    strip_leading_section_number,
)
from app.services.document_analysis_services import (
    extract_main_sections,
    load_paragraphs,
)
from app.services.heading_corrections import (
    _build_style_level_map,
    _heading_level,
    apply_heading_corrections,
)
from app.services.jutlp_validator import _matches_section
from app.services.output_generation_samfix import (
    HEADING_1_TEXT_NORMALIZATIONS,
    build_edited_document,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _doc(tmp_path, headings):
    """headings: list of (text, style)."""
    path = tmp_path / "doc.docx"
    doc = Document()
    for text, style in headings:
        doc.add_paragraph(text, style=style)
    doc.save(str(path))
    return str(path)


def _edits(docx_path):
    """Return list of (deleted_text, inserted_text) tracked pairs."""
    with zipfile.ZipFile(docx_path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    pairs = []
    for p in root.iter(f"{WQ}p"):
        dels = "".join(t.text or "" for t in p.iter(f"{WQ}delText"))
        ins = "".join(
            t.text or ""
            for ins_el in p.findall(f"{WQ}ins")
            for t in ins_el.iter(f"{WQ}t")
        )
        if dels or ins:
            pairs.append((dels, ins))
    return pairs


# ── strip_leading_section_number ─────────────────────────────────────────────

def test_strip_leading_number_variants():
    assert strip_leading_section_number("1. Introduction") == "Introduction"
    assert strip_leading_section_number("2. Literature Review") == "Literature Review"
    assert strip_leading_section_number("1.2 Background") == "Background"
    assert strip_leading_section_number("(3) Method") == "Method"
    # No false positives:
    assert strip_leading_section_number("3D Printing") == "3D Printing"
    assert strip_leading_section_number("Results") == "Results"
    # A heading that is only a number is left untouched (never emptied):
    assert strip_leading_section_number("3.") == "3."


# ── validator number tolerance ───────────────────────────────────────────────

def test_matches_section_tolerates_leading_number():
    assert _matches_section("1. Introduction", "Introduction")
    assert _matches_section("2. Literature Review", "Literature", ["Literature Review"])


def test_extract_main_sections_strips_numbers(tmp_path):
    src = _doc(tmp_path, [
        ("1. Introduction", "Heading 1"),
        ("2. Literature Review", "Heading 1"),
        ("3. Method", "Heading 1"),
    ])
    names = extract_main_sections(load_paragraphs(src))
    assert names == ["Introduction", "Literature Review", "Method"]


# ── heading-corrections tracked pass ─────────────────────────────────────────

def test_numbered_heading_is_stripped(tmp_path):
    src = _doc(tmp_path, [("1. Introduction", "Heading 1")])
    out = tmp_path / "out.docx"
    next_id, actions = apply_heading_corrections(src, str(out), 1)
    assert next_id == 3  # one del + one ins consumed
    assert actions == [{"rule": "heading_correction",
                        "from": "1. Introduction", "to": "Introduction"}]
    assert _edits(str(out)) == [("1. Introduction", "Introduction")]


def test_findings_is_renamed_to_results(tmp_path):
    src = _doc(tmp_path, [("Findings", "Heading 1")])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    assert _edits(str(out)) == [("Findings", "Results")]


def test_numbered_alias_stripped_and_renamed_in_one_edit(tmp_path):
    src = _doc(tmp_path, [("2. Literature Review", "Heading 1")])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    assert _edits(str(out)) == [("2. Literature Review", "Literature")]


def test_combined_results_and_discussion_not_renamed(tmp_path):
    src = _doc(tmp_path, [("Results and Discussion", "Heading 1")])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    assert actions == []
    assert _edits(str(out)) == []


def test_canonical_headings_untouched(tmp_path):
    src = _doc(tmp_path, [
        ("Introduction", "Heading 1"),
        ("Method", "Heading 1"),
        ("Results", "Heading 1"),
    ])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    assert actions == []
    assert _edits(str(out)) == []


def test_subheading_number_stripped_but_not_renamed(tmp_path):
    src = _doc(tmp_path, [("3.1 Sample", "Heading 2")])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    # number stripped, but no H2 alias renaming
    assert _edits(str(out)) == [("3.1 Sample", "Sample")]


def test_build_edited_document_strips_heading_numbers(tmp_path):
    src = tmp_path / "input.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("Example title")
    doc.add_paragraph("Sam Person^a")
    doc.add_paragraph("a RMIT University, Australia")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This abstract paragraph gives enough text for a simple test.")
    doc.add_paragraph("Keywords")
    doc.add_paragraph("copy editing; testing")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text.")
    doc.add_paragraph(
        "1.1 the Scarcity of Longitudinal Evaluation in Engineering Education",
        style="Heading 2",
    )
    doc.add_paragraph("More body text.")
    doc.save(src)

    with patch("app.services.output_generation_samfix.call_llm_json") as mock_llm:
        mock_llm.side_effect = Exception("skip llm")
        with patch("app.services.output_generation_samfix.build_body_edit_plan") as mock_body:
            mock_body.return_value = {"action": "none", "edits": []}
            with patch("app.services.output_generation_samfix.build_editorial_review_comment_plan") as mock_review:
                mock_review.return_value = {"action": "none", "comments": []}
                build_edited_document(str(src), str(out))

    edits = _edits(str(out))
    assert ("1. Introduction", "Introduction") in edits
    assert (
        "1.1 The Scarcity of Longitudinal Evaluation in Engineering Education",
        "The Scarcity of Longitudinal Evaluation in Engineering Education",
    ) in edits


# ── expanded alias coverage ──────────────────────────────────────────────────

@pytest.mark.parametrize("heading,expected", [
    ("Methodology", "Method"),
    ("Research Methodology", "Method"),
    ("Materials and Methods", "Method"),
    ("Related Work", "Literature"),
    ("Review of the Literature", "Literature"),
    ("Theoretical Background", "Literature"),
    ("Empirical Findings", "Results"),
    ("Study Results", "Results"),
    ("Concluding Remarks", "Conclusion"),
    ("Summary and Conclusion", "Conclusion"),
    ("Conclusion and Recommendations", "Conclusion"),
    ("Bibliography", "References"),
    ("Works Cited", "References"),
    ("Acknowledgments", "Acknowledgements"),
])
def test_expanded_aliases_rename_to_canonical(tmp_path, heading, expected):
    src = _doc(tmp_path, [(heading, "Heading 1")])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    assert _edits(str(out)) == [(heading, expected)]


@pytest.mark.parametrize("heading", ["Results and Discussion", "Findings and Discussion"])
def test_combined_sections_never_renamed(tmp_path, heading):
    """Headings merging two canonical sections must be left for the author to
    split — the validator flags them separately."""
    src = _doc(tmp_path, [(heading, "Heading 1")])
    out = tmp_path / "out.docx"
    _, actions = apply_heading_corrections(src, str(out))
    assert actions == []
    assert _edits(str(out)) == []


# ── cross-map consistency (single source of truth) ───────────────────────────

def test_every_rename_alias_is_accepted_by_the_validator():
    """A wording we auto-rename must also be recognised as that section being
    present, so the validator never reports it missing before the rename."""
    for alias_norm, canonical in SECTION_RENAME_MAP.items():
        aliases = CANONICAL_STRUCTURE["section_aliases"].get(canonical, [])
        assert _matches_section(alias_norm, canonical, aliases), (alias_norm, canonical)


def test_sam_silent_rename_uses_the_same_map():
    """Sam's Heading 1 normalisation and the tracked pass cover identical
    wordings (both derive from SECTION_RENAME_MAP)."""
    assert set(HEADING_1_TEXT_NORMALIZATIONS) == set(SECTION_RENAME_MAP)
    for key, canonical in SECTION_RENAME_MAP.items():
        assert HEADING_1_TEXT_NORMALIZATIONS[key] == canonical


# ── numeric style-id resolution (converted-doc headings) ─────────────────────

_NUMERIC_STYLES = (
    '<?xml version="1.0"?>'
    '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:style w:type="paragraph" w:styleId="3"><w:name w:val="heading 1"/></w:style>'
    '<w:style w:type="paragraph" w:styleId="4"><w:name w:val="heading 2"/></w:style>'
    '<w:style w:type="paragraph" w:styleId="25"><w:name w:val="header"/></w:style>'
    '<w:style w:type="paragraph" w:styleId="X">'
    '<w:name w:val="MainHead"/><w:pPr><w:outlineLvl w:val="0"/></w:pPr></w:style>'
    '</w:styles>'
).encode()


def test_numeric_style_ids_resolve_to_heading_level():
    # Google-Docs-style numeric ids ("3" == heading 1) must resolve, otherwise
    # the rename/number-strip pass skips every heading (the reported bug).
    level_map = _build_style_level_map(_NUMERIC_STYLES)
    assert level_map["3"] == 1
    assert level_map["4"] == 2
    assert level_map["X"] == 1            # via outlineLvl fallback
    assert "25" not in level_map          # a non-heading style
    assert _heading_level("3", level_map) == 1
    assert _heading_level("4", level_map) == 2
    assert _heading_level("25", level_map) is None
    # friendly form still works without a map
    assert _heading_level("Heading1") == 1
    assert _heading_level("Heading 2") == 2
