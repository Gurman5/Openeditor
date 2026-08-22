import zipfile
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from app.services.body_llm_edits import (
    MAX_EDITS_PER_PARAGRAPH,
    _dedupe_edits,
    _llm_edits_for_paragraph,
    _propagate_repeated_edits,
    _validate_edit,
    build_body_edit_plan,
)
from app.services.output_generation_samfix import (
    FIGURE_NUMBER_STYLE_ID,
    FIGURE_TITLE_STYLE_ID,
    HEADING_1_STYLE_ID,
    HEADING_2_STYLE_ID,
    QUOTE_STYLE_ID,
    TABLE_EMPHASIS_STYLE_ID,
    TABLE_NOTE_STYLE_ID,
    TABLE_NUMBER_STYLE_ID,
    TABLE_TEXT_STYLE_ID,
    TABLE_TITLE_STYLE_ID,
    _apply_body_and_reference_style_fixes,
    _apply_body_edit_plan,
    _apply_intra_paragraph_tracked_replace,
    _apply_tracked_table_formatting,
    _remove_bold_italic_from_normal_body_paragraphs,
)

PARAGRAPH = "Researchers utilise various methods to analyse data, and they utilize those methods often."


def test_validate_edit_accepts_substring_match():
    edit = {"find": "utilize", "replace": "utilise", "reason": "AU spelling"}
    assert _validate_edit(edit, PARAGRAPH) is True


def test_validate_edit_rejects_when_find_absent():
    edit = {"find": "synthesize", "replace": "synthesise", "reason": "AU spelling"}
    assert _validate_edit(edit, PARAGRAPH) is False


def test_validate_edit_rejects_overlong_replace():
    edit = {
        "find": "utilize",
        "replace": "use in a more formal and academic tone to match the style",
        "reason": "rewrite",
    }
    assert _validate_edit(edit, PARAGRAPH) is False


def test_validate_edit_rejects_long_find():
    edit = {
        "find": "utilise various methods to analyse data they",
        "replace": "use these",
        "reason": "too long",
    }
    assert _validate_edit(edit, PARAGRAPH) is False


def test_validate_edit_rejects_identical():
    edit = {"find": "utilize", "replace": "utilize", "reason": "no change"}
    assert _validate_edit(edit, PARAGRAPH) is False


def test_validate_edit_rejects_au_to_us_spelling():
    para = "Researchers utilise a licence to analyse organised practice with generalisability."
    assert _validate_edit(
        {"find": "utilise", "replace": "utilize", "reason": "US spelling"},
        para,
    ) is False
    assert _validate_edit(
        {"find": "licence", "replace": "license", "reason": "US spelling"},
        para,
    ) is False
    assert _validate_edit(
        {"find": "generalisability", "replace": "generalizability", "reason": "US spelling"},
        para,
    ) is False


def test_validate_edit_rejects_empty_strings():
    assert _validate_edit({"find": "", "replace": "x", "reason": ""}, PARAGRAPH) is False
    assert _validate_edit({"find": "utilize", "replace": "", "reason": ""}, PARAGRAPH) is False


def test_validate_edit_rejects_hyphen_removal_same_case():
    """The LLM may not strip a hyphen from a compound word."""
    para = "An evidences-based approach was used."
    edit = {"find": "evidences-based", "replace": "evidencesbased", "reason": "x"}
    assert _validate_edit(edit, para) is False


def test_validate_edit_rejects_hyphen_removal_with_case_change():
    """Regression: `Evidences-Based → evidencesbased` slipped past a
    case-sensitive strip-and-compare. The fix lower-cases both sides
    before stripping ``-`` so case-flipping hyphen-strip edits are
    rejected too."""
    para = "Evidences-Based research is preferred."
    edit = {"find": "Evidences-Based", "replace": "evidencesbased", "reason": "x"}
    assert _validate_edit(edit, para) is False


def test_validate_edit_rejects_hyphen_removal_two_stage():
    para = "The two-stage interview process worked."
    edit = {"find": "two-stage", "replace": "twostage", "reason": "x"}
    assert _validate_edit(edit, para) is False


def test_validate_edit_accepts_legitimate_hyphen_keeping_edit():
    """An edit that keeps the hyphen and fixes a real typo must still pass."""
    para = "The two-stage interveiw process worked."
    edit = {"find": "interveiw", "replace": "interview", "reason": "typo"}
    assert _validate_edit(edit, para) is True


def test_validate_edit_rejects_edits_inside_direct_quotation():
    """JUTLP policy: text inside `"..."` retains the source's spelling."""
    para = 'The participant said "the data was incomplete" during the interview.'
    # `was` sits inside the quoted span — must be rejected even though
    # `was → were` is otherwise a legitimate grammar edit.
    edit = {"find": "was", "replace": "were", "reason": "subj-verb"}
    assert _validate_edit(edit, para) is False


def test_validate_edit_accepts_edits_outside_quotation():
    """Same paragraph, unquoted region: a meaning-preserving edit must
    still be accepted. We use a typo fix rather than a tense/auxiliary
    flip because those are already rejected by an unrelated guard."""
    para = 'The participant said "the data was incomplete" but our analyssis was complete.'
    edit = {"find": "analyssis", "replace": "analysis", "reason": "typo"}
    assert _validate_edit(edit, para) is True


def test_validate_edit_accepts_unquoted_repeat_after_quoted_repeat():
    para = 'The participant said "Russell group" before the Russell group example.'
    edit = {"find": "Russell group", "replace": "Russell Group", "reason": "caps"}
    assert _validate_edit(edit, para) is True


def test_validate_edit_rejects_edits_inside_curly_quotation():
    para = "Smith claims “the methodology was flawed” throughout."
    edit = {"find": "was", "replace": "were", "reason": "subj-verb"}
    assert _validate_edit(edit, para) is False


def test_dedupe_edits_drops_duplicates():
    edits = [
        {"find": "a", "replace": "b", "reason": "x"},
        {"find": "a", "replace": "b", "reason": "y"},
        {"find": "a", "replace": "c", "reason": "z"},
    ]
    result = _dedupe_edits(edits)
    assert len(result) == 2


@patch("app.services.body_llm_edits.call_llm_json")
def test_llm_edits_for_paragraph_caps_edits(mock_call):
    mock_call.return_value = {
        "content": {
            "edits": [
                {"find": f"token{i}", "replace": f"repl{i}", "reason": "r"}
                for i in range(10)
            ]
        }
    }
    paragraph = " ".join(f"token{i}" for i in range(10))
    accepted = _llm_edits_for_paragraph(paragraph)
    assert len(accepted) <= MAX_EDITS_PER_PARAGRAPH


@patch("app.services.body_llm_edits.call_llm_json")
def test_llm_edits_for_paragraph_drops_invalid(mock_call):
    mock_call.return_value = {
        "content": {
            "edits": [
                {"find": "utilize", "replace": "utilise", "reason": "AU"},
                {"find": "not-in-paragraph", "replace": "x", "reason": "bad"},
                {
                    "find": "analyse",
                    "replace": "analyse thoroughly and with great care",
                    "reason": "too long",
                },
            ]
        }
    }
    accepted = _llm_edits_for_paragraph(PARAGRAPH)
    assert len(accepted) == 1
    assert accepted[0]["find"] == "utilize"


@patch("app.services.body_llm_edits.call_llm_json")
def test_llm_edits_for_paragraph_handles_llm_error(mock_call):
    from app.services.body_llm_edits import LLMError

    mock_call.side_effect = LLMError("boom")
    assert _llm_edits_for_paragraph(PARAGRAPH) == []


def test_propagate_repeated_edits_adds_matching_body_paragraph():
    edits = [
        {
            "paragraph_index": 1,
            "find": "Russell group",
            "replace": "Russell Group",
            "reason": "Capitalisation",
        },
    ]
    paragraphs = [
        SimpleNamespace(index=1, text="The Russell group universities were discussed."),
        SimpleNamespace(index=2, text="Students from a Russell group institution responded."),
        SimpleNamespace(index=3, text="No matching phrase here."),
    ]

    propagated = _propagate_repeated_edits(edits, paragraphs)

    assert propagated == [
        {
            "paragraph_index": 1,
            "find": "Russell group",
            "replace": "Russell Group",
            "reason": "Capitalisation",
        },
        {
            "paragraph_index": 2,
            "find": "Russell group",
            "replace": "Russell Group",
            "reason": "Capitalisation",
        },
    ]


def test_propagate_repeated_edits_skips_quote_only_body_paragraph():
    edits = [
        {
            "paragraph_index": 1,
            "find": "Russell group",
            "replace": "Russell Group",
            "reason": "Capitalisation",
        },
    ]
    paragraphs = [
        SimpleNamespace(index=1, text="The Russell group universities were discussed."),
        SimpleNamespace(index=2, text='A participant said "Russell group" verbatim.'),
    ]

    propagated = _propagate_repeated_edits(edits, paragraphs)

    assert propagated == edits


@patch("app.services.body_llm_edits._llm_edits_for_paragraph")
@patch("app.services.body_llm_edits.load_paragraphs")
def test_body_edit_plan_propagates_accepted_exact_repeat(mock_load_paragraphs, mock_llm_edits):
    def fake_llm_edits(paragraph_text):
        if "Russell group universities" not in paragraph_text:
            return []
        return [
            {
                "find": "Russell group",
                "replace": "Russell Group",
                "reason": "Capitalisation",
            },
        ]

    mock_load_paragraphs.return_value = [
        SimpleNamespace(index=0, style="Heading 1", text="Introduction", is_empty=False),
        SimpleNamespace(
            index=1,
            style="Normal",
            text="The Russell group universities were discussed in relation to teaching practice.",
            is_empty=False,
        ),
        SimpleNamespace(
            index=2,
            style="Normal",
            text="Students from a Russell group institution reported similar experiences.",
            is_empty=False,
        ),
    ]
    mock_llm_edits.side_effect = fake_llm_edits

    plan = build_body_edit_plan("unused.docx")

    assert plan["edits"] == [
        {
            "paragraph_index": 1,
            "find": "Russell group",
            "replace": "Russell Group",
            "reason": "Capitalisation",
        },
        {
            "paragraph_index": 2,
            "find": "Russell group",
            "replace": "Russell Group",
            "reason": "Capitalisation",
        },
    ]


def _paragraph_with_runs(run_texts, run_properties_xml=""):
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraph = etree.Element(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p",
        nsmap=nsmap,
    )
    for run_text in run_texts:
        run = etree.SubElement(
            paragraph,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r",
        )
        if run_properties_xml:
            run.append(etree.fromstring(run_properties_xml))
        text_element = etree.SubElement(
            run,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t",
        )
        text_element.text = run_text
    return paragraph


def test_intra_paragraph_tracked_replace_single_run():
    paragraph = _paragraph_with_runs(["We utilize these methods."])
    success = _apply_intra_paragraph_tracked_replace(paragraph, "utilize", "utilise", 42)

    assert success is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    del_elements = paragraph.findall(f"{{{W}}}del")
    ins_elements = paragraph.findall(f"{{{W}}}ins")
    assert len(del_elements) == 1
    assert len(ins_elements) == 1

    del_text = del_elements[0].find(f".//{{{W}}}delText").text
    ins_text = ins_elements[0].find(f".//{{{W}}}t").text
    assert del_text == "utilize"
    assert ins_text == "utilise"

    # Before/after text survive in surrounding runs
    run_texts = [
        t.text for t in paragraph.findall(f"{{{W}}}r/{{{W}}}t")
    ]
    full_kept = "".join(t or "" for t in run_texts)
    assert full_kept == "We  these methods."


def test_intra_paragraph_tracked_replace_reports_false_when_missing():
    paragraph = _paragraph_with_runs(["We utilize these methods."])
    assert _apply_intra_paragraph_tracked_replace(
        paragraph, "synthesize", "synthesise", 1
    ) is False


def test_body_edit_plan_applies_all_repeated_occurrences_in_paragraph(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("The Russell group includes another Russell group example.")
    doc.save(input_path)

    changed = _apply_body_edit_plan(
        str(input_path),
        str(output_path),
        {
            "action": "apply_body_edits",
            "edits": [
                {
                    "paragraph_index": 1,
                    "find": "Russell group",
                    "replace": "Russell Group",
                    "reason": "Capitalisation",
                },
            ],
        },
    )

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    deletions = [el.text for el in root.findall(f".//{{{W}}}delText")]
    insertions = [
        "".join(t.text or "" for t in ins.findall(f".//{{{W}}}t"))
        for ins in root.findall(f".//{{{W}}}ins")
    ]

    assert deletions == ["Russell group", "Russell group"]
    assert insertions == ["Russell Group", "Russell Group"]


def test_body_edit_plan_skips_repeated_occurrence_inside_direct_quote(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(
        'A participant said "Russell group" before the Russell group example.'
    )
    doc.save(input_path)

    changed = _apply_body_edit_plan(
        str(input_path),
        str(output_path),
        {
            "action": "apply_body_edits",
            "edits": [
                {
                    "paragraph_index": 1,
                    "find": "Russell group",
                    "replace": "Russell Group",
                    "reason": "Capitalisation",
                },
            ],
        },
    )

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    deletions = [el.text for el in root.findall(f".//{{{W}}}delText")]
    insertions = [
        "".join(t.text or "" for t in ins.findall(f".//{{{W}}}t"))
        for ins in root.findall(f".//{{{W}}}ins")
    ]

    assert deletions == ["Russell group"]
    assert insertions == ["Russell Group"]


def test_body_style_fix_changes_required_body_styles(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction")
    doc.add_paragraph("Intro body.")
    doc.add_paragraph("Method")
    doc.add_paragraph("Research Design")
    doc.add_paragraph("Participants")
    doc.add_paragraph("Measures")
    doc.add_paragraph("Procedure")
    doc.add_paragraph("Analysis")
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Participant overview")
    doc.add_paragraph("Results")
    doc.add_paragraph("Figure 1")
    doc.add_paragraph("Model output")
    doc.add_paragraph("Discussion")
    doc.add_paragraph("Practical Implications")
    doc.add_paragraph("Theoretical Implications")
    doc.add_paragraph("Limitations and Future Research")
    doc.add_paragraph("Acknowledgements")
    doc.add_paragraph("Conclusions:")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Smith, J. (2025). Example title.")
    doc.save(input_path)

    changed = _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)

    expected_styles = {
        "Introduction": HEADING_1_STYLE_ID,
        "Method": HEADING_1_STYLE_ID,
        "Research Design": HEADING_2_STYLE_ID,
        "Participants": HEADING_2_STYLE_ID,
        "Measures": HEADING_2_STYLE_ID,
        "Procedure": HEADING_2_STYLE_ID,
        "Analysis": HEADING_2_STYLE_ID,
        "Table 1": TABLE_NUMBER_STYLE_ID,
        "Participant overview": TABLE_TITLE_STYLE_ID,
        "Results": HEADING_1_STYLE_ID,
        "Figure 1": FIGURE_NUMBER_STYLE_ID,
        "Model output": FIGURE_TITLE_STYLE_ID,
        "Discussion": HEADING_1_STYLE_ID,
        "Practical Implications": HEADING_2_STYLE_ID,
        "Theoretical Implications": HEADING_2_STYLE_ID,
        "Limitations and Future Research": HEADING_2_STYLE_ID,
        "Acknowledgements": HEADING_1_STYLE_ID,
        "Conclusions:": HEADING_1_STYLE_ID,
    }
    found_styles = {}

    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t"))
        if text in expected_styles:
            p_style = para.find(f"{WQ}pPr/{WQ}pStyle")
            assert p_style is not None
            found_styles[text] = p_style.get(f"{WQ}val")

    assert found_styles == expected_styles


def test_body_style_fix_changes_bold_body_subheading_to_heading2(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Intro body.")
    subheading = doc.add_paragraph()
    subheading.add_run("Theoretical Framework").bold = True
    doc.add_paragraph("This paragraph belongs under the subheading.")
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    changed = _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t"))
        if text == "Theoretical Framework":
            assert para.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == HEADING_2_STYLE_ID
            break
    else:
        raise AssertionError("Theoretical Framework paragraph not found")


def test_body_style_fix_does_not_change_bold_sentence_to_heading2(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    paragraph = doc.add_paragraph()
    paragraph.add_run("This is a bold body sentence.").bold = True
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t"))
        if text == "This is a bold body sentence.":
            assert para.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") != HEADING_2_STYLE_ID
            break


def test_body_style_fix_keeps_table_title_out_of_heading2(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Table 1")
    table_title = doc.add_paragraph()
    table_title.add_run("Participant demographics").bold = True
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t"))
        if text == "Participant demographics":
            assert para.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_TITLE_STYLE_ID
            break


def test_body_style_fix_changes_long_quote_to_quote_style(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    quote = " ".join(["quoted"] * 41)

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph(f"“{quote}“ (Smith, 2025, p. 12).")
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    assert _apply_body_and_reference_style_fixes(str(input_path), str(output_path)) is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
        styles = etree.fromstring(z.read("word/styles.xml"))

    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t"))
        if quote in text:
            assert para.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == QUOTE_STYLE_ID
            break
    else:
        raise AssertionError("Long quote paragraph not found")

    quote_style = next(s for s in styles.findall(f"{WQ}style") if s.get(f"{WQ}styleId") == QUOTE_STYLE_ID)
    assert quote_style.find(f"{WQ}pPr/{WQ}ind").get(f"{WQ}left") == "720"


def test_body_style_fix_does_not_quote_style_mixed_paragraph(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    quote = " ".join(["quoted"] * 41)

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph(f'The authors explain “{quote}“ (Smith, 2025, p. 12).').style = "Normal"
    doc.add_paragraph("References").style = "Heading 1"
    doc.save(input_path)

    _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    for para in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in para.iter(f"{WQ}t"))
        if quote in text:
            assert para.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") != QUOTE_STYLE_ID
            break
    else:
        raise AssertionError("Mixed long quote paragraph not found")


def test_remove_bold_italic_only_changes_normal_body_paragraphs(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    body_run = doc.add_paragraph().add_run("Body text")
    body_run.bold = True
    body_run.italic = True
    heading_run = doc.add_paragraph("Method", style="Heading 2").runs[0]
    heading_run.bold = True
    heading_run.italic = True
    caption_run = doc.add_paragraph("Figure 1. Example", style="Caption").runs[0]
    caption_run.bold = True
    caption_run.italic = True
    doc.add_paragraph("References").style = "Heading 1"
    ref_run = doc.add_paragraph().add_run("Smith, J. (2025). Example title.")
    ref_run.bold = True
    ref_run.italic = True
    doc.save(input_path)

    changed = _remove_bold_italic_from_normal_body_paragraphs(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    paragraphs = list(root.iter(f"{WQ}p"))
    body_run_pr = paragraphs[1].find(f"{WQ}r/{WQ}rPr")
    heading_run_pr = paragraphs[2].find(f"{WQ}r/{WQ}rPr")
    caption_run_pr = paragraphs[3].find(f"{WQ}r/{WQ}rPr")
    ref_run_pr = paragraphs[5].find(f"{WQ}r/{WQ}rPr")

    assert body_run_pr.find(f"{WQ}b") is None
    assert body_run_pr.find(f"{WQ}i") is None

    for r_pr in (heading_run_pr, caption_run_pr, ref_run_pr):
        assert r_pr.find(f"{WQ}b").get(f"{WQ}val") is None
        assert r_pr.find(f"{WQ}i").get(f"{WQ}val") is None


def test_table_formatting_applies_template_styles_directly(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Table 1")
    doc.add_paragraph("student engagement outcomes")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Heading 1"
    table.cell(0, 1).text = "Heading 2"
    table.cell(1, 0).text = "Row heading"
    data_run = table.cell(1, 1).paragraphs[0].add_run("Body value")
    data_run.bold = True
    data_run.italic = True
    data_run.font.name = "Times New Roman"
    data_run.font.size = Pt(9)
    table.cell(1, 1).paragraphs[0].paragraph_format.space_before = Pt(12)
    table.cell(1, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
    table.cell(2, 0).text = "Total responses"
    table.cell(2, 1).text = "42"
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "A6A6A6")
    table.cell(0, 0)._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph("Note. Values are counts.")
    for cell in table._cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(input_path)

    changed = _apply_tracked_table_formatting(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    table_number, table_title, table_note = root.findall(f"{WQ}body/{WQ}p")[:3]
    assert table_number.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_NUMBER_STYLE_ID
    assert table_number.find(f"{WQ}pPr/{WQ}pPrChange") is None
    assert table_title.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_TITLE_STYLE_ID
    assert "".join(t.text or "" for t in table_title.iter(f"{WQ}t")) == "Student Engagement Outcomes"
    assert table_title.find(f"{WQ}pPr/{WQ}pPrChange") is None
    assert table_note.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_NOTE_STYLE_ID

    tbl = root.find(f".//{WQ}tbl")
    tbl_pr = tbl.find(f"{WQ}tblPr")
    assert tbl_pr.find(f"{WQ}tblStyle").get(f"{WQ}val") == "TableGrid"
    assert tbl_pr.find(f"{WQ}tblPrChange") is None
    assert tbl_pr.find(f"{WQ}tblBorders/{WQ}insideV").get(f"{WQ}val") == "none"
    assert tbl_pr.find(f"{WQ}tblBorders/{WQ}insideH").get(f"{WQ}val") == "none"
    assert tbl_pr.find(f"{WQ}tblBorders/{WQ}top").get(f"{WQ}val") == "single"
    assert tbl_pr.find(f"{WQ}tblBorders/{WQ}bottom").get(f"{WQ}val") == "single"

    rows = tbl.findall(f"{WQ}tr")
    header_cell = rows[0].findall(f"{WQ}tc")[0]
    row_heading_cell = rows[1].findall(f"{WQ}tc")[0]
    data_cell = rows[1].findall(f"{WQ}tc")[1]
    total_label_cell = rows[2].findall(f"{WQ}tc")[0]
    total_value_cell = rows[2].findall(f"{WQ}tc")[1]

    assert header_cell.find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_EMPHASIS_STYLE_ID
    assert row_heading_cell.find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_TEXT_STYLE_ID
    assert data_cell.find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_TEXT_STYLE_ID
    assert total_label_cell.find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_EMPHASIS_STYLE_ID
    assert total_value_cell.find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_EMPHASIS_STYLE_ID
    assert header_cell.find(f"{WQ}tcPr/{WQ}tcPrChange") is None
    assert header_cell.find(f"{WQ}tcPr/{WQ}shd") is None
    assert header_cell.find(f".//{WQ}pPr/{WQ}jc") is None
    assert data_cell.find(f".//{WQ}pPr/{WQ}jc") is None
    assert data_cell.find(f".//{WQ}pPr/{WQ}spacing") is None
    assert header_cell.find(f".//{WQ}pPr/{WQ}pPrChange") is None
    assert data_cell.find(f".//{WQ}pPr/{WQ}pPrChange") is None
    assert header_cell.find(f"{WQ}tcPr/{WQ}tcBorders/{WQ}bottom").get(f"{WQ}val") == "single"
    assert header_cell.find(f"{WQ}tcPr/{WQ}tcBorders/{WQ}top") is None
    assert data_cell.find(f"{WQ}tcPr/{WQ}tcBorders/{WQ}bottom").get(f"{WQ}val") == "nil"
    assert total_value_cell.find(f"{WQ}tcPr/{WQ}tcBorders/{WQ}top").get(f"{WQ}val") == "nil"

    data_r_pr = data_cell.find(f".//{WQ}rPr")
    assert data_r_pr.find(f"{WQ}b") is None
    assert data_r_pr.find(f"{WQ}i") is None
    assert data_r_pr.find(f"{WQ}rFonts") is None
    assert data_r_pr.find(f"{WQ}sz") is None
    assert data_r_pr.find(f"{WQ}rPrChange") is None


def test_table_formatting_splits_combined_caption_paragraphs(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Table 1. mean rubric scores by cohort and dimension")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Header"
    doc.add_paragraph("Table 2: attribution distribution at week 5")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Header"
    doc.add_paragraph("Table A1. appendix outcome distribution")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Header"
    doc.save(input_path)

    changed = _apply_tracked_table_formatting(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    body_paragraphs = root.findall(f"{WQ}body/{WQ}p")
    texts = ["".join(t.text or "" for t in p.iter(f"{WQ}t")) for p in body_paragraphs]
    styles = [p.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") for p in body_paragraphs]

    assert texts[:6] == [
        "Table 1.",
        "Mean Rubric Scores by Cohort and Dimension",
        "Table 2:",
        "Attribution Distribution at Week 5",
        "Table A1.",
        "Appendix Outcome Distribution",
    ]
    assert styles[:6] == [
        TABLE_NUMBER_STYLE_ID,
        TABLE_TITLE_STYLE_ID,
        TABLE_NUMBER_STYLE_ID,
        TABLE_TITLE_STYLE_ID,
        TABLE_NUMBER_STYLE_ID,
        TABLE_TITLE_STYLE_ID,
    ]
    assert body_paragraphs[0].find(f"{WQ}r/{WQ}rPr") is None
    assert body_paragraphs[1].find(f"{WQ}r/{WQ}rPr") is None


def test_table_formatting_does_not_split_body_sentence_before_paragraph(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Table 1. presents the variables used in the next analysis.")
    doc.add_paragraph("This paragraph separates the prose from the table.")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Header"
    doc.save(input_path)

    _apply_tracked_table_formatting(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    body_paragraphs = root.findall(f"{WQ}body/{WQ}p")
    text = "".join(t.text or "" for t in body_paragraphs[0].iter(f"{WQ}t"))
    assert text == "Table 1. presents the variables used in the next analysis."
    assert body_paragraphs[0].find(f"{WQ}pPr/{WQ}pStyle") is None


def test_table_formatting_does_not_split_prose_like_caption_before_table(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Table 1. presents the variables used in the next analysis.")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Header"
    doc.save(input_path)

    _apply_tracked_table_formatting(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    body_paragraphs = root.findall(f"{WQ}body/{WQ}p")
    text = "".join(t.text or "" for t in body_paragraphs[0].iter(f"{WQ}t"))
    assert text == "Table 1. presents the variables used in the next analysis."
    assert body_paragraphs[0].find(f"{WQ}pPr/{WQ}pStyle") is None


def test_body_style_fix_splits_combined_figure_caption_before_image(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Figure 1 Average OSCE overall grade (%) comparison based on AI Usage")
    doc.add_picture("app/static/oapa-logo.png")
    doc.save(input_path)

    changed = _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    body_paragraphs = root.findall(f"{WQ}body/{WQ}p")
    texts = ["".join(t.text or "" for t in p.iter(f"{WQ}t")) for p in body_paragraphs]
    styles = [
        p.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val")
        if p.find(f"{WQ}pPr/{WQ}pStyle") is not None else ""
        for p in body_paragraphs
    ]

    assert texts[1:3] == [
        "Figure 1",
        "Average OSCE Overall Grade (%) Comparison Based on AI Usage",
    ]
    assert styles[1:3] == [FIGURE_NUMBER_STYLE_ID, FIGURE_TITLE_STYLE_ID]
    assert body_paragraphs[3].find(f".//{WQ}drawing") is not None


def test_body_style_fix_uses_real_figure_style_ids_for_split_caption(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Figure 1")
    doc.add_paragraph("Average OSCE Overall Grade (%) Comparison Based on AI Usage")
    doc.add_picture("app/static/oapa-logo.png")
    doc.save(input_path)

    changed = _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    assert changed is True
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
        styles_root = etree.fromstring(z.read("word/styles.xml"))

    style_ids = {
        style.get(f"{WQ}styleId")
        for style in styles_root.findall(f"{WQ}style")
    }
    body_paragraphs = root.findall(f"{WQ}body/{WQ}p")
    styles = [
        p.find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val")
        if p.find(f"{WQ}pPr/{WQ}pStyle") is not None else ""
        for p in body_paragraphs
    ]

    assert styles[1:3] == [FIGURE_NUMBER_STYLE_ID, FIGURE_TITLE_STYLE_ID]
    assert styles[1] in style_ids
    assert styles[2] in style_ids
    assert " " not in styles[1]
    assert " " not in styles[2]


def test_body_style_fix_does_not_split_prose_like_figure_reference_before_image(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Figure 1 shows average OSCE overall grade by AI usage.")
    doc.add_picture("app/static/oapa-logo.png")
    doc.save(input_path)

    _apply_body_and_reference_style_fixes(str(input_path), str(output_path))

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    body_paragraphs = root.findall(f"{WQ}body/{WQ}p")
    text = "".join(t.text or "" for t in body_paragraphs[1].iter(f"{WQ}t"))
    style = body_paragraphs[1].find(f"{WQ}pPr/{WQ}pStyle").get(f"{WQ}val")
    assert text == "Figure 1 shows average OSCE overall grade by AI usage."
    assert style != FIGURE_NUMBER_STYLE_ID


def test_table_formatting_removes_vertical_borders_from_wide_unmerged_tables(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=5)
    table.cell(0, 0).text = "Header"
    table.cell(1, 0).text = "Row"
    symbol_run = table.cell(1, 1).paragraphs[0].add_run("✓")
    symbol_run.font.name = "Segoe UI Symbol"
    doc.save(input_path)

    changed = _apply_tracked_table_formatting(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    tbl = root.find(f".//{WQ}tbl")
    assert tbl.find(f"{WQ}tblPr/{WQ}tblPrChange") is None
    assert tbl.find(f"{WQ}tblPr/{WQ}tblBorders/{WQ}top").get(f"{WQ}val") == "single"
    assert tbl.find(f"{WQ}tblPr/{WQ}tblBorders/{WQ}bottom").get(f"{WQ}val") == "single"
    assert tbl.find(f"{WQ}tblPr/{WQ}tblBorders/{WQ}insideH").get(f"{WQ}val") == "none"
    assert tbl.find(f"{WQ}tblPr/{WQ}tblBorders/{WQ}insideV").get(f"{WQ}val") == "none"
    assert tbl.find(f".//{WQ}tcPr/{WQ}tcPrChange") is None
    rows = tbl.findall(f"{WQ}tr")
    assert rows[0].find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_EMPHASIS_STYLE_ID
    assert rows[1].find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_TEXT_STYLE_ID
    assert tbl.find(f".//{WQ}rPr/{WQ}rFonts").get(f"{WQ}ascii") == "Segoe UI Symbol"


def test_table_formatting_demotes_body_emphasis_unless_total(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document("app/domain/JUTLP Template 2026.docx")
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Header"
    table.cell(1, 1).text = "Second header"
    table.cell(1, 1).paragraphs[0].style = "Table Emphasis"
    table.cell(2, 1).text = "Value"
    doc.save(input_path)

    changed = _apply_tracked_table_formatting(str(input_path), str(output_path))

    assert changed is True

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WQ = f"{{{W}}}"
    with zipfile.ZipFile(output_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))

    tbl = root.findall(f".//{WQ}tbl")[-1]
    second_header = tbl.findall(f"{WQ}tr")[1].findall(f"{WQ}tc")[1]
    assert second_header.find(f".//{WQ}pPr/{WQ}pStyle").get(f"{WQ}val") == TABLE_TEXT_STYLE_ID


# Removed: test_table_formatting_replaces_table_style_definitions_with_template
# relied on a real-world manuscript fixture (Zhou+et+al+Final-5.docx) that
# was never committed. The three other table-formatting tests above cover
# tracked-change application, wide-border preservation, and not demoting
# existing emphasis. If the style-definition-replacement edge case ever
# regresses, build a synthetic fixture inline rather than re-importing the
# Zhou manuscript.
