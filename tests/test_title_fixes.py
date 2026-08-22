"""Tests for title fixes: acronym-safe title casing and not restyling the
article title as an affiliation."""

from docx import Document

from app.services.output_generation_samfix import (
    _to_title_case_title,
    abstractFound,
)

# ── acronym-safe title casing ────────────────────────────────────────────────

def test_paper_acronyms_preserved_in_mixed_case_title():
    # A lone all-caps word in a normally-cased title is an acronym (NFSC, OBE)
    # and must NOT be lowered to "Nfsc"/"Obe".
    out = _to_title_case_title(
        "A Three-Year Evaluation of the NFSC Pedagogy and OBE Outcomes"
    )
    assert "NFSC" in out
    assert "OBE" in out
    assert "Nfsc" not in out


def test_shouted_all_caps_title_is_title_cased():
    # An all-caps title is "shouted": title-case it, preserving only allow-list
    # acronyms (regression for the existing behaviour).
    assert _to_title_case_title("THE IMPACT OF AI: AN APA STUDY IN STEM") == (
        "The Impact of AI: An APA Study in STEM"
    )


def test_title_case_treats_first_word_after_section_number_as_first_word():
    assert _to_title_case_title(
        "1.1 the Scarcity of Longitudinal Evaluation in Engineering Education"
    ) == "1.1 The Scarcity of Longitudinal Evaluation in Engineering Education"


# ── title is never restyled as an affiliation ────────────────────────────────

def test_title_not_flagged_as_affiliation(tmp_path):
    """The title "A Three-Year ..." starts with "A ", which the affiliation
    heuristic mistakes for an affiliation marker. The title's paragraph index
    must not be queued for an AuthorAffiliations restyle."""
    doc = Document()
    doc.add_paragraph("A Three-Year Evaluation of the NFSC Pedagogy")  # title (Normal)
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("Background: this is the abstract body text.")
    doc.add_paragraph("Introduction", style="Heading 1")
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    state = abstractFound(str(path))
    assert state["affiliations_found"] is False
    assert state["affiliations_end_pos"] == -1
    assert 0 not in state["affiliation_style_fix_positions"]
