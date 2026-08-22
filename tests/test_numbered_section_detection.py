"""Regression tests for section/subsection detection on numbered or aliased
headings — the Phase-1 validators run before the heading-correction pass
removes numbers, so detection must tolerate "3. Methodology" / "3.1 Research
Design" headings. (Previously these made the Method/Discussion subsection
checks report present subsections as missing.)"""

from docx import Document

from app.services.document_analysis_services import (
    extract_subsections,
    get_section_bounds,
    load_paragraphs,
)
from app.services.jutlp_validator import validate


def _doc(tmp_path):
    doc = Document()
    doc.add_paragraph("A Study Title")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Intro body text goes here for context.")
    doc.add_paragraph("3. Methodology", style="Heading 1")   # numbered alias
    doc.add_paragraph("3.1 Research Design", style="Heading 2")
    doc.add_paragraph("Design body.")
    doc.add_paragraph("3.2 Participants", style="Heading 2")
    doc.add_paragraph("Participants body.")
    doc.add_paragraph("4. Results", style="Heading 1")
    doc.add_paragraph("Results body.")
    path = tmp_path / "doc.docx"
    doc.save(str(path))
    return str(path)


def test_section_bounds_match_numbered_alias_heading(tmp_path):
    paras = load_paragraphs(_doc(tmp_path))
    # "3. Methodology" must resolve to the Method section.
    assert get_section_bounds(paras, "Method") is not None


def test_subsections_found_under_numbered_section(tmp_path):
    paras = load_paragraphs(_doc(tmp_path))
    subs = extract_subsections(paras, "Method")
    assert any("Research Design" in s for s in subs)
    assert any("Participants" in s for s in subs)


def test_present_subsection_not_flagged_missing(tmp_path):
    results = validate(_doc(tmp_path))["results"]
    met = {r["rule_id"]: r for r in results if r["rule_id"].startswith("MET00")}
    # Research Design + Participants are present (numbered) -> must pass.
    rd = next(r for r in met.values() if "Research Design" in r["message"])
    pt = next(r for r in met.values() if "Participants" in r["message"])
    assert rd["status"] == "pass"
    assert pt["status"] == "pass"
