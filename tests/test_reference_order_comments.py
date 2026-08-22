"""Tests for the reference-list alphabetical-order check."""

import zipfile

from docx import Document
from lxml import etree

from app.services.reference_order_comments import (
    _first_inversion,
    apply_reference_order_comments,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _refs_doc(tmp_path, entries, heading="References"):
    doc = Document()
    if heading is not None:
        doc.add_paragraph(heading, style="Heading 1")
    for e in entries:
        doc.add_paragraph(e, style="Normal")
    path = tmp_path / "doc.docx"
    doc.save(str(path))
    return str(path)


def _comment_texts(path):
    with zipfile.ZipFile(path) as z:
        if "word/comments.xml" not in z.namelist():
            return []
        root = etree.fromstring(z.read("word/comments.xml"))
    return ["".join(t.text or "" for t in c.iter(f"{WQ}t")) for c in root.findall(f"{WQ}comment")]


_ORDERED = [
    "Adams, B. (2020). A study. Journal A.",
    "Brown, C. (2019). Another study. Journal B.",
    "Carter, D. (2021). Third study. Journal C.",
]
_SHUFFLED = [
    "Brown, C. (2019). Another study. Journal B.",
    "Adams, B. (2020). A study. Journal A.",
    "Carter, D. (2021). Third study. Journal C.",
]


# ── detection ────────────────────────────────────────────────────────────────

def test_inversion_none_when_ordered():
    assert _first_inversion(_ORDERED) is None


def test_inversion_detected_when_out_of_order():
    assert _first_inversion(_SHUFFLED) == ("Brown", "Adams")


def test_same_first_author_is_not_an_inversion():
    # Same surname (different years) is a tie, not an ordering error.
    entries = [
        "Smith, J. (2018). Early. Journal.",
        "Smith, J. (2020). Later. Journal.",
        "Wong, K. (2019). Other. Journal.",
    ]
    assert _first_inversion(entries) is None


# ── comment pass ─────────────────────────────────────────────────────────────

def test_ordered_list_gets_no_comment(tmp_path):
    out = tmp_path / "out.docx"
    _, actions = apply_reference_order_comments(_refs_doc(tmp_path, _ORDERED), str(out))
    assert actions == []
    assert _comment_texts(str(out)) == []


def test_out_of_order_list_gets_one_comment(tmp_path):
    out = tmp_path / "out.docx"
    _, actions = apply_reference_order_comments(_refs_doc(tmp_path, _SHUFFLED), str(out))
    assert len(actions) == 1
    text = _comment_texts(str(out))[0]
    assert "alphabetical" in text.lower()
    assert "Adams" in text and "Brown" in text


def test_too_few_entries_no_comment(tmp_path):
    out = tmp_path / "out.docx"
    _, actions = apply_reference_order_comments(
        _refs_doc(tmp_path, _SHUFFLED[:2]), str(out)
    )
    assert actions == []


def test_no_references_heading_no_comment(tmp_path):
    out = tmp_path / "out.docx"
    _, actions = apply_reference_order_comments(
        _refs_doc(tmp_path, _SHUFFLED, heading=None), str(out)
    )
    assert actions == []
