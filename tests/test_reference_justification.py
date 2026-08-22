"""Reference entries must be justified (w:jc=both), not left-aligned.

Two layers of protection:
  1. The canonical JUTLP template (the .docx the bot reads its styles from)
     carries jc=both for the APA 7 reference style.
  2. The body/reference style pass replaces an author-supplied reference style
     definition — which commonly sets jc=left — with the canonical justified
     one, so a document that already defines the style still ends up justified.
"""

import zipfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from app.services.output_generation_samfix import (
    _TEMPLATE_PATH,
    _apply_body_and_reference_style_fixes,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"
_REF_STYLE_NAME = "APA 7 Reference List Entry"


def _ref_style_jc(styles_xml: bytes) -> str | None:
    """Return the w:jc value of the APA 7 reference style, or None."""
    root = etree.fromstring(styles_xml)
    for style_el in root.findall(f"{WQ}style"):
        name_el = style_el.find(f"{WQ}name")
        if name_el is None or name_el.get(f"{WQ}val") != _REF_STYLE_NAME:
            continue
        jc = style_el.find(f"{WQ}pPr/{WQ}jc")
        return jc.get(f"{WQ}val") if jc is not None else None
    return None


def test_template_reference_style_is_justified():
    with zipfile.ZipFile(_TEMPLATE_PATH, "r") as z:
        assert _ref_style_jc(z.read("word/styles.xml")) == "both"


def test_left_aligned_author_reference_style_is_replaced_with_justified(tmp_path):
    src = tmp_path / "refs.docx"
    out = tmp_path / "refs_out.docx"

    doc = Document()
    # Author defines the reference style themselves, left-aligned.
    ref_style = doc.styles.add_style(_REF_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph(
        "Smith, J. (2020). A study of things. Journal of Things, 1(1), 1-10.",
        style=_REF_STYLE_NAME,
    )
    doc.save(str(src))

    # Sanity: the author definition starts left-aligned.
    with zipfile.ZipFile(src, "r") as z:
        assert _ref_style_jc(z.read("word/styles.xml")) == "left"

    _apply_body_and_reference_style_fixes(str(src), str(out))

    with zipfile.ZipFile(out, "r") as z:
        assert _ref_style_jc(z.read("word/styles.xml")) == "both"
