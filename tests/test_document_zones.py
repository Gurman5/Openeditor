import zipfile

from docx import Document
from lxml import etree

from app.services.decimal_corrections import apply_decimal_corrections
from app.services.document_zones import iter_paragraphs_with_zone, para_plain_text
from app.services.number_word_corrections import apply_number_word_corrections


def _make_doc(path, paragraphs):
    doc = Document()
    for text, style in paragraphs:
        doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    doc.save(path)


def _zones(path):
    with zipfile.ZipFile(path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    return [(para_plain_text(p), zone) for p, zone in iter_paragraphs_with_zone(root)]


def test_unstyled_intro_and_reference_aliases_set_zones(tmp_path):
    path = tmp_path / "doc.docx"
    _make_doc(path, [
        ("Abstract text.", None),
        ("1. Introduction.", None),
        ("Body text.", None),
        ("Works Cited", None),
        ("Smith, J. (2024). Reference entry.", None),
    ])

    assert _zones(path) == [
        ("Abstract text.", "front"),
        ("1. Introduction.", "body"),
        ("Body text.", "body"),
        ("Works Cited", "outside"),
        ("Smith, J. (2024). Reference entry.", "outside"),
    ]


def test_unstyled_references_without_intro_starts_as_body_then_exits(tmp_path):
    path = tmp_path / "doc.docx"
    _make_doc(path, [
        ("Body before a missing Introduction.", None),
        ("References", None),
        ("Smith, J. (2024). Reference entry.", None),
    ])

    assert _zones(path) == [
        ("Body before a missing Introduction.", "body"),
        ("References", "outside"),
        ("Smith, J. (2024). Reference entry.", "outside"),
    ]


def test_body_sentence_mentioning_references_is_not_boundary(tmp_path):
    path = tmp_path / "doc.docx"
    _make_doc(path, [
        ("Introduction", None),
        ("References to 5 participants are included in the analysis.", None),
    ])

    assert _zones(path) == [
        ("Introduction", "body"),
        ("References to 5 participants are included in the analysis.", "body"),
    ]


def test_number_words_skip_unstyled_references_section(tmp_path):
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, [
        ("Introduction", None),
        ("We surveyed 5 students in the study.", None),
        ("References", None),
        ("Smith, J. (2024). 6 studentss in a title.", None),
    ])

    _, actions = apply_number_word_corrections(str(src), str(out), 1)

    assert [(a["original"], a["replacement"]) for a in actions] == [("5", "five")]


def test_decimal_checks_skip_unstyled_acknowledgements_and_refs(tmp_path):
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, [
        ("Introduction", None),
        ("In the sample, 21,7% selected the option.", None),
        ("Acknowledgements", None),
        ("We thank 42,1% of contributors.", None),
        ("Reference List", None),
        ("Smith, J. (2024). A 36,1% finding.", None),
    ])

    _, actions = apply_decimal_corrections(str(src), str(out), 1)

    assert [a["value"] for a in actions if a["check"] == "comma_decimal"] == ["21,7"]
