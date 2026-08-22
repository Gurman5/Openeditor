import zipfile

from docx import Document
from lxml import etree

from app.services.language_corrections import (
    WQ,
    apply_au_spelling_corrections,
    summarize_spelling_correction_repeats,
)


def _docx_xml(path, part):
    with zipfile.ZipFile(path, "r") as z:
        return etree.fromstring(z.read(part))


def test_spelling_summary_groups_repeated_fix_once():
    corrections = [
        {"original": "analyze", "replacement": "analyse"},
        {"original": "analyze", "replacement": "analyse"},
        {"original": "color", "replacement": "colour"},
    ]

    summary = summarize_spelling_correction_repeats(corrections)

    assert summary == [
        {"original": "analyze", "replacement": "analyse", "grouped_repeats": 1},
        {"original": "color", "replacement": "colour", "grouped_repeats": 0},
    ]


def test_spelling_summary_groups_case_variants_of_same_fix():
    corrections = [
        {"original": "Analyze", "replacement": "Analyse"},
        {"original": "analyze", "replacement": "analyse"},
    ]

    summary = summarize_spelling_correction_repeats(corrections)

    assert summary == [
        {"original": "Analyze", "replacement": "Analyse", "grouped_repeats": 1},
    ]


def test_repeated_spelling_fix_gets_one_author_query_comment(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("We analyze the data and then analyze the pattern.")
    doc.save(source)

    _, corrections = apply_au_spelling_corrections(str(source), str(output))

    assert corrections == [
        {"original": "analyze", "replacement": "analyse"},
        {"original": "analyze", "replacement": "analyse"},
    ]

    document_root = _docx_xml(output, "word/document.xml")
    comments_root = _docx_xml(output, "word/comments.xml")
    comments = comments_root.findall(f"{WQ}comment")
    comment_refs = document_root.findall(f".//{WQ}commentReference")
    insertions = [
        "".join(t.text or "" for t in ins.findall(f".//{WQ}t"))
        for ins in document_root.findall(f".//{WQ}ins")
    ]

    assert insertions == ["analyse", "analyse"]
    assert len(comments) == 1
    assert len(comment_refs) == 1
    comment_text = " ".join(t.text or "" for t in comments[0].iter(f"{WQ}t"))
    assert 'US spelling "analyze" was changed to Australian English "analyse".' in comment_text
    assert "1 later occurrence of the same spelling issue was fixed" in comment_text


def test_generalizability_is_changed_to_australian_spelling(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("The generalizability of the model was tested.")
    doc.save(source)

    _, corrections = apply_au_spelling_corrections(str(source), str(output))

    assert corrections == [
        {"original": "generalizability", "replacement": "generalisability"},
    ]


def test_us_spelling_inside_direct_quotation_is_skipped(tmp_path):
    """A US-spelled word INSIDE quoted text must NOT be rewritten —
    direct quotations retain the source's original spelling."""
    source = tmp_path / "src.docx"
    output = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph(
        'The author wrote "we analyze the data" in the original paper.'
    )
    doc.save(source)

    _, corrections = apply_au_spelling_corrections(str(source), str(output))
    assert corrections == []


def test_us_spelling_outside_quote_still_fixed_when_quote_present(tmp_path):
    """An unquoted US spelling in the same paragraph as a quoted one
    must still be corrected."""
    source = tmp_path / "src.docx"
    output = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph(
        'The author wrote "we analyze the data" but we instead analyze further.'
    )
    doc.save(source)

    _, corrections = apply_au_spelling_corrections(str(source), str(output))
    # Exactly one fix — the second `analyze` outside the quote.
    assert corrections == [{"original": "analyze", "replacement": "analyse"}]
