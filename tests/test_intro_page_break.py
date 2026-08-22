"""Tests for the front-page / body page break before the Introduction."""

import zipfile

from docx import Document
from lxml import etree

from app.services.output_generation_samfix import _apply_intro_page_break

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _intro_has_page_break(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    for p in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in p.iter(f"{WQ}t"))
        if "Introduction" in text:
            pPr = p.find(f"{WQ}pPr")
            return pPr is not None and pPr.find(f"{WQ}pageBreakBefore") is not None
    return False


def _doc(tmp_path, intro_heading):
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("Abstract body text.", style="Normal")
    doc.add_paragraph("Keywords: a; b; c", style="Normal")
    doc.add_paragraph(intro_heading, style="Heading 1")
    doc.add_paragraph("Intro body.", style="Normal")
    src = tmp_path / "in.docx"
    doc.save(str(src))
    return str(src)


def test_page_break_inserted_before_plain_introduction(tmp_path):
    src = _doc(tmp_path, "Introduction")
    out = tmp_path / "out.docx"
    assert _apply_intro_page_break(src, str(out)) is True
    assert _intro_has_page_break(str(out))


def test_page_break_inserted_before_numbered_introduction(tmp_path):
    """Regression: a numbered "1. Introduction" heading must still get the
    front-page/body page break (text match previously required exactly
    "introduction")."""
    src = _doc(tmp_path, "1. Introduction")
    out = tmp_path / "out.docx"
    assert _apply_intro_page_break(src, str(out)) is True
    assert _intro_has_page_break(str(out))
