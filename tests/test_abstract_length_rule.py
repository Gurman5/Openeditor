"""Tests for the JUTLP abstract length rule: the abstract must fit lines 7–23
of the front page (max 17 lines) AND be at most 250 words."""

from docx import Document

from app.domain.canonical_jultp_template import CANONICAL_STRUCTURE
from app.services.document_analysis_services import (
    estimate_line_count,
    extract_abstract,
    load_paragraphs,
)
from app.services.jutlp_validator import check_abstract
from app.services.output_generation_samfix import abstractFormatCheck, abstractFound

_FRONT = CANONICAL_STRUCTURE["front_page"]


def _parsed(found=True, word_count=0, line_count=0):
    return {"front_page": {"abstract": {
        "found": found, "word_count": word_count, "line_count": line_count,
    }}}


def _fp008(results):
    return next(r for r in results if r["rule_id"] == "FP008")


# ── canonical wiring ─────────────────────────────────────────────────────────

def test_rule_constants_match_template():
    assert _FRONT["abstract_max_words"] == 250
    assert _FRONT["abstract_line_region"] == [7, 23]
    # lines 7–23 inclusive == 17 lines
    assert _FRONT["abstract_max_lines"] == 23 - 7 + 1 == 17


# ── line estimator ───────────────────────────────────────────────────────────

def test_estimate_line_count_uses_words_and_hard_breaks():
    wpl = _FRONT["abstract_words_per_line_estimate"]
    assert estimate_line_count("", wpl) == 0
    # 250 words at the template estimate ≈ the 17-line limit
    assert estimate_line_count(" ".join(["w"] * 250), wpl) == 17
    # hard line breaks dominate when text is padded with manual breaks
    assert estimate_line_count("a\n" * 19 + "a", wpl) == 20


# ── validator FP008 ──────────────────────────────────────────────────────────

def test_abstract_within_limits_passes():
    r = _fp008(check_abstract(_parsed(word_count=200, line_count=14)))
    assert r["status"] == "pass"


def test_abstract_at_the_limit_passes():
    r = _fp008(check_abstract(_parsed(word_count=250, line_count=17)))
    assert r["status"] == "pass"


def test_abstract_over_word_limit_fails():
    r = _fp008(check_abstract(_parsed(word_count=300, line_count=20)))
    assert r["status"] == "fail"
    assert "300 words" in r["message"]
    assert "250" in r["message"]


def test_abstract_over_line_limit_under_word_limit_fails():
    # 200 words but spread over many lines (manual breaks) — within the word
    # limit yet overflowing the lines 7–23 region.
    r = _fp008(check_abstract(_parsed(word_count=200, line_count=21)))
    assert r["status"] == "fail"
    assert "lines 7" in r["message"] or "7–23" in r["message"]


def test_abstract_not_found_fails():
    r = _fp008(check_abstract(_parsed(found=False)))
    assert r["status"] == "fail"


# ── extract_abstract measures non-template (Normal) abstracts ────────────────

def test_normal_styled_abstract_is_measured(tmp_path):
    """Regression: an abstract that is NOT styled 'Front Page Text' (authors
    routinely leave it Normal) must still be counted — it was previously read
    as 0 words and never flagged."""
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph(" ".join(["word"] * 260), style="Normal")
    doc.add_paragraph("Keywords: alpha; beta; gamma", style="Normal")
    doc.add_paragraph("Introduction", style="Heading 1")
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    ab = extract_abstract(load_paragraphs(str(path)))
    assert ab["found"] is True
    assert ab["paragraph_count"] == 1          # keyword line excluded
    assert ab["word_count"] == 260             # over the 250 limit
    assert ab["line_count"] > _FRONT["abstract_max_lines"]


def test_abstract_stops_at_introduction_heading(tmp_path):
    """The body (after Introduction) must not be swallowed into the abstract."""
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("Short abstract sentence.", style="Normal")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(" ".join(["body"] * 500), style="Normal")
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    ab = extract_abstract(load_paragraphs(str(path)))
    assert ab["paragraph_count"] == 1
    assert ab["word_count"] < 10               # only the abstract sentence


# ── Sam's abstract-length comment is standalone (survives comment de-dup) ─────

def _long_abstract_doc(tmp_path):
    doc = Document()
    doc.add_paragraph("A Study of Things")           # title
    doc.add_paragraph("Abstract", style="Heading 1")  # abstract heading (by text)
    doc.add_paragraph(" ".join(["word"] * 260))       # 260-word body (over limit)
    doc.add_paragraph("Keywords")
    doc.add_paragraph("Introduction", style="Heading 1")
    path = tmp_path / "doc.docx"
    doc.save(str(path))
    return str(path)


def test_abstract_length_comment_is_a_standalone_comment(tmp_path):
    """The length flag must be its OWN comment (not folded into the merge/style
    bundle) so the pipeline's Sam-comment de-dup can't drop it with the bundle."""
    path = _long_abstract_doc(tmp_path)
    plan = abstractFormatCheck(path, abstractFound(path))

    length_comments = [
        c for c in plan.get("comments", [])
        if c["message"].startswith("Abstract is too long")
    ]
    assert len(length_comments) == 1
    msg = length_comments[0]["message"]
    # standalone: not concatenated with the merge / style messages
    assert "Abstract merged" not in msg
    assert "Tracked style change" not in msg
    assert "250 words" in msg
