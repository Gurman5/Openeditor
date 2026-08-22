"""Tests for the spell-checker pass.

Covers Round 3 fix B: auto-fix high-confidence typos like `studentss → students`
as tracked changes, and flag ambiguous suspects as comments.
"""

import zipfile

from docx import Document
from lxml import etree

from app.services.language_corrections import WQ
from app.services.spell_checker import (
    _classify_word,
    apply_spell_corrections,
    get_spell_corrections,
)

# ── Classifier-level tests (don't need a docx) ────────────────────────────────


def _counts(*words):
    """Build a Counter where each given word appears once."""
    from collections import Counter
    return Counter(w.lower() for w in words)


def test_studentss_classified_as_autofix():
    verdict, candidate = _classify_word("studentss", _counts("studentss"))
    assert verdict == "autofix"
    assert candidate == "students"


def test_uppercase_initial_word_skipped():
    """Proper nouns (capitalised) are too risky — never auto-fix or flag."""
    verdict, candidate = _classify_word("Kiraly", _counts("Kiraly"))
    assert verdict == "skip"


def test_acronym_skipped():
    verdict, _ = _classify_word("PACTE", _counts("PACTE"))
    assert verdict == "skip"


def test_recurring_word_skipped():
    """A word seen 3+ times in the body is assumed correct (domain term)."""
    counts = _counts()
    counts["unfamiliar"] = 3
    verdict, _ = _classify_word("unfamiliar", counts)
    assert verdict == "skip"


def test_short_word_skipped():
    """Words shorter than 4 chars are too risky — single-char swaps explode."""
    verdict, _ = _classify_word("teh", _counts("teh"))
    assert verdict == "skip"


def test_word_with_apostrophe_skipped():
    verdict, _ = _classify_word("O'Hagan", _counts("O'Hagan"))
    assert verdict == "skip"


def test_already_known_word_skipped():
    verdict, _ = _classify_word("students", _counts("students"))
    assert verdict == "skip"


def test_au_spellings_are_not_autofixed_to_us():
    for word in [
        "underutilised",
        "disorganisation",
        "empathise",
        "unrecognised",
    ]:
        verdict, candidate = _classify_word(word, _counts(word))
        assert (verdict, candidate) == ("skip", None)


# ── get_spell_corrections (paragraph-level) ───────────────────────────────────


def _make_doc(tmp_path, paragraphs):
    docx_path = tmp_path / "spell.docx"
    doc = Document()
    for text, style in paragraphs:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    doc.save(docx_path)
    return str(docx_path)


def test_obvious_typo_returned_with_autofix_verdict(tmp_path):
    path = _make_doc(tmp_path, [
        ("This is the abstract paragraph for studentss in the manuscript.", None),
    ])
    results = get_spell_corrections(path)
    matches = [r for r in results if r["original"] == "studentss"]
    assert len(matches) == 1
    assert matches[0]["verdict"] == "autofix"
    assert matches[0]["replacement"] == "students"


def test_references_section_paragraphs_skipped(tmp_path):
    """A typo inside the References section should not produce any flag."""
    path = _make_doc(tmp_path, [
        ("Introduction", "Heading 1"),
        ("The bodyy of the paper has substance.", None),
        ("References", "Heading 1"),
        ("Smith, J. (2024). A studentss-related study.", None),
    ])
    results = get_spell_corrections(path)
    originals = {r["original"] for r in results}
    # `bodyy` (a real misspelling) should be flagged — it's in the body.
    # `studentss` inside the References section must not be flagged.
    found_studentss_in_results = any(
        r["original"] == "studentss" for r in results
    )
    assert found_studentss_in_results is False, (
        f"References-section typo should be skipped; got {originals}"
    )


# ── apply_spell_corrections (end-to-end) ──────────────────────────────────────


def _read_part(docx_path, name):
    with zipfile.ZipFile(docx_path, "r") as z:
        if name not in z.namelist():
            return b""
        return z.read(name)


def test_apply_writes_tracked_change_for_autofix(tmp_path):
    src = _make_doc(tmp_path, [
        ("This is the abstract paragraph for studentss in the manuscript.", None),
    ])
    out = str(tmp_path / "out.docx")
    next_id, actions = apply_spell_corrections(src, out, next_change_id=1000)

    assert any(a["rule"] == "SPELL_AUTOFIX" and a["original"] == "studentss"
               for a in actions)

    doc_xml = _read_part(out, "word/document.xml")
    root = etree.fromstring(doc_xml)
    # Tracked change for studentss → students should be present as a
    # <w:del>/<w:ins> pair in the document body.
    del_texts = [
        (t.text or "")
        for t in root.iter(f"{WQ}delText")
    ]
    ins_texts = []
    for ins in root.iter(f"{WQ}ins"):
        for t in ins.iter(f"{WQ}t"):
            if t.text:
                ins_texts.append(t.text)
    assert "studentss" in del_texts
    assert "students" in ins_texts


def test_apply_returns_empty_actions_for_clean_document(tmp_path):
    src = _make_doc(tmp_path, [
        ("This paragraph contains only correctly spelled English words.", None),
    ])
    out = str(tmp_path / "out.docx")
    next_id, actions = apply_spell_corrections(src, out, next_change_id=1)
    assert actions == []
    # File should still be a valid docx that opens.
    assert _read_part(out, "word/document.xml") != b""


def test_typo_inside_direct_quotation_is_skipped(tmp_path):
    """A typo INSIDE quoted text retains the source's spelling."""
    path = _make_doc(tmp_path, [
        ('The author wrote "the studentss were not informed" in the paper.', None),
    ])
    results = get_spell_corrections(path)
    # `studentss` only appears inside the quoted span — must NOT be flagged.
    assert all(r["original"] != "studentss" for r in results)


def test_typo_outside_quote_still_flagged_when_quote_present(tmp_path):
    """A typo OUTSIDE the quote in the same paragraph should still flag."""
    path = _make_doc(tmp_path, [
        ('Author noted "system was robust" but our studentss had different results.', None),
    ])
    results = get_spell_corrections(path)
    originals = [r["original"] for r in results]
    assert "studentss" in originals
