"""Tests for the reference-list hanging-indent correction."""

import zipfile

from docx import Document
from docx.shared import Pt
from lxml import etree

from app.services.reference_indent_corrections import (
    _build_style_name_map,
    apply_reference_indent_corrections,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _ref_entry_inds(path):
    """Return the w:ind element (attrib dict) of each reference entry."""
    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    inds, in_refs = [], False
    for p in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip()
        if text.lower() == "references":
            in_refs = True
            continue
        if in_refs and text:
            pPr = p.find(f"{WQ}pPr")
            ind = pPr.find(f"{WQ}ind") if pPr is not None else None
            inds.append(dict(ind.attrib) if ind is not None else {})
    return inds


def _refs_doc(tmp_path, first_line_indent_pt=18):
    doc = Document()
    doc.add_paragraph("References", style="Heading 1")
    for r in ["Adams, B. (2020). A study. Journal.",
              "Brown, C. (2019). Another. Journal."]:
        p = doc.add_paragraph(r)
        if first_line_indent_pt is not None:
            # Simulate the wrong format: first line indented (not hanging).
            p.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    path = tmp_path / "in.docx"
    doc.save(str(path))
    return str(path)


def test_reference_gets_left_and_hanging_indent(tmp_path):
    src = _refs_doc(tmp_path)
    out = tmp_path / "out.docx"
    _, made = apply_reference_indent_corrections(src, str(out), 1)
    assert len(made) == 2
    for ind in _ref_entry_inds(str(out)):
        assert ind.get(f"{WQ}left") == "720"
        assert ind.get(f"{WQ}hanging") == "720"
        # the wrong first-line indent must be removed
        assert f"{WQ}firstLine" not in ind


def test_already_correct_indent_not_touched(tmp_path):
    # Apply once to get the correct hanging indent, then again → no new change.
    src = _refs_doc(tmp_path)
    once = tmp_path / "once.docx"
    apply_reference_indent_corrections(src, str(once), 1)
    twice = tmp_path / "twice.docx"
    _, made = apply_reference_indent_corrections(str(once), str(twice), 1)
    assert made == []


def test_fixes_indent_when_pprchange_already_exists(tmp_path):
    """Regression: Sam restyles references (adding a w:pPrChange) but leaves the
    wrong direct firstLine. The pass must still rewrite the indent — reusing the
    existing pPrChange (only one allowed) — not skip the paragraph."""
    src = _refs_doc(tmp_path)
    with zipfile.ZipFile(src) as z:
        root = etree.fromstring(z.read("word/document.xml"))
        members = {i.filename: z.read(i.filename) for i in z.infolist()}
    # Inject a pre-existing pPrChange on each reference paragraph (as Sam would).
    body = root.find(f"{WQ}body")
    in_refs = False
    for p in body.findall(f"{WQ}p"):
        text = "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip().lower()
        if text == "references":
            in_refs = True
            continue
        if in_refs and text:
            pPr = p.find(f"{WQ}pPr")
            change = etree.SubElement(pPr, f"{WQ}pPrChange")
            change.set(f"{WQ}id", "1")
            change.set(f"{WQ}author", "Sam")
            change.set(f"{WQ}date", "2024-01-01T00:00:00Z")
            change.append(etree.SubElement(change, f"{WQ}pPr"))
    members["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    forced = tmp_path / "forced.docx"
    with zipfile.ZipFile(str(forced), "w", zipfile.ZIP_DEFLATED) as zout:
        for n, d in members.items():
            zout.writestr(n, d)

    out = tmp_path / "out.docx"
    _, made = apply_reference_indent_corrections(str(forced), str(out), 1)
    assert len(made) == 2
    for ind in _ref_entry_inds(str(out)):
        assert ind.get(f"{WQ}left") == "720"
        assert ind.get(f"{WQ}hanging") == "720"
        assert f"{WQ}firstLine" not in ind
    # still exactly one pPrChange per reference paragraph (no duplicate)
    with zipfile.ZipFile(str(out)) as z:
        root2 = etree.fromstring(z.read("word/document.xml"))
    body2 = root2.find(f"{WQ}body")
    in_refs = False
    for p in body2.findall(f"{WQ}p"):
        text = "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip().lower()
        if text == "references":
            in_refs = True
            continue
        if in_refs and text:
            assert len(p.find(f"{WQ}pPr").findall(f"{WQ}pPrChange")) == 1


def test_style_name_map_resolves_heading(tmp_path):
    # The References section must be found even though detection is style-based:
    # the resolver maps the heading's style id to its "heading 1" name.
    src = _refs_doc(tmp_path)
    with zipfile.ZipFile(src) as z:
        name_map = _build_style_name_map(z)
    assert any(v == "heading 1" for v in name_map.values())
