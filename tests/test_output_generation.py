import os
import zipfile
from types import SimpleNamespace

import pytest
from docx import Document
from lxml import etree

from app.services.jutlp_validator import validate
from app.services.output_generation import (
    _canonical_insert_index,
    _format_author_query_text,
    _make_comment_element,
    _section_for_rule,
    _section_heading_present,
    generate_commented_docx,
)

PACK = "tests/jutlp_sample_docx_test_pack"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def test_author_query_comment_text_is_numbered():
    assert _format_author_query_text(3, "Please revise this section.") == (
        "Author Query 3. Please revise this section."
    )


def test_author_query_comment_text_is_not_double_prefixed():
    assert _format_author_query_text(4, "Author Query 2. Existing message.") == (
        "Author Query 2. Existing message."
    )


def test_author_query_prefix_is_bold_in_comment_xml():
    comment = _make_comment_element(3, "Please revise this section.")
    runs = comment.findall(f".//{WQ}r")
    prefix_run = next(
        run for run in runs
        if "".join(t.text or "" for t in run.findall(f"{WQ}t")).startswith("Author Query 3.")
    )

    assert prefix_run.find(f"{WQ}rPr/{WQ}b") is not None


def test_comment_xml_does_not_add_visible_id_before_author_query():
    comment = _make_comment_element(24, "Reference issue.")
    text = "".join(t.text or "" for t in comment.iter(f"{WQ}t"))

    assert text == "Author Query 24. Reference issue."


def count_comments_in_docx(path: str) -> int:
    with zipfile.ZipFile(path, "r") as z:
        if "word/comments.xml" not in z.namelist():
            return 0
        xml = z.read("word/comments.xml")
    tree = etree.fromstring(xml)
    return len(tree.findall(f"{WQ}comment"))


def has_comments_relationship(path: str) -> bool:
    with zipfile.ZipFile(path, "r") as z:
        rels_xml = z.read("word/_rels/document.xml.rels")
    tree = etree.fromstring(rels_xml)
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    return any(el.get("Type") == rel_type for el in tree)


def has_comments_content_type(path: str) -> bool:
    with zipfile.ZipFile(path, "r") as z:
        ct_xml = z.read("[Content_Types].xml")
    tree = etree.fromstring(ct_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    return any(
        el.get("PartName") == "/word/comments.xml"
        for el in tree.findall(f"{{{ns}}}Override")
    )


def get_document_text(path: str) -> str:
    with zipfile.ZipFile(path, "r") as z:
        xml = z.read("word/document.xml")
    tree = etree.fromstring(xml)
    return " ".join(t.text or "" for t in tree.iter(f"{WQ}t"))


def get_comment_anchor_texts(path: str) -> dict[str, str]:
    with zipfile.ZipFile(path, "r") as z:
        doc_root = etree.fromstring(z.read("word/document.xml"))
        comments_root = etree.fromstring(z.read("word/comments.xml"))

    body = doc_root.find(f"{WQ}body")
    anchors = {}
    for para in body:
        if para.tag != f"{WQ}p":
            continue
        text = " ".join("".join(t.text or "" for t in para.iter(f"{WQ}t")).split())
        for marker in para.findall(f"{WQ}commentRangeStart"):
            anchors[marker.get(f"{WQ}id")] = text

    return {
        " ".join("".join(t.text or "" for t in c.iter(f"{WQ}t")).split()): anchors.get(c.get(f"{WQ}id"), "")
        for c in comments_root.findall(f"{WQ}comment")
    }


class TestValidDocNoComments:
    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        out = str(tmp_path_factory.mktemp("out") / "01_out.docx")
        report = validate(f"{PACK}/01_valid_identified.docx")
        generate_commented_docx(f"{PACK}/01_valid_identified.docx", out, report)
        return out

    def test_output_file_created(self, output):
        assert os.path.exists(output)

    def test_no_comments_added(self, output):
        assert count_comments_in_docx(output) == 0



class TestOneFailureOneComment:
    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        out = str(tmp_path_factory.mktemp("out") / "02_out.docx")
        report = validate(f"{PACK}/02_missing_method_subsection.docx")
        generate_commented_docx(f"{PACK}/02_missing_method_subsection.docx", out, report)
        return out

    def test_output_file_created(self, output):
        assert os.path.exists(output)

    def test_comments_added(self, output):
        # The fixture renames Participants → "Sampling", so it now draws TWO
        # complementary comments: MET002 (required 'Participants' missing) and
        # the MET900 advisory (Method has a non-template subheading 'Sampling').
        assert count_comments_in_docx(output) == 2

    def test_comment_text_covers_missing_and_extra(self, output):
        with zipfile.ZipFile(output, "r") as z:
            text = " ".join(t.text or "" for t in etree.fromstring(
                z.read("word/comments.xml")).iter(f"{WQ}t"))
        assert "Participants" in text  # missing required subsection
        assert "not in the JUTLP expected set" in text  # extra-subheading advisory

    def test_comments_relationship_present(self, output):
        assert has_comments_relationship(output)

    def test_comments_content_type_present(self, output):
        assert has_comments_content_type(output)



class TestTwoFailuresTwoComments:
    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        out = str(tmp_path_factory.mktemp("out") / "03_out.docx")
        report = validate(f"{PACK}/03_front_page_issues.docx")
        generate_commented_docx(f"{PACK}/03_front_page_issues.docx", out, report)
        return out

    def test_output_file_created(self, output):
        assert os.path.exists(output)

    def test_two_comments_added(self, output):
        assert count_comments_in_docx(output) == 2


class TestFiveFailuresFiveComments:
    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        out = str(tmp_path_factory.mktemp("out") / "04_out.docx")
        report = validate(f"{PACK}/04_structure_and_endmatter_issues.docx")
        generate_commented_docx(f"{PACK}/04_structure_and_endmatter_issues.docx", out, report)
        return out

    def test_output_file_created(self, output):
        assert os.path.exists(output)

    def test_failing_rules_become_comments(self, output):
        # Fixture has 8 issues from validate(): 5 fail (SEC005 combined,
        # DIS001/002/003 missing subsections, STY001 Guidance Notes) and
        # 3 warn (SPE001, SPE003, STY003). The comment-emission filter in
        # generate_commented_docx only converts fail-status results (and
        # REFE-prefixed warns) into Word comments — the 3 warn-status
        # structural results stay in the side-panel UI rather than
        # cluttering the document. So 5 comments, matching the fail count.
        assert count_comments_in_docx(output) == 5

    def test_comments_relationship_present(self, output):
        assert has_comments_relationship(output)

    def test_comments_content_type_present(self, output):
        assert has_comments_content_type(output)

    def test_comment_text_references_discussion(self, output):
        with zipfile.ZipFile(output, "r") as z:
            xml = z.read("word/comments.xml")
        tree = etree.fromstring(xml)
        all_text = " ".join(t.text or "" for t in tree.iter(f"{WQ}t"))
        assert "Discussion" in all_text

    def test_comment_text_references_guidance_notes(self, output):
        with zipfile.ZipFile(output, "r") as z:
            xml = z.read("word/comments.xml")
        tree = etree.fromstring(xml)
        all_text = " ".join(t.text or "" for t in tree.iter(f"{WQ}t"))
        assert "Guidance Notes" in all_text

    def _h1_paras(self, output):
        """Return ordered list of (text, has_ins, [commentRangeStart ids]) for
        every Heading-1 paragraph in the output document."""
        with zipfile.ZipFile(output, "r") as z:
            root = etree.fromstring(z.read("word/document.xml"))
        body = root.find(f"{WQ}body")
        rows = []
        for p in body:
            if p.tag != f"{WQ}p":
                continue
            pStyle = p.find(f"{WQ}pPr/{WQ}pStyle")
            sid = pStyle.get(f"{WQ}val") if pStyle is not None else ""
            if sid.replace(" ", "").lower() != "heading1":
                continue
            text = "".join(t.text or "" for t in p.iter(f"{WQ}t"))
            has_ins = p.find(f"{WQ}ins") is not None
            crs = [c.get(f"{WQ}id") for c in p.findall(f"{WQ}commentRangeStart")]
            rows.append((text, has_ins, crs))
        return rows

    def test_discussion_stub_heading_is_inserted(self, output):
        """The combined 'Results and Discussion' fixture has no standalone
        Discussion heading, so a tracked Heading-1 'Discussion' stub is
        inserted (wrapped in <w:ins>)."""
        rows = self._h1_paras(output)
        discussion = [r for r in rows if r[0].strip() == "Discussion"]
        assert len(discussion) == 1
        assert discussion[0][1] is True  # tracked-inserted

    def test_discussion_comments_anchor_to_stub_not_results(self, output):
        """SEC005 + DIS001-003 comments anchor inside the inserted Discussion
        stub, NOT on the Results heading."""
        rows = self._h1_paras(output)
        discussion = next(r for r in rows if r[0].strip() == "Discussion")
        results = next(r for r in rows if r[0].strip() == "Results")
        # The stub carries the missing-section/subsection comment anchors.
        assert len(discussion[2]) >= 4
        # The Results heading must NOT carry those anchors.
        assert results[2] == []

    def test_discussion_stub_positioned_before_references(self, output):
        """The inserted Discussion stub sits after Results and before
        References in canonical order."""
        rows = self._h1_paras(output)
        texts = [r[0].strip() for r in rows]
        assert "Discussion" in texts and "References" in texts
        assert texts.index("Results") < texts.index("Discussion")
        assert texts.index("Discussion") < texts.index("References")


class TestNonTemplateSubheadingComment:
    """A MET900/DIS900 advisory (warn) must surface as a Word comment anchored
    at its section heading."""

    def _doc(self, tmp_path):
        path = tmp_path / "in.docx"
        doc = Document()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Intro body.")
        doc.add_paragraph("Method", style="Heading 1")
        doc.add_paragraph("Interviews", style="Heading 2")
        doc.add_paragraph("Method body.")
        doc.add_paragraph("References", style="Heading 1")
        doc.add_paragraph("Smith, J. (2020). A study.")
        doc.save(str(path))
        return str(path)

    def test_met900_warn_becomes_comment(self, tmp_path):
        src = self._doc(tmp_path)
        out = str(tmp_path / "out.docx")
        report = {"results": [{
            "rule_id": "MET900", "status": "warn",
            "message": "Method contains subheading(s) not in the JUTLP expected set: 'Interviews'.",
        }]}
        generate_commented_docx(src, out, report)

        texts = get_comment_anchor_texts(out)
        match = next((c for c in texts if "not in the JUTLP expected set" in c), None)
        assert match is not None, "MET900 advisory did not become a comment"
        # Anchored on the Method heading.
        assert texts[match].strip() == "Method"


class TestPresentUnstyledDiscussionNoStub:
    """A Discussion section that is present but not yet Heading-1 styled (the
    author marked it visually; the restyle pass applies Heading 1 later) must
    NOT get a spurious duplicate stub, and its subsection comment must anchor
    on the real Discussion paragraph — not on an inserted stub."""

    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        src = str(tmp_path_factory.mktemp("in") / "unstyled_discussion.docx")
        doc = Document()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Intro sentence one. Sentence two. Sentence three.")
        doc.add_paragraph("Results", style="Heading 1")
        doc.add_paragraph("Results sentence one. Two. Three.")
        # Discussion present, but styled bold-Normal (NOT Heading 1).
        disc = doc.add_paragraph()
        disc.add_run("Discussion").bold = True
        for sub in ("Practical Implications", "Limitations and Future Research"):
            sp = doc.add_paragraph()
            sp.add_run(sub).bold = True
            doc.add_paragraph("A body sentence under the subsection.")
        doc.add_paragraph("References", style="Heading 1")
        doc.add_paragraph("Smith, J. (2020). A study. Journal, 1(1), 1-10.")
        doc.save(src)

        out = str(tmp_path_factory.mktemp("out") / "unstyled_discussion_out.docx")
        report = validate(src)
        generate_commented_docx(src, out, report)
        return out

    def _body_paras(self, output):
        with zipfile.ZipFile(output, "r") as z:
            root = etree.fromstring(z.read("word/document.xml"))
        body = root.find(f"{WQ}body")
        return [p for p in body if p.tag == f"{WQ}p"]

    def test_discussion_was_detected(self, output):
        # The fixture is only meaningful if the validator flagged the missing
        # Discussion subsection (proving Discussion itself was recognised).
        with zipfile.ZipFile(output, "r") as z:
            comments = z.read("word/comments.xml").decode("utf-8")
        assert "Theoretical Implications" in comments

    def test_no_inserted_discussion_stub(self, output):
        """No tracked-inserted 'Discussion' heading paragraph should exist."""
        for p in self._body_paras(output):
            text = "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip()
            is_inserted = p.find(f"{WQ}ins") is not None
            assert not (text == "Discussion" and is_inserted), (
                "a spurious tracked Discussion stub was inserted"
            )

    def test_only_one_discussion_paragraph(self, output):
        discussions = [
            p for p in self._body_paras(output)
            if "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip() == "Discussion"
        ]
        assert len(discussions) == 1

    def test_comment_anchors_on_real_discussion_paragraph(self, output):
        """The Discussion-subsection comment must anchor on the original
        (non-inserted) Discussion paragraph."""
        for p in self._body_paras(output):
            if not p.findall(f"{WQ}commentRangeStart"):
                continue
            text = "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip()
            if text == "Discussion":
                assert p.find(f"{WQ}ins") is None  # the real one, not a stub
                return
        raise AssertionError("no comment anchored on the Discussion paragraph")


class TestReferenceVisualIndicators:
    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        out = str(tmp_path_factory.mktemp("out") / "05_ref_visuals_out.docx")
        report = validate(f"{PACK}/01_valid_identified.docx")
        ref_results = [
            {
                "rule_id": "CREF001",
                "status": "pass",
                "message": "Citation matched reference list entry.",
            },
            {
                "rule_id": "CREF002",
                "status": "fail",
                "message": "Citation missing from reference list.",
            },
        ]
        generate_commented_docx(
            f"{PACK}/01_valid_identified.docx",
            out,
            report,
            ref_results=ref_results,
        )
        return out



class TestCrefPassIsNotCommented:
    @pytest.fixture(scope="class")
    def output(self, tmp_path_factory):
        out = str(tmp_path_factory.mktemp("out") / "cref_filter_out.docx")
        report = validate(f"{PACK}/01_valid_identified.docx")
        ref_results = [
            {"rule_id": "CREF001", "status": "pass", "message": "Entry 1 verified"},
            {"rule_id": "CREF002", "status": "fail", "message": "Entry 2 failed verification"},
        ]
        generate_commented_docx(
            f"{PACK}/01_valid_identified.docx",
            out,
            report,
            ref_results=ref_results,
            llm_results=None,
        )
        return out

    def test_failing_cref_gets_detail_and_summary_comments(self, output):
        assert count_comments_in_docx(output) == 2
        comments = get_comment_anchor_texts(output)
        assert any("Reference list issue summary:" in text for text in comments)
        assert any("Entry 2 failed verification" in text for text in comments)


def test_reference_summary_comment_anchors_to_references_heading_for_cons_warn(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text.")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Alpha, A. (2023). First reference.")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {"results": []},
        ref_results=[
            {
                "rule_id": "CONS001",
                "status": "warn",
                "message": "1 reference is not cited in the body text.",
            }
        ],
    )

    anchors = get_comment_anchor_texts(str(output_path))
    summary = next(
        text for text in anchors
        if "Reference list issue summary:" in text
    )
    assert "1 citation-consistency warning" in summary
    assert anchors[summary] == "References"


def test_cref_comments_anchor_to_unstyled_reference_entries(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body text cites Beta (2024).")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Alpha, A. (2023). First reference.")
    doc.add_paragraph("Beta, B. (2024). Second reference.")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {"results": []},
        ref_results=[{"rule_id": "CREF002", "status": "fail", "message": "Entry 2 failed"}],
    )

    anchors = get_comment_anchor_texts(str(output_path))
    target = next(text for text in anchors if "Entry 2 failed" in text)
    assert anchors[target].startswith("Beta, B. (2024).")


def test_doit_comment_anchors_at_correct_reference_entry(tmp_path):
    """A `DOIT005`-style result should drop a Word comment anchored at
    reference entry 5 — same dispatch as CREF, just a different prefix."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body.")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Alpha, A. (2023). First ref. 10.1111/aaa")
    doc.add_paragraph("Beta, B. (2024). Second ref. 10.2222/bbb")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {"results": []},
        ref_results=[
            {
                "rule_id": "DOIT002",
                "status": "fail",
                "message": "Entry 2: DOI not found in CrossRef (10.2222/bbb)",
            }
        ],
    )

    anchors = get_comment_anchor_texts(str(output_path))
    target = next(
        (k for k in anchors if "10.2222/bbb" in k),
        None,
    )
    assert target is not None
    assert anchors[target].startswith("Beta, B. (2024).")


def test_doit_suffixed_rule_id_anchors_at_same_entry(tmp_path):
    """`DOIT002_2` (the second DOI inside reference 2) must anchor at
    the same reference paragraph as `DOIT002`."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body.")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Alpha, A. (2023). First ref.")
    doc.add_paragraph("Beta, B. (2024). Second ref. 10.x/a 10.x/b")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {"results": []},
        ref_results=[
            {"rule_id": "DOIT002_1", "status": "fail", "message": "DOI #1 bad"},
            {"rule_id": "DOIT002_2", "status": "fail", "message": "DOI #2 bad"},
        ],
    )

    anchors = get_comment_anchor_texts(str(output_path))
    matched = [v for k, v in anchors.items() if "DOI #" in k]
    # Both comments should land on the same Beta reference paragraph.
    assert len(matched) == 2
    assert all(v.startswith("Beta, B. (2024).") for v in matched)


def test_href_comment_anchors_at_correct_reference_entry(tmp_path):
    """Regression: the pre-existing `HREF***` rule (hyperlink-DOI check)
    previously fell through the anchor dispatch and was silently dropped
    from the docx. After the fix it lands at the referenced entry."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Body.")
    doc.add_paragraph("References").style = "Heading 1"
    doc.add_paragraph("Alpha, A. (2023). First ref.")
    doc.add_paragraph("Beta, B. (2024). Second ref.")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {"results": []},
        ref_results=[
            {
                "rule_id": "HREF002",
                "status": "fail",
                "message": "Entry 2: hyperlinked DOI does not resolve",
            }
        ],
    )

    anchors = get_comment_anchor_texts(str(output_path))
    target = next(
        (k for k in anchors if "hyperlinked DOI" in k),
        None,
    )
    assert target is not None
    assert anchors[target].startswith("Beta, B. (2024).")


def test_duplicate_comments_for_same_anchor_are_collapsed(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Example title")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Abstract body.")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {
            "results": [
                {
                    "rule_id": "FP004",
                    "status": "fail",
                    "message": "Please check use of acronyms.",
                },
                {
                    "rule_id": "FP007",
                    "status": "fail",
                    "message": "Please check use of acronyms.",
                },
            ]
        },
    )

    assert count_comments_in_docx(str(output_path)) == 1
    anchors = get_comment_anchor_texts(str(output_path))
    assert list(anchors) == ["Author Query 1. Please check use of acronyms."]
    assert anchors["Author Query 1. Please check use of acronyms."] == "Abstract"


def test_andor_occurrences_are_grouped_into_one_comment(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("Students and/or teachers responded.")
    doc.add_paragraph("Parents and/or guardians also responded.")
    doc.save(input_path)

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {
            "results": [
                {
                    "rule_id": "ANDOR001",
                    "status": "fail",
                    "message": "Avoid 'and/or'.",
                    "para_idx": 1,
                    "anchor_phrase": "and/or",
                },
                {
                    "rule_id": "ANDOR002",
                    "status": "fail",
                    "message": "Avoid 'and/or'.",
                    "para_idx": 2,
                    "anchor_phrase": "and/or",
                },
            ]
        },
    )

    assert count_comments_in_docx(str(output_path)) == 1
    anchors = get_comment_anchor_texts(str(output_path))
    comment_text = next(iter(anchors))
    assert comment_text.startswith("Author Query 1. Avoid 'and/or'.")
    assert "2 occurrences" in comment_text
    assert anchors[comment_text] == "Students and/or teachers responded."


def test_llm_acronym_notes_are_grouped_into_one_comment(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("AI supports LMS use.")
    doc.add_paragraph("Introduction").style = "Heading 1"
    doc.add_paragraph("AI tools support LMS platforms.")
    doc.save(input_path)

    llm_results = {
        "notes": [
            {
                "severity": "high",
                "message": "Please check acronym use for accessibility.",
                "suggestion": "Expand AI on first use.",
                "section": "Abstract",
                "quote": "AI",
            },
            {
                "severity": "high",
                "message": "Please check acronym use for accessibility.",
                "suggestion": "Expand LMS on first use.",
                "section": "Introduction",
                "quote": "LMS",
            },
        ]
    }

    generate_commented_docx(
        str(input_path),
        str(output_path),
        {"results": []},
        llm_results=llm_results,
    )

    assert count_comments_in_docx(str(output_path)) == 1
    comment_text = next(iter(get_comment_anchor_texts(str(output_path))))
    assert comment_text.startswith("Author Query 1. Please check acronym use")
    assert "2 similar acronym notes were grouped here" in comment_text


# ── Missing-section stub helper unit tests ───────────────────────────────────

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ_NS = f"{{{W_NS}}}"


def _rec(text, style="Heading 1"):
    """Minimal ParagraphRecord-like stub for heading-presence checks."""
    return SimpleNamespace(text=text, style=style, is_empty=(text.strip() == ""))


def _h1_para(text):
    """Build a Heading-1 lxml <w:p> for insert-index tests."""
    p = etree.Element(f"{WQ_NS}p")
    pPr = etree.SubElement(p, f"{WQ_NS}pPr")
    pStyle = etree.SubElement(pPr, f"{WQ_NS}pStyle")
    pStyle.set(f"{WQ_NS}val", "Heading 1")
    r = etree.SubElement(p, f"{WQ_NS}r")
    t = etree.SubElement(r, f"{WQ_NS}t")
    t.text = text
    return p


def test_section_for_rule_mapping():
    assert _section_for_rule("SEC005") == "Discussion"
    assert _section_for_rule("DIS002") == "Discussion"
    assert _section_for_rule("MET003") == "Method"
    assert _section_for_rule("SEC009") is None   # ordering rule, no section
    assert _section_for_rule("CONS001") is None


def test_section_heading_present_exact_and_combined():
    # Combined "Results and Discussion" counts as Results, NOT Discussion.
    recs = [_rec("Introduction"), _rec("Results and Discussion"), _rec("Conclusion")]
    assert _section_heading_present("Results", recs) is True
    assert _section_heading_present("Discussion", recs) is False
    # A standalone Discussion heading IS present.
    recs2 = [_rec("Results"), _rec("Discussion")]
    assert _section_heading_present("Discussion", recs2) is True


def test_canonical_insert_index_after_present_predecessor():
    # Results present, Discussion missing → insert right after Results block
    # (before Conclusion).
    paras = [_h1_para("Introduction"), _h1_para("Results"), _h1_para("Conclusion")]
    idx = _canonical_insert_index("Discussion", paras)
    assert idx == 2  # before Conclusion


def test_canonical_insert_index_before_successor_when_no_predecessor():
    # Only References present; insert Discussion before it.
    paras = [_h1_para("References")]
    idx = _canonical_insert_index("Discussion", paras)
    assert idx == 0


def test_canonical_insert_index_consecutive_missing():
    # Results present, both Discussion and Conclusion missing. After a
    # Discussion stub is inserted, Conclusion should land after it.
    paras = [_h1_para("Results"), _h1_para("References")]
    disc_idx = _canonical_insert_index("Discussion", paras)
    assert disc_idx == 1  # before References
    paras.insert(disc_idx, _h1_para("Discussion"))
    concl_idx = _canonical_insert_index("Conclusion", paras)
    # Conclusion goes after the freshly-inserted Discussion, before References.
    assert concl_idx == 2
