"""Tests for the table keep-together pass."""

import zipfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

from app.services.table_keep_together import apply_table_keep_together

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"


def _rows(docx_path: str):
    with zipfile.ZipFile(docx_path, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    tbl = root.find(f".//{WQ}tbl")
    return tbl.findall(f"{WQ}tr") if tbl is not None else []


def _has(el, tag):
    return el is not None and el.find(f"{WQ}{tag}") is not None


def _basic_table_doc(tmp_path, *, with_caption=False):
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    if with_caption:
        # Define the caption styles the bot uses.
        for name in ("Table Number", "Table Title"):
            try:
                doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            except Exception:
                pass
        doc.add_paragraph("Table 4", style="Table Number")
        doc.add_paragraph("Descriptive statistics", style="Table Title")
    table = doc.add_table(rows=3, cols=2)
    for r in range(3):
        for c in range(2):
            table.rows[r].cells[c].text = f"r{r}c{c}"
    doc.save(str(path))
    return str(path)


def test_all_rows_get_cantsplit(tmp_path):
    src = _basic_table_doc(tmp_path)
    out = tmp_path / "out.docx"

    actions = apply_table_keep_together(src, str(out))

    assert len(actions) == 1
    rows = _rows(str(out))
    assert rows
    for row in rows:
        assert _has(row.find(f"{WQ}trPr"), "cantSplit")


def test_only_first_row_is_tblheader(tmp_path):
    src = _basic_table_doc(tmp_path)
    out = tmp_path / "out.docx"

    apply_table_keep_together(src, str(out))

    rows = _rows(str(out))
    assert _has(rows[0].find(f"{WQ}trPr"), "tblHeader")
    for row in rows[1:]:
        assert not _has(row.find(f"{WQ}trPr"), "tblHeader")


def test_caption_paragraphs_get_keepnext(tmp_path):
    src = _basic_table_doc(tmp_path, with_caption=True)
    out = tmp_path / "out.docx"

    apply_table_keep_together(src, str(out))

    with zipfile.ZipFile(out, "r") as z:
        root = etree.fromstring(z.read("word/document.xml"))
    caption_keepnext = 0
    for p in root.iter(f"{WQ}p"):
        text = "".join(t.text or "" for t in p.iter(f"{WQ}t")).strip()
        if text in ("Table 4", "Descriptive statistics"):
            ppr = p.find(f"{WQ}pPr")
            if ppr is not None and ppr.find(f"{WQ}keepNext") is not None:
                caption_keepnext += 1
    assert caption_keepnext == 2


def test_idempotent(tmp_path):
    """Running twice produces no second set of changes and no duplicate props."""
    src = _basic_table_doc(tmp_path)
    mid = tmp_path / "mid.docx"
    out = tmp_path / "out.docx"

    apply_table_keep_together(src, str(mid))
    actions2 = apply_table_keep_together(str(mid), str(out))

    assert actions2 == []
    for row in _rows(str(out)):
        tr_pr = row.find(f"{WQ}trPr")
        assert len(tr_pr.findall(f"{WQ}cantSplit")) == 1


def test_trpr_child_order_is_schema_valid(tmp_path):
    """cantSplit must precede tblHeader, and trPr must be tr's first child."""
    src = _basic_table_doc(tmp_path)
    out = tmp_path / "out.docx"
    apply_table_keep_together(src, str(out))

    first_row = _rows(str(out))[0]
    assert etree.QName(first_row[0]).localname == "trPr"
    names = [etree.QName(c).localname for c in first_row.find(f"{WQ}trPr")]
    assert "cantSplit" in names and "tblHeader" in names
    assert names.index("cantSplit") < names.index("tblHeader")


def test_no_tables_is_noop(tmp_path):
    path = tmp_path / "notable.docx"
    doc = Document()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Just prose.")
    doc.save(str(path))
    out = tmp_path / "out.docx"

    assert apply_table_keep_together(str(path), str(out)) == []
    assert out.exists()


def test_existing_docx_still_opens(tmp_path):
    src = _basic_table_doc(tmp_path, with_caption=True)
    out = tmp_path / "out.docx"
    apply_table_keep_together(src, str(out))
    # Must remain a valid, openable document.
    assert Document(str(out))
