"""Author-only output generation scratch workflow with LLM naming checks."""

import copy
import os
import re
import shutil
import zipfile

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from app.domain.canonical_jultp_template import (
    CANONICAL_STRUCTURE,
    SECTION_RENAME_MAP,
    strip_leading_section_number,
)
from app.services.acronym_store import load_acronyms
from app.services.ai.llm_client import call_llm_json
from app.services.body_llm_edits import build_body_edit_plan
from app.services.document_analysis_services import get_front_page, load_paragraphs
from app.services.document_styling_fixes import extract_template_styles
from app.services.editorial_review_comments import build_editorial_review_comment_plan
from app.services.heading_corrections import apply_heading_corrections
from app.services.language_corrections import AU_CORRECTIONS as _AU_CORRECTIONS
from app.services.narrative_citation_corrections import apply_narrative_citation_corrections
from app.services.output_filename import build_output_filename_from_author_line
from app.services.quotation_utils import find_quote_spans
from app.services.run_font_corrections import apply_run_font_corrections
from app.services.timestamps import now_sydney_iso


def _fix_au_spellings_in_text(text: str) -> str:
    """Apply AU_CORRECTIONS to any US spellings found in text, preserving capitalisation."""
    def _replace(m: re.Match) -> str:
        word = m.group(0)
        au = _AU_CORRECTIONS.get(word.lower())
        if au is None:
            return word
        if word.isupper():
            return au.upper()
        if word[0].isupper():
            return au[0].upper() + au[1:]
        return au
    return re.sub(r'\b[A-Za-z]+\b', _replace, text)

#PROMPT AND SCHEMA PAIRS -------------------------------------------------------------------------------
# These parts are for the LLM parts of the checks


##AUTHOR
AUTHOR_SYSTEM_PROMPT = "You are a strict naming pattern checker. Return valid JSON only."

AUTHOR_NAMING_PATTERN = """Rewrite the authors line to match this exact JUTLP pattern:
Name I. Surname^a, Name Surname^b, and Name I. Surname^c

Hard rules:
- Keep the same authors and original order.
- Use commas between authors and use ", and " before the final author.
- Ensure each author ends with exactly one affiliation marker using ^letter format (for example: ^a, ^b, ^c).
- Preserve titles exactly where present in the source and normalise punctuation/case (for example: Dr., Prof., Mrs., Mx., Rev., Capt., Assoc. Prof.).
- Fix obvious OCR punctuation errors (no "...", no double punctuation, no digits inside names).
- Use proper case for names (for example: Patsie Polly).
- Do not include affiliation institution text in the authors line.
- Keep output to a single line only.

Return JSON only with this schema:
{
  "is_author_naming_correct": true or false,
  "corrected_authors_line": "string",
  "reason": "short explanation"
}
"""

AUTHOR_NAMING_PATTERN_SCHEMA = {
    "name": "author_naming_pattern_check",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "is_author_naming_correct": {"type": "boolean"},
            "corrected_authors_line": {"type": "string"},
            "reason": {"type": "string"},
        },
        "required": [
            "is_author_naming_correct",
            "corrected_authors_line",
            "reason",
        ],
        "additionalProperties": False,
    },
}
##AFFILIATION
AFFILIATION_MATCH_SYSTEM_PROMPT = "You check affiliation marker consistency. Return valid JSON only."

AFFILIATION_MATCH_PROMPT = """Check if author affiliation markers match affiliation section markers.

How to read markers:
- Author markers are suffixes in authors line.
  They may be shown as caret markers (^a, ^b, ^1), as superscript glyphs (ᵃ, ᵇ, ¹),
  or as plain trailing markers attached to names (for example: Smitha, Jonesb).
- Affiliation section markers are the leading marker at the start of each affiliation line
  (for example: aUniversity..., bSchool..., 1Institute...).
  These may also appear as superscript at the start (for example: ᵃUniversity...).
- Note markers (* and #) can be valid author-note markers when matching note lines
  like *corresponding author or #equal contributing authors.
- Treat visually different marker styles as the same marker (a == ^a == ᵃ).
- If all authors share one institution marker (for example only marker a), this is valid.

Return JSON only with this schema:
{
  "is_marker_mapping_consistent": true or false,
  "missing_in_affiliations": ["marker", "..."],
  "missing_in_authors": ["marker", "..."],
  "reason": "short explanation"
}
"""

AFFILIATION_MATCH_SCHEMA = {
    "name": "affiliation_marker_match",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "is_marker_mapping_consistent": {"type": "boolean"},
            "missing_in_affiliations": {
                "type": "array",
                "items": {"type": "string"},
            },
            "missing_in_authors": {
                "type": "array",
                "items": {"type": "string"},
            },
            "reason": {"type": "string"},
        },
        "required": [
            "is_marker_mapping_consistent",
            "missing_in_affiliations",
            "missing_in_authors",
            "reason",
        ],
        "additionalProperties": False,
    },
}

AFFILIATION_SECTION_FOUND_SYSTEM_PROMPT = "You check if front-page author affiliations are present. Return valid JSON only."

AFFILIATION_SECTION_FOUND_PROMPT = """Check the front-page paragraphs listed below and decide if an author affiliations section exists.

Rules:
- An affiliations section usually has marker lines like 1University..., 2School..., aInstitute...
- Ignore note lines starting with * or #.
- Focus only on paragraphs before the Abstract heading.

Return JSON only with this schema:
{
  "affiliations_found": true or false,
  "affiliations_end_index": integer,
  "reason": "short explanation"
}
"""

AFFILIATION_SECTION_FOUND_SCHEMA = {
    "name": "affiliation_section_found",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "affiliations_found": {"type": "boolean"},
            "affiliations_end_index": {"type": "integer"},
            "reason": {"type": "string"},
        },
        "required": [
            "affiliations_found",
            "affiliations_end_index",
            "reason",
        ],
        "additionalProperties": False,
    },
}

##ABSTRACT
ABSTRACT_COMPONENTS_SYSTEM_PROMPT = "You check if an abstract covers required components. Return valid JSON only."

ABSTRACT_COMPONENTS_PROMPT = """Check whether this abstract paragraph includes all required parts:
- problem statement
- theoretical framework (if used)
- method
- key findings
- implications

Return JSON only with this schema:
{
  "is_abstract_components_complete": true or false,
  "missing_components": ["component", "..."],
  "suggested_inclusions": ["short suggestion", "short suggestion"],
  "reason": "short explanation"
}
"""

ABSTRACT_COMPONENTS_SCHEMA = {
    "name": "abstract_components_check",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "is_abstract_components_complete": {"type": "boolean"},
            "missing_components": {
                "type": "array",
                "items": {"type": "string"},
            },
            "suggested_inclusions": {
                "type": "array",
                "items": {"type": "string"},
                "minItems": 2,
                "maxItems": 2,
            },
            "reason": {"type": "string"},
        },
        "required": [
            "is_abstract_components_complete",
            "missing_components",
            "suggested_inclusions",
            "reason",
        ],
        "additionalProperties": False,
    },
}

##CITATION
CITATION_FORMAT_SYSTEM_PROMPT = "You check if a citation line is complete. Return valid JSON only."

CITATION_FORMAT_PROMPT = """Check whether this citation includes the expected parts:
- author list
- year in parentheses
- article title
- journal title
- volume(issue)
- DOI or URL

Return JSON only with this schema:
{
  "is_citation_complete": true or false,
  "missing_parts": ["part", "..."],
  "suggested_fix": "short suggestion",
  "reason": "short explanation"
}
"""

CITATION_FORMAT_SCHEMA = {
    "name": "citation_format_check",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "is_citation_complete": {"type": "boolean"},
            "missing_parts": {
                "type": "array",
                "items": {"type": "string"},
            },
            "suggested_fix": {"type": "string"},
            "reason": {"type": "string"},
        },
        "required": [
            "is_citation_complete",
            "missing_parts",
            "suggested_fix",
            "reason",
        ],
        "additionalProperties": False,
    },
}

##TITLE SECTION -------------------------------------------------------------------------------------------
#Deterministic check for Title, will be first text on the front page - if found, check it uses style :Article Style
#and then check it's in title case. if not, change the title into title case and insert in into the document as a tracked change, replacing the original title.
#then, check if the title is over two lines long, and check if its over 15 words long. If it is, get the LLM to write a comment in effect saying the title can be  No More Than Two Lines and/or Typically 15 Words in Total (depending on which error is flagged - and then have a short suggestions section suggesting some fixes)

_TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "domain", "JUTLP Template 2026.docx")
_TEMPLATE_STYLE_NAMES = None
_TEMPLATE_STYLE_ID_MAP = None
_TEMPLATE_STYLE_XML_MAP = None
_STYLE_ID_ATTR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId"


def _load_template_style_names():
    global _TEMPLATE_STYLE_NAMES
    global _TEMPLATE_STYLE_ID_MAP
    global _TEMPLATE_STYLE_XML_MAP

    if (
        _TEMPLATE_STYLE_NAMES is not None
        and _TEMPLATE_STYLE_ID_MAP is not None
        and _TEMPLATE_STYLE_XML_MAP is not None
    ):
        return _TEMPLATE_STYLE_NAMES

    try:
        style_map = extract_template_styles(_TEMPLATE_PATH)
        _TEMPLATE_STYLE_XML_MAP = style_map
        _TEMPLATE_STYLE_NAMES = set(style_map.keys())
        _TEMPLATE_STYLE_ID_MAP = {}

        for style_name, style_xml in style_map.items():
            style_id = style_name
            try:
                style_el = etree.fromstring(style_xml.encode("utf-8"))
                xml_style_id = style_el.get(_STYLE_ID_ATTR)
                if xml_style_id is not None and xml_style_id.strip() != "":
                    style_id = xml_style_id
            except Exception:
                style_id = style_name
            _TEMPLATE_STYLE_ID_MAP[style_name] = style_id
    except Exception:
        _TEMPLATE_STYLE_NAMES = set()
        _TEMPLATE_STYLE_ID_MAP = {}
        _TEMPLATE_STYLE_XML_MAP = {}

    return _TEMPLATE_STYLE_NAMES


def _resolve_style_name(default_name, alias_names=None):
    if alias_names is None:
        alias_names = []

    names = _load_template_style_names()

    if default_name in names:
        return default_name

    for alias in alias_names:
        if alias in names:
            return alias

    return default_name


def _resolve_style_id(style_name):
    _load_template_style_names()
    if _TEMPLATE_STYLE_ID_MAP is None:
        return style_name

    style_id = _TEMPLATE_STYLE_ID_MAP.get(style_name)
    if style_id is None:
        return style_name
    if str(style_id).strip() == "":
        return style_name
    return style_id


TITLE_REQUIRED_STYLE = _resolve_style_name("Article Title", ["Article Style"])
TITLE_STYLE_ALIASES = ["Article Title", "Article Style"]
if TITLE_REQUIRED_STYLE not in TITLE_STYLE_ALIASES:
    TITLE_STYLE_ALIASES.append(TITLE_REQUIRED_STYLE)
AUTHORS_REQUIRED_STYLE = _resolve_style_name("Authors")
AFFILIATIONS_REQUIRED_STYLE = _resolve_style_name("Author Affiliations")
REFERENCE_ENTRY_REQUIRED_STYLE = _resolve_style_name("APA 7 Reference List Entry")
TITLE_REQUIRED_STYLE_ID = _resolve_style_id(TITLE_REQUIRED_STYLE)
TITLE_LEFT_ALIGNED_STYLE_IDS = frozenset({
    "Article Title", "ArticleTitle", TITLE_REQUIRED_STYLE, TITLE_REQUIRED_STYLE_ID,
})
AUTHORS_REQUIRED_STYLE_ID = _resolve_style_id(AUTHORS_REQUIRED_STYLE)
AFFILIATIONS_REQUIRED_STYLE_ID = _resolve_style_id(AFFILIATIONS_REQUIRED_STYLE)
DEFAULT_AUTHORS_LINE = "First Author^a, Second Author^a, and Third Author^b"
DEFAULT_AUTHOR_AFFILIATIONS = ("a Affiliation 1; b Affiliation 2",)
REFERENCE_ENTRY_REQUIRED_STYLE_ID = _resolve_style_id(REFERENCE_ENTRY_REQUIRED_STYLE)
HEADING_1_STYLE_ID = _resolve_style_id("Heading 1")
HEADING_2_STYLE_ID = _resolve_style_id("Heading 2")
HEADING_3_STYLE_ID = _resolve_style_id("Heading 3")
TABLE_NUMBER_STYLE_ID = _resolve_style_id("Table Number")
TABLE_TITLE_STYLE_ID = _resolve_style_id("Table Title")
TABLE_GRID_STYLE_ID = _resolve_style_id("Table Grid")
TABLE_TEXT_STYLE_ID = _resolve_style_id("Table Text")
TABLE_EMPHASIS_STYLE_ID = _resolve_style_id("Table Emphasis")
TABLE_NOTE_STYLE_ID = _resolve_style_id("Table Note")
QUOTE_STYLE_ID = _resolve_style_id("Quote")
TABLE_TEMPLATE_STYLE_NAMES = [
    "Table Grid",
    "Table Text",
    "Table Text Char",
    "Table Emphasis",
    "Table Emphasis Char",
    "Table Number",
    "Table Number Char",
    "Table Title",
    "Figure/Table Number",
    "Figure/Table Title",
    "Figure/Table Notes",
    "Figure/Table Notes Char",
    "Table Note",
    "Table Note Char",
]
FIGURE_NUMBER_STYLE = _resolve_style_name("Figure Number", ["Figure/Table Number"])
FIGURE_TITLE_STYLE = _resolve_style_name("Figure Title", ["Figure/Table Title"])
FIGURE_NUMBER_STYLE_ID = _resolve_style_id(FIGURE_NUMBER_STYLE)
FIGURE_TITLE_STYLE_ID = _resolve_style_id(FIGURE_TITLE_STYLE)
HEADING_KEEP_NEXT_STYLE_IDS = frozenset({
    "Heading 1", "Heading 2", "Heading 3",
    "Heading1", "Heading2", "Heading3",
    HEADING_1_STYLE_ID, HEADING_2_STYLE_ID, HEADING_3_STYLE_ID,
})
TITLE_MAX_LINES = 2
TITLE_MAX_WORDS = 15

TITLE_LENGTH_SYSTEM_PROMPT = "You are a concise academic copy editor. Return valid JSON only."

TITLE_LENGTH_PROMPT = """Write one short Word comment for this title issue.

Rules:
- Mention only the triggered limits:
  - No More Than Two Lines
  - Typically 15 Words in Total
- Include a short "Suggestions:" section with two quick fixes.
- Also provide exactly two improved title options based on the same fix advice.
- Keep the tone practical and direct.
- Keep it under 90 words.

Return JSON only with this schema:
{
  "comment": "string",
  "suggested_titles": ["string", "string"]
}
"""

TITLE_LENGTH_COMMENT_SCHEMA = {
    "name": "title_length_comment",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "comment": {"type": "string"},
            "suggested_titles": {
                "type": "array",
                "items": {"type": "string"},
                "minItems": 2,
                "maxItems": 2,
            },
        },
        "required": ["comment", "suggested_titles"],
        "additionalProperties": False,
    },
}


#title helper checks -------------------------------------------------------------------------------
#count words in a title-like string
def _count_title_words(text):
    return len(re.findall(r"\b\w+(?:[-']\w+)*\b", text))


#count non-empty visual lines in a title-like string
def _count_title_lines(text):
    lines = text.splitlines()
    non_empty_count = 0

    for line in lines:
        if line.strip() != "":
            non_empty_count += 1

    if non_empty_count == 0 and text.strip() != "":
        return 1

    return non_empty_count


#normalise title text into Title Case with small-word handling
def _to_title_case_title(text):
    """Return ``text`` rewritten in academic title case.

    Rules applied (APA 7 / Chicago-flavoured):

    1. The first word of the title is always capitalised.
    2. The first word after a colon is always capitalised.
    3. "Small" words (``a, an, and, as, at, but, by, for, in, nor, of, on,
       or, the, to, up, via, with``) are lower-cased at any other position.
    4. Known acronyms (case-insensitive match against the acronym allow-list
       plus the canonical set ``AI / APA / EF / UNSW``) are restored to
       their canonical casing, so both ``ai`` and ``AI`` become ``AI``.
    5. Hyphenated compounds are split on the hyphen and rules 1–4 are
       applied to each segment independently; small words *inside* a
       compound stay lower-cased — e.g. ``state-of-the-art`` →
       ``State-of-the-Art`` (not ``State-Of-The-Art``).
    6. Any other word: first character upper, rest lower.

    The previous implementation processed each space-separated word as a
    single unit, which mangled hyphenated compounds like ``AI-Mediated``
    into ``Ai-mediated`` (the full word wasn't ``isupper()``, so the
    acronym-preservation branch missed it, and the fallback lower-cased
    everything after the first letter). Splitting on internal hyphens
    fixes that case while preserving the original behaviour for every
    other shape of title.
    """
    small_words = {
        "a", "an", "and", "as", "at", "but", "by", "for", "in", "nor",
        "of", "on", "or", "the", "to", "up", "via", "with",
    }
    acronyms = {a.upper(): a for a in load_acronyms()}
    acronyms.update({"AI": "AI", "APA": "APA", "EF": "EF", "UNSW": "UNSW"})

    # Is the title "shouted" (written predominantly in capitals)? If so, every
    # word is upper-case, so we must NOT treat all-caps words as acronyms —
    # title-case them and rely on the allow-list for genuine acronyms. But in a
    # normally-cased title a lone all-caps word (e.g. NFSC, OBE) IS an acronym
    # and must be preserved rather than lowered to "Nfsc".
    _alpha_words = [w for w in re.findall(r"[A-Za-z]{2,}", text)]
    _allcaps_words = [w for w in _alpha_words if w.isupper()]
    shouted = bool(_alpha_words) and len(_allcaps_words) / len(_alpha_words) >= 0.5

    def _fix_segment(segment: str, allow_small: bool) -> str:
        """Apply rules 3, 4, 6 to a single hyphen-less segment.

        ``allow_small`` controls whether rule 3 fires for this segment —
        callers pass False for the first word of the title, the first word
        after a colon, and (for hyphenated compounds) the first segment of
        a compound whose containing word itself sits at one of those
        positions.
        """
        if not segment:
            return segment
        upper_key = segment.upper()
        lower = segment.lower()
        if upper_key in acronyms:
            return acronyms[upper_key]
        # Preserve all-caps acronyms not in the allow-list (e.g. NFSC, OBE,
        # JUTLP) — but only in a normally-cased title. In a shouted (all-caps)
        # title every word is upper-case, so this would defeat title-casing.
        if (
            not shouted
            and segment == upper_key
            and segment.isalpha()
            and 2 <= len(segment) <= 6
        ):
            return segment
        if segment in {"Indigenous", "indigenous"}:
            return segment
        if allow_small and lower in small_words:
            return lower
        if len(segment) == 1:
            return segment.upper()
        return segment[0].upper() + segment[1:].lower()

    lines = text.split("\n")
    fixed_lines = []

    for line in lines:
        words = line.split(" ")
        fixed_words = []
        cap_next = False
        seen_title_word = False

        i = 0
        while i < len(words):
            word = words[i]
            if word == "":
                fixed_words.append(word)
                i += 1
                continue

            prefix = ""
            suffix = ""
            core = word

            while len(core) > 0 and core[0].isalnum() is False:
                prefix = prefix + core[0]
                core = core[1:]

            while len(core) > 0 and core[-1].isalnum() is False:
                suffix = core[-1] + suffix
                core = core[:-1]

            if core == "":
                fixed_words.append(word)
                if ":" in word:
                    cap_next = True
                i += 1
                continue

            # Word-level eligibility for the small-word rule: only when the
            # word isn't the title's first alphabetic word AND isn't
            # immediately after a colon. Leading section numbers like "1.1"
            # should not make the following "the" look like a middle word.
            word_allows_small = (seen_title_word and not cap_next)

            if "-" in core:
                # Hyphenated compound: capitalise each segment independently.
                # The first segment inherits the outer ``allow_small`` flag
                # so a compound that opens the title (``"AI-Mediated Era"``)
                # still capitalises its first segment correctly. Subsequent
                # segments always carry ``allow_small=True`` so small words
                # nested inside the compound stay lowercase
                # (``"State-of-the-Art"``).
                segments = core.split("-")
                fixed_segments = [
                    _fix_segment(seg, word_allows_small if j == 0 else True)
                    for j, seg in enumerate(segments)
                ]
                fixed_core = "-".join(fixed_segments)
            else:
                fixed_core = _fix_segment(core, word_allows_small)

            fixed_words.append(prefix + fixed_core + suffix)
            if any(char.isalpha() for char in core):
                seen_title_word = True
            cap_next = ":" in suffix
            i += 1

        fixed_lines.append(" ".join(fixed_words))

    return "\n".join(fixed_lines)


#boolean check for if a title is already in Title Case
def _is_title_case_title(text):
    candidate = _to_title_case_title(text)
    return candidate == text


#build safe fallback title options when LLM is unavailable
def _build_title_option_candidates(title_text):
    clean = title_text.replace("\n", " ")
    clean = re.sub(r"\s+", " ", clean).strip()
    if clean == "":
        return [
            "Improving Feedback Quality in Introductory Statistics",
            "Improving Introductory Statistics Feedback Through Reflective Practice",
        ]

    title_case = _to_title_case_title(clean)
    words = [w for w in title_case.split(" ") if w.strip() != ""]

    option_one_words = words[:12]
    option_two_words = words[:10]

    option_one = " ".join(option_one_words).strip(" ,:;-")
    option_two = " ".join(option_two_words).strip(" ,:;-")

    if option_one == "":
        option_one = title_case
    if option_two == "":
        option_two = title_case
    if option_two == option_one:
        option_two = option_one + " in Higher Education"

    return [option_one, option_two]


#quick heuristic to detect when front-page text looks like an authors line
def _looks_like_authors_line(text):
    low = text.lower()
    has_author_separators = ", and " in low or low.count(",") >= 2
    has_person_title = _has_known_title(low)
    has_author_markers = "^" in text or "*" in text or "#" in text

    if has_author_separators is True:
        return True
    if has_person_title is True:
        return True
    if has_author_markers is True:
        return True

    return False


#find best anchor paragraph for placing a missing-title comment
def _find_anchor_above(paragraphs, index):
    i = index - 1
    while i >= 0:
        if paragraphs[i].is_empty is True:
            return i
        i -= 1

    if index > 0:
        return index - 1
    return index


#fallback title length comment (deterministic wording)
def _build_title_length_comment_fallback(title_text, line_count, word_count, too_many_lines, too_many_words):
    message = "Title guidance: "

    if too_many_lines is True and too_many_words is True:
        message += "Titles should be No More Than Two Lines and Typically 15 Words in Total."
    elif too_many_lines is True:
        message += "Titles should be No More Than Two Lines."
    elif too_many_words is True:
        message += "Titles should be Typically 15 Words in Total."

    message += " Current title is "
    message += str(line_count)
    message += " line(s) and "
    message += str(word_count)
    message += " word(s)."

    message += "\nSuggestions:\n"
    message += "- Remove extra filler words.\n"
    message += "- Keep one core idea and move detail to the abstract."
    options = _build_title_option_candidates(title_text)
    message += "\nSuggested Title Options:\n"
    message += "- " + options[0] + "\n"
    message += "- " + options[1]

    return message


#LLM title length comment with deterministic fallback
def _build_title_length_comment_with_llm(title_text, line_count, word_count, too_many_lines, too_many_words):
    triggered_lines = []

    if too_many_lines is True:
        triggered_lines.append("- No More Than Two Lines")
    if too_many_words is True:
        triggered_lines.append("- Typically 15 Words in Total")

    trigger_text = "\n".join(triggered_lines)

    user_prompt = (
        TITLE_LENGTH_PROMPT
        + "\n\nTitle:\n"
        + title_text
        + "\n\nCurrent counts:\n"
        + "- lines: "
        + str(line_count)
        + "\n- words: "
        + str(word_count)
        + "\n\nTriggered limits:\n"
        + trigger_text
    )

    try:
        llm_response = call_llm_json(
            system_prompt=TITLE_LENGTH_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            response_schema=TITLE_LENGTH_COMMENT_SCHEMA,
        )
        content = llm_response["content"]
        comment = content["comment"].strip()
        options = content["suggested_titles"]

        if comment != "" and len(options) >= 2:
            final_message = comment
            final_message += "\nSuggested Title Options:\n"
            final_message += "- " + options[0].strip() + "\n"
            final_message += "- " + options[1].strip()
            return final_message
    except Exception:
        pass

    return _build_title_length_comment_fallback(
        title_text=title_text,
        line_count=line_count,
        word_count=word_count,
        too_many_lines=too_many_lines,
        too_many_words=too_many_words,
    )


#find title candidate and capture all title-related state
def titleFound(docxpath):
    paragraphs = load_paragraphs(docxpath)
    front = get_front_page(paragraphs)

    styled_title = None
    for p in front:
        if p.is_empty is True:
            continue

        for allowed_style in TITLE_STYLE_ALIASES:
            if p.style == allowed_style:
                styled_title = p
                break

        if styled_title is not None:
            break

    first_non_empty = None
    for p in front:
        if p.is_empty is False:
            first_non_empty = p
            break

    if styled_title is not None:
        title_text = styled_title.text
        corrected_title = _fix_au_spellings_in_text(_to_title_case_title(title_text))

        style_ok = False
        for allowed_style in TITLE_STYLE_ALIASES:
            if styled_title.style == allowed_style:
                style_ok = True
                break

        return {
            "title_found": True,
            "anchor_pos": styled_title.index,
            "title_text": title_text,
            "title_style": styled_title.style,
            "style_ok": style_ok,
            "line_count": _count_title_lines(title_text),
            "word_count": _count_title_words(title_text),
            "is_title_case": _is_title_case_title(title_text),
            "corrected_title": corrected_title,
            "missing_reason": "",
        }

    for i, p in enumerate(front):
        if _normalised_label_text(p.text) not in {"title", "article title", "proposed article title"}:
            continue
        for next_p in front[i + 1:]:
            if next_p.is_empty is True:
                continue
            title_text = next_p.text
            corrected_title = _fix_au_spellings_in_text(_to_title_case_title(title_text))
            return {
                "title_found": True,
                "anchor_pos": next_p.index,
                "title_text": title_text,
                "title_style": next_p.style,
                "style_ok": next_p.style in TITLE_STYLE_ALIASES,
                "line_count": _count_title_lines(title_text),
                "word_count": _count_title_words(title_text),
                "is_title_case": _is_title_case_title(title_text),
                "corrected_title": corrected_title,
                "missing_reason": "",
            }

    if first_non_empty is None:
        return {
            "title_found": False,
            "anchor_pos": 0,
            "title_text": "",
            "title_style": "",
            "style_ok": False,
            "line_count": 0,
            "word_count": 0,
            "is_title_case": True,
            "corrected_title": "",
            "missing_reason": "No non-empty front-page text",
        }

    if _looks_like_authors_line(first_non_empty.text) is True:
        anchor_pos = _find_anchor_above(paragraphs, first_non_empty.index)
        return {
            "title_found": False,
            "anchor_pos": anchor_pos,
            "title_text": "",
            "title_style": "",
            "style_ok": False,
            "line_count": 0,
            "word_count": 0,
            "is_title_case": True,
            "corrected_title": "",
            "missing_reason": "First front-page text looks like authors, not a title",
        }

    style_ok = False
    for allowed_style in TITLE_STYLE_ALIASES:
        if first_non_empty.style == allowed_style:
            style_ok = True
            break

    title_text = first_non_empty.text
    corrected_title = _fix_au_spellings_in_text(_to_title_case_title(title_text))

    return {
        "title_found": True,
        "anchor_pos": first_non_empty.index,
        "title_text": title_text,
        "title_style": first_non_empty.style,
        "style_ok": style_ok,
        "line_count": _count_title_lines(title_text),
        "word_count": _count_title_words(title_text),
        "is_title_case": _is_title_case_title(title_text),
        "corrected_title": corrected_title,
        "missing_reason": "",
    }


#build action plan for title checks and title comments/tracked changes
def titleFormatCheck(docxpath, title_state):
    if title_state["title_found"] is False:
        message = "Missing title on front page. Add the article title above the authors paragraph using style 'Article Title'."
        return {
            "action": "add_title_comment",
            "reason": "Title not found on front page",
            "anchor_pos": title_state["anchor_pos"],
            "title_text": "",
            "corrected_title": "",
            "needs_title_case_fix": False,
            "too_many_lines": False,
            "too_many_words": False,
            "style_ok": False,
            "line_count": 0,
            "word_count": 0,
            "message": message,
        }

    style_ok = title_state["style_ok"]
    line_count = title_state["line_count"]
    word_count = title_state["word_count"]

    too_many_lines = line_count > TITLE_MAX_LINES
    too_many_words = word_count > TITLE_MAX_WORDS

    needs_title_case_fix = False
    if title_state["is_title_case"] is False:
        needs_title_case_fix = True

    comment_parts = []

    if style_ok is False:
        style_message = "Title style should be '" + TITLE_REQUIRED_STYLE + "'. "
        style_message += "Current style is '" + title_state["title_style"] + "'."
        comment_parts.append(style_message)

    if too_many_lines is True or too_many_words is True:
        length_message = _build_title_length_comment_with_llm(
            title_text=title_state["title_text"],
            line_count=line_count,
            word_count=word_count,
            too_many_lines=too_many_lines,
            too_many_words=too_many_words,
        )
        comment_parts.append(length_message)

    if needs_title_case_fix is True:
        title_case_message = "Title converted to Title Case so it matches JUTLP title formatting."
        comment_parts.append(title_case_message)

    message = "\n\n".join(comment_parts)

    action = "none"
    if needs_title_case_fix is True:
        action = "replace_title_and_comment"
    elif message.strip() != "":
        action = "add_title_comment"

    return {
        "action": action,
        "reason": "Title section checked",
        "anchor_pos": title_state["anchor_pos"],
        "title_text": title_state["title_text"],
        "title_style": title_state.get("title_style", ""),
        "corrected_title": title_state["corrected_title"],
        "needs_title_case_fix": needs_title_case_fix,
        "too_many_lines": too_many_lines,
        "too_many_words": too_many_words,
        "style_ok": style_ok,
        "line_count": line_count,
        "word_count": word_count,
        "message": message,
    }


#wrapper to run title find + title format planning
def build_title_check_plan(docxpath):
    title_state = titleFound(docxpath)
    return titleFormatCheck(docxpath, title_state)




#AUTHORS SECTION ----------------------------------------------------------------------------------------------
#1. authors founc ##fixes the issue of inccorect 'missing authors heading' comment
##deterministic check for isAuthors style found, if not, no authors. if no authors, store the location where they shoudl be in the doc (xml position probably best)
#under the title secion on the front page and put a comment there saying missing authors
#else, return the part of the document (xml position probably best)just return 'authors found' message in the console then move on to style checking

auth_start_pos = None
auth_end_pos = None


# WordprocessingML namespace used for manual XML edits inside .docx files.
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
RQ = f"{{{R}}}"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
PACKAGE_REL_Q = f"{{{PACKAGE_REL_NS}}}"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WPQ = f"{{{WP}}}"

COMMENT_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
COMMENT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
FOOTER_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
)
FOOTER_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
)
HEADER_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
)
HEADER_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
)

AUTHOR = "CopyEditor AI"
INITIALS = "AI"
DATE = now_sydney_iso()

TRACKED_CHANGE_AUTHOR = "CopyEditor AI"
TRACKED_CHANGE_DATE = now_sydney_iso()
RUN_AFFILIATION_CHECK_IN_AUTHORS = True

# REGEXES
TITLE_PREFIX_FLEX_REGEX = (
    r"(?:asst\.?\s+prof\.?|assistant professor|assoc\.?\s+prof\.?|associate professor|"
    r"a/?prof\.?|professor emeritus|emeritus professor|air commodore|air cdr|"
    r"dr\.?|doctor|prof\.?|professor|mr\.?|mrs\.?|ms\.?|miss|mx\.?|sir|dame|lady|lord|"
    r"rev\.?|reverend|fr\.?|father|rabbi|imam|sheikh|shaikh|"
    r"hon\.?|honourable|judge|justice|"
    r"capt\.?|captain|cmdr\.?|commander|col\.?|colonel|lt\.?|lieutenant|"
    r"maj\.?|major|gen\.?|general|sgt\.?|sergeant|cpl\.?|corporal|adm\.?|admiral)"
)

TITLE_PREFIX_CANONICAL_REGEX = (
    r"(?:Asst\. Prof\.|Assistant Professor|Assoc\. Prof\.|Associate Professor|"
    r"Professor Emeritus|Emeritus Professor|Air Commodore|Air Cdr|"
    r"Dr\.|Prof\.|Professor|Mr\.|Mrs\.|Ms\.|Miss|Mx\.|Sir|Dame|Lady|Lord|"
    r"Rev\.|Reverend|Fr\.|Father|Rabbi|Imam|Sheikh|Shaikh|"
    r"Hon\.|Honourable|Judge|Justice|"
    r"Capt\.|Captain|Cmdr\.|Commander|Col\.|Colonel|Lt\.|Lieutenant|"
    r"Maj\.|Major|Gen\.|General|Sgt\.|Sergeant|Cpl\.|Corporal|Adm\.|Admiral)"
)

TITLE_CANONICAL_MAP = {
    "asst prof": "Asst. Prof.",
    "assistant professor": "Assistant Professor",
    "assoc prof": "Assoc. Prof.",
    "associate professor": "Associate Professor",
    "aprof": "Assoc. Prof.",
    "professor emeritus": "Professor Emeritus",
    "emeritus professor": "Emeritus Professor",
    "air commodore": "Air Commodore",
    "air cdr": "Air Cdr",
    "dr": "Dr.",
    "doctor": "Dr.",
    "prof": "Prof.",
    "professor": "Professor",
    "mr": "Mr.",
    "mrs": "Mrs.",
    "ms": "Ms.",
    "miss": "Miss",
    "mx": "Mx.",
    "sir": "Sir",
    "dame": "Dame",
    "lady": "Lady",
    "lord": "Lord",
    "rev": "Rev.",
    "reverend": "Reverend",
    "fr": "Fr.",
    "father": "Father",
    "rabbi": "Rabbi",
    "imam": "Imam",
    "sheikh": "Sheikh",
    "shaikh": "Shaikh",
    "hon": "Hon.",
    "honourable": "Honourable",
    "judge": "Judge",
    "justice": "Justice",
    "capt": "Capt.",
    "captain": "Captain",
    "cmdr": "Cmdr.",
    "commander": "Commander",
    "col": "Col.",
    "colonel": "Colonel",
    "lt": "Lt.",
    "lieutenant": "Lieutenant",
    "maj": "Maj.",
    "major": "Major",
    "gen": "Gen.",
    "general": "General",
    "sgt": "Sgt.",
    "sergeant": "Sergeant",
    "cpl": "Cpl.",
    "corporal": "Corporal",
    "adm": "Adm.",
    "admiral": "Admiral",
}


#author helper checks -------------------------------------------------------------------------------
#normalise a title prefix to a canonical lookup key
def _normalise_title_key(title_text):
    key = title_text.strip().lower()
    key = key.replace(".", "")
    key = key.replace("/", "")
    key = re.sub(r"\s+", " ", key)
    return key


#detect known person title patterns in free text
def _has_known_title(text):
    return re.search(
        r"(^|[\s,(])(" + TITLE_PREFIX_FLEX_REGEX + r")(\s|$)",
        text,
        flags=re.IGNORECASE,
    ) is not None


def _has_trailing_person_title(author_text):
    text = author_text.strip()
    text = re.sub(r"(?:\^[a-z])+\s*$", "", text, flags=re.IGNORECASE).strip()
    return re.search(
        r"(^|\s)(" + TITLE_PREFIX_FLEX_REGEX + r")\.?$",
        text,
        flags=re.IGNORECASE,
    ) is not None


def _normalised_label_text(text):
    return _normalise_heading_text(text).rstrip(":").strip()


def _looks_like_author_name_text(text):
    words = re.findall(r"[A-Z][A-Za-z'’-]+", text)
    return len(words) >= 2 and _count_title_words(text) <= 30


def _looks_like_author_list_text(text):
    clean = text.strip()
    if clean == "" or _count_title_words(clean) > 45:
        return False
    if _has_known_title(clean) and _looks_like_author_name_text(clean):
        return True
    if ", and " not in clean.lower() and clean.count(",") < 2:
        return False
    candidates = [
        part.strip()
        for part in re.split(r",\s+and\s+|,\s*|;\s*", clean)
        if part.strip() != ""
    ]
    return len(candidates) >= 2 and all(_looks_like_author_name_text(part) for part in candidates[:3])


#fallback content-based authors detection when style is missing
def _find_content_authors(front, title_idx):
    abstract_idx = None
    for p in front:
        if p.text.strip().lower() == "abstract":
            abstract_idx = p.index
            break

    if abstract_idx is None:
        abstract_idx = 10**9

    candidates = [p for p in front if p.index > title_idx and p.index < abstract_idx]
    for i, p in enumerate(candidates):
        label = _normalised_label_text(p.text)
        if label in {"author", "authors"}:
            for next_p in candidates[i + 1:]:
                if next_p.is_empty is True:
                    continue
                return next_p.index

    for p in candidates:
        if p.index <= title_idx:
            continue
        if p.index >= abstract_idx:
            break
        if p.is_empty is True:
            continue
        if _normalised_label_text(p.text) in {
            "abstract", "keywords", "citation", "introduction", "practitioner notes", "practioner notes",
            "proposed article title", "rationale", "research questions",
        }:
            break
        if _looks_like_author_list_text(p.text):
            return p.index

    return None


#extract first non-empty line from a multi-line authors block
def _extract_authors_line_only(text):
    if text is None:
        return ""

    clean = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = clean.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line != "":
            return line
        i += 1

    return ""


def _superscript_markers_to_carets(text):
    superscripts = {v: k for k, v in SUPERSCRIPT_CHAR_MAP.items()}
    output = []
    for ch in text or "":
        if ch in superscripts:
            output.append("^" + superscripts[ch])
        else:
            output.append(ch)
    return "".join(output)


def _author_affiliation_marker_count(author_text):
    return len(re.findall(r"\^([a-z0-9])", _superscript_markers_to_carets(author_text).lower()))


def _authors_with_multiple_affiliation_markers(authors_line):
    authors = []
    for author in _split_author_candidates(authors_line or ""):
        if _author_affiliation_marker_count(author) > 1:
            authors.append(author.strip())
    return authors


def _run_has_superscript(run):
    r_pr = run._element.find(f"{WQ}rPr")
    if r_pr is None:
        return False
    vert_align = r_pr.find(f"{WQ}vertAlign")
    return vert_align is not None and vert_align.get(f"{WQ}val") == "superscript"


def _extract_authors_line_with_run_markers(docxpath, start_pos, end_pos):
    if start_pos is None or end_pos is None:
        return ""
    doc = DocxDocument(docxpath)
    parts = []
    for i in range(start_pos, min(end_pos + 1, len(doc.paragraphs))):
        text = ""
        for run in doc.paragraphs[i].runs:
            run_text = run.text or ""
            if _run_has_superscript(run):
                for ch in run_text:
                    if ch.lower() in "abcdefghijklmnopqrstuvwxyz0123456789":
                        text += "^" + ch.lower()
                    else:
                        text += ch
            else:
                text += _superscript_markers_to_carets(run_text)
        if text.strip() != "":
            parts.append(text)
    return _extract_authors_line_only("\n".join(parts))


#strict deterministic check for expected authors naming pattern
def _author_naming_pattern_valid(line):
    if line is None:
        return False

    line = line.strip()
    if line == "":
        return False

    if ", and " not in line:
        return False

    parts = line.rsplit(", and ", 1)
    if len(parts) != 2:
        return False

    head = parts[0].strip()
    last = parts[1].strip()
    if head == "" or last == "":
        return False

    authors = head.split(", ")
    authors.append(last)
    if len(authors) < 2:
        return False

    author_regex = re.compile(
        r"^(?:" + TITLE_PREFIX_CANONICAL_REGEX + r" )?[A-Z][a-z]+(?: [A-Z][a-z]+)*(?: [A-Z]\.)? [A-Z][a-zA-Z'-]+(?:\^[a-z])+$"
    )

    for author in authors:
        text = author.strip()
        if text == "":
            return False
        if ".." in text:
            return False
        if re.search(r"\d", text):
            return False

        if _has_trailing_person_title(text):
            return False

        if "^" not in text:
            return False
        if _author_affiliation_marker_count(text) != 1:
            return False
        if re.search(r"\^[a-z](?:\^[a-z])*$", text) is None:
            return False
        if author_regex.fullmatch(text) is None:
            return False

    return True


#extract leading person title from an author string
def _extract_leading_title(author_text):
    clean = author_text.replace("...", ". ").replace("..", ". ")
    match = re.match(
        r"^\s*(" + TITLE_PREFIX_FLEX_REGEX + r")\s+",
        clean,
        flags=re.IGNORECASE,
    )
    if match is None:
        return ""
    key = _normalise_title_key(match.group(1))
    if key in TITLE_CANONICAL_MAP:
        return TITLE_CANONICAL_MAP[key]
    return ""


#title-case single token helper
def _to_title_case_word(word):
    if word == "":
        return ""
    if len(word) == 1:
        return word.upper()
    return word[0].upper() + word[1:].lower()


#normalise raw author token into clean display form
def _normalise_author_core(author_text):
    text = author_text
    title = _extract_leading_title(text)
    text = text.replace("...", ". ")
    text = text.replace("..", ". ")
    text = text.replace("*", " ")
    text = text.replace("#", " ")
    text = re.sub(r"\d", "", text)
    text = re.sub(
        r"^\s*(" + TITLE_PREFIX_FLEX_REGEX + r")\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"[^A-Za-z\s'-]", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" ,")

    words = text.split(" ")
    fixed_words = []
    for word in words:
        if word == "":
            continue
        if len(word) == 2 and word.endswith(".") and word[0].isalpha():
            fixed_words.append(word[0].upper() + ".")
        else:
            fixed_words.append(_to_title_case_word(word))

    core = " ".join(fixed_words).strip()
    if title != "":
        if core == "":
            return title
        return title + " " + core
    return core


#split author candidates from a comma/and separated line
def _split_author_candidates(line):
    text = line.strip()
    text = re.sub(r"\s+", " ", text)

    if ", and " in text:
        head, last = text.rsplit(", and ", 1)
        candidates = [part.strip() for part in head.split(",") if part.strip() != ""]
        if last.strip() != "":
            candidates.append(last.strip())
        return candidates

    if " and " in text:
        head, last = text.rsplit(" and ", 1)
        candidates = [part.strip() for part in head.split(",") if part.strip() != ""]
        if last.strip() != "":
            candidates.append(last.strip())
        return candidates

    return [part.strip() for part in text.split(",") if part.strip() != ""]


#count how many author candidates include a leading title
def _count_titled_authors(line):
    candidates = _split_author_candidates(line)
    count = 0
    for candidate in candidates:
        if _extract_leading_title(candidate) != "":
            count += 1
    return count


#deterministic rewrite into expected authors pattern
def _force_author_naming_pattern(line):
    candidates = _split_author_candidates(line)
    if len(candidates) < 2:
        return line.strip()

    fixed_authors = []
    letter_code = ord("a")

    for candidate in candidates:
        core = _normalise_author_core(candidate)
        if core == "":
            continue

        suffixes = re.findall(r"\^([a-z])", candidate.lower())
        if len(suffixes) == 0:
            suffix = "^" + chr(letter_code)
            letter_code += 1
        else:
            suffix = ""
            for marker in suffixes:
                suffix += "^" + marker

        fixed_authors.append(core + suffix)

    if len(fixed_authors) < 2:
        return line.strip()

    if len(fixed_authors) == 2:
        return fixed_authors[0] + ", and " + fixed_authors[1]

    return ", ".join(fixed_authors[:-1]) + ", and " + fixed_authors[-1]


#find authors paragraph range and return authors section state
def authorFound(docxpath):
    global auth_start_pos, auth_end_pos

    paragraphs = load_paragraphs(docxpath)
    front = get_front_page(paragraphs)

    # 1) find title index for fallback anchor
    title_idx = None
    for p in front:
        if p.style == "Article Title" and p.is_empty is False:
            title_idx = p.index
            break
    if title_idx is None:
        title_idx = 0

    # 2) strict style detection
    styled_authors = []
    for p in front:
        if p.style == AUTHORS_REQUIRED_STYLE and p.is_empty is False:
            styled_authors.append(p)

    if len(styled_authors) > 0:
        auth_start_pos = styled_authors[0].index
        auth_end_pos = styled_authors[-1].index
        return {
            "authors_present": True,
            "authors_style_ok": True,
            "anchor_pos": auth_start_pos,
            "title_idx": title_idx,
        }

    # 3) fallback content detection (if style is wrong but names are there)
    content_author_idx = _find_content_authors(front, title_idx)
    if content_author_idx is not None:
        auth_start_pos = content_author_idx
        auth_end_pos = content_author_idx
        return {
            "authors_present": True,
            "authors_style_ok": False,
            "anchor_pos": auth_start_pos,
            "title_idx": title_idx,
        }

    auth_start_pos = None
    auth_end_pos = None
    return {
        "authors_present": False,
        "authors_style_ok": False,
        "anchor_pos": title_idx,
        "title_idx": title_idx,
    }


def _next_default_marker(marker):
    if marker.isdigit():
        return str(int(marker) + 1)
    if len(marker) == 1 and marker.isalpha() and marker.lower() != "z":
        return chr(ord(marker.lower()) + 1)
    return "b"


def _default_authors_line_for_markers(markers):
    if len(markers) == 0:
        return DEFAULT_AUTHORS_LINE
    first = markers[0]
    second = markers[1] if len(markers) > 1 else _next_default_marker(first)
    return f"First Author^{first}, Second Author^{first}, and Third Author^{second}"


def _extract_affiliation_markers_from_text(text):
    markers = []
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        marker = _extract_affiliation_prefix_marker(line)
        if marker != "" and marker not in markers:
            markers.append(marker)
    for marker in re.findall(r"(?:^|[;,\n])\s*([a-zA-Z0-9])\s+", text):
        marker = marker.lower()
        if marker not in markers:
            markers.append(marker)
    for marker in re.findall(r"(?:^|[;,\n])\s*([a-zA-Z])(?=[A-Z])", text):
        marker = marker.lower()
        if marker not in markers:
            markers.append(marker)
    return markers


def _find_existing_affiliations_for_missing_authors(docxpath, title_idx):
    positions = []
    markers = []
    for p in get_front_page(load_paragraphs(docxpath)):
        low = p.text.strip().lower()
        if p.index <= title_idx or p.is_empty is True:
            continue
        if low in ["abstract", "keywords", "citation", "introduction", "practitioner notes", "practioner notes"]:
            break
        if p.style == AFFILIATIONS_REQUIRED_STYLE or _paragraph_has_affiliations_block(p.text) is True:
            positions.append(p.index)
            for marker in _extract_affiliation_markers_from_text(p.text):
                if marker != "" and marker not in markers:
                    markers.append(marker)
    if len(positions) == 0:
        return None
    return {"first_pos": positions[0], "last_pos": positions[-1], "markers": markers}


def _append_affiliation_before_notes(text, affiliation):
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n") if line.strip() != ""]
    if len(lines) == 0:
        return affiliation
    target = 0
    for i, line in enumerate(lines):
        if line.startswith("*") or line.startswith("#"):
            break
        target = i
    separator = " " if lines[target].rstrip().endswith(";") else "; "
    lines[target] = lines[target].rstrip() + separator + affiliation
    return "\n".join(lines)


def _short_comment_text(text, limit=350):
    clean = re.sub(r"\s+", " ", (text or "")).strip()
    return clean if len(clean) <= limit else clean[:limit].rstrip() + "..."


def _build_author_naming_query(current_authors_line, corrected_authors_line):
    message = (
        "Check the author line formatting. Use author names only, separated by commas, "
        "with an affiliation marker after each author surname, e.g. "
        "'First Author^a, Second Author^a, and Third Author^b'."
    )
    if current_authors_line.strip() != "":
        message += "\n\nCurrent author line: " + _short_comment_text(current_authors_line)
    if corrected_authors_line.strip() != "" and corrected_authors_line.strip() != current_authors_line.strip():
        message += "\n\nSuggested author line: " + _short_comment_text(corrected_authors_line)
    return message


#2. if authors found, check styling and postion
#Correct format of authors:
#AutName I. Surnamea, Name Surnameb, and Name I. Surnamec (Style: Authors)
#use LLM: if incorrect format, get LLM to rewrite in correct format, then get that LLM (ChatGPT) output and insert it into the document as a tracked change at the xml position previously identified
#correct position of authors is directly underneath article title, using style spacing rather than a blank paragraph


def authorFormatCheck(docxpath, author_state):
    global auth_start_pos, auth_end_pos

    if author_state["authors_present"] is False:
        existing_affiliations = _find_existing_affiliations_for_missing_authors(
            docxpath,
            author_state["title_idx"],
        )
        if existing_affiliations is not None:
            markers = existing_affiliations["markers"]
            default_affiliations = ()
            plan = {
                "action": "insert_default_authors",
                "message": "Default author placeholders inserted because authors were not found.",
                "anchor_pos": author_state["anchor_pos"],
                "insert_pos": existing_affiliations["first_pos"],
                "reason": "Authors not found on front page",
                "corrected_authors_line": _default_authors_line_for_markers(markers),
                "default_affiliations": default_affiliations,
                "authors_style_ok": True,
                "authors_current_style": "",
            }
            if len(markers) == 1:
                plan["append_affiliation_pos"] = existing_affiliations["last_pos"]
                plan["append_affiliation_text"] = f"{_next_default_marker(markers[0])} Affiliation 2"
            return plan

        return {
            "action": "insert_default_authors",
            "message": "Default author placeholders inserted because authors were not found.",
            "anchor_pos": author_state["anchor_pos"],
            "reason": "Authors not found on front page",
            "corrected_authors_line": DEFAULT_AUTHORS_LINE,
            "default_affiliations": DEFAULT_AUTHOR_AFFILIATIONS,
            "authors_style_ok": True,
            "authors_current_style": "",
        }

    paragraphs = load_paragraphs(docxpath)
    _authors_style_ok = author_state.get("authors_style_ok", True)
    _authors_current_style = (
        paragraphs[auth_start_pos].style
        if auth_start_pos is not None and auth_start_pos < len(paragraphs)
        else ""
    )

    # join author text from auth_start_pos..auth_end_pos
    current_authors_block = ""
    current_authors_line = ""
    if auth_start_pos is not None and auth_end_pos is not None:
        parts = []
        i = auth_start_pos
        while i <= auth_end_pos:
            if i < len(paragraphs):
                text = paragraphs[i].text
                if text.strip() != "":
                    parts.append(text)
            i += 1
        current_authors_block = "\n".join(parts)
        current_authors_line = _extract_authors_line_only(current_authors_block)
        run_authors_line = _extract_authors_line_with_run_markers(docxpath, auth_start_pos, auth_end_pos)
        if run_authors_line != "":
            current_authors_line = run_authors_line

    multiple_affiliation_authors = _authors_with_multiple_affiliation_markers(current_authors_line)
    if len(multiple_affiliation_authors) > 0:
        return {
            "action": "add_authors_naming_comment",
            "reason": "Author has multiple affiliation markers",
            "anchor_pos": auth_start_pos,
            "corrected_authors_line": current_authors_line,
            "message": (
                "Each author should have one primary affiliation only. "
                "Please keep one affiliation marker per author. Check: "
                + ", ".join(multiple_affiliation_authors)
                + "."
            ),
            "authors_style_ok": _authors_style_ok,
            "authors_current_style": _authors_current_style,
        }

    # deterministic first, only call LLM if current line is not valid
    llm_failed = False
    llm_attempted = False
    initial_naming_valid = _author_naming_pattern_valid(current_authors_line)
    naming_pattern_valid = initial_naming_valid

    result = {
        "is_author_naming_correct": True,
        "corrected_authors_line": current_authors_line,
        "reason": "Authors naming already valid",
    }
    corrected_authors_line = current_authors_line

    if naming_pattern_valid is False:
        llm_attempted = True
        result = {
            "is_author_naming_correct": False,
            "corrected_authors_line": current_authors_line,
            "reason": "",
        }

        try:
            llm_response = call_llm_json(
                system_prompt=AUTHOR_SYSTEM_PROMPT,
                user_prompt=AUTHOR_NAMING_PATTERN + "\n\nInput authors line:\n" + current_authors_line,
                response_schema=AUTHOR_NAMING_PATTERN_SCHEMA,
            )
            result = llm_response["content"]
        except Exception as e:
            llm_failed = True
            result["reason"] = "LLM failed: " + str(e)

        corrected_authors_line = _extract_authors_line_only(result["corrected_authors_line"])
        if corrected_authors_line == "":
            corrected_authors_line = current_authors_line

        naming_pattern_valid = _author_naming_pattern_valid(corrected_authors_line)
        source_titled_count = _count_titled_authors(current_authors_line)
        output_titled_count = _count_titled_authors(corrected_authors_line)
        if source_titled_count > 0 and output_titled_count < source_titled_count:
            naming_pattern_valid = False

    # affiliations marker check will be handled in the affiliations section later
    affiliation_check = _run_affiliation_validation_in_authors(
        current_authors_line,
        current_authors_block,
    )

    if naming_pattern_valid is False:
        reason = "Author naming pattern is incorrect"
        if llm_attempted is True:
            reason = "Author naming pattern is still incorrect after LLM rewrite"
        if llm_failed is True:
            reason = "LLM failed and author naming pattern is incorrect"

        message = _build_author_naming_query(current_authors_line, corrected_authors_line)
        if affiliation_check is not None:
            message += "\n\n" + affiliation_check["message"]
            reason += " + affiliation mismatch"

        return {
            "action": "add_authors_naming_comment",
            "reason": reason,
            "anchor_pos": auth_start_pos,
            "corrected_authors_line": corrected_authors_line,
            "message": message,
            "authors_style_ok": _authors_style_ok,
            "authors_current_style": _authors_current_style,
        }

    if affiliation_check is not None:
        return {
            "action": "add_affiliation_mismatch_comment",
            "reason": affiliation_check["reason"],
            "anchor_pos": auth_start_pos,
            "corrected_authors_line": corrected_authors_line,
            "message": affiliation_check["message"],
            "authors_style_ok": _authors_style_ok,
            "authors_current_style": _authors_current_style,
        }

    # check if position of authors is correct: directly underneath article title
    title_idx = author_state["title_idx"]
    expected_author_idx = title_idx + 1
    position_correct = auth_start_pos == expected_author_idx

    naming_needs_change = False
    if initial_naming_valid is False:
        naming_needs_change = True
    if corrected_authors_line.strip() != current_authors_line.strip():
        naming_needs_change = True

    # if result.is_author_style == False
    if naming_needs_change is True:
        return {
            "action": "replace_in_place",
            "reason": result["reason"],
            "anchor_pos": auth_start_pos,
            "corrected_authors_line": corrected_authors_line,
            "replace_first_line_only": True,
            "message": "Author format corrected, inserted tracked changes",
            "authors_style_ok": _authors_style_ok,
            "authors_current_style": _authors_current_style,
        }

    # naming is correct, then only position may need move
    if position_correct is False:
        return {
            "action": "move_only",
            "reason": "Author naming is correct but position is wrong",
            "anchor_pos": auth_start_pos,
            "target_pos": expected_author_idx,
            "corrected_authors_line": corrected_authors_line,
            "message": "Author position corrected, inserted tracked changes",
            "authors_style_ok": _authors_style_ok,
            "authors_current_style": _authors_current_style,
        }

    # Style wrong but everything else is fine — still need a tracked style fix.
    if _authors_style_ok is False:
        return {
            "action": "fix_author_style",
            "reason": "Authors paragraph style is incorrect",
            "anchor_pos": auth_start_pos,
            "corrected_authors_line": corrected_authors_line,
            "message": (
                f"Author paragraph style is '{_authors_current_style}' — "
                f"should be '{AUTHORS_REQUIRED_STYLE}'. Tracked style change applied."
            ),
            "authors_style_ok": False,
            "authors_current_style": _authors_current_style,
        }

    return {
        "action": "none",
        "reason": result["reason"],
        "anchor_pos": auth_start_pos,
        "corrected_authors_line": corrected_authors_line,
        "message": "Authors naming and position already correct",
        "authors_style_ok": True,
        "authors_current_style": _authors_current_style,
    }


def build_author_check_plan(docxpath):
    author_state = authorFound(docxpath)
    return authorFormatCheck(docxpath, author_state)


#AFFLILIATIONS SECTION ----------------------------------------------------------------------------------------------
#markers and lines parsing for affiliation checking

#split affiliations lines from authors block (skip first author line)
def _extract_authors_affiliations_section_lines(text):
    clean = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = clean.split("\n")

    affiliations_lines = []
    i = 1
    while i < len(lines):
        line = lines[i].strip()
        if line != "":
            affiliations_lines.append(line)
        i += 1

    return affiliations_lines


#extract affiliation markers from authors line (^a style only)
def _extract_author_affiliation_markers(authors_line):
    markers = set()
    if authors_line is None:
        return markers

    found = re.findall(r"\^([a-z0-9])", authors_line.lower())
    for marker in found:
        markers.add(marker)

    return markers


def _default_affiliations_text_for_authors(docxpath):
    global auth_start_pos, auth_end_pos

    author_state = authorFound(docxpath)
    if author_state["authors_present"] is False or auth_start_pos is None:
        return ""

    authors_line = _extract_authors_line_with_run_markers(docxpath, auth_start_pos, auth_end_pos)
    if authors_line == "":
        paragraphs = load_paragraphs(docxpath)
        parts = []
        end_pos = auth_end_pos if auth_end_pos is not None else auth_start_pos
        for i in range(auth_start_pos, min(end_pos + 1, len(paragraphs))):
            if paragraphs[i].text.strip() != "":
                parts.append(paragraphs[i].text)
        authors_line = _extract_authors_line_only("\n".join(parts))

    markers = sorted(_extract_author_affiliation_markers(authors_line))
    if len(markers) == 0:
        markers = ["a", "b"]

    return "; ".join(f"{marker} Affiliation {i}" for i, marker in enumerate(markers, start=1))


#extract prefix marker from one affiliation line
def _extract_affiliation_prefix_marker(affiliation_line):
    text = affiliation_line.strip()
    if text == "":
        return ""

    if text[0] == "*" or text[0] == "#":
        return ""

    if text[0].isdigit():
        return text[0]

    if len(text) >= 2:
        if text[0].isalpha() and text[1].isupper():
            return text[0].lower()
        if text[0].isalpha() and text[1].isspace():
            return text[0].lower()

    if len(text) == 1 and text[0].isalpha():
        return text[0].lower()

    return ""


#collect all affiliation prefix markers from affiliations lines
def _extract_affiliation_markers(affiliations_lines):
    markers = set()

    for line in affiliations_lines:
        marker = _extract_affiliation_prefix_marker(line)
        if marker != "":
            markers.add(marker)

    return markers


#superscript marker formatting for affiliation letters
#append plain text run to XML paragraph/run container
def _append_plain_text_run(parent_element, text):
    if text == "":
        return

    run = etree.SubElement(parent_element, f"{WQ}r")
    run_text = etree.SubElement(run, f"{WQ}t")
    run_text.text = text
    if text[0].isspace() or text[-1].isspace():
        run_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _append_plain_text_run_with_size(parent_element, text, size_half_points):
    if text == "":
        return

    run = etree.SubElement(parent_element, f"{WQ}r")
    run_properties = etree.SubElement(run, f"{WQ}rPr")

    run_size = etree.SubElement(run_properties, f"{WQ}sz")
    run_size.set(f"{WQ}val", str(size_half_points))

    run_size_cs = etree.SubElement(run_properties, f"{WQ}szCs")
    run_size_cs.set(f"{WQ}val", str(size_half_points))

    run_text = etree.SubElement(run, f"{WQ}t")
    run_text.text = text
    if text[0].isspace() or text[-1].isspace():
        run_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


SUPERSCRIPT_CHAR_MAP = {
    "0": "⁰",
    "1": "¹",
    "2": "²",
    "3": "³",
    "4": "⁴",
    "5": "⁵",
    "6": "⁶",
    "7": "⁷",
    "8": "⁸",
    "9": "⁹",
    "a": "ᵃ",
    "b": "ᵇ",
    "c": "ᶜ",
    "d": "ᵈ",
    "e": "ᵉ",
    "f": "ᶠ",
    "g": "ᵍ",
    "h": "ʰ",
    "i": "ⁱ",
    "j": "ʲ",
    "k": "ᵏ",
    "l": "ˡ",
    "m": "ᵐ",
    "n": "ⁿ",
    "o": "ᵒ",
    "p": "ᵖ",
    "r": "ʳ",
    "s": "ˢ",
    "t": "ᵗ",
    "u": "ᵘ",
    "v": "ᵛ",
    "w": "ʷ",
    "x": "ˣ",
    "y": "ʸ",
    "z": "ᶻ",
}


#convert ^marker tokens to superscript unicode text
def _caret_markers_to_superscript_text(text):
    if text is None:
        return ""

    output = []
    i = 0

    while i < len(text):
        ch = text[i]
        if ch == "^" and i + 1 < len(text):
            marker = text[i + 1]
            lower = marker.lower()
            if lower in SUPERSCRIPT_CHAR_MAP:
                output.append(SUPERSCRIPT_CHAR_MAP[lower])
            else:
                output.append(marker)
            i += 2
            continue

        output.append(ch)
        i += 1

    return "".join(output)


#append text after converting any caret markers to superscript
def _append_text_with_superscript_markers(parent_element, text):
    _append_plain_text_run(parent_element, _caret_markers_to_superscript_text(text))


#affiliation validation helpers
#build readable mismatch message from affiliation check result
def _build_affiliation_mismatch_message(affiliation_result):
    missing_in_affiliations = []
    for marker in affiliation_result["missing_in_affiliations"]:
        text = str(marker).strip().lower()
        if text != "":
            missing_in_affiliations.append(text)
    missing_in_affiliations = sorted(set(missing_in_affiliations))

    missing_in_authors = []
    for marker in affiliation_result["missing_in_authors"]:
        text = str(marker).strip().lower()
        if text != "":
            missing_in_authors.append(text)
    missing_in_authors = sorted(set(missing_in_authors))

    parts = []
    if len(missing_in_affiliations) > 0:
        parts.append(
            "Author marker(s) "
            + ", ".join(missing_in_affiliations)
            + " are missing from the affiliations section."
        )
    if len(missing_in_authors) > 0:
        parts.append(
            "Affiliation marker(s) "
            + ", ".join(missing_in_authors)
            + " are not used by any author."
        )
    if len(parts) == 0:
        parts.append("Author and affiliation markers do not match.")

    if affiliation_result["reason"].strip() != "":
        parts.append(affiliation_result["reason"].strip())

    return " ".join(parts)


#run affiliation consistency check (LLM) from source author line + affiliation lines
def _run_affiliation_validation_in_authors(source_authors_line, current_authors_block):
    if RUN_AFFILIATION_CHECK_IN_AUTHORS is False:
        return None

    affiliations_lines = _extract_authors_affiliations_section_lines(current_authors_block)
    affiliation_text = "\n".join(affiliations_lines)
    affiliation_result = {
        "is_marker_mapping_consistent": True,
        "missing_in_affiliations": [],
        "missing_in_authors": [],
        "reason": "",
    }

    try:
        affiliation_llm_response = call_llm_json(
            system_prompt=AFFILIATION_MATCH_SYSTEM_PROMPT,
            user_prompt=(
                AFFILIATION_MATCH_PROMPT
                + "\n\nAuthors line:\n"
                + source_authors_line
                + "\n\nAffiliations section lines:\n"
                + affiliation_text
            ),
            response_schema=AFFILIATION_MATCH_SCHEMA,
        )
        affiliation_result = affiliation_llm_response["content"]
    except Exception:
        return None

    if affiliation_result["is_marker_mapping_consistent"] is True:
        return None

    return {
        "reason": "Author and affiliation markers do not match",
        "message": _build_affiliation_mismatch_message(affiliation_result),
    }



##ABSTRACT SECTION
#find heading 'Abstract' - deterministic
#abstract section follows the author affiliations section - deterministic
#Abstract heading must use style: 'Heading Front Page' - deterministic
#underneath the 'Abstract' heading there is a paragraph
#paragraph checks are:
# - a single paragraphm unstructured - derterministic
# - includes problem statement, any theoretical framework used, method, key findings, and implications - LLM check (if any of those are missing, comment that its missing and add a couple of suggested inclusions)
# - IMPORTANT: The abstract should be no longer than lines 7 to 23 and a maximum of 250. - deterministic
# - paragraph uses (Style: Front Page Text) - deterministic
# THEN -> the Practitioner notes section of the abstract, whcih must be underneath the abstract paragraph with a blank line in between it and the paragraph:
# - must contain heading '"Practitioner Notes" in style  "Heading Front Page" - deterministic
# - directly underneath that heading, there must be 5 short statements which are summary points,(style: Practitioner Notes) Each statement should be short and no more two lines of text each. - deterministic

ABSTRACT_REQUIRED_HEADING_STYLE = _resolve_style_name("Heading Front Page")
ABSTRACT_REQUIRED_PARAGRAPH_STYLE = _resolve_style_name("Front Page Text")
ABSTRACT_REQUIRED_HEADING_STYLE_ID = _resolve_style_id(ABSTRACT_REQUIRED_HEADING_STYLE)
ABSTRACT_REQUIRED_PARAGRAPH_STYLE_ID = _resolve_style_id(ABSTRACT_REQUIRED_PARAGRAPH_STYLE)
# Abstract front-page length rule (single source of truth in the canonical
# template): the abstract must fit lines 7–23 of the front page (= 17 lines)
# and be at most 250 words.
_ABSTRACT_RULE = CANONICAL_STRUCTURE["front_page"]
ABSTRACT_REGION_START_LINE, ABSTRACT_REGION_END_LINE = _ABSTRACT_RULE["abstract_line_region"]
ABSTRACT_MAX_LINES = _ABSTRACT_RULE["abstract_max_lines"]
ABSTRACT_MAX_WORDS = _ABSTRACT_RULE["abstract_max_words"]
PRACTITIONER_REQUIRED_HEADING_STYLE = _resolve_style_name("Heading Front Page")
PRACTITIONER_REQUIRED_NOTE_STYLE = _resolve_style_name("Practitioner Notes")
PRACTITIONER_REQUIRED_HEADING_STYLE_ID = _resolve_style_id(PRACTITIONER_REQUIRED_HEADING_STYLE)
PRACTITIONER_REQUIRED_NOTE_STYLE_ID = _resolve_style_id(PRACTITIONER_REQUIRED_NOTE_STYLE)
PRACTITIONER_INSERTED_HEADING_SIZE_HALF_POINTS = 28
PRACTITIONER_REQUIRED_NOTE_COUNT = 5
PRACTITIONER_MAX_LINES_PER_NOTE = 2
ABSTRACT_WORDS_PER_LINE_ESTIMATE = _ABSTRACT_RULE["abstract_words_per_line_estimate"]


#find practitioner notes heading from front page
def _find_practitioner_heading_in_front(front):
    for p in front:
        low = p.text.strip().lower()
        if low == "practitioner notes" or low == "practioner notes":
            return p
    return None


#heuristic check for if line looks like affiliation institution line
def _looks_like_affiliation_line_for_front_page(text):
    line = text.strip()
    if line == "":
        return False
    if line[0] == "*" or line[0] == "#":
        return False

    marker = _extract_affiliation_prefix_marker(line)
    if marker == "":
        return False

    if len(line) <= 2:
        return False

    low = line.lower()
    hint_words = [
        "university",
        "school",
        "faculty",
        "department",
        "institute",
        "college",
        "centre",
        "center",
        "campus",
        "unsw",
        "rmit",
        "sydney",
        "canberra",
        "australia",
    ]
    for word in hint_words:
        if word in low:
            return True

    remaining = line[1:].strip()
    if "," in remaining:
        return True
    if len(remaining.split()) >= 2:
        return True

    return False


#check whether paragraph contains an affiliations block
def _paragraph_has_affiliations_block(text):
    if text is None:
        return False

    clean = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = clean.split("\n")

    i = 0
    while i < len(lines):
        if _looks_like_affiliation_line_for_front_page(lines[i]) is True:
            return True
        i += 1

    return False


#build compact front-page dump for LLM affiliation-location check
def _build_front_page_affiliations_dump(front, abstract_heading_index):
    lines = []

    for p in front:
        if abstract_heading_index is not None and p.index >= abstract_heading_index:
            break

        row = "index=" + str(p.index)
        row += " | style=" + str(p.style)
        row += " | text=" + p.text.replace("\n", " \\n ")
        lines.append(row)

    return "\n".join(lines)


#ask LLM to identify where affiliation section ends
def _find_affiliations_end_pos_with_llm(front, abstract_heading_index, skip_index=None):
    dump_text = _build_front_page_affiliations_dump(front, abstract_heading_index)
    if dump_text.strip() == "":
        return None

    valid_indexes = set()
    for p in front:
        if abstract_heading_index is not None and p.index >= abstract_heading_index:
            break
        if p.index == skip_index:
            continue
        valid_indexes.add(p.index)

    try:
        llm_response = call_llm_json(
            system_prompt=AFFILIATION_SECTION_FOUND_SYSTEM_PROMPT,
            user_prompt=AFFILIATION_SECTION_FOUND_PROMPT + "\n\nFront-page paragraphs:\n" + dump_text,
            response_schema=AFFILIATION_SECTION_FOUND_SCHEMA,
        )
        content = llm_response["content"]
        if content["affiliations_found"] is False:
            return None

        guessed_index = int(content["affiliations_end_index"])
        if guessed_index in valid_indexes:
            return guessed_index
    except Exception:
        pass

    return None


#find end of affiliations section with style check + heuristics + LLM fallback
def _find_affiliations_end_pos(front, abstract_heading_index, skip_index=None):
    affiliations_end_pos = None

    for p in front:
        if abstract_heading_index is not None and p.index >= abstract_heading_index:
            break
        if p.index == skip_index:
            continue
        if p.style == AFFILIATIONS_REQUIRED_STYLE and p.is_empty is False:
            affiliations_end_pos = p.index

    if affiliations_end_pos is not None:
        return affiliations_end_pos

    for p in front:
        if abstract_heading_index is not None and p.index >= abstract_heading_index:
            break
        if p.index == skip_index:
            continue
        if p.is_empty is True:
            continue
        if _paragraph_has_affiliations_block(p.text) is True:
            affiliations_end_pos = p.index

    if affiliations_end_pos is not None:
        return affiliations_end_pos

    return _find_affiliations_end_pos_with_llm(front, abstract_heading_index, skip_index)


#marks where abstract body scanning should stop
def _is_abstract_body_end_marker(paragraph):
    low = paragraph.text.strip().lower()
    if low == "practioner notes" or low == "practitioner notes":
        return True
    if low == "keywords" or low == "keyword":
        return True
    if low.startswith("keywords:"):
        return True
    if low == "citation":
        return True
    if low.startswith("citation:"):
        return True
    if paragraph.style == "Heading 1" and low == "introduction":
        return True
    return False


#line counting for abstract paragraph (supports wrapped-line estimate)
def _count_abstract_lines(text):
    plain_line_count = _count_title_lines(text)
    if text.strip() == "":
        return 0

    if "\n" in text:
        return plain_line_count

    word_count = _count_title_words(text)
    if word_count == 0:
        return plain_line_count

    estimated_line_count = (word_count + ABSTRACT_WORDS_PER_LINE_ESTIMATE - 1) // ABSTRACT_WORDS_PER_LINE_ESTIMATE
    if estimated_line_count < plain_line_count:
        return plain_line_count
    return estimated_line_count


#LLM abstract content coverage check (problem/method/findings/etc)
def _check_abstract_components_with_llm(abstract_text):
    check_result = {
        "checked": False,
        "is_complete": True,
        "missing_components": [],
        "suggested_inclusions": [],
        "reason": "",
    }

    if abstract_text.strip() == "":
        return check_result

    try:
        llm_response = call_llm_json(
            system_prompt=ABSTRACT_COMPONENTS_SYSTEM_PROMPT,
            user_prompt=ABSTRACT_COMPONENTS_PROMPT + "\n\nAbstract paragraph:\n" + abstract_text,
            response_schema=ABSTRACT_COMPONENTS_SCHEMA,
        )
        content = llm_response["content"]

        missing_components = []
        for item in content["missing_components"]:
            text = str(item).strip()
            if text != "":
                missing_components.append(text)

        suggested_inclusions = []
        for item in content["suggested_inclusions"]:
            text = str(item).strip()
            if text != "":
                suggested_inclusions.append(text)

        check_result["checked"] = True
        check_result["is_complete"] = content["is_abstract_components_complete"]
        check_result["missing_components"] = missing_components
        check_result["suggested_inclusions"] = suggested_inclusions
        check_result["reason"] = str(content["reason"]).strip()
    except Exception:
        pass

    return check_result


#build readable abstract-component-missing message with suggestions
def _build_missing_abstract_components_message(abstract_check):
    missing_text = ", ".join(abstract_check["missing_components"])
    message = "Abstract content is missing required part(s): " + missing_text + "."

    suggestions = abstract_check["suggested_inclusions"]
    if len(suggestions) < 2:
        suggestions = [
            "Add one sentence that clearly states the study problem and context.",
            "Add one sentence that states method, key findings, and implications.",
        ]

    message += "\nSuggested inclusions:\n"
    message += "- " + suggestions[0]
    message += "\n- " + suggestions[1]

    if abstract_check["reason"] != "":
        message += "\nReason: " + abstract_check["reason"]

    return message


#collect full abstract section state from front page
def abstractFound(docxpath):
    paragraphs = load_paragraphs(docxpath)
    front = get_front_page(paragraphs)

    # The first non-empty front-page paragraph is the article title. It must
    # never be reclassified as an affiliation: the affiliation heuristic treats
    # a leading "A " (the article) as an affiliation marker, so a title like
    # "A Three-Year Evaluation of ..." would otherwise be restyled to
    # AuthorAffiliations, clobbering the Article Title style the title pass set.
    title_index = next((p.index for p in front if not p.is_empty), None)

    abstract_heading = None
    for p in front:
        if p.text.strip().lower() == "abstract":
            abstract_heading = p
            break

    abstract_heading_index = None
    if abstract_heading is not None:
        abstract_heading_index = abstract_heading.index

    affiliations_end_pos = _find_affiliations_end_pos(front, abstract_heading_index, title_index)
    affiliation_style_fix_positions = []
    affiliation_paragraph_positions = []
    if affiliations_end_pos is not None:
        direct_space_before_positions = _find_direct_space_before_positions(docxpath)
        for p in front:
            if abstract_heading_index is not None and p.index >= abstract_heading_index:
                break
            if p.index > affiliations_end_pos:
                break
            if p.is_empty is True:
                continue
            if p.index == title_index:
                continue  # never restyle the article title as an affiliation
            if p.style == AFFILIATIONS_REQUIRED_STYLE:
                affiliation_paragraph_positions.append(p.index)
                if p.index in direct_space_before_positions:
                    affiliation_style_fix_positions.append(p.index)
                continue
            if _paragraph_has_affiliations_block(p.text) is True:
                affiliation_paragraph_positions.append(p.index)
                affiliation_style_fix_positions.append(p.index)

    affiliation_after_zero_positions = []
    if len(affiliation_paragraph_positions) > 1:
        affiliation_after_zero_positions = affiliation_paragraph_positions[:-1]

    anchor_pos = 0
    if affiliations_end_pos is not None:
        anchor_pos = affiliations_end_pos
    else:
        for p in front:
            if p.is_empty is False:
                anchor_pos = p.index
                break

    practitioner_heading = _find_practitioner_heading_in_front(front)

    abstract_all_body = []
    abstract_body_paragraphs = []
    if abstract_heading is not None:
        for p in front:
            if p.index <= abstract_heading.index:
                continue
            if _is_abstract_body_end_marker(p) is True:
                break
            abstract_all_body.append(p)
            if p.is_empty is False:
                abstract_body_paragraphs.append(p)

    # Merge positions: all paragraphs from first to last non-empty (inclusive of empty ones between)
    abstract_merge_positions = []
    if len(abstract_body_paragraphs) >= 1:
        first_pos = abstract_body_paragraphs[0].index
        last_pos = abstract_body_paragraphs[-1].index
        for p in abstract_all_body:
            if first_pos <= p.index <= last_pos:
                abstract_merge_positions.append(p.index)

    abstract_paragraph = None
    if len(abstract_body_paragraphs) > 0:
        abstract_paragraph = abstract_body_paragraphs[0]

    abstract_line_count = 0
    abstract_word_count = 0
    abstract_paragraph_style_ok = False
    abstract_text = ""
    abstract_paragraph_pos = 0
    if abstract_paragraph is not None:
        abstract_text = " ".join(p.text for p in abstract_body_paragraphs)
        abstract_paragraph_pos = abstract_paragraph.index
        abstract_line_count = _count_abstract_lines(abstract_text)
        abstract_word_count = _count_title_words(abstract_text)
        if abstract_paragraph.style == ABSTRACT_REQUIRED_PARAGRAPH_STYLE:
            abstract_paragraph_style_ok = True

    blank_line_between_abstract_and_practitioner = False
    if abstract_paragraph is not None and practitioner_heading is not None:
        if practitioner_heading.index > abstract_paragraph.index:
            i = abstract_paragraph.index + 1
            while i < practitioner_heading.index:
                if i < len(paragraphs) and paragraphs[i].is_empty is True:
                    blank_line_between_abstract_and_practitioner = True
                    break
                i += 1

    practitioner_notes = []
    practitioner_non_note_lines_count = 0
    practitioner_non_note_positions = []
    practitioner_notes_directly_under_heading = False
    if practitioner_heading is not None:
        i = practitioner_heading.index + 1
        while i < len(paragraphs):
            p = paragraphs[i]
            low = p.text.strip().lower()
            if low == "keywords":
                break
            if p.style == "Heading 1" and low == "introduction":
                break

            if p.is_empty is False:
                if practitioner_notes_directly_under_heading is False:
                    if p.style == PRACTITIONER_REQUIRED_NOTE_STYLE:
                        practitioner_notes_directly_under_heading = True

                if p.style == PRACTITIONER_REQUIRED_NOTE_STYLE:
                    practitioner_notes.append(p)
                else:
                    practitioner_non_note_lines_count += 1
                    practitioner_non_note_positions.append(p.index)
            i += 1

    practitioner_notes_too_long_count = 0
    for p in practitioner_notes:
        if _count_title_lines(p.text) > PRACTITIONER_MAX_LINES_PER_NOTE:
            practitioner_notes_too_long_count += 1

    abstract_heading_found = abstract_heading is not None
    abstract_heading_style_ok = False
    abstract_heading_pos = 0
    if abstract_heading is not None:
        abstract_heading_pos = abstract_heading.index
        if abstract_heading.style == ABSTRACT_REQUIRED_HEADING_STYLE:
            abstract_heading_style_ok = True

    practitioner_heading_found = practitioner_heading is not None
    practitioner_heading_style_ok = False
    practitioner_heading_pos = 0
    if practitioner_heading is not None:
        practitioner_heading_pos = practitioner_heading.index
        if practitioner_heading.style == PRACTITIONER_REQUIRED_HEADING_STYLE:
            practitioner_heading_style_ok = True

    follows_affiliations = False
    if abstract_heading is not None and affiliations_end_pos is not None:
        if abstract_heading.index > affiliations_end_pos:
            follows_affiliations = True

    return {
        "anchor_pos": anchor_pos,
        "affiliations_found": affiliations_end_pos is not None,
        "affiliations_end_pos": affiliations_end_pos if affiliations_end_pos is not None else -1,
        "affiliation_style_fix_positions": affiliation_style_fix_positions,
        "affiliation_after_zero_positions": affiliation_after_zero_positions,
        "abstract_found": abstract_heading_found,
        "abstract_heading_pos": abstract_heading_pos,
        "abstract_heading_style_ok": abstract_heading_style_ok,
        "abstract_follows_affiliations": follows_affiliations,
        "abstract_paragraph_found": abstract_paragraph is not None,
        "abstract_paragraph_pos": abstract_paragraph_pos,
        "abstract_paragraph_count": len(abstract_body_paragraphs),
        "abstract_body_paragraph_positions": [p.index for p in abstract_body_paragraphs],
        "abstract_merge_positions": abstract_merge_positions,
        "abstract_paragraph_style_ok": abstract_paragraph_style_ok,
        "abstract_line_count": abstract_line_count,
        "abstract_word_count": abstract_word_count,
        "abstract_text": abstract_text,
        "practitioner_heading_found": practitioner_heading_found,
        "practitioner_heading_pos": practitioner_heading_pos,
        "practitioner_heading_style_ok": practitioner_heading_style_ok,
        "blank_line_between_abstract_and_practitioner": blank_line_between_abstract_and_practitioner,
        "practitioner_notes_directly_under_heading": practitioner_notes_directly_under_heading,
        "practitioner_note_count": len(practitioner_notes),
        "practitioner_non_note_lines_count": practitioner_non_note_lines_count,
        "practitioner_non_note_positions": practitioner_non_note_positions,
        "practitioner_notes_too_long_count": practitioner_notes_too_long_count,
    }


_DEIDENTIFIED_FILENAME_RE = re.compile(
    # Match deidentified markers only when they sit at a token boundary in the
    # FILE name (not anywhere along the path). Otherwise a temp directory like
    # `pytest-of-…/test_deidentified_xx/…` triggers a false positive on the
    # test directory's name.
    r"(?:^|[\s_\-+])(?:deidentified|de[-_]identified|anonymi[sz]ed|blinded|blind[-_]review)(?=[\s_\-+.]|$)",
    re.IGNORECASE,
)

# Placeholder text authors commonly insert when they strip identity
# information from a manuscript before peer review. Covers
# `[Authors removed for review]`, `[Anonymised]`, `[Affiliations
# removed]`, `[Author affiliations withheld for blind review]`, etc.
_DEIDENTIFIED_PLACEHOLDER_RE = re.compile(
    r"\[[^\]]*?(?:removed|redacted|blinded|anonymi[sz]ed|withheld|for[ -]review)[^\]]*?\]",
    re.IGNORECASE,
)

# Numbered placeholder authors use as a stand-in for real names, e.g.
# `[Author 1]`, `[Author 2]`, sometimes `[Affiliation 1]`. Matched inside
# square brackets so an actual citation like "(Author, 2024)" doesn't
# trip it.
_DEIDENTIFIED_NUMBERED_AUTHOR_RE = re.compile(
    r"\[\s*(?:author|affiliation)\s*\d+\s*\]",
    re.IGNORECASE,
)


def _front_page_paragraphs_for_deid(docxpath):
    """Read up to the first ~15 front-page paragraphs of ``docxpath``.

    Used by the deidentified-manuscript content check. Best-effort — any
    error reading the file results in an empty list (filename signal
    remains the fallback). Capped at 15 paragraphs so a malformed doc
    can't slow the check down.
    """
    try:
        from app.services.document_analysis_services import (
            get_front_page,
            load_paragraphs,
        )
        paragraphs = load_paragraphs(str(docxpath))
        front = get_front_page(paragraphs)
        return front[:15]
    except Exception:
        return []


def _is_deidentified_manuscript(docxpath, abstract_state):
    """Return True if the manuscript was deliberately stripped of identity.

    Three signals, any one is sufficient. The filename check is the
    fast-path so the second docx read only happens when needed:

    1. The filename contains a deidentification marker
       (``deidentified``, ``anonymised``, ``blinded``, etc.).
    2. The Authors paragraph(s) are styled but empty/whitespace — the
       author replaced their names with blank lines.
    3. Any front-page paragraph contains a placeholder pattern like
       ``[Authors removed for review]``, ``[Anonymised]``,
       ``[Author 1]``, etc.

    Returning True silences the "Author affiliations section was not
    found" query that otherwise fires for every deidentified submission.
    """
    try:
        # Match only against the basename — running the regex on the full
        # path produces false positives when a parent directory happens to
        # contain "deid" (e.g. pytest tmp dirs).
        from pathlib import Path as _Path
        basename = _Path(str(docxpath)).name
        if _DEIDENTIFIED_FILENAME_RE.search(basename):
            return True
    except Exception:
        pass

    for p in _front_page_paragraphs_for_deid(docxpath):
        # Empty paragraph styled as Authors or Author Affiliations — a
        # blanked-out identity line.
        if p.style in ("Authors", "Author Affiliations") and not p.text.strip():
            return True
        text = p.text or ""
        if not text.strip():
            continue
        if _DEIDENTIFIED_PLACEHOLDER_RE.search(text):
            return True
        if _DEIDENTIFIED_NUMBERED_AUTHOR_RE.search(text):
            return True

    return False


#build abstract action plan with separated heading/body anchored comments
def abstractFormatCheck(docxpath, abstract_state):
    heading_comment_parts = []
    body_comment_parts = []
    # The abstract length flag is emitted as its OWN comment (not folded into
    # the body bundle) so the downstream Sam-comment de-duplication can't drop
    # it just because the bundled merge/style text overlaps an LLM abstract note.
    length_comment = ""
    anchor_pos = abstract_state["anchor_pos"]
    _merge_abstract = False
    style_fixes = []
    insert_default_affiliations = False
    default_affiliations_text = ""
    insert_practitioner_stub = False
    ensure_practitioner_section = False
    practitioner_missing_comment = ""

    if abstract_state["abstract_found"] is False:
        message = "Missing 'Abstract' heading on the front page after author affiliations."
        deidentified = _is_deidentified_manuscript(docxpath, abstract_state)
        default_affiliations_insert_pos = abstract_state["anchor_pos"]
        if abstract_state["affiliations_found"] is False and not deidentified:
            default_affiliations_text = _default_affiliations_text_for_authors(docxpath)
            if default_affiliations_text != "":
                insert_default_affiliations = True
                if auth_end_pos is not None:
                    default_affiliations_insert_pos = auth_end_pos + 1
        practitioner_insert_after_pos = abstract_state["anchor_pos"]
        if abstract_state["affiliations_found"] is True:
            practitioner_insert_after_pos = abstract_state["affiliations_end_pos"]
        elif insert_default_affiliations is True:
            practitioner_insert_after_pos = default_affiliations_insert_pos
        if abstract_state["practitioner_heading_found"] is False:
            insert_practitioner_stub = True
            ensure_practitioner_section = True
            practitioner_missing_comment = (
                "Practitioner Notes section not found.\n\n"
                "Tracked change applied: 'Practitioner Notes' heading and one placeholder dot point have been inserted after the front-page author/affiliation block because the Abstract heading is also missing.\n\n"
                "Please add the Abstract above it, then replace the placeholder with 5 concise practitioner notes (max 2 lines each) using the 'Practitioner Notes' paragraph style."
            )
        elif abstract_state["practitioner_note_count"] == 0:
            ensure_practitioner_section = True
            practitioner_missing_comment = (
                "Practitioner Notes heading found but no note dot points are present.\n\n"
                "Tracked change applied: one placeholder dot point has been inserted under the heading.\n\n"
                "Please replace the placeholder with 5 concise practitioner notes (max 2 lines each) using the 'Practitioner Notes' paragraph style."
            )
        return {
            "action": "add_abstract_comment",
            "reason": "Abstract heading not found",
            "anchor_pos": anchor_pos,
            "message": message,
            "insert_default_affiliations": insert_default_affiliations,
            "default_affiliations_text": default_affiliations_text,
            "default_affiliations_insert_pos": default_affiliations_insert_pos,
            "insert_practitioner_stub": insert_practitioner_stub,
            "ensure_practitioner_section": ensure_practitioner_section,
            "practitioner_missing_comment": practitioner_missing_comment,
            "practitioner_insert_after_pos": practitioner_insert_after_pos,
        }

    anchor_pos = abstract_state["abstract_heading_pos"]
    deidentified = _is_deidentified_manuscript(docxpath, abstract_state)

    if abstract_state["affiliations_found"] is False:
        # Suppress the missing-affiliations query on deliberately deidentified
        # manuscripts (blinded peer review). The editor already knows the
        # affiliations are intentionally absent — flagging it adds noise.
        if not deidentified:
            default_affiliations_text = _default_affiliations_text_for_authors(docxpath)
            if default_affiliations_text != "":
                insert_default_affiliations = True
            else:
                heading_comment_parts.append("Author affiliations section was not found, so Abstract position cannot be verified.")
    elif abstract_state["abstract_follows_affiliations"] is False:
        heading_comment_parts.append("Abstract should follow the author affiliations section on the front page.")

    affiliation_fix_positions = []
    for pos in abstract_state.get("affiliation_style_fix_positions", []):
        if pos not in affiliation_fix_positions:
            affiliation_fix_positions.append(pos)
    for pos in abstract_state.get("affiliation_after_zero_positions", []):
        if pos not in affiliation_fix_positions:
            affiliation_fix_positions.append(pos)

    if len(affiliation_fix_positions) > 0:
        after_zero_positions = set(abstract_state.get("affiliation_after_zero_positions", []))
        for pos in affiliation_fix_positions:
            style_fix = {
                "paragraph_index": pos,
                "new_style": AFFILIATIONS_REQUIRED_STYLE_ID,
            }
            if pos in after_zero_positions:
                style_fix["clear_after_spacing"] = True
            style_fixes.append(
                style_fix
            )
        heading_comment_parts.append(
            "Tracked formatting change applied: Author affiliations changed to style '"
            + AFFILIATIONS_REQUIRED_STYLE
            + "' and spacing between affiliation lines was tightened."
        )

    if abstract_state["abstract_heading_style_ok"] is False:
        style_fixes.append(
            {
                "paragraph_index": abstract_state["abstract_heading_pos"],
                "new_style": ABSTRACT_REQUIRED_HEADING_STYLE_ID,
            }
        )
        heading_comment_parts.append(
            "Tracked style change applied: Abstract heading changed to style '"
            + ABSTRACT_REQUIRED_HEADING_STYLE
            + "'."
        )

    if abstract_state["abstract_paragraph_found"] is False:
        body_comment_parts.append("Add one abstract paragraph directly under the Abstract heading.")
    else:
        if len(abstract_state.get("abstract_merge_positions", [])) > 1:
            n = len(abstract_state["abstract_merge_positions"])
            body_comment_parts.insert(
                0,
                f"Abstract merged: {n} separate paragraphs combined into one continuous paragraph "
                f"as required by JUTLP style. Paragraph breaks removed as tracked changes — "
                f"please accept or reject."
            )
            # Signal that we need to apply merge tracked changes
            _merge_abstract = True

        if abstract_state["abstract_paragraph_style_ok"] is False:
            style_fixes.append(
                {
                    "paragraph_index": abstract_state["abstract_paragraph_pos"],
                    "new_style": ABSTRACT_REQUIRED_PARAGRAPH_STYLE_ID,
                }
            )
            body_comment_parts.append(
                "Tracked style change applied: Abstract paragraph changed to style '"
                + ABSTRACT_REQUIRED_PARAGRAPH_STYLE
                + "'."
            )

        # JUTLP rule: the abstract must fit lines 7–23 of the front page
        # (≈ ABSTRACT_MAX_LINES lines) AND be ≤ ABSTRACT_MAX_WORDS words. This is
        # a maximum, so only flag when the abstract is too long.
        line_count = abstract_state["abstract_line_count"]
        word_count = abstract_state["abstract_word_count"]
        if line_count > ABSTRACT_MAX_LINES or word_count > ABSTRACT_MAX_WORDS:
            length_message = (
                "Abstract is too long. Current: "
                + str(word_count) + " word(s), ~" + str(line_count) + " line(s). "
                + "JUTLP limit: a maximum of " + str(ABSTRACT_MAX_WORDS)
                + " words, and it must fit within lines "
                + str(ABSTRACT_REGION_START_LINE) + " to " + str(ABSTRACT_REGION_END_LINE)
                + " of the front page (about " + str(ABSTRACT_MAX_LINES)
                + " lines) so all summary content stays before the Introduction. "
                + "Please shorten the abstract."
            )
            length_comment = length_message

        abstract_component_check = _check_abstract_components_with_llm(abstract_state["abstract_text"])
        if abstract_component_check["checked"] is True:
            if abstract_component_check["is_complete"] is False:
                body_comment_parts.append(_build_missing_abstract_components_message(abstract_component_check))

    if abstract_state["practitioner_heading_found"] is False:
        insert_practitioner_stub = True
        ensure_practitioner_section = True
        practitioner_missing_comment = (
            "Practitioner Notes section not found.\n\n"
            "Tracked change applied: 'Practitioner Notes' heading and one placeholder dot point have been inserted under the abstract.\n\n"
            "Please replace the placeholder with 5 concise practitioner notes (max 2 lines each) using the 'Practitioner Notes' paragraph style."
        )
    else:
        if abstract_state["practitioner_heading_style_ok"] is False:
            ensure_practitioner_section = True
            style_fixes.append(
                {
                    "paragraph_index": abstract_state["practitioner_heading_pos"],
                    "new_style": PRACTITIONER_REQUIRED_HEADING_STYLE_ID,
                }
            )
            body_comment_parts.append(
                "Tracked style change applied: Practitioner Notes heading changed to style '"
                + PRACTITIONER_REQUIRED_HEADING_STYLE
                + "'."
            )

        if abstract_state["blank_line_between_abstract_and_practitioner"] is False:
            body_comment_parts.append("Keep one blank line between abstract paragraph and Practitioner Notes heading.")

        if abstract_state["practitioner_notes_directly_under_heading"] is False and abstract_state["practitioner_note_count"] > 0:
            body_comment_parts.append("Practitioner Notes statements should start directly under the heading.")

        if abstract_state["practitioner_note_count"] == 0:
            ensure_practitioner_section = True
            practitioner_missing_comment = (
                "Practitioner Notes heading found but no note dot points are present.\n\n"
                "Tracked change applied: one placeholder dot point has been inserted under the heading.\n\n"
                "Please replace the placeholder with 5 concise practitioner notes (max 2 lines each) using the 'Practitioner Notes' paragraph style."
            )
        elif abstract_state["practitioner_note_count"] != PRACTITIONER_REQUIRED_NOTE_COUNT:
            body_comment_parts.append(
                "Practitioner Notes should contain exactly 5 short statements. Found "
                + str(abstract_state["practitioner_note_count"])
                + "."
            )

        if abstract_state["practitioner_non_note_lines_count"] > 0:
            for pos in abstract_state.get("practitioner_non_note_positions", []):
                style_fixes.append(
                    {
                        "paragraph_index": pos,
                        "new_style": PRACTITIONER_REQUIRED_NOTE_STYLE_ID,
                    }
                )
            body_comment_parts.append(
                "Tracked style change applied: Practitioner Notes statements changed to style '"
                + PRACTITIONER_REQUIRED_NOTE_STYLE
                + "'."
            )

        if abstract_state["practitioner_notes_too_long_count"] > 0:
            body_comment_parts.append(
                str(abstract_state["practitioner_notes_too_long_count"])
                + " Practitioner Notes statement(s) are over two lines; keep each statement to max two lines."
            )

    comments = []

    if len(heading_comment_parts) > 0:
        comments.append(
            {
                "anchor_pos": abstract_state["abstract_heading_pos"],
                "message": "\n\n".join(heading_comment_parts),
            }
        )

    body_anchor_pos = abstract_state["abstract_heading_pos"]
    if abstract_state["abstract_paragraph_found"] is True:
        body_anchor_pos = abstract_state["abstract_paragraph_pos"]

    if len(body_comment_parts) > 0:
        comments.append(
            {
                "anchor_pos": body_anchor_pos,
                "message": "\n\n".join(body_comment_parts),
            }
        )

    # Standalone abstract-length comment (kept separate so it survives the
    # Sam-comment de-duplication even when the body bundle is removed).
    if length_comment:
        comments.append(
            {
                "anchor_pos": body_anchor_pos,
                "message": length_comment,
            }
        )

    merge_positions = abstract_state.get("abstract_merge_positions", []) if _merge_abstract else []

    if len(comments) == 0 and not _merge_abstract:
        return {
            "action": "none",
            "reason": "Abstract section checked",
            "anchor_pos": anchor_pos,
            "message": "",
            "style_fixes": style_fixes,
            "insert_default_affiliations": insert_default_affiliations,
            "default_affiliations_text": default_affiliations_text,
            "default_affiliations_insert_pos": abstract_state["abstract_heading_pos"],
            "insert_practitioner_stub": insert_practitioner_stub,
            "ensure_practitioner_section": ensure_practitioner_section,
            "practitioner_missing_comment": practitioner_missing_comment,
            "practitioner_insert_after_pos": abstract_state["abstract_paragraph_pos"],
        }

    base_action = "add_abstract_comment" if len(comments) == 1 else "add_abstract_comments"
    if _merge_abstract:
        base_action = "merge_abstract_paragraphs"

    return {
        "action": base_action,
        "reason": "Abstract section has issues",
        "anchor_pos": comments[0]["anchor_pos"] if comments else anchor_pos,
        "message": comments[0]["message"] if comments else "",
        "comments": comments,
        "merge_paragraph_positions": merge_positions,
        "style_fixes": style_fixes,
        "insert_default_affiliations": insert_default_affiliations,
        "default_affiliations_text": default_affiliations_text,
        "default_affiliations_insert_pos": abstract_state["abstract_heading_pos"],
        "insert_practitioner_stub": insert_practitioner_stub,
        "ensure_practitioner_section": ensure_practitioner_section,
        "practitioner_missing_comment": practitioner_missing_comment,
        "practitioner_insert_after_pos": abstract_state["abstract_paragraph_pos"],
    }


#wrapper to run abstract find + abstract format planning
def build_abstract_check_plan(docxpath):
    abstract_state = abstractFound(docxpath)
    return abstractFormatCheck(docxpath, abstract_state)


##KEYWORDS SECTION
#find heading 'Keywords' - deterministic
#keywords section follows Practitioner Notes section - deterministic
#Keywords heading must use style: 'Heading Front Page' - deterministic
#underneath the 'Keywords' heading there is one line of keywords text
#keywords checks are:
# - single keywords paragraph - deterministic
# - keywords line uses style 'Front Page Text' - deterministic
# - maximum 5 keywords - deterministic
# - one line only - deterministic
# - no abbreviations - deterministic

KEYWORDS_REQUIRED_HEADING_STYLE = _resolve_style_name("Heading Front Page")
KEYWORDS_REQUIRED_TEXT_STYLE = _resolve_style_name("Front Page Text")
KEYWORDS_REQUIRED_HEADING_STYLE_ID = _resolve_style_id(KEYWORDS_REQUIRED_HEADING_STYLE)
KEYWORDS_REQUIRED_TEXT_STYLE_ID = _resolve_style_id(KEYWORDS_REQUIRED_TEXT_STYLE)
KEYWORDS_MAX_COUNT = 5
KEYWORDS_MAX_LINES = 1


##CITATION SECTION
#find heading 'Citation' - deterministic
#citation text belongs in the first page footer, not in the body before Introduction

CITATION_REQUIRED_HEADING_STYLE = _resolve_style_name("Heading Front Page")
CITATION_REQUIRED_TEXT_STYLE = _resolve_style_name("Front Page Text")
CITATION_REQUIRED_HEADING_STYLE_ID = _resolve_style_id(CITATION_REQUIRED_HEADING_STYLE)
CITATION_REQUIRED_TEXT_STYLE_ID = _resolve_style_id(CITATION_REQUIRED_TEXT_STYLE)
CITATION_FOOTER_STYLE = _resolve_style_name("Citation")
CITATION_FOOTER_STYLE_ID = _resolve_style_id(CITATION_FOOTER_STYLE)


#shared heading text normaliser for front-page heading checks
def _normalise_front_heading_text(text):
    clean = text.strip().lower()
    clean = re.sub(r"\s+", " ", clean)
    return clean


#check heading word variants like keyword/keywords/citation/citations
def _matches_front_heading_text(text, heading_word):
    low = _normalise_front_heading_text(text)
    if low == heading_word:
        return True
    if low == heading_word + "s":
        return True
    if low.startswith(heading_word + ":"):
        return True
    if low.startswith(heading_word + "s:"):
        return True
    if low.startswith(heading_word + " "):
        return True
    if low.startswith(heading_word + "s "):
        return True
    if low.startswith(heading_word + "("):
        return True
    if low.startswith(heading_word + "s("):
        return True
    return False


def _extract_inline_keywords_text(text):
    match = re.match(r"^\s*keywords?\s*[:\-]\s*(.+)$", text, flags=re.IGNORECASE)
    if match is not None:
        return match.group(1).strip()

    match = re.match(r"^\s*keywords?\s+(.+)$", text, flags=re.IGNORECASE)
    if match is not None:
        return match.group(1).strip()

    return ""


#find keywords heading on front page
def _find_keywords_heading_in_front(front):
    for p in front:
        if _matches_front_heading_text(p.text, "keywords") is True:
            return p
    return None


#find citation heading on front page
def _find_citation_heading_in_front(front):
    for p in front:
        if _matches_front_heading_text(p.text, "citation") is True:
            return p
    return None


#find abstract heading on front page
def _find_abstract_heading_in_front(front):
    for p in front:
        if _matches_front_heading_text(p.text, "abstract") is True:
            return p
    return None


#find first-level Introduction heading on front page
def _find_introduction_heading_in_front(front):
    for p in front:
        low = _normalise_front_heading_text(p.text)
        if p.style == "Heading 1" and low == "introduction":
            return p
    return None


#stop markers while scanning keywords body block
def _is_keywords_body_end_marker(paragraph):
    if _matches_front_heading_text(paragraph.text, "citation") is True:
        return True
    if paragraph.style == "Guidance Notes":
        return True
    # Break on ANY heading, not just `Heading 1 + "Introduction"`. The
    # original check was too narrow: manuscripts whose Introduction is
    # missing the proper Heading 1 style (or whose front matter has a
    # subheading before Introduction) would fall through and the loop
    # would walk into body prose, treating it as the keywords line.
    if paragraph.style in {
        "Heading 1", "Heading1",
        "Heading 2", "Heading2",
        "Heading 3", "Heading3",
        "Heading 4", "Heading4",
    }:
        return True
    return False


def _looks_like_keywords_line(paragraph) -> bool:
    """Return True when ``paragraph`` plausibly contains a keyword list.

    Used to reject false positives where the Keywords heading has no
    real body and the scan would otherwise pick up the first paragraph
    of the Introduction (or any other non-keyword prose) and restyle it
    as keywords. A real keywords paragraph is:

    - already styled as the keywords body style (``Front Page Text``), OR
    - short (under 240 characters — five keywords plus separators easily
      fit) AND contains at least one separator (``,`` ``;`` or ``/``)
      AND does NOT end with a sentence terminator (``.`` ``!`` ``?``),
      which is the prose-sentence shape we explicitly want to exclude.
    """
    if paragraph is None:
        return False
    if paragraph.style == KEYWORDS_REQUIRED_TEXT_STYLE:
        return True
    text = (paragraph.text or "").strip()
    if not text or len(text) > 240:
        return False
    if text[-1] in ".!?":
        return False
    # Authors sometimes write `Keywords: foo, bar` on a single line — the
    # split pass handles that case via `_extract_inline_keywords_text`;
    # here we only need to recognise the body-line shape.
    if any(sep in text for sep in (",", ";", "/")):
        return True
    return False


#stop markers while scanning citation body block
def _is_citation_body_end_marker(paragraph):
    if paragraph.style == "Guidance Notes":
        return True
    low = _normalise_front_heading_text(paragraph.text)
    if paragraph.style == "Heading 1" and low == "introduction":
        return True
    return False


#find blank line anchor after a given paragraph index
def _find_blank_line_after_index(paragraphs, start_index, stop_index):
    i = start_index + 1
    while i < len(paragraphs):
        if stop_index is not None and i >= stop_index:
            break
        if paragraphs[i].is_empty is True:
            return i
        i += 1
    return None


#split keywords line into individual keyword items
def _extract_keywords_list(text):
    if text is None:
        return []

    clean = text.replace("\n", " ")
    clean = re.sub(r"\s+", " ", clean).strip()
    clean = re.sub(r"^(keywords?|keyword)\s*:\s*", "", clean, flags=re.IGNORECASE)

    parts = clean.split(",")
    if len(parts) == 1 and ";" in clean:
        parts = clean.split(";")

    keywords = []
    for part in parts:
        item = part.strip().strip(".")
        if item != "":
            keywords.append(item)

    return keywords


#simple abbreviation detector for keywords rule
def _keyword_looks_like_abbreviation(keyword):
    text = keyword.strip()
    if text == "":
        return False

    if "." in text:
        return True

    if re.search(r"\b[A-Z]{2,6}\b", text) is not None:
        return True

    tokens = re.findall(r"[A-Za-z]+", text)
    for token in tokens:
        if len(token) >= 2 and len(token) <= 6:
            if token.upper() == token:
                return True

    return False


#collect full keywords section state from front page
def keywordsFound(docxpath):
    paragraphs = load_paragraphs(docxpath)
    front = get_front_page(paragraphs)

    practitioner_heading = _find_practitioner_heading_in_front(front)
    keywords_heading = _find_keywords_heading_in_front(front)
    citation_heading = _find_citation_heading_in_front(front)
    intro_heading = _find_introduction_heading_in_front(front)

    citation_paragraph = None
    if citation_heading is not None:
        citation_body_paragraphs = []
        for p in front:
            if p.index <= citation_heading.index:
                continue
            if _is_citation_body_end_marker(p) is True:
                break
            if p.is_empty is False:
                citation_body_paragraphs.append(p)
        if len(citation_body_paragraphs) > 0:
            citation_paragraph = citation_body_paragraphs[0]

    anchor_pos = 0
    if keywords_heading is not None:
        anchor_pos = keywords_heading.index
    elif citation_paragraph is not None:
        anchor_pos = citation_paragraph.index
        stop_index = None
        if intro_heading is not None:
            stop_index = intro_heading.index
        blank_anchor = _find_blank_line_after_index(paragraphs, citation_paragraph.index, stop_index)
        if blank_anchor is not None:
            anchor_pos = blank_anchor
    elif citation_heading is not None:
        anchor_pos = citation_heading.index
    elif practitioner_heading is not None:
        anchor_pos = practitioner_heading.index
    else:
        for p in front:
            if p.is_empty is False:
                anchor_pos = p.index
                break

    keyword_body_paragraphs = []
    if keywords_heading is not None:
        for p in front:
            if p.index <= keywords_heading.index:
                continue
            if _is_keywords_body_end_marker(p) is True:
                break
            if p.is_empty is False:
                keyword_body_paragraphs.append(p)

    keywords_paragraph = None
    # Only accept the first candidate when it actually LOOKS like a
    # keyword list. Previously the loop unconditionally promoted
    # ``keyword_body_paragraphs[0]`` to ``keywords_paragraph``, which
    # caused the Introduction's first body sentence to be restyled as
    # keywords when (a) the Keywords heading had no real content and
    # (b) `get_front_page` couldn't locate a proper Heading 1
    # Introduction boundary (it then falls back to the whole document).
    for candidate in keyword_body_paragraphs:
        if _looks_like_keywords_line(candidate):
            keywords_paragraph = candidate
            break

    keywords_text = ""
    keywords_line_count = 0
    keywords_style_ok = False
    keywords_paragraph_pos = 0
    keywords_items = []
    abbreviation_items = []
    inline_keywords_text = ""

    if keywords_heading is not None:
        inline_keywords_text = _extract_inline_keywords_text(keywords_heading.text)
    keywords_heading_text = ""
    if keywords_heading is not None:
        keywords_heading_text = keywords_heading.text

    if keywords_paragraph is not None:
        keywords_text = keywords_paragraph.text
        keywords_paragraph_pos = keywords_paragraph.index
        keywords_line_count = _count_title_lines(keywords_text)
        if keywords_paragraph.style == KEYWORDS_REQUIRED_TEXT_STYLE:
            keywords_style_ok = True

        keywords_items = _extract_keywords_list(keywords_text)
        for item in keywords_items:
            if _keyword_looks_like_abbreviation(item) is True:
                abbreviation_items.append(item)
    elif inline_keywords_text != "":
        keywords_text = inline_keywords_text
        keywords_line_count = _count_title_lines(inline_keywords_text)
        keywords_items = _extract_keywords_list(inline_keywords_text)
        for item in keywords_items:
            if _keyword_looks_like_abbreviation(item) is True:
                abbreviation_items.append(item)

    keywords_follows_practitioner = False
    if keywords_heading is not None and practitioner_heading is not None:
        if keywords_heading.index > practitioner_heading.index:
            keywords_follows_practitioner = True

    keywords_heading_style_ok = False
    keywords_heading_pos = 0
    if keywords_heading is not None:
        keywords_heading_pos = keywords_heading.index
        if keywords_heading.style == KEYWORDS_REQUIRED_HEADING_STYLE:
            keywords_heading_style_ok = True

    return {
        "anchor_pos": anchor_pos,
        "practitioner_heading_found": practitioner_heading is not None,
        "keywords_found": keywords_heading is not None,
        "keywords_heading_pos": keywords_heading_pos,
        "keywords_heading_style_ok": keywords_heading_style_ok,
        "keywords_follows_practitioner": keywords_follows_practitioner,
        "keywords_paragraph_found": keywords_paragraph is not None,
        "keywords_paragraph_pos": keywords_paragraph_pos,
        "keywords_paragraph_count": len(keyword_body_paragraphs),
        "keywords_paragraph_style_ok": keywords_style_ok,
        "keywords_line_count": keywords_line_count,
        "keywords_count": len(keywords_items),
        "abbreviation_items": abbreviation_items,
        "keywords_items": keywords_items,
        "inline_keywords_text": inline_keywords_text,
        "keywords_heading_text": keywords_heading_text,
    }


#extract title + abstract and ask the LLM for 5 keyword candidates.
#Returns ``[]`` on any failure path so callers can fall back to the
#empty-stub + "please add 5 keywords" comment without crashing.
def _maybe_generate_keywords_for_missing_section(docxpath):
    try:
        # Local imports so unit tests can monkeypatch
        # ``output_generation_samfix.generate_keywords`` directly without
        # pulling the network-bound LLM client into module import time.
        from app.services.document_analysis_services import (
            extract_abstract,
            load_paragraphs,
        )
        from app.services.keywords_generation import generate_keywords

        paragraphs = load_paragraphs(docxpath)
        title_parts = [
            p.text.strip()
            for p in paragraphs
            if p.style == "Article Title" and p.text.strip()
        ]
        title = " ".join(title_parts)
        abstract = extract_abstract(paragraphs)
        abstract_text = " ".join(
            p.text for p in abstract.get("paragraphs", []) if p.text
        )
        return generate_keywords(title, abstract_text)
    except Exception:
        # Never let the keyword-generation path break the rest of the
        # front-matter pipeline. Best-effort by design.
        return []


#build action plan for keywords checks and comments
def keywordsFormatCheck(docxpath, keywords_state):
    comment_parts = []
    style_fixes = []
    anchor_pos = keywords_state["anchor_pos"]

    if keywords_state["keywords_found"] is False:
        # Generate 5 candidate keywords from the manuscript's title +
        # abstract so the editor receives a draft list as tracked text
        # rather than an empty heading + "please add" comment. Falls back
        # to the old behaviour when the LLM is unavailable or every
        # candidate fails the post-validation gate.
        generated = _maybe_generate_keywords_for_missing_section(docxpath)
        if generated:
            message = (
                "Tracked change applied: inserted 'Keywords' heading and "
                "AI-suggested keywords under Practitioner Notes — please "
                "review and edit to match your paper's contribution."
            )
            missing_comment = (
                "AI-suggested keywords inserted as tracked changes — please "
                "review and edit to match your paper's contribution."
            )
        else:
            message = "Tracked change applied: inserted 'Keywords' heading under Practitioner Notes."
            message += "\n\nKeywords missing, please add up to 5 keywords pertaining to the manuscript."
            missing_comment = "Keywords missing, please add up to 5 keywords pertaining to the manuscript."
        return {
            "action": "none",
            "reason": "Keywords heading not found",
            "anchor_pos": anchor_pos,
            "message": message,
            "style_fixes": style_fixes,
            "ensure_keywords_section": True,
            "missing_keywords_comment": missing_comment,
            "generated_keywords": generated,
        }

    anchor_pos = keywords_state["keywords_heading_pos"]

    if keywords_state["practitioner_heading_found"] is True:
        if keywords_state["keywords_follows_practitioner"] is False:
            comment_parts.append("Keywords section should follow the Practitioner Notes section.")

    if keywords_state["keywords_heading_style_ok"] is False:
        style_fixes.append(
            {
                "paragraph_index": keywords_state["keywords_heading_pos"],
                "new_style": KEYWORDS_REQUIRED_HEADING_STYLE_ID,
            }
        )
        comment_parts.append(
            "Tracked style change applied: Keywords heading changed to style '"
            + KEYWORDS_REQUIRED_HEADING_STYLE
            + "'."
        )

    if keywords_state.get("inline_keywords_text", "") != "":
        comment_parts.append("Tracked change applied: split combined Keywords line into heading and keywords line.")
    elif keywords_state["keywords_paragraph_found"] is False:
        comment_parts.append("Add one keywords line directly under the Keywords heading.")
    else:
        if keywords_state["keywords_paragraph_count"] != 1:
            comment_parts.append(
                "Keywords should be in a single paragraph. Found "
                + str(keywords_state["keywords_paragraph_count"])
                + " paragraph(s)."
            )

        if keywords_state["keywords_paragraph_style_ok"] is False:
            style_fixes.append(
                {
                    "paragraph_index": keywords_state["keywords_paragraph_pos"],
                    "new_style": KEYWORDS_REQUIRED_TEXT_STYLE_ID,
                }
            )
            comment_parts.append(
                "Tracked style change applied: Keywords line changed to style '"
                + KEYWORDS_REQUIRED_TEXT_STYLE
                + "'."
            )

        if keywords_state["keywords_line_count"] > KEYWORDS_MAX_LINES:
            comment_parts.append("Keywords should be one line only.")

        keyword_count = keywords_state["keywords_count"]
        if keyword_count == 0:
            comment_parts.append("Keywords line is empty. Add up to 5 keywords.")
        elif keyword_count > KEYWORDS_MAX_COUNT:
            comment_parts.append(
                "Keywords should be maximum 5 items. Found "
                + str(keyword_count)
                + "."
            )

        abbreviation_items = keywords_state["abbreviation_items"]
        if len(abbreviation_items) > 0:
            comment_parts.append(
                "Keywords should not use abbreviations. Found: " + ", ".join(abbreviation_items) + "."
            )

    if len(comment_parts) == 0:
        return {
            "action": "none",
            "reason": "Keywords section checked",
            "anchor_pos": anchor_pos,
            "message": "",
            "style_fixes": style_fixes,
        }

    message = "\n\n".join(comment_parts)
    return {
        "action": "add_keywords_comment",
        "reason": "Keywords section has issues",
        "anchor_pos": anchor_pos,
        "message": message,
        "style_fixes": style_fixes,
        "split_inline_keywords": keywords_state.get("inline_keywords_text", "") != "",
        "inline_keywords_text": keywords_state.get("inline_keywords_text", ""),
        "keywords_heading_text": keywords_state.get("keywords_heading_text", ""),
        "keywords_heading_pos": keywords_state["keywords_heading_pos"],
    }


#wrapper to run keywords find + keywords format planning
def build_keywords_check_plan(docxpath):
    keywords_state = keywordsFound(docxpath)
    return keywordsFormatCheck(docxpath, keywords_state)


#LLM citation component coverage check
def _check_citation_components_with_llm(citation_text):
    check_result = {
        "checked": False,
        "is_complete": True,
        "missing_parts": [],
        "suggested_fix": "",
        "reason": "",
    }

    if citation_text.strip() == "":
        return check_result

    try:
        llm_response = call_llm_json(
            system_prompt=CITATION_FORMAT_SYSTEM_PROMPT,
            user_prompt=CITATION_FORMAT_PROMPT + "\n\nCitation text:\n" + citation_text,
            response_schema=CITATION_FORMAT_SCHEMA,
        )
        content = llm_response["content"]

        missing_parts = []
        for item in content["missing_parts"]:
            text = str(item).strip()
            if text != "":
                missing_parts.append(text)

        check_result["checked"] = True
        check_result["is_complete"] = content["is_citation_complete"]
        check_result["missing_parts"] = missing_parts
        check_result["suggested_fix"] = str(content["suggested_fix"]).strip()
        check_result["reason"] = str(content["reason"]).strip()
    except Exception:
        pass

    return check_result


#deterministic fallback for citation component coverage
def _check_citation_components_deterministic(citation_text):
    missing_parts = []
    text = citation_text.strip()

    if "," not in text:
        missing_parts.append("author list")

    if re.search(r"\(\d{4}[a-z]?\)", text) is None:
        missing_parts.append("year in parentheses")

    if re.search(r"\b\d+\s*\(\d+\)", text) is None:
        missing_parts.append("volume(issue)")

    if re.search(r"(https?://|doi\\.org|doi:)", text, flags=re.IGNORECASE) is None:
        missing_parts.append("DOI or URL")

    if re.search(r"\bjournal\b", text, flags=re.IGNORECASE) is None:
        missing_parts.append("journal title")

    if re.search(r"\(\d{4}[a-z]?\)\.?\s+.+", text) is None:
        missing_parts.append("article title")

    return {
        "checked": True,
        "is_complete": len(missing_parts) == 0,
        "missing_parts": missing_parts,
        "suggested_fix": "Include all required citation parts in one line: authors, year, title, journal, volume(issue), and DOI/URL.",
        "reason": "Used deterministic citation component checks.",
    }


#build readable missing-citation-components comment message
def _build_missing_citation_message(citation_check):
    message = "Citation appears incomplete. Missing part(s): " + ", ".join(citation_check["missing_parts"]) + "."

    if citation_check["suggested_fix"] != "":
        message += "\nSuggested fix: " + citation_check["suggested_fix"]

    if citation_check["reason"] != "":
        message += "\nReason: " + citation_check["reason"]

    return message


#collect full citation section state from front page
def citationFound(docxpath):
    paragraphs = load_paragraphs(docxpath)
    front = get_front_page(paragraphs)

    citation_heading = _find_citation_heading_in_front(front)
    abstract_heading = _find_abstract_heading_in_front(front)
    intro_heading = _find_introduction_heading_in_front(front)

    anchor_pos = 0
    for p in front:
        if p.is_empty is False:
            anchor_pos = p.index
            break

    citation_heading_pos = 0
    citation_heading_style_ok = False
    if citation_heading is not None:
        citation_heading_pos = citation_heading.index
        anchor_pos = citation_heading_pos
        if citation_heading.style == CITATION_REQUIRED_HEADING_STYLE:
            citation_heading_style_ok = True

    citation_body_paragraphs = []
    if citation_heading is not None:
        for p in front:
            if p.index <= citation_heading.index:
                continue
            if _is_citation_body_end_marker(p) is True:
                break
            if p.is_empty is False:
                citation_body_paragraphs.append(p)

    citation_paragraph = None
    if len(citation_body_paragraphs) > 0:
        citation_paragraph = citation_body_paragraphs[0]

    citation_text = ""
    citation_line_count = 0
    citation_paragraph_style_ok = False
    citation_paragraph_pos = 0
    if citation_paragraph is not None:
        citation_text_parts = []
        for paragraph in citation_body_paragraphs:
            clean_text = paragraph.text.strip()
            if clean_text != "":
                citation_text_parts.append(clean_text)
        citation_text = " ".join(citation_text_parts)
        citation_line_count = _count_title_lines(citation_text)
        citation_paragraph_pos = citation_paragraph.index
        if citation_paragraph.style == CITATION_REQUIRED_TEXT_STYLE:
            citation_paragraph_style_ok = True

    citation_follows_abstract = True
    if citation_heading is not None and abstract_heading is not None:
        if citation_heading.index <= abstract_heading.index:
            citation_follows_abstract = False

    citation_before_intro = True
    if citation_heading is not None and intro_heading is not None:
        if citation_heading.index >= intro_heading.index:
            citation_before_intro = False

    return {
        "anchor_pos": anchor_pos,
        "citation_found": citation_heading is not None,
        "citation_heading_pos": citation_heading_pos,
        "citation_heading_style_ok": citation_heading_style_ok,
        "citation_paragraph_found": citation_paragraph is not None,
        "citation_paragraph_pos": citation_paragraph_pos,
        "citation_body_positions": [p.index for p in citation_body_paragraphs],
        "citation_paragraph_count": len(citation_body_paragraphs),
        "citation_paragraph_style_ok": citation_paragraph_style_ok,
        "citation_line_count": citation_line_count,
        "citation_text": citation_text,
        "abstract_found": abstract_heading is not None,
        "citation_follows_abstract": citation_follows_abstract,
        "intro_found": intro_heading is not None,
        "citation_before_intro": citation_before_intro,
    }


#build action plan for citation checks and comments
def citationFormatCheck(docxpath, citation_state):
    if citation_state["citation_found"] is False:
        return {
            "action": "none",
            "reason": "Citation section not found in body",
            "anchor_pos": citation_state["anchor_pos"],
            "message": "",
            "style_fixes": [],
            "citation_text": "",
            "citation_remove_positions": [],
        }

    anchor_pos = citation_state["citation_heading_pos"]
    if citation_state["citation_paragraph_found"] is True:
        anchor_pos = citation_state["citation_paragraph_pos"]

    # Citation should not stay in the document body.
    # Store the body positions so they can be removed after the text is copied to the footer.
    remove_positions = [citation_state["citation_heading_pos"]]
    for paragraph_pos in citation_state.get("citation_body_positions", []):
        remove_positions.append(paragraph_pos)

    return {
        "action": "move_citation_to_footer",
        "reason": "Citation belongs in the footer",
        "anchor_pos": anchor_pos,
        "message": "",
        "style_fixes": [],
        "citation_text": citation_state["citation_text"],
        "citation_remove_positions": remove_positions,
    }


#wrapper to run citation find + citation format planning
def build_citation_check_plan(docxpath):
    citation_state = citationFound(docxpath)
    return citationFormatCheck(docxpath, citation_state)
#----------------------------------------------------------------------------------------------------------------------------------------------
#DOCUMENT BODY
#----------------------------------------------------------------------------------------------------------------------------------------------

##global rules to follow:

    #- #body text must use font size 11pt, justified, Arial font with 1.15 line-spacing and should not be indented.
    #any text not following this style should be corrected as a tracked change to size 11pt, justified, Arial font with 1.15 line-spacing and should not be indented

    #Quotes less than 40 words should have quotation marks around the words be incorporated in-line with a paragraph without additional formatting, such as "quotation here with up to 40 words" (Author, 2025).
    #Quotes longer than 40 words should be formatted as block quotations using the Quote style.
    #in processing, split them: under-40 inline quotes = deterministic-ish, over-40 block quote style = deterministic, APA correctness details = LLM/manual.



    # Refer to APA guidance for further details about quotes. Common errors with quotes are the use of italics for formatting and the use of quotation marks in block quotes.
    # Author name and page number would be included at the end such as this. (Author & Author, 2025, p. 123) (Style: Quote)



##TABLES (if present)
# avoid tables that split across pages and avoid leaving more than a quarter of a page blank




#REQUIRED SECTIONS:

##INTRODUCTION SECTION ------------------------------------------------------------------------------------------------------------------------
#Must have heading: "Introduction" at the top of page 2
#Introduction exists as Heading 1 and is the first body section heading.






##LITERATURE SECTION ---------------------------------------------------------------------------------------------------------------------------











##RESULTS SECTION ------------------------------------------------------------------------------------------------------------------------------

# if tables or figures are present in the document,
# then they must be reffered to in this section, if they are not reffered to,
# please create a comment anchored to paragraph instructing them to do so - deterministic
#








##DISCUSSION SECTION ----------------------------------------------------------------------------------------------------------------------------

# required subheadings in this section: Practical Implications, Theoretical Implications, and Limitations and Future Research.
# subheadings use style: "Heading 2"
# deterministic (Heading 2 + exact names).

#
#


#document body constants -------------------------------------------------------------------------------
BODY_REQUIRED_FONT_NAME = "Arial"
BODY_REQUIRED_FONT_SIZE_PT = 11.0
BODY_REQUIRED_FONT_SIZE_XML = "22"
BODY_REQUIRED_LINE_SPACING = 1.15
BODY_REQUIRED_LINE_SPACING_XML = "276"
BODY_REQUIRED_SPACE_AFTER_XML = "120"
BODY_REQUIRED_STYLE = "Normal"

# Styles that are never body prose — skip when checking paragraph development.
_PARA_DEV_SKIP_STYLES = frozenset({
    "Article Title", "Authors", "Author Affiliations",
    "Heading Front Page", "Front Page Text", "Practitioner Notes",
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Figure Number", "Figure Title",
    "Table Number", "Table Title", "Table Text", "Table Emphasis",
    "APA 7 Reference List Entry", "APA7ReferenceListEntry",
    "APA7 Reference List Entry", "APAReferenceListEntry", "Reference List Entry",
    "Quote", "Guidance Notes", "CommentText", "Caption",
    "Default Paragraph Style",
})

METHOD_REQUIRED_SUBHEADINGS = [
    "Research Design",
    "Participants",
    "Measures",
    "Procedure",
    "Analysis",
]

DISCUSSION_REQUIRED_SUBHEADINGS = [
    "Practical Implications",
    "Theoretical Implications",
    "Limitations and Future Research",
]

# Synonym map: each required canonical subheading maps to a list of
# equivalent heading texts authors commonly use. Single source of truth
# lives in ``canonical_jultp_template.CANONICAL_STRUCTURE['subsection_aliases']``
# so the validator (per-rule pass/fail) and this module (consolidated
# missing-subsection comment) stay aligned. The alias here is just a
# local reference for backwards compatibility with code that imported the
# old name.
_SUBSECTION_ALIASES: dict[str, list[str]] = CANONICAL_STRUCTURE.get(
    "subsection_aliases", {}
)


def _matches_required_subsection(found_normalised_set, canonical_label):
    """Return True if any alias (including the canonical) appears in the set.

    ``found_normalised_set`` is a set of heading texts already normalised
    via ``_normalise_heading_for_match`` (lower-cased, whitespace-collapsed,
    trailing colon/period stripped).
    """
    canonical_norm = _normalise_heading_for_match(canonical_label)
    if canonical_norm in found_normalised_set:
        return True
    for alias in _SUBSECTION_ALIASES.get(canonical_label, []):
        if _normalise_heading_for_match(alias) in found_normalised_set:
            return True
    return False


def _all_alias_norms(canonical_label):
    """Return the set of all accepted normalised forms for a canonical label."""
    aliases = {_normalise_heading_for_match(canonical_label)}
    for alias in _SUBSECTION_ALIASES.get(canonical_label, []):
        aliases.add(_normalise_heading_for_match(alias))
    return aliases

CENTER_ALIGNED_HEADING_TEXTS = [
    "introduction",
    "literature",
    "literature review",
    "method",
    "methods",
    "results",
    "discussion",
    "conclusion",
    "conclusions",
    "acknowledgements",
    "acknowledgments",
]

LEFT_ALIGNED_HEADING_TEXTS = [
    "practioner notes",
    "practitioner notes",
    "keywords",
]

INLINE_QUOTE_MAX_WORDS = 40
DOT_POINT_PATTERN = re.compile(r"^\s*(?:•|◦|○|●|▪|▫|·|\*|-)\s+\S")
QUOTE_MARK_PATTERN = re.compile(r'["“”]')
QUOTE_SEGMENT_PATTERN = re.compile(r'["“”](.*?)["“”]')


#document body helper checks -------------------------------------------------------------------------------
#normalise heading text for case-insensitive exact matching
def _normalise_heading_text(text):
    clean = text.strip().lower()
    clean = re.sub(r"\s+", " ", clean)
    return clean


#find first Heading 1 paragraph index for a target section title
def _find_heading1_index(paragraphs, heading_text):
    target = _normalise_heading_text(heading_text)

    # Exact match first
    for p in paragraphs:
        if p.is_empty is True:
            continue
        if p.style != "Heading 1":
            continue
        if _normalise_heading_text(p.text) == target:
            return p.index

    # Substring fallback — handles combined headings like "Results and Discussion"
    for p in paragraphs:
        if p.is_empty is True:
            continue
        if p.style != "Heading 1":
            continue
        if target in _normalise_heading_text(p.text):
            return p.index

    return None


def _heading1_indices_after(paragraphs, start_index):
    if start_index is None:
        return []
    return [
        p.index
        for p in paragraphs
        if p.index > start_index and p.style == "Heading 1" and p.is_empty is False
    ]


def _nth_heading1_after(paragraphs, start_index, number):
    headings = _heading1_indices_after(paragraphs, start_index)
    if number <= 0 or len(headings) < number:
        return None
    return headings[number - 1]


def _find_combined_results_discussion_index(paragraphs):
    combined_names = {
        "results and discussion",
        "findings and discussion",
    }
    for p in paragraphs:
        if p.is_empty is True or p.style != "Heading 1":
            continue
        clean = _normalise_heading_text(p.text).rstrip(":.")
        if clean in combined_names:
            return p.index
    return None


#find next Heading 1 index after a section start
def _find_next_heading1_index(paragraphs, start_index):
    for p in paragraphs:
        if p.index <= start_index:
            continue
        if p.style == "Heading 1" and p.is_empty is False:
            return p.index
    return len(paragraphs)


#find first Heading 1 index in the document
def _find_first_heading1_index(paragraphs):
    for p in paragraphs:
        if p.style == "Heading 1" and p.is_empty is False:
            return p.index
    return None


#find first non-empty paragraph index in the document
def _find_first_non_empty_index(paragraphs):
    for p in paragraphs:
        if p.is_empty is False:
            return p.index
    return 0


#add one body comment item anchored to a paragraph
def _append_body_comment(comment_items, anchor_pos, message):
    text = message.strip()
    if text == "":
        return
    for item in comment_items:
        if item.get("message") == text and abs(item.get("anchor_pos", 0) - anchor_pos) <= 3:
            return

    comment_items.append(
        {
            "anchor_pos": anchor_pos,
            "message": text,
        }
    )


def _is_missing_required_subheading_issue(issue):
    message = issue.get("message", "")
    return "is missing required Heading 2 subheading(s):" in message


#alignment check for body text paragraphs
def _body_alignment_ok(doc_paragraph):
    alignment = doc_paragraph.alignment
    if alignment is None:
        return True  # inherits from Normal style (expected to be justified)
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return True
    return False


#line spacing check for body text paragraphs
def _body_line_spacing_ok(doc_paragraph):
    spacing = doc_paragraph.paragraph_format.line_spacing
    if spacing is None:
        return True

    if isinstance(spacing, float) or isinstance(spacing, int):
        value = float(spacing)
        if abs(value - BODY_REQUIRED_LINE_SPACING) <= 0.05:
            return True
        return False

    return True


def _body_spacing_ok(doc_paragraph):
    before = doc_paragraph.paragraph_format.space_before
    after = doc_paragraph.paragraph_format.space_after
    before_ok = before is None or abs(float(before.pt)) <= 0.1
    after_ok = after is None or abs(float(after.pt) - 6.0) <= 0.1
    return before_ok and after_ok


#helper for checking if indent length is effectively zero
def _is_zero_indent_length(length_value):
    if length_value is None:
        return True

    try:
        pt = length_value.pt
        if pt is None:
            return True
        if abs(float(pt)) <= 0.1:
            return True
        return False
    except Exception:
        return True


#indent check for body text paragraphs
def _body_indent_ok(doc_paragraph):
    left_indent_ok = _is_zero_indent_length(doc_paragraph.paragraph_format.left_indent)
    first_line_indent_ok = _is_zero_indent_length(doc_paragraph.paragraph_format.first_line_indent)

    if left_indent_ok is False:
        return False
    if first_line_indent_ok is False:
        return False
    return True


#font checks for body text paragraphs
def _body_font_ok(doc_paragraph):
    font_name_ok = True
    font_size_ok = True

    for run in doc_paragraph.runs:
        if run.text is None:
            continue
        if run.text.strip() == "":
            continue

        if run.font.name is not None:
            run_name = run.font.name.strip().lower()
            if run_name != BODY_REQUIRED_FONT_NAME.lower():
                font_name_ok = False

        if run.font.size is not None:
            run_size = run.font.size.pt
            if run_size is not None:
                if abs(float(run_size) - BODY_REQUIRED_FONT_SIZE_PT) > 0.2:
                    font_size_ok = False

    return font_name_ok, font_size_ok


##PARAGRAPH DEVELOPMENT CHECKS ---------------------------------------------------------------------------------------------------------------

_BODY_PARA_MIN_WORDS = 35
_BODY_PARA_MIN_SENTENCES = 3
_BODY_PARA_MAX_REPORTED = 5  # cap to avoid flooding editor with comments

# Academic abbreviations that end in a period but are NOT sentence boundaries
_SENT_ABB = re.compile(
    r"\b(?:e\.g|i\.e|etc|cf|vs|et\s+al|[Ff]ig|[Ff]igs|pp?|[Vv]ol|ed|eds|[Nn]o|"
    r"Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec|"
    r"[Dd]r|[Pp]rof|[Mm]r|[Mm]rs|[Mm]s|[Jj]r|[Ss]r)\."
)
# Single uppercase initial (e.g. "J. Smith")
_SENT_INITIAL = re.compile(r"(?<!\w)[A-Z]\.")

# Nouns that make "This/These/Such + noun" a self-referencing (not back-referencing) opener
_SELF_REF_NOUNS = frozenset({
    "study", "paper", "article", "research", "review",
    "chapter", "section", "analysis", "report", "work", "investigation",
})


def _count_body_sentences(text):
    masked = _SENT_ABB.sub(lambda m: m.group(0)[:-1] + "\x00", text)
    masked = _SENT_INITIAL.sub(lambda m: m.group(0)[:-1] + "\x00", masked)
    # Allow both uppercase and lowercase sentence starters — handles brand names like ePortfolio
    count = len(re.findall(r"[.!?](?=\s+[A-Za-z(]|\s*$)", masked))
    return max(1, count)


def _has_continuation_opener(text):
    m = re.match(r"^(This|These|Such)\s+(\w+)", text.strip(), re.IGNORECASE)
    if m is None:
        return False, ""
    second = m.group(2).lower()
    if second in _SELF_REF_NOUNS:
        return False, ""
    return True, m.group(1)


def _is_skippable_para(text):
    stripped = text.strip()
    if not stripped:
        return True
    last_char = stripped[-1]
    first_char = stripped[0]
    # Lead-in sentences (introduce a quote or list — intentionally incomplete)
    if last_char == ":":
        return True
    # Unclosed/fragment paragraphs — not complete prose
    if last_char not in ".!?":
        return True
    # Embedded block quotes styled as Normal (paragraph opens with a quotation mark)
    # Includes straight and curly variants: " " " ' ' '
    if first_char in ('"', '"', '"', "'", ''', '''):
        return True
    return False


def _collect_paragraph_development_issues(
    paragraph_records,
    body_start_index,
    body_end_index,
    ack_index=None,
    method_index=None,
):
    issues = []

    if body_end_index is None:
        body_end_index = len(paragraph_records)

    ack_section_end = None
    if ack_index is not None:
        ack_section_end = _find_next_heading1_index(paragraph_records, ack_index)

    method_section_end = None
    if method_index is not None:
        method_section_end = _find_next_heading1_index(paragraph_records, method_index)

    i = body_start_index
    while i < body_end_index and i < len(paragraph_records):
        p = paragraph_records[i]
        i += 1

        if p.is_empty is True:
            continue
        style = p.style or ""
        if style in _PARA_DEV_SKIP_STYLES:
            continue
        if "heading" in style.lower() or "list" in style.lower():
            continue
        if ack_index is not None and ack_section_end is not None:
            if ack_index <= p.index < ack_section_end:
                continue
        if method_index is not None and method_section_end is not None:
            if method_index <= p.index < method_section_end:
                continue

        text = p.text or ""
        if _is_skippable_para(text):
            continue

        word_count = _count_title_words(text)
        if word_count < _BODY_PARA_MIN_WORDS:
            continue

        sentence_count = _count_body_sentences(text)
        if sentence_count >= _BODY_PARA_MIN_SENTENCES:
            continue

        is_continuation, opener_word = _has_continuation_opener(text)

        if sentence_count == 1:
            parts = [
                f"This is not a real paragraph ({word_count} words). "
                "A single sentence does not constitute a paragraph. "
                "Consider merging it with an adjacent paragraph or expanding it."
            ]
        else:
            parts = [
                f"This paragraph has only {sentence_count} sentences ({word_count} words) and is too short. "
                "Consider merging it with an adjacent paragraph or expanding it."
            ]
        if is_continuation:
            parts.append(
                f"The paragraph opens with '{opener_word}', which may indicate it is "
                "continuing the previous paragraph's argument rather than establishing "
                "its own claim. Consider revising the opening to state a clear topic sentence."
            )

        issues.append({
            "anchor_pos": p.index,
            "message": "\n".join(parts),
        })

    if len(issues) > _BODY_PARA_MAX_REPORTED:
        remaining = len(issues) - _BODY_PARA_MAX_REPORTED
        summary = {
            "anchor_pos": issues[_BODY_PARA_MAX_REPORTED]["anchor_pos"],
            "message": (
                f"{remaining} additional paragraph(s) in this document may also be "
                "underdeveloped (fewer than 3 sentences). Please review the full manuscript "
                "for paragraph development and ensure each paragraph contains a topic sentence, "
                "supporting evidence or analysis, and a concluding synthesis."
            ),
        }
        issues = issues[:_BODY_PARA_MAX_REPORTED] + [summary]

    return issues


#collect style/formatting issues for normal body paragraphs
def _collect_body_formatting_issues(doc, paragraph_records, body_start_index, body_end_index):
    issues = []

    if body_end_index is None:
        body_end_index = len(paragraph_records)

    i = body_start_index
    while i < body_end_index and i < len(paragraph_records) and i < len(doc.paragraphs):
        p = paragraph_records[i]

        if p.is_empty is False:
            if p.style == BODY_REQUIRED_STYLE:
                doc_p = doc.paragraphs[i]

                alignment_ok = _body_alignment_ok(doc_p)
                line_spacing_ok = _body_line_spacing_ok(doc_p)
                spacing_ok = _body_spacing_ok(doc_p)
                indent_ok = _body_indent_ok(doc_p)
                font_name_ok, font_size_ok = _body_font_ok(doc_p)

                if alignment_ok is False or line_spacing_ok is False or spacing_ok is False or indent_ok is False or font_name_ok is False or font_size_ok is False:
                    issues.append(
                        {
                            "paragraph_index": p.index,
                            "font_name": font_name_ok is False,
                            "font_size": font_size_ok is False,
                            "alignment": alignment_ok is False,
                            "line_spacing": line_spacing_ok is False,
                            "spacing": spacing_ok is False,
                            "indent": indent_ok is False,
                        }
                    )

        i += 1

    return issues


#extract inline quoted segments between straight/curly quote marks
def _extract_quoted_segments(text):
    segments = []
    for match in QUOTE_SEGMENT_PATTERN.finditer(text):
        segment = match.group(1).strip()
        if segment != "":
            segments.append(segment)
    return segments


#quote formatting checks (deterministic-ish split by short vs long quote)
def _collect_quote_formatting_issues(paragraph_records, body_start_index, body_end_index):
    issues = []

    if body_end_index is None:
        body_end_index = len(paragraph_records)

    i = body_start_index
    while i < body_end_index and i < len(paragraph_records):
        p = paragraph_records[i]
        if p.is_empty is True:
            i += 1
            continue

        style = p.style
        text = p.text
        has_quote_marks = QUOTE_MARK_PATTERN.search(text) is not None
        quote_segments = _extract_quoted_segments(text)
        paragraph_words = _count_title_words(text)
        has_page_citation = re.search(r"\([^)]+,\s*\d{4}[a-z]?,\s*p{1,2}\.?\s*\d+[^)]*\)", text, flags=re.IGNORECASE) is not None

        if style == "Quote":
            if paragraph_words <= INLINE_QUOTE_MAX_WORDS:
                issues.append(
                    {
                        "anchor_pos": p.index,
                        "message": "Quote style used for a short quote (40 words or less). Keep short quotes inline with quotation marks instead of block Quote style.",
                    }
                )
            if has_quote_marks is True:
                issues.append(
                    {
                        "anchor_pos": p.index,
                        "message": "Block quotes (over 40 words, style 'Quote') should not use quotation marks.",
                    }
                )

        if style != "Quote":
            long_inline_found = False
            for segment in quote_segments:
                if _count_title_words(segment) > INLINE_QUOTE_MAX_WORDS:
                    long_inline_found = True
                    break

            if long_inline_found is True:
                issues.append(
                    {
                        "anchor_pos": p.index,
                        "message": "Found an inline quoted segment over 40 words. Convert it to a block quote using style 'Quote'.",
                    }
                )

            if has_page_citation is True and has_quote_marks is False:
                issues.append(
                    {
                        "anchor_pos": p.index,
                        "message": "This looks like a direct quote citation with page number, but quotation marks are missing for inline quote format.",
                    }
                )

        i += 1

    return issues


#detect disallowed dot-point style content in manuscript body
def _collect_dot_point_issues(doc, paragraph_records, body_start_index, body_end_index):
    if body_end_index is None:
        body_end_index = len(paragraph_records)

    first_anchor = None
    count = 0
    in_practitioner_notes = False

    i = body_start_index
    while i < body_end_index and i < len(paragraph_records) and i < len(doc.paragraphs):
        p = paragraph_records[i]
        if p.is_empty is True:
            i += 1
            continue

        text = p.text or ""
        low = text.strip().lower()
        style = p.style or ""

        if low == "practioner notes" or low == "practitioner notes":
            in_practitioner_notes = True
            i += 1
            continue
        if low == "keywords":
            in_practitioner_notes = False
        if in_practitioner_notes is True:
            i += 1
            continue

        is_dot_point = False
        if DOT_POINT_PATTERN.search(text) is not None:
            is_dot_point = True
        if "list" in style.lower():
            is_dot_point = True

        doc_p = doc.paragraphs[i]
        pPr = doc_p._p.pPr
        if pPr is not None and pPr.find(f"{WQ}numPr") is not None:
            is_dot_point = True

        if is_dot_point:
            if first_anchor is None:
                first_anchor = p.index
            count += 1

        i += 1

    if count == 0:
        return []

    if count == 1:
        message = (
            "Dot-point lists are not permitted in academic journal manuscripts outside "
            "the Practitioner Notes section. Please convert this dot-point list to "
            "continuous prose."
        )
    else:
        message = (
            f"Dot-point lists were found in {count} locations in this document. "
            "Dot-point lists are not permitted in academic journal manuscripts outside "
            "the Practitioner Notes section. Please convert all dot-point lists to "
            "continuous prose."
        )

    return [{"anchor_pos": first_anchor, "message": message}]


#collect unique Figure/Table labels found in the document
def _collect_table_figure_labels(paragraph_records):
    labels = []
    seen = set()

    for p in paragraph_records:
        if p.is_empty is True:
            continue

        matches = re.findall(r"\b(Figure|Table)\s+(\d+)\b", p.text, flags=re.IGNORECASE)
        for kind, number in matches:
            label = kind.title() + " " + number
            if label not in seen:
                seen.add(label)
                labels.append(
                    {
                        "label": label,
                        "anchor_pos": p.index,
                    }
                )

    return labels


#check if figures/tables are referenced in Results section when present
def _collect_results_reference_issues(paragraph_records, results_start_index, labels):
    issues = []

    if len(labels) == 0:
        return issues

    if results_start_index is None:
        first_anchor = labels[0]["anchor_pos"]
        issues.append(
            {
                "anchor_pos": first_anchor,
                "message": "Figures/tables are present but Results section heading is missing, so figure/table references in Results cannot be verified.",
            }
        )
        return issues

    results_end_index = _find_next_heading1_index(paragraph_records, results_start_index)

    results_text_parts = []
    i = results_start_index
    while i < results_end_index:
        p = paragraph_records[i]
        if p.is_empty is False:
            results_text_parts.append(p.text.lower())
        i += 1
    results_text = " ".join(results_text_parts)

    missing_labels = []
    for item in labels:
        label_text = item["label"].lower()
        if label_text not in results_text:
            missing_labels.append(item["label"])

    if len(missing_labels) > 0:
        issues.append(
            {
                "anchor_pos": results_start_index,
                "message": "Figures/tables should be referred to in the Results section. Missing reference(s): " + ", ".join(missing_labels) + ".",
            }
        )

    return issues


#required Heading 1 section checks (Introduction/Literature/Results/Discussion)
def _collect_required_section_issues(paragraph_records, intro_index, literature_index, method_index, results_index, discussion_index, first_heading1_index):
    issues = []

    first_non_empty_index = _find_first_non_empty_index(paragraph_records)
    section_anchor = first_non_empty_index
    if first_heading1_index is not None:
        section_anchor = first_heading1_index

    combined_results_discussion_index = _find_combined_results_discussion_index(paragraph_records)

    def _nearest_anchor(*candidates):
        for candidate in candidates:
            if candidate is not None:
                return candidate
        return section_anchor

    def _heading_between(start_index, end_index, prefer_last=False):
        if start_index is None or end_index is None:
            return None
        candidates = []
        for p in paragraph_records:
            if p.index <= start_index or p.index >= end_index:
                continue
            if p.style == "Heading 1" and p.is_empty is False:
                candidates.append(p.index)
        if len(candidates) == 0:
            return None
        if prefer_last is True:
            return candidates[-1]
        return candidates[0]

    if intro_index is None:
        issues.append(
            {
                "anchor_pos": section_anchor,
                "message": "Missing required Heading 1 section: Introduction.",
            }
        )
    else:
        if first_heading1_index is not None and intro_index != first_heading1_index:
            issues.append(
                {
                    "anchor_pos": intro_index,
                    "message": "Introduction should be the first Heading 1 body section heading.",
                }
            )

    if literature_index is None:
        literature_anchor = _nth_heading1_after(paragraph_records, intro_index, 1)
        if literature_anchor is None:
            literature_anchor = _nearest_anchor(intro_index)
        issues.append(
            {
                "anchor_pos": literature_anchor,
                "message": CANONICAL_STRUCTURE["missing_section_queries"]["Literature"],
            }
        )
    elif intro_index is not None and literature_index < intro_index:
        issues.append(
            {
                "anchor_pos": literature_index,
                "message": "Literature should appear after Introduction.",
            }
        )

    if method_index is None:
        next_known = _nearest_anchor(results_index, discussion_index)
        method_anchor = _heading_between(literature_index, next_known)
        if method_anchor is None:
            method_anchor = _nearest_anchor(literature_index, intro_index)
        issues.append(
            {
                "anchor_pos": method_anchor,
                "message": CANONICAL_STRUCTURE["missing_section_queries"]["Method"],
            }
        )
    elif literature_index is not None and method_index < literature_index:
        issues.append(
            {
                "anchor_pos": method_index,
                "message": "Method should appear after Literature.",
            }
        )

    if results_index is None:
        results_anchor = _nth_heading1_after(paragraph_records, method_index, 1)
        if results_anchor is None:
            results_anchor = _nearest_anchor(method_index, literature_index, intro_index)
            between_anchor = _heading_between(results_anchor, discussion_index, prefer_last=True)
            if between_anchor is not None:
                results_anchor = between_anchor
        issues.append(
            {
                "anchor_pos": results_anchor,
                "message": CANONICAL_STRUCTURE["missing_section_queries"]["Results"],
            }
        )

    if combined_results_discussion_index is not None:
        issues.append(
            {
                "anchor_pos": combined_results_discussion_index,
                "message": CANONICAL_STRUCTURE["combined_results_discussion_query"],
            }
        )

    if discussion_index is None:
        issues.append(
            {
                "anchor_pos": _nearest_anchor(results_index, method_index, literature_index, intro_index),
                "message": "Missing required Heading 1 section: Discussion.",
            }
        )

    return issues


#method subsection checks for required Heading 2 names
def _collect_method_subheading_issues(paragraph_records, method_index):
    issues = []

    if method_index is None:
        return issues

    section_end = _find_next_heading1_index(paragraph_records, method_index)

    found_subheading_normalised = set()
    i = method_index + 1
    while i < section_end:
        p = paragraph_records[i]
        if p.is_empty is False and p.style == "Heading 2":
            found_subheading_normalised.add(_normalise_heading_for_match(p.text))
        i += 1

    missing = []
    for required in METHOD_REQUIRED_SUBHEADINGS:
        if not _matches_required_subsection(found_subheading_normalised, required):
            missing.append(required)

    if len(missing) > 0:
        issues.append(
            {
                "anchor_pos": method_index,
                "message": "Method is missing required Heading 2 subheading(s): " + ", ".join(missing) + ".",
            }
        )

    return issues


#normalise a heading for matching — strips trailing punctuation on top of the base normalise
def _normalise_heading_for_match(text):
    return _normalise_heading_text(text).rstrip(".:;")


#discussion subsection checks for required Heading 2 names
def _collect_discussion_subheading_issues(paragraph_records, discussion_index):
    issues = []

    if discussion_index is None:
        return issues

    section_end = _find_next_heading1_index(paragraph_records, discussion_index)

    # Track found subheadings: normalised name -> paragraph index
    found_heading2 = {}
    i = discussion_index + 1
    while i < section_end:
        p = paragraph_records[i]
        if p.is_empty is False and p.style == "Heading 2":
            normalised = _normalise_heading_for_match(p.text)
            found_heading2[normalised] = p.index
        i += 1

    found_set = set(found_heading2.keys())

    missing = []
    for req_label in DISCUSSION_REQUIRED_SUBHEADINGS:
        if not _matches_required_subsection(found_set, req_label):
            missing.append(req_label)

    if len(missing) > 0:
        issues.append(
            {
                "anchor_pos": discussion_index,
                "message": "Discussion is missing required Heading 2 subheading(s): " + ", ".join(missing) + ".",
            }
        )

    # Ordering check: required subheadings that are present must appear in template order
    present_positions = []
    for req_label in DISCUSSION_REQUIRED_SUBHEADINGS:
        for alias_norm in _all_alias_norms(req_label):
            if alias_norm in found_heading2:
                present_positions.append(found_heading2[alias_norm])
                break
    if present_positions != sorted(present_positions):
        issues.append(
            {
                "anchor_pos": discussion_index,
                "message": (
                    "Discussion subheadings are out of order. Required order: "
                    + ", ".join(DISCUSSION_REQUIRED_SUBHEADINGS) + "."
                ),
            }
        )

    # Style check: required subheadings present but with wrong style.
    # Union all alias normalisations so a heading like "Limitations" still
    # triggers the style check when it should have been "Heading 2".
    all_required_norms = set()
    for req_label in DISCUSSION_REQUIRED_SUBHEADINGS:
        all_required_norms |= _all_alias_norms(req_label)

    i = discussion_index + 1
    while i < section_end:
        p = paragraph_records[i]
        if p.is_empty is False:
            normalised = _normalise_heading_for_match(p.text)
            if normalised in all_required_norms and p.style != "Heading 2":
                issues.append(
                    {
                        "anchor_pos": p.index,
                        "message": "Discussion required subheading '" + p.text + "' should use style 'Heading 2'.",
                    }
                )
        i += 1


    return issues


##ACKNOWLEDGEMENTS SECTION -------------------------------------------------------------------------------------------------------------------

# required elements (checked by keyword presence across all ack paragraphs):
#   conflict of interest, funding, ethics, AI use, CRediT author contributions
# CRediT comment is anchored at the CRediT paragraph, not the section heading.

_ACK_COI_KEYWORDS = ["conflict of interest", "conflicts of interest", "competing interest"]
_ACK_FUNDING_KEYWORDS = ["funding", "funded", "financial support", "grant"]
_ACK_ETHICS_KEYWORDS = ["ethical", "ethics"]
_ACK_AI_KEYWORDS = ["artificial intelligence", "generative ai", "ai tool", "language model", "copilot", "chatgpt", "gpt-"]
_ACK_CREDIT_KEYWORDS = ["credit contribution", "credit author", "author contribution", "authorship contribution", "credit:"]
_NUMBERED_AUTHOR_BODY_RE = re.compile(r"\bAuthors?\s+\d+\b", re.IGNORECASE)
_NUMBERED_AUTHOR_BODY_COMMENT = (
    "Use author surnames rather than numbered placeholders such as 'Author 1' "
    "or 'Author 2' when referring to authors in the manuscript body or CRediT "
    "contribution statement."
)


def _ack_contains(text, keywords):
    lower = text.lower()
    for kw in keywords:
        if kw in lower:
            return True
    return False


def _collect_acknowledgements_issues(paragraph_records, ack_index):
    issues = []

    if ack_index is None:
        return issues

    section_end = _find_next_heading1_index(paragraph_records, ack_index)

    ack_paras = []
    i = ack_index + 1
    while i < section_end:
        p = paragraph_records[i]
        if p.is_empty is False:
            ack_paras.append(p)
        i += 1

    if len(ack_paras) == 0:
        issues.append({
            "anchor_pos": ack_index,
            "message": (
                "Acknowledgements section is empty. Add statements for: "
                "conflict of interest, funding, ethics, AI use, and CRediT author contributions."
            ),
        })
        return issues

    combined_text = " ".join(p.text for p in ack_paras)
    first_anchor = ack_paras[0].index

    if not _ack_contains(combined_text, _ACK_COI_KEYWORDS):
        issues.append({
            "anchor_pos": first_anchor,
            "message": "Acknowledgements is missing a conflict of interest statement.",
        })

    if not _ack_contains(combined_text, _ACK_FUNDING_KEYWORDS):
        issues.append({
            "anchor_pos": first_anchor,
            "message": "Acknowledgements is missing a funding statement.",
        })

    if not _ack_contains(combined_text, _ACK_ETHICS_KEYWORDS):
        issues.append({
            "anchor_pos": first_anchor,
            "message": "Acknowledgements is missing an ethics statement.",
        })

    if not _ack_contains(combined_text, _ACK_AI_KEYWORDS):
        issues.append({
            "anchor_pos": first_anchor,
            "message": "Acknowledgements is missing an AI use statement.",
        })

    # CRediT: anchor comment at the CRediT paragraph itself, or at last para if absent
    credit_para = None
    for p in ack_paras:
        if _ack_contains(p.text, _ACK_CREDIT_KEYWORDS):
            credit_para = p
            break

    if credit_para is None:
        credit_anchor = ack_paras[-1].index
        issues.append({
            "anchor_pos": credit_anchor,
            "message": (
                "Acknowledgements is missing CRediT author contributions. "
                "Add a paragraph listing each author's role (see https://credit.niso.org/)."
            ),
        })

    return issues


def _collect_numbered_author_placeholder_issues(paragraph_records, body_start_index, body_end_index):
    issues = []

    i = body_start_index
    while i < body_end_index and i < len(paragraph_records):
        p = paragraph_records[i]
        i += 1
        if p.is_empty is True or p.style in _PARA_DEV_SKIP_STYLES:
            continue
        if _NUMBERED_AUTHOR_BODY_RE.search(p.text or "") is None:
            continue
        issues.append({
            "anchor_pos": p.index,
            "message": _NUMBERED_AUTHOR_BODY_COMMENT,
        })

    return issues


#collect document-body state used for body-level checks
def documentBodyFound(docxpath):
    paragraphs = load_paragraphs(docxpath)

    first_heading1_index = _find_first_heading1_index(paragraphs)
    intro_index = _find_heading1_index(paragraphs, "Introduction")
    literature_index = _find_heading1_index(paragraphs, "Literature")
    method_index = _find_heading1_index(paragraphs, "Method")
    results_index = _find_heading1_index(paragraphs, "Results")
    discussion_index = _find_heading1_index(paragraphs, "Discussion")
    ack_index = _find_heading1_index(paragraphs, "Acknowledgements")
    if ack_index is None:
        ack_index = _find_heading1_index(paragraphs, "Acknowledgments")
    references_index = _find_heading1_index(paragraphs, "References")

    labels = _collect_table_figure_labels(paragraphs)

    anchor_pos = _find_first_non_empty_index(paragraphs)
    if first_heading1_index is not None:
        anchor_pos = first_heading1_index

    return {
        "anchor_pos": anchor_pos,
        "paragraphs": paragraphs,
        "first_heading1_index": first_heading1_index,
        "intro_index": intro_index,
        "literature_index": literature_index,
        "method_index": method_index,
        "results_index": results_index,
        "discussion_index": discussion_index,
        "ack_index": ack_index,
        "references_index": references_index,
        "table_figure_labels": labels,
    }


#build document-body plan with anchored comments for body rules
def documentBodyFormatCheck(docxpath, body_state):
    paragraphs = body_state["paragraphs"]
    first_heading1_index = body_state["first_heading1_index"]
    intro_index = body_state["intro_index"]
    literature_index = body_state["literature_index"]
    method_index = body_state["method_index"]
    results_index = body_state["results_index"]
    discussion_index = body_state["discussion_index"]
    ack_index = body_state.get("ack_index")
    references_index = body_state["references_index"]
    labels = body_state["table_figure_labels"]

    comment_items = []

    required_section_issues = _collect_required_section_issues(
        paragraph_records=paragraphs,
        intro_index=intro_index,
        literature_index=literature_index,
        method_index=method_index,
        results_index=results_index,
        discussion_index=discussion_index,
        first_heading1_index=first_heading1_index,
    )
    for issue in required_section_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    subheading_issues = []
    other_subheading_issues = []

    method_issues = _collect_method_subheading_issues(paragraphs, method_index)
    discussion_issues = _collect_discussion_subheading_issues(paragraphs, discussion_index)
    for issue in method_issues + discussion_issues:
        if _is_missing_required_subheading_issue(issue) is True:
            subheading_issues.append(issue)
        else:
            other_subheading_issues.append(issue)

    if len(subheading_issues) > 0:
        anchor_pos = subheading_issues[0]["anchor_pos"]
        message = "\n\n".join(issue["message"] for issue in subheading_issues)
        _append_body_comment(comment_items, anchor_pos, message)

    for issue in other_subheading_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    ack_issues = _collect_acknowledgements_issues(paragraphs, ack_index)
    for issue in ack_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    results_reference_issues = _collect_results_reference_issues(
        paragraph_records=paragraphs,
        results_start_index=results_index,
        labels=labels,
    )
    for issue in results_reference_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    body_start_index = 0
    if first_heading1_index is not None:
        body_start_index = first_heading1_index
    body_end_index = len(paragraphs)
    if references_index is not None and references_index > body_start_index:
        body_end_index = references_index

    quote_issues = _collect_quote_formatting_issues(paragraphs, body_start_index, body_end_index)
    for issue in quote_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    if first_heading1_index is not None:
        numbered_author_issues = _collect_numbered_author_placeholder_issues(
            paragraphs, body_start_index, body_end_index
        )
        for issue in numbered_author_issues:
            _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    doc = DocxDocument(docxpath)

    dot_point_issues = _collect_dot_point_issues(doc, paragraphs, body_start_index, body_end_index)
    for issue in dot_point_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    development_issues = _collect_paragraph_development_issues(
        paragraphs, body_start_index, body_end_index, ack_index=ack_index, method_index=method_index
    )
    for issue in development_issues:
        _append_body_comment(comment_items, issue["anchor_pos"], issue["message"])

    format_fixes = _collect_body_formatting_issues(doc, paragraphs, body_start_index, body_end_index)

    comment_items.sort(key=lambda x: x["anchor_pos"])

    if len(comment_items) == 0 and len(format_fixes) == 0:
        return {
            "action": "none",
            "reason": "Document body checked",
            "anchor_pos": body_state["anchor_pos"],
            "message": "",
            "format_fixes": [],
        }

    if len(comment_items) == 0:
        return {
            "action": "apply_document_body_formatting",
            "reason": "Document body formatting has issues",
            "anchor_pos": format_fixes[0]["paragraph_index"],
            "message": "",
            "comments": [],
            "format_fixes": format_fixes,
        }

    if len(comment_items) == 1:
        return {
            "action": "add_document_body_comment",
            "reason": "Document body has issues",
            "anchor_pos": comment_items[0]["anchor_pos"],
            "message": comment_items[0]["message"],
            "comments": comment_items,
            "format_fixes": format_fixes,
        }

    return {
        "action": "add_document_body_comments",
        "reason": "Document body has issues",
        "anchor_pos": comment_items[0]["anchor_pos"],
        "message": comment_items[0]["message"],
        "comments": comment_items,
        "format_fixes": format_fixes,
    }


#wrapper to run body find + body format planning
def build_document_body_check_plan(docxpath):
    body_state = documentBodyFound(docxpath)
    return documentBodyFormatCheck(docxpath, body_state)


FRONT_PAGE_BANNER_COMMENT = (
    "Check whether there is a picture/banner at the top of the first page. "
    "If not, paste in the JUTLP banner."
)
FRONT_PAGE_TEXT_BOX_COMMENT = (
    "Check whether there is a text box on the first page. "
    "If not, insert the JUTLP text box."
)
IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
_IMAGE_CONTENT_TYPES = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "tif": "image/tiff",
    "tiff": "image/tiff",
    "emf": "image/x-emf",
    "wmf": "image/x-wmf",
}


def _xml_local_name(element):
    return etree.QName(element).localname


def _paragraph_visible_text(paragraph_element):
    return "".join(t.text or "" for t in paragraph_element.iter(f"{WQ}t")).strip()


def _paragraph_direct_text(paragraph_element):
    return "".join(t.text or "" for t in paragraph_element.findall(f"{WQ}r/{WQ}t")).strip()


def _paragraph_has_picture(paragraph_element):
    return len(_collect_image_rel_ids(paragraph_element)) > 0


def _element_has_textbox(element):
    for child in element.iter():
        if _xml_local_name(child) in {"txbxContent", "textbox"}:
            return True
    return False


def _collect_image_rel_ids(paragraph_element):
    rel_ids = []
    for element in paragraph_element.iter():
        if _xml_local_name(element) not in {"blip", "imagedata"}:
            continue
        for attr in (f"{RQ}embed", f"{RQ}link", f"{RQ}id"):
            rel_id = element.get(attr)
            if rel_id and rel_id not in rel_ids:
                rel_ids.append(rel_id)
    return rel_ids


def _collect_relationship_ids(element):
    rel_ids = []
    for child in element.iter():
        for attr, val in child.attrib.items():
            if attr.startswith(RQ) and val and val not in rel_ids:
                rel_ids.append(val)
    return rel_ids


def _normalise_media_target(target):
    clean = (target or "").replace("\\", "/").lstrip("/")
    if clean.startswith("word/"):
        clean = clean[5:]
    return clean


def _next_media_name(existing_media, extension):
    index = 1
    while True:
        name = f"image{index}.{extension}"
        if f"word/media/{name}" not in existing_media:
            return name
        index += 1


def _ensure_image_content_types(ct_xml, extensions):
    if not extensions:
        return ct_xml
    tree = etree.fromstring(ct_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    existing = {
        el.get("Extension", "")
        for el in tree.findall(f"{{{ns}}}Default")
    }
    for ext in sorted(extensions):
        if ext in existing:
            continue
        content_type = _IMAGE_CONTENT_TYPES.get(ext)
        if not content_type:
            continue
        default = etree.SubElement(tree, f"{{{ns}}}Default")
        default.set("Extension", ext)
        default.set("ContentType", content_type)
    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _front_page_body_paras(body_paras):
    front_page = []
    for paragraph in body_paras:
        text = _paragraph_visible_text(paragraph)
        style_element = paragraph.find(f"{WQ}pPr/{WQ}pStyle")
        style_id = style_element.get(f"{WQ}val") if style_element is not None else ""
        if style_id in {HEADING_1_STYLE_ID, "Heading1", "Heading 1"} and _normalise_heading_text(text) == "introduction":
            break
        front_page.append(paragraph)
    return front_page


def _find_first_picture_paragraph(paragraphs):
    for paragraph in paragraphs:
        if _paragraph_has_picture(paragraph):
            return paragraph
    return None


def _word_part_path(target):
    clean = (target or "").replace("\\", "/").lstrip("/")
    if clean.startswith("word/"):
        return clean
    return "word/" + clean


def _rels_path_for_part(part_path):
    folder, name = os.path.split(part_path)
    return f"{folder}/_rels/{name}.rels"


def _find_header_picture_part(template_zip, template_rels_root):
    for rel in template_rels_root.findall(f"{PACKAGE_REL_Q}Relationship"):
        if not rel.get("Type", "").endswith("/header"):
            continue
        part_path = _word_part_path(rel.get("Target", ""))
        rels_path = _rels_path_for_part(part_path)
        if part_path not in template_zip.namelist() or rels_path not in template_zip.namelist():
            continue
        header_root = etree.fromstring(template_zip.read(part_path))
        if _find_first_picture_paragraph(header_root.iter(f"{WQ}p")) is None:
            continue
        header_rels_root = etree.fromstring(template_zip.read(rels_path))
        return header_root, header_rels_root
    return None, None


def _first_section_properties(doc_root):
    return next(doc_root.iter(f"{WQ}sectPr"), None)


def _doc_has_first_page_header_picture(docxpath, doc_root):
    try:
        sect_pr = _first_section_properties(doc_root)
        if sect_pr is None:
            return False

        with zipfile.ZipFile(docxpath, "r") as z:
            if "word/_rels/document.xml.rels" not in z.namelist():
                return False
            rels_root = etree.fromstring(z.read("word/_rels/document.xml.rels"))
            header_type = "first" if sect_pr.find(f"{WQ}titlePg") is not None else "default"
            for header_ref in sect_pr.findall(f"{WQ}headerReference"):
                if header_ref.get(f"{WQ}type", "") != header_type:
                    continue
                target = _find_relationship_target(rels_root, header_ref.get(f"{RQ}id", ""))
                part_path = _normalise_docx_target(target)
                if part_path in z.namelist():
                    header_root = etree.fromstring(z.read(part_path))
                    return _find_first_picture_paragraph(header_root.iter(f"{WQ}p")) is not None
    except Exception:
        return False
    return False


def _next_header_name(names):
    index = 1
    while True:
        name = "word/header" + str(index) + ".xml"
        if name not in names:
            return name
        index += 1


def _find_first_header_ref(sect_pr):
    for header_ref in sect_pr.findall(f"{WQ}headerReference"):
        if header_ref.get(f"{WQ}type", "") == "first":
            return header_ref
    return None


def _ensure_first_header(doc_root, rels_root, names):
    body = doc_root.find(f"{WQ}body")
    if body is None:
        return ""
    sect_pr = _first_section_properties(doc_root)
    if sect_pr is None:
        sect_pr = etree.SubElement(body, f"{WQ}sectPr")

    header_ref = _find_first_header_ref(sect_pr)
    if header_ref is not None:
        target = _find_relationship_target(rels_root, header_ref.get(f"{RQ}id", ""))
        if target != "":
            if sect_pr.find(f"{WQ}titlePg") is None:
                etree.SubElement(sect_pr, f"{WQ}titlePg")
            return _normalise_docx_target(target)

    header_name = _next_header_name(names)
    rel_id = _next_relationship_id(rels_root)
    new_rel = etree.SubElement(rels_root, f"{PACKAGE_REL_Q}Relationship")
    new_rel.set("Id", rel_id)
    new_rel.set("Type", HEADER_REL_TYPE)
    new_rel.set("Target", header_name.replace("word/", ""))

    new_ref = etree.Element(f"{WQ}headerReference")
    new_ref.set(f"{WQ}type", "first")
    new_ref.set(f"{RQ}id", rel_id)
    sect_pr.insert(0, new_ref)
    if sect_pr.find(f"{WQ}titlePg") is None:
        etree.SubElement(sect_pr, f"{WQ}titlePg")
    return header_name


def _patch_header_content_type(ct_xml, header_name):
    tree = etree.fromstring(ct_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    part_name = "/" + header_name
    for override in tree.findall(f"{{{ns}}}Override"):
        if override.get("PartName") == part_name:
            return ct_xml
    new_override = etree.SubElement(tree, f"{{{ns}}}Override")
    new_override.set("PartName", part_name)
    new_override.set("ContentType", HEADER_CONTENT_TYPE)
    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _strip_paragraph_to_textbox(paragraph_element):
    kept = False
    for child in list(paragraph_element):
        if child.tag == f"{WQ}pPr":
            continue
        if child.tag == f"{WQ}r" and _element_has_textbox(child):
            kept = True
            continue
        paragraph_element.remove(child)
    if kept is False:
        return None
    return paragraph_element


def _make_textbox_wrap_around_text(paragraph_element):
    """Make a front-page floating textbox use *square* text-wrapping so the
    surrounding content (the abstract) flows around it instead of being hidden
    beneath it.

    The JUTLP template box uses ``wrapSquare`` (text on both sides). An earlier
    revision forced ``wrapNone`` here, which makes the box float ON TOP of the
    abstract and obscure whatever sits under it. Square wrapping keeps the box
    in place while pushing the text aside, so nothing is blocked.
    """
    changed = False
    for anchor in paragraph_element.iter(f"{WPQ}anchor"):
        # Keep the box in the text layer (not drawn behind the text), so it is
        # never hidden and never hides text.
        if anchor.get("behindDoc") not in (None, "0"):
            anchor.set("behindDoc", "0")
            changed = True
        # Give wrapped text a small breathing margin around the box.
        for name, value in (
            ("distT", "36195"), ("distB", "36195"),
            ("distL", "114300"), ("distR", "114300"),
        ):
            if anchor.get(name) in (None, "0"):
                anchor.set(name, value)
                changed = True
        # Replace any non-square wrap mode (notably wrapNone, the overlay mode)
        # with wrapSquare; leave an existing wrapSquare untouched.
        has_square = False
        for child in list(anchor):
            if child.tag in {
                f"{WPQ}wrapNone",
                f"{WPQ}wrapTight",
                f"{WPQ}wrapThrough",
                f"{WPQ}wrapTopAndBottom",
            }:
                anchor.remove(child)
                changed = True
            elif child.tag == f"{WPQ}wrapSquare":
                has_square = True
        if not has_square:
            wrap_square = etree.Element(f"{WPQ}wrapSquare")
            wrap_square.set("wrapText", "bothSides")
            doc_pr = anchor.find(f"{WPQ}docPr")
            if doc_pr is not None:
                anchor.insert(list(anchor).index(doc_pr), wrap_square)
            else:
                anchor.append(wrap_square)
            changed = True

    # Legacy VML textboxes: <w10:wrap type="none"> is the overlay mode; "square"
    # wraps the surrounding text around the box.
    for element in paragraph_element.iter():
        if _xml_local_name(element) == "wrap" and element.get("type") in (None, "none"):
            element.set("type", "square")
            changed = True

    return changed


def _normalise_front_page_textbox_layout(input_path, output_path):
    doc_root, body_paras = _load_body_paras_raw(input_path)
    front_page_paras = _front_page_body_paras(body_paras)

    abstract_para = None
    for para in front_page_paras:
        if _normalise_heading_text(_paragraph_direct_text(para)) == "abstract":
            abstract_para = para
            break
    if abstract_para is None:
        return False

    body = doc_root.find(f"{WQ}body")
    changed = False
    for para in list(front_page_paras):
        if _element_has_textbox(para) is False:
            continue

        if _make_textbox_wrap_around_text(para) is True:
            changed = True

        if para is abstract_para:
            continue

        # Only re-anchor textboxes from empty dedicated-anchor paragraphs.
        # Content paragraphs that have a floating textbox anchored to them
        # (e.g. AuthorAffiliations with a journal-info box) must keep their
        # anchor — moving it would corrupt the paragraph's visible text.
        if _paragraph_direct_text(para) != "":
            continue

        for child in list(para):
            if child.tag != f"{WQ}pPr" and _element_has_textbox(child) is True:
                abstract_para.append(child)
                changed = True

        if body is not None and _element_has_textbox(para) is False and _paragraph_direct_text(para) == "":
            body.remove(para)
            changed = True

    if changed is False:
        return False

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


def _insert_front_page_banner_from_template(input_path, output_path, insert_index):
    if not os.path.exists(_TEMPLATE_PATH):
        return False

    try:
        with zipfile.ZipFile(_TEMPLATE_PATH, "r") as template_zip:
            if "word/document.xml" not in template_zip.namelist():
                return False
            if "word/_rels/document.xml.rels" not in template_zip.namelist():
                return False
            template_rels_xml = template_zip.read("word/_rels/document.xml.rels")

            template_rels_root = etree.fromstring(template_rels_xml)

            header_root, banner_rels_root = _find_header_picture_part(
                template_zip,
                template_rels_root,
            )
            if header_root is None:
                return False

            header_copy = copy.deepcopy(header_root)
            rel_ids = _collect_image_rel_ids(header_copy)
            if not rel_ids:
                return False

            rel_targets = {}
            for rel in banner_rels_root.findall(f"{PACKAGE_REL_Q}Relationship"):
                rel_id = rel.get("Id")
                if rel_id in rel_ids:
                    rel_targets[rel_id] = rel.get("Target", "")
            if not rel_targets:
                return False

            with zipfile.ZipFile(input_path, "r") as zin:
                names = zin.namelist()
                doc_xml = zin.read("word/document.xml")
                rels_xml = zin.read("word/_rels/document.xml.rels")
                ct_xml = zin.read("[Content_Types].xml")
                existing_media = {
                    name for name in names if name.startswith("word/media/")
                }

            doc_root = etree.fromstring(doc_xml)
            body = doc_root.find(f"{WQ}body")
            if body is None:
                return False

            rels_root = etree.fromstring(rels_xml)
            header_name = _ensure_first_header(doc_root, rels_root, names)
            if header_name == "":
                return False
            header_rels_path = _rels_path_for_part(header_name)
            header_rels_root = etree.Element(
                f"{PACKAGE_REL_Q}Relationships",
                nsmap={None: PACKAGE_REL_NS},
            )
            new_media_files = {}
            image_exts = set()
            rid_map = {}

            for old_rid, target in rel_targets.items():
                clean_target = _normalise_media_target(target)
                if clean_target == "":
                    continue
                ext = os.path.splitext(clean_target)[1].lstrip(".").lower() or "png"
                image_exts.add(ext)
                new_name = _next_media_name(existing_media, ext)
                existing_media.add(f"word/media/{new_name}")

                new_rid = _next_relationship_id(header_rels_root)
                new_rel = etree.SubElement(header_rels_root, f"{PACKAGE_REL_Q}Relationship")
                new_rel.set("Id", new_rid)
                new_rel.set("Type", IMAGE_REL_TYPE)
                new_rel.set("Target", f"media/{new_name}")
                rid_map[old_rid] = new_rid

                template_media_path = f"word/{clean_target}"
                if template_media_path in template_zip.namelist():
                    new_media_files[f"word/media/{new_name}"] = template_zip.read(
                        template_media_path
                    )

            if not rid_map:
                return False

            for element in header_copy.iter():
                old_embed = element.get(f"{RQ}embed")
                if old_embed in rid_map:
                    element.set(f"{RQ}embed", rid_map[old_embed])
                old_link = element.get(f"{RQ}link")
                if old_link in rid_map:
                    element.set(f"{RQ}link", rid_map[old_link])
                old_id = element.get(f"{RQ}id")
                if old_id in rid_map:
                    element.set(f"{RQ}id", rid_map[old_id])

            new_doc_xml = etree.tostring(
                doc_root, xml_declaration=True, encoding="UTF-8", standalone=True
            )
            new_rels_xml = etree.tostring(
                rels_root, xml_declaration=True, encoding="UTF-8", standalone=True
            )
            new_header_xml = etree.tostring(
                header_copy, xml_declaration=True, encoding="UTF-8", standalone=True
            )
            new_header_rels_xml = etree.tostring(
                header_rels_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone=True,
            )
            new_ct_xml = _patch_header_content_type(
                _ensure_image_content_types(ct_xml, image_exts),
                header_name,
            )

            tmp_path = output_path + ".banner.tmp"
            with zipfile.ZipFile(input_path, "r") as zin, zipfile.ZipFile(
                tmp_path, "w", zipfile.ZIP_DEFLATED
            ) as zout:
                for item in zin.infolist():
                    if item.filename == "word/document.xml":
                        zout.writestr(item, new_doc_xml)
                    elif item.filename == "word/_rels/document.xml.rels":
                        zout.writestr(item, new_rels_xml)
                    elif item.filename == header_name:
                        zout.writestr(item, new_header_xml)
                    elif item.filename == header_rels_path:
                        zout.writestr(item, new_header_rels_xml)
                    elif item.filename == "[Content_Types].xml":
                        zout.writestr(item, new_ct_xml)
                    else:
                        zout.writestr(item, zin.read(item.filename))
                if header_name not in names:
                    zout.writestr(header_name, new_header_xml)
                if header_rels_path not in names:
                    zout.writestr(header_rels_path, new_header_rels_xml)
                for media_path, data in new_media_files.items():
                    zout.writestr(media_path, data)
            os.replace(tmp_path, output_path)
            return True
    except Exception:
        return False


def _insert_front_page_text_box_from_template(input_path, output_path, insert_index):
    if not os.path.exists(_TEMPLATE_PATH):
        return False

    try:
        with zipfile.ZipFile(_TEMPLATE_PATH, "r") as template_zip:
            if "word/document.xml" not in template_zip.namelist():
                return False
            if "word/_rels/document.xml.rels" not in template_zip.namelist():
                return False
            template_doc_xml = template_zip.read("word/document.xml")
            template_rels_xml = template_zip.read("word/_rels/document.xml.rels")

            template_root = etree.fromstring(template_doc_xml)
            template_rels_root = etree.fromstring(template_rels_xml)
            template_body = template_root.find(f"{WQ}body")
            template_paras = [
                el for el in (template_body if template_body is not None else [])
                if el.tag == f"{WQ}p"
            ]

            template_front = _front_page_body_paras(template_paras)
            textbox_para = next(
                (p for p in template_front if _element_has_textbox(p)),
                None,
            )
            if textbox_para is None:
                textbox_para = next(
                    (p for p in template_paras if _element_has_textbox(p)),
                    None,
                )
            if textbox_para is None:
                return False

            textbox_copy = copy.deepcopy(textbox_para)
            textbox_copy = _strip_paragraph_to_textbox(textbox_copy)
            if textbox_copy is None:
                return False
            _make_textbox_wrap_around_text(textbox_copy)

            textbox_rel_ids = _collect_relationship_ids(textbox_copy)

        with zipfile.ZipFile(input_path, "r") as zin:
            doc_xml = zin.read("word/document.xml")
            rels_xml = zin.read("word/_rels/document.xml.rels")

        doc_root = etree.fromstring(doc_xml)
        body = doc_root.find(f"{WQ}body")
        if body is None:
            return False

        rels_root = etree.fromstring(rels_xml)
        rid_map = {}

        if textbox_rel_ids:
            for old_rid in textbox_rel_ids:
                template_rel = template_rels_root.find(
                    f"{PACKAGE_REL_Q}Relationship[@Id='{old_rid}']"
                )
                if template_rel is None:
                    continue
                new_rid = _next_relationship_id(rels_root)
                new_rel = etree.SubElement(rels_root, f"{PACKAGE_REL_Q}Relationship")
                new_rel.set("Id", new_rid)
                new_rel.set("Type", template_rel.get("Type", ""))
                new_rel.set("Target", template_rel.get("Target", ""))
                target_mode = template_rel.get("TargetMode")
                if target_mode:
                    new_rel.set("TargetMode", target_mode)
                rid_map[old_rid] = new_rid

        if rid_map:
            for element in textbox_copy.iter():
                for attr, val in list(element.attrib.items()):
                    if attr.startswith(RQ) and val in rid_map:
                        element.set(attr, rid_map[val])

        paras = [el for el in body if el.tag == f"{WQ}p"]
        target_para = None
        for para in paras:
            if _normalise_heading_text(_paragraph_visible_text(para)) == "abstract":
                target_para = para
                break

        if target_para is not None:
            for child in list(textbox_copy):
                if child.tag != f"{WQ}pPr":
                    target_para.append(copy.deepcopy(child))
        elif insert_index < len(paras):
            target_para = paras[insert_index]
            insert_at = list(body).index(target_para)
            body.insert(insert_at, textbox_copy)
        else:
            sect_pr = body.find(f"{WQ}sectPr")
            if sect_pr is not None:
                body.insert(list(body).index(sect_pr), textbox_copy)
            else:
                body.append(textbox_copy)

        new_doc_xml = etree.tostring(
            doc_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        new_rels_xml = etree.tostring(
            rels_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        tmp_path = output_path + ".textbox.tmp"
        with zipfile.ZipFile(input_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_doc_xml)
                elif item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item, new_rels_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp_path, output_path)
        return True
    except Exception:
        return False


def build_front_page_asset_check_plan(docxpath):
    doc_root, body_paras = _load_body_paras_raw(docxpath)
    front_page_paras = _front_page_body_paras(body_paras)

    anchor_pos = 0
    for i, paragraph in enumerate(body_paras):
        if _paragraph_visible_text(paragraph) != "" or _paragraph_has_picture(paragraph):
            anchor_pos = i
            break

    first_content_para = None
    for paragraph in front_page_paras:
        if _paragraph_visible_text(paragraph) != "" or _paragraph_has_picture(paragraph):
            first_content_para = paragraph
            break

    has_top_picture = (
        first_content_para is not None and _paragraph_has_picture(first_content_para)
    ) or _doc_has_first_page_header_picture(docxpath, doc_root)
    has_first_page_textbox = _element_has_textbox(doc_root)
    if front_page_paras:
        has_first_page_textbox = any(_element_has_textbox(p) for p in front_page_paras)

    comments = []
    if has_top_picture is False:
        comments.append({"anchor_pos": anchor_pos, "message": FRONT_PAGE_BANNER_COMMENT})
    if has_first_page_textbox is False:
        comments.append({"anchor_pos": anchor_pos, "message": FRONT_PAGE_TEXT_BOX_COMMENT})

    if not comments:
        return {
            "action": "none",
            "reason": "Front-page picture/banner and text box found",
            "anchor_pos": anchor_pos,
            "comments": [],
        }

    return {
        "action": "add_front_page_asset_comments",
        "reason": "Front-page picture/banner or text box missing",
        "anchor_pos": anchor_pos,
        "comments": comments,
    }


def _apply_front_page_asset_plan(input_path, output_path, plan):
    if plan.get("action") != "add_front_page_asset_comments":
        return False
    comment_items = plan.get("comments", [])
    if not comment_items:
        return False

    current_input = input_path
    applied_any = False

    for item in comment_items:
        message = (item.get("message") or "").strip()
        if message == "":
            continue

        if message == FRONT_PAGE_BANNER_COMMENT:
            inserted = _insert_front_page_banner_from_template(
                current_input,
                output_path,
                item.get("anchor_pos", 0),
            )
            if inserted is True:
                current_input = output_path
                applied_any = True
                continue

        if message == FRONT_PAGE_TEXT_BOX_COMMENT:
            inserted = _insert_front_page_text_box_from_template(
                current_input,
                output_path,
                item.get("anchor_pos", 0),
            )
            if inserted is True:
                current_input = output_path
                applied_any = True
                continue

        _write_single_comment_docx(
            input_path=current_input,
            output_path=output_path,
            paragraph_index=item.get("anchor_pos", 0),
            message=message,
        )
        current_input = output_path
        applied_any = True

    return applied_any










#OUTPUT GENERATION SECTION ----------------------------------------------------------------------------------------------------------------------
#3. apply tracked changes ----------------------------------------------------------------------------------------------


def _load_body_paras_raw(input_path: str):
    """Return (doc_root, body_paras) by parsing document.xml directly from the zip.

    Unlike DocxDocument(), this preserves mc:AlternateContent, VML shapes, and
    all other XML that python-docx may strip when it builds its object model.
    The returned body_paras list has the same indices as DocxDocument().paragraphs
    (direct <w:p> children of <w:body> only).
    """
    with zipfile.ZipFile(input_path, "r") as z:
        doc_xml = z.read("word/document.xml")
    doc_root = etree.fromstring(doc_xml)
    body = doc_root.find(f"{WQ}body")
    body_paras = [el for el in (body if body is not None else []) if el.tag == f"{WQ}p"]
    return doc_root, body_paras


def _save_doc_xml_to_zip(input_path: str, output_path: str, new_doc_xml: bytes) -> None:
    """Repack the docx zip replacing only word/document.xml; all other entries are copied as-is."""
    tmp = output_path + ".tmp"
    with zipfile.ZipFile(input_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            zout.writestr(item, new_doc_xml if item.filename == "word/document.xml"
                          else zin.read(item.filename))
    os.replace(tmp, output_path)


def _remove_line_numbers(docx_path):
    doc_root, _ = _load_body_paras_raw(docx_path)
    changed = False
    for el in doc_root.findall(f".//{WQ}lnNumType"):
        el.getparent().remove(el)
        changed = True
    if changed:
        new_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        _save_doc_xml_to_zip(docx_path, docx_path, new_xml)
    return changed


def _save_styles_xml_to_zip(input_path: str, output_path: str, new_styles_xml: bytes) -> None:
    """Repack the docx zip replacing only word/styles.xml; all other entries are copied as-is."""
    tmp = output_path + ".tmp"
    saw_styles = False

    with zipfile.ZipFile(input_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/styles.xml":
                zout.writestr(item, new_styles_xml)
                saw_styles = True
            else:
                zout.writestr(item, zin.read(item.filename))

        if saw_styles is False:
            zout.writestr("word/styles.xml", new_styles_xml)

    os.replace(tmp, output_path)


def _style_matches_template(style_el, template_style_el):
    try:
        parser = etree.XMLParser(remove_blank_text=True)
        style_xml = etree.tostring(style_el)
        template_xml = etree.tostring(template_style_el)
        style_c14n = etree.tostring(etree.fromstring(style_xml, parser), method="c14n")
        template_c14n = etree.tostring(
            etree.fromstring(template_xml, parser), method="c14n"
        )
        return style_c14n == template_c14n
    except Exception:
        return False


def _ensure_template_styles_available(
    input_path, output_path, required_style_names, replace_existing=False
):
    if required_style_names is None:
        return False
    if len(required_style_names) == 0:
        return False

    _load_template_style_names()
    if _TEMPLATE_STYLE_XML_MAP is None:
        return False

    with zipfile.ZipFile(input_path, "r") as z:
        try:
            styles_xml = z.read("word/styles.xml")
        except Exception:
            return False

    styles_root = etree.fromstring(styles_xml)

    existing_style_names = {}
    existing_style_ids = {}
    for style_el in styles_root.findall(f"{WQ}style"):
        style_id = style_el.get(f"{WQ}styleId")
        if style_id is not None and style_id.strip() != "":
            existing_style_ids[style_id] = style_el

        name_el = style_el.find(f"{WQ}name")
        if name_el is not None:
            style_name = name_el.get(f"{WQ}val")
            if style_name is not None and style_name.strip() != "":
                existing_style_names[style_name] = style_el

    added_any = False
    for style_name in required_style_names:
        style_xml = _TEMPLATE_STYLE_XML_MAP.get(style_name)
        if style_xml is None:
            continue

        try:
            style_el = etree.fromstring(style_xml.encode("utf-8"))
        except Exception:
            continue

        wanted_style_id = style_el.get(f"{WQ}styleId", _resolve_style_id(style_name))
        existing_el = existing_style_names.get(style_name)
        if existing_el is None:
            existing_el = existing_style_ids.get(wanted_style_id)

        if existing_el is not None:
            if replace_existing is False or _style_matches_template(existing_el, style_el):
                continue
            insert_at = list(styles_root).index(existing_el)
            styles_root.remove(existing_el)
            styles_root.insert(insert_at, style_el)
            added_any = True
            existing_style_ids[wanted_style_id] = style_el
            existing_style_names[style_name] = style_el
            continue

        styles_root.append(style_el)
        added_any = True

        new_style_id = style_el.get(f"{WQ}styleId")
        if new_style_id is not None and new_style_id.strip() != "":
            existing_style_ids[new_style_id] = style_el

        name_el = style_el.find(f"{WQ}name")
        if name_el is not None:
            added_name = name_el.get(f"{WQ}val")
            if added_name is not None and added_name.strip() != "":
                existing_style_names[added_name] = style_el

    if added_any is False:
        return False

    new_styles_xml = etree.tostring(styles_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_styles_xml_to_zip(input_path, output_path, new_styles_xml)
    return True


def _ensure_normal_style_body_rpr(input_path, output_path):
    """Pin the body font + size on the Normal style's run properties.

    The template's Normal style inherits its 11pt size from the *template's*
    docDefaults. When that style is copied into a manuscript whose own
    docDefaults carry no size, the body silently falls back to Word's 10pt
    default (the reported "Arial size 10" bug). Setting rFonts=Arial and
    sz/szCs=22 (11pt) directly on the Normal style guarantees every paragraph
    that inherits Normal renders as 11pt Arial regardless of docDefaults.
    """
    with zipfile.ZipFile(input_path, "r") as z:
        names = z.namelist()
        if "word/styles.xml" not in names:
            return _passthrough_copy(input_path, output_path)
        styles_xml = z.read("word/styles.xml")

    root = etree.fromstring(styles_xml)
    normal = None
    for style_el in root.findall(f"{WQ}style"):
        name_el = style_el.find(f"{WQ}name")
        if name_el is not None and (name_el.get(f"{WQ}val", "") or "").strip().lower() == "normal":
            normal = style_el
            break
    if normal is None:
        return _passthrough_copy(input_path, output_path)

    r_pr = normal.find(f"{WQ}rPr")
    if r_pr is None:
        r_pr = etree.SubElement(normal, f"{WQ}rPr")

    changed = False
    r_fonts = r_pr.find(f"{WQ}rFonts")
    if r_fonts is None:
        r_fonts = etree.SubElement(r_pr, f"{WQ}rFonts")
    for attr in ("ascii", "hAnsi", "cs"):
        if r_fonts.get(f"{WQ}{attr}") != BODY_REQUIRED_FONT_NAME:
            r_fonts.set(f"{WQ}{attr}", BODY_REQUIRED_FONT_NAME)
            changed = True
    for tag in ("sz", "szCs"):
        el = r_pr.find(f"{WQ}{tag}")
        if el is None:
            el = etree.SubElement(r_pr, f"{WQ}{tag}")
        if el.get(f"{WQ}val") != BODY_REQUIRED_FONT_SIZE_XML:
            el.set(f"{WQ}val", BODY_REQUIRED_FONT_SIZE_XML)
            changed = True

    if not changed:
        return _passthrough_copy(input_path, output_path)

    new_styles = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    tmp = output_path + ".normalrpr.tmp"
    try:
        with zipfile.ZipFile(input_path, "r") as zin, zipfile.ZipFile(
            tmp, "w", zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                if item.filename == "word/styles.xml":
                    zout.writestr(item, new_styles)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp, output_path)
    except Exception:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise
    return True


def _passthrough_copy(input_path, output_path):
    if input_path != output_path:
        with zipfile.ZipFile(input_path, "r") as zin, zipfile.ZipFile(
            output_path, "w", zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                zout.writestr(item, zin.read(item.filename))
    return False


def _apply_normal_style_fix(input_path, output_path):
    applied = _ensure_template_styles_available(
        input_path,
        output_path,
        ["Normal"],
        replace_existing=True,
    )
    # Pin Arial 11pt onto the Normal style so body text that inherits Normal is
    # not left at Word's 10pt default after the template style is copied in.
    source = output_path if applied else input_path
    pinned = _ensure_normal_style_body_rpr(source, output_path)
    return applied or pinned


def _resolve_style_id_from_doc(input_path, style_name, fallback_style_id):
    try:
        with zipfile.ZipFile(input_path, "r") as z:
            styles_xml = z.read("word/styles.xml")
        styles_root = etree.fromstring(styles_xml)
    except Exception:
        return fallback_style_id

    for style_el in styles_root.findall(f"{WQ}style"):
        name_el = style_el.find(f"{WQ}name")
        if name_el is None:
            continue
        found_name = name_el.get(f"{WQ}val", "")
        if found_name != style_name:
            continue
        style_id = style_el.get(f"{WQ}styleId")
        if style_id is not None and style_id.strip() != "":
            return style_id

    return fallback_style_id


#remove all paragraph text/runs but keep paragraph style/properties
def _remove_paragraph_content_keep_style(paragraph_element):
    children = list(paragraph_element)
    for child in children:
        if child.tag != f"{WQ}pPr":
            paragraph_element.remove(child)


#append tracked delete+insert pair into one paragraph
def _append_tracked_replace(paragraph_element, old_text, new_text, change_id):
    nsmap = {"w": W}

    delete_element = etree.Element(f"{WQ}del", nsmap=nsmap)
    delete_element.set(f"{WQ}id", str(change_id))
    delete_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
    delete_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)
    delete_run = etree.SubElement(delete_element, f"{WQ}r")
    delete_text = etree.SubElement(delete_run, f"{WQ}delText")
    delete_text.text = old_text
    if old_text != "":
        if old_text[0].isspace() or old_text[-1].isspace():
            delete_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    insert_element = etree.Element(f"{WQ}ins", nsmap=nsmap)
    insert_element.set(f"{WQ}id", str(change_id + 1))
    insert_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
    insert_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)
    _append_text_with_superscript_markers(insert_element, new_text)

    paragraph_element.append(delete_element)
    paragraph_element.append(insert_element)


def _append_plain_text_run(paragraph_element, text):
    run = etree.SubElement(paragraph_element, f"{WQ}r")
    text_el = etree.SubElement(run, f"{WQ}t")
    text_el.text = text
    if text != "" and (text[0].isspace() or text[-1].isspace()):
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _apply_title_text_tracked_changes(para_el, old_text, new_text):
    pattern = r"\w+(?:[-']\w+)*|\s+|[^\w\s]+"
    old_parts = re.findall(pattern, old_text)
    new_parts = re.findall(pattern, new_text)
    if len(old_parts) != len(new_parts):
        return False

    _remove_paragraph_content_keep_style(para_el)
    change_id = 1
    for old, new in zip(old_parts, new_parts):
        if old != new and old.strip() != "":
            _append_tracked_replace(para_el, old, new, change_id)
            change_id += 2
        else:
            _append_plain_text_run(para_el, new)
    return True


# Heading style ids (both the resolved template ids and the common raw
# variants) — when a paragraph is restyled TO one of these we must clear the
# body formatting it carried, otherwise the heading renders at body size.
_HEADING_STYLE_IDS_FOR_RESTYLE = frozenset({
    HEADING_1_STYLE_ID, HEADING_2_STYLE_ID, HEADING_3_STYLE_ID,
    "Heading1", "Heading2", "Heading3",
    "Heading 1", "Heading 2", "Heading 3",
})


def _strip_conflicting_direct_format_for_heading(para_element):
    """Remove direct formatting that would override a heading style.

    A paragraph that was body text before being restyled to a heading carries
    direct run font sizes (``w:sz``) and paragraph spacing/indent from its
    body life. Those direct properties win over the heading style, so the
    "heading" renders at body size with body spacing — visibly different from a
    natively-styled heading. Stripping them lets the heading style govern.

    Paragraph spacing/indent removal is captured by the surrounding
    ``w:pPrChange`` (which already recorded the original ``w:pPr``); run font
    sizes are direct character properties — removing them is safe on accept
    (heading size applies) and on reject (the runs fall back to the restored
    style's size, same as before).
    """
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is not None:
        for tag in ("spacing", "ind"):
            for el in pPr.findall(f"{WQ}{tag}"):
                pPr.remove(el)
        mark_rpr = pPr.find(f"{WQ}rPr")
        if mark_rpr is not None:
            for tag in ("sz", "szCs"):
                for el in mark_rpr.findall(f"{WQ}{tag}"):
                    mark_rpr.remove(el)
    for run in para_element.iter(f"{WQ}r"):
        rpr = run.find(f"{WQ}rPr")
        if rpr is None:
            continue
        for tag in ("sz", "szCs"):
            for el in rpr.findall(f"{WQ}{tag}"):
                rpr.remove(el)


#apply a paragraph style change as a tracked change using <w:pPrChange>
def _apply_tracked_style_change(para_element, new_style_name, old_style_name, change_id):
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is None:
        old_pPr = etree.Element(f"{WQ}pPr")
        pPr = etree.Element(f"{WQ}pPr")
        para_element.insert(0, pPr)
    else:
        old_pPr = copy.deepcopy(pPr)
        old_existing = old_pPr.find(f"{WQ}pPrChange")
        if old_existing is not None:
            old_pPr.remove(old_existing)

    # Remove any pre-existing pPrChange to avoid nesting
    existing = pPr.find(f"{WQ}pPrChange")
    if existing is not None:
        pPr.remove(existing)

    # Read old style from existing pStyle (may differ from what the caller passed)
    pStyle = pPr.find(f"{WQ}pStyle")
    if pStyle is not None:
        old_style_name = pStyle.get(f"{WQ}val", old_style_name)
    else:
        pStyle = etree.Element(f"{WQ}pStyle")
        pPr.insert(0, pStyle)

    old_pStyle = old_pPr.find(f"{WQ}pStyle")
    if old_pStyle is None and old_style_name:
        old_pStyle = etree.Element(f"{WQ}pStyle")
        old_pStyle.set(f"{WQ}val", old_style_name)
        old_pPr.insert(0, old_pStyle)

    # Set the proposed (new) style
    pStyle.set(f"{WQ}val", new_style_name)

    # Append pPrChange recording the original paragraph properties.
    pPrChange = etree.SubElement(pPr, f"{WQ}pPrChange")
    pPrChange.set(f"{WQ}id", str(change_id))
    pPrChange.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
    pPrChange.set(f"{WQ}date", TRACKED_CHANGE_DATE)
    pPrChange.append(old_pPr)

    # When promoting a body paragraph to a heading, drop the body's direct run
    # sizes and paragraph spacing so the heading style's formatting actually
    # shows (otherwise the heading renders at body size / spacing). Done after
    # pPrChange so the original properties are preserved for accept/reject.
    if new_style_name in _HEADING_STYLE_IDS_FOR_RESTYLE:
        _strip_conflicting_direct_format_for_heading(para_element)


def _paragraph_has_direct_space_before(para_element):
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is None:
        return False

    spacing = pPr.find(f"{WQ}spacing")
    if spacing is None:
        return False

    before_names = [
        f"{WQ}before",
        f"{WQ}beforeLines",
        f"{WQ}beforeAutospacing",
    ]
    for before_name in before_names:
        if before_name in spacing.attrib:
            return True

    return False


def _remove_direct_space_before(para_element):
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is None:
        return False

    spacing = pPr.find(f"{WQ}spacing")
    if spacing is None:
        return False

    changed = False
    before_names = [
        f"{WQ}before",
        f"{WQ}beforeLines",
        f"{WQ}beforeAutospacing",
    ]
    for before_name in before_names:
        if before_name in spacing.attrib:
            del spacing.attrib[before_name]
            changed = True

    if len(spacing.attrib) == 0:
        pPr.remove(spacing)

    return changed


def _direct_space_after_is_zero(para_element):
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is None:
        return False

    spacing = pPr.find(f"{WQ}spacing")
    if spacing is None:
        return False

    after_value = spacing.get(f"{WQ}after")
    return after_value == "0"


def _set_direct_space_after_zero(para_element):
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is None:
        pPr = etree.Element(f"{WQ}pPr")
        para_element.insert(0, pPr)

    spacing = pPr.find(f"{WQ}spacing")
    if spacing is None:
        spacing = etree.SubElement(pPr, f"{WQ}spacing")

    changed = False
    if spacing.get(f"{WQ}after") != "0":
        spacing.set(f"{WQ}after", "0")
        changed = True

    after_names = [
        f"{WQ}afterLines",
        f"{WQ}afterAutospacing",
    ]
    for after_name in after_names:
        if after_name in spacing.attrib:
            del spacing.attrib[after_name]
            changed = True

    return changed


def _find_direct_space_before_positions(input_path):
    positions = set()

    try:
        loaded = _load_body_paras_raw(input_path)
        body_paras = loaded[1]
        for i, para_element in enumerate(body_paras):
            if _paragraph_has_direct_space_before(para_element) is True:
                positions.add(i)
    except Exception:
        pass

    return positions


def _read_raw_paragraph_style_name(para_element):
    pPr = para_element.find(f"{WQ}pPr")
    if pPr is None:
        return ""

    pStyle = pPr.find(f"{WQ}pStyle")
    if pStyle is None:
        return ""

    value = pStyle.get(f"{WQ}val", "")
    if value is None:
        return ""

    return value


def _apply_tracked_style_fixes(input_path, output_path, style_fixes, change_id_start):
    if style_fixes is None:
        return False

    if len(style_fixes) == 0:
        return False

    doc_root, body_paras = _load_body_paras_raw(input_path)
    changed = False
    change_id = change_id_start
    seen_positions = set()

    for fix in style_fixes:
        paragraph_index = fix.get("paragraph_index")
        new_style_name = str(fix.get("new_style", "")).strip()

        if paragraph_index is None:
            continue
        if new_style_name == "":
            continue
        if paragraph_index in seen_positions:
            continue
        if paragraph_index < 0 or paragraph_index >= len(body_paras):
            continue

        para_el = body_paras[paragraph_index]
        old_style_name = _read_raw_paragraph_style_name(para_el)
        style_changed = False
        spacing_changed = False
        tracked_change_added = False

        if old_style_name != new_style_name:
            _apply_tracked_style_change(para_el, new_style_name, old_style_name, change_id)
            change_id += 1
            style_changed = True
            tracked_change_added = True

        if new_style_name == AFFILIATIONS_REQUIRED_STYLE_ID:
            if _paragraph_has_direct_space_before(para_el) is True:
                if tracked_change_added is False:
                    _apply_tracked_style_change(para_el, new_style_name, old_style_name, change_id)
                    change_id += 1
                    tracked_change_added = True
                spacing_changed = _remove_direct_space_before(para_el)
            if fix.get("clear_after_spacing") is True:
                if _direct_space_after_is_zero(para_el) is False:
                    if tracked_change_added is False:
                        _apply_tracked_style_change(para_el, new_style_name, old_style_name, change_id)
                        change_id += 1
                        tracked_change_added = True
                    if _set_direct_space_after_zero(para_el) is True:
                        spacing_changed = True

        if style_changed is True or spacing_changed is True:
            changed = True
        seen_positions.add(paragraph_index)

    if changed is False:
        return False

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


#build new inserted authors paragraph with Authors style
def _make_inserted_authors_paragraph(text, change_id):
    nsmap = {"w": W}
    paragraph_element = etree.Element(f"{WQ}p", nsmap=nsmap)

    paragraph_properties = etree.SubElement(paragraph_element, f"{WQ}pPr")
    paragraph_style = etree.SubElement(paragraph_properties, f"{WQ}pStyle")
    paragraph_style.set(f"{WQ}val", AUTHORS_REQUIRED_STYLE_ID)

    insert_element = etree.SubElement(paragraph_element, f"{WQ}ins")
    insert_element.set(f"{WQ}id", str(change_id))
    insert_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
    insert_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)
    _append_text_with_superscript_markers(insert_element, text)

    return paragraph_element


def _make_plain_styled_paragraph(text, style_id):
    nsmap = {"w": W}
    paragraph_element = etree.Element(f"{WQ}p", nsmap=nsmap)
    paragraph_properties = etree.SubElement(paragraph_element, f"{WQ}pPr")
    paragraph_style = etree.SubElement(paragraph_properties, f"{WQ}pStyle")
    paragraph_style.set(f"{WQ}val", style_id)
    _append_text_with_superscript_markers(paragraph_element, text)
    return paragraph_element


def _append_text_with_line_breaks(parent_element, text):
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            run = etree.SubElement(parent_element, f"{WQ}r")
            etree.SubElement(run, f"{WQ}br")
        _append_text_with_superscript_markers(parent_element, line)


def _make_inserted_styled_paragraph(text, style_id, change_id, font_size_half_points=None):
    nsmap = {"w": W}
    paragraph_element = etree.Element(f"{WQ}p", nsmap=nsmap)

    paragraph_properties = etree.SubElement(paragraph_element, f"{WQ}pPr")
    paragraph_style = etree.SubElement(paragraph_properties, f"{WQ}pStyle")
    paragraph_style.set(f"{WQ}val", style_id)

    insert_element = etree.SubElement(paragraph_element, f"{WQ}ins")
    insert_element.set(f"{WQ}id", str(change_id))
    insert_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
    insert_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)

    if text != "":
        text_with_markers = _caret_markers_to_superscript_text(text)
        if font_size_half_points is None:
            _append_plain_text_run(insert_element, text_with_markers)
        else:
            _append_plain_text_run_with_size(
                insert_element,
                text_with_markers,
                font_size_half_points,
            )

    return paragraph_element


#insert XML paragraph at target document index safely
def _insert_paragraph_at_index(doc, target_index, paragraph_element):
    paragraphs = doc.paragraphs

    if target_index < 0:
        target_index = 0

    if target_index < len(paragraphs):
        paragraphs[target_index]._element.addprevious(paragraph_element)
        return

    body = doc.element.body
    body_children = list(body)
    if len(body_children) > 0:
        last_child = body_children[-1]
        if last_child.tag == f"{WQ}sectPr":
            last_child.addprevious(paragraph_element)
            return

    body.append(paragraph_element)


def _remove_blank_paragraph_at_index(doc, paragraph_index):
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        return False
    paragraph = doc.paragraphs[paragraph_index]
    if paragraph.text.strip() != "":
        return False
    paragraph_element = paragraph._element
    if paragraph_element.find(f".//{WQ}drawing") is not None:
        return False
    p_pr = paragraph_element.find(f"{WQ}pPr")
    if p_pr is not None and p_pr.find(f"{WQ}sectPr") is not None:
        return False
    parent = paragraph_element.getparent()
    if parent is None:
        return False
    parent.remove(paragraph_element)
    return True


def _insert_raw_body_paragraph(input_path, output_path, target_index, paragraph_element):
    doc_root, paragraphs = _load_body_paras_raw(input_path)
    body = doc_root.find(f"{WQ}body")
    if body is None:
        shutil.copy2(input_path, output_path)
        return

    if target_index < 0:
        target_index = 0

    if target_index < len(paragraphs):
        paragraphs[target_index].addprevious(paragraph_element)
    else:
        sect_pr = body.find(f"{WQ}sectPr")
        if sect_pr is not None:
            sect_pr.addprevious(paragraph_element)
        else:
            body.append(paragraph_element)

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)


def _shift_abstract_plan_positions_after_insert(abstract_plan, insert_at):
    for key in ("anchor_pos", "practitioner_insert_after_pos"):
        value = abstract_plan.get(key)
        if isinstance(value, int) and value >= insert_at:
            abstract_plan[key] = value + 1

    shifted_positions = []
    for pos in abstract_plan.get("merge_paragraph_positions", []):
        shifted_positions.append(pos + 1 if pos >= insert_at else pos)
    if shifted_positions:
        abstract_plan["merge_paragraph_positions"] = shifted_positions

    for item in abstract_plan.get("comments", []):
        value = item.get("anchor_pos")
        if isinstance(value, int) and value >= insert_at:
            item["anchor_pos"] = value + 1


def _format_author_query_text(comment_id, text):
    prefix, query_text = _author_query_parts(comment_id, text)
    return prefix + query_text


def _author_query_parts(comment_id, text):
    clean_text = (text or "").strip()
    match = re.match(r"^(Author Query\s+\d+\.\s*)(.*)$", clean_text, flags=re.IGNORECASE)
    if match is not None:
        return match.group(1), match.group(2)
    return f"Author Query {comment_id}. ", clean_text


def _append_author_query_runs(parent, comment_id, text):
    prefix, query_text = _author_query_parts(comment_id, text)

    prefix_run = etree.SubElement(parent, f"{WQ}r")
    prefix_r_pr = etree.SubElement(prefix_run, f"{WQ}rPr")
    etree.SubElement(prefix_r_pr, f"{WQ}b")
    prefix_t = etree.SubElement(prefix_run, f"{WQ}t")
    prefix_t.text = prefix
    prefix_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    text_run = etree.SubElement(parent, f"{WQ}r")
    t = etree.SubElement(text_run, f"{WQ}t")
    t.text = query_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


#create a single Word comment XML element
def _make_comment_element(comment_id, text):
    nsmap = {"w": W}
    comment = etree.Element(f"{WQ}comment", nsmap=nsmap)
    comment.set(f"{WQ}id", str(comment_id))
    comment.set(f"{WQ}author", AUTHOR)
    comment.set(f"{WQ}date", DATE)
    comment.set(f"{WQ}initials", INITIALS)

    p = etree.SubElement(comment, f"{WQ}p")
    p_pr = etree.SubElement(p, f"{WQ}pPr")
    p_style = etree.SubElement(p_pr, f"{WQ}pStyle")
    p_style.set(f"{WQ}val", "CommentText")

    ref_run = etree.SubElement(p, f"{WQ}r")
    ref_r_pr = etree.SubElement(ref_run, f"{WQ}rPr")
    ref_r_style = etree.SubElement(ref_r_pr, f"{WQ}rStyle")
    ref_r_style.set(f"{WQ}val", "CommentReference")
    etree.SubElement(ref_run, f"{WQ}annotationRef")

    _append_author_query_runs(p, comment_id, text)
    return comment


#attach comment range markers to target paragraph
def _inject_comment_markers(para_element, comment_id):
    nsmap = {"w": W}

    start = etree.Element(f"{WQ}commentRangeStart", nsmap=nsmap)
    start.set(f"{WQ}id", str(comment_id))

    end = etree.Element(f"{WQ}commentRangeEnd", nsmap=nsmap)
    end.set(f"{WQ}id", str(comment_id))

    ref_run = etree.Element(f"{WQ}r", nsmap=nsmap)
    r_pr = etree.SubElement(ref_run, f"{WQ}rPr")
    r_style = etree.SubElement(r_pr, f"{WQ}rStyle")
    r_style.set(f"{WQ}val", "CommentReference")
    comment_ref = etree.SubElement(ref_run, f"{WQ}commentReference")
    comment_ref.set(f"{WQ}id", str(comment_id))

    insert_at = 0
    children = list(para_element)
    while insert_at < len(children) and children[insert_at].tag == f"{WQ}commentRangeStart":
        insert_at += 1

    para_element.insert(insert_at, start)
    para_element.append(end)
    para_element.append(ref_run)


#ensure comments relationship exists in document rels
def _patch_rels(rels_xml):
    tree = etree.fromstring(rels_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    for rel in tree:
        if rel.get("Type") == COMMENT_REL_TYPE:
            return rels_xml

    new_rel = etree.SubElement(tree, f"{{{ns}}}Relationship")
    new_rel.set("Id", "rIdComments")
    new_rel.set("Type", COMMENT_REL_TYPE)
    new_rel.set("Target", "comments.xml")

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


#ensure comments content-type override exists
def _patch_content_types(ct_xml):
    tree = etree.fromstring(ct_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    part_name = "/word/comments.xml"

    for override in tree.findall(f"{{{ns}}}Override"):
        if override.get("PartName") == part_name:
            return ct_xml

    new_override = etree.SubElement(tree, f"{{{ns}}}Override")
    new_override.set("PartName", part_name)
    new_override.set("ContentType", COMMENT_CONTENT_TYPE)

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


#write one Word comment into docx at target paragraph index
def _write_single_comment_docx(input_path, output_path, paragraph_index, message):
    # Read raw bytes directly — preserves images, mc:AlternateContent, VML shapes.
    with zipfile.ZipFile(input_path, "r") as z:
        names        = z.namelist()
        doc_xml      = z.read("word/document.xml")
        rels_xml     = z.read("word/_rels/document.xml.rels")
        ct_xml       = z.read("[Content_Types].xml")
        has_comments = "word/comments.xml" in names
        comments_xml = z.read("word/comments.xml") if has_comments else None

    doc_root  = etree.fromstring(doc_xml)
    body      = doc_root.find(f"{WQ}body")
    all_paras = [el for el in (body if body is not None else []) if el.tag == f"{WQ}p"]

    if len(all_paras) == 0:
        shutil.copy2(input_path, output_path)
        return

    target = max(0, min(paragraph_index, len(all_paras) - 1))

    if comments_xml:
        comments_root = etree.fromstring(comments_xml)
        existing_ids  = [int(el.get(f"{WQ}id", 0)) for el in comments_root.findall(f"{WQ}comment")]
        next_id = max(existing_ids, default=0) + 1
    else:
        comments_root = etree.Element(f"{WQ}comments", nsmap={"w": W})
        next_id = 1

    comments_root.append(_make_comment_element(next_id, message))
    _inject_comment_markers(all_paras[target], next_id)

    new_doc_xml      = etree.tostring(doc_root,      xml_declaration=True, encoding="UTF-8", standalone=True)
    new_comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    new_rels_xml     = _patch_rels(rels_xml)
    new_ct_xml       = _patch_content_types(ct_xml)

    tmp_path = output_path + ".tmp"
    with zipfile.ZipFile(input_path, "r") as zin, \
         zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, new_doc_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, new_rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, new_ct_xml)
            elif item.filename == "word/comments.xml":
                zout.writestr(item, new_comments_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
        if not has_comments:
            zout.writestr("word/comments.xml", new_comments_xml)

    os.replace(tmp_path, output_path)


def _comment_ids_in_document_order(doc_root):
    seen = set()
    ordered_ids = []
    for element in doc_root.iter():
        if element.tag not in (f"{WQ}commentRangeStart", f"{WQ}commentReference"):
            continue
        raw_id = element.get(f"{WQ}id")
        if raw_id is None or raw_id in seen:
            continue
        seen.add(raw_id)
        ordered_ids.append(raw_id)
    return ordered_ids


def _ensure_bold_run(run_element):
    r_pr = run_element.find(f"{WQ}rPr")
    if r_pr is None:
        r_pr = etree.Element(f"{WQ}rPr")
        run_element.insert(0, r_pr)
    if r_pr.find(f"{WQ}b") is None:
        etree.SubElement(r_pr, f"{WQ}b")


def _insert_author_query_prefix_run(paragraph_element, before_run, prefix_text):
    prefix_run = etree.Element(f"{WQ}r", nsmap=paragraph_element.nsmap)
    prefix_r_pr = etree.SubElement(prefix_run, f"{WQ}rPr")
    etree.SubElement(prefix_r_pr, f"{WQ}b")
    prefix_t = etree.SubElement(prefix_run, f"{WQ}t")
    prefix_t.text = prefix_text
    prefix_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    insert_at = list(paragraph_element).index(before_run)
    paragraph_element.insert(insert_at, prefix_run)


def _renumber_comment_element_author_query(comment_element, query_number):
    prefix_text = f"Author Query {query_number}. "
    paragraph = comment_element.find(f"{WQ}p")
    if paragraph is None:
        return

    for run in paragraph.findall(f"{WQ}r"):
        text_nodes = run.findall(f"{WQ}t")
        for text_node in text_nodes:
            text = text_node.text or ""
            match = re.match(r"^(Author Query\s+\d+\.\s*)(.*)$", text, flags=re.IGNORECASE)
            if match is None:
                continue
            _ensure_bold_run(run)
            text_node.text = prefix_text + match.group(2)
            text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            return

    runs = paragraph.findall(f"{WQ}r")
    for run in runs:
        text_nodes = run.findall(f"{WQ}t")
        if not text_nodes:
            continue
        text = text_nodes[0].text or ""
        if text.strip().isdigit():
            continue
        _insert_author_query_prefix_run(paragraph, run, prefix_text)
        return


def _renumber_author_queries_by_anchor_order(docx_path):
    with zipfile.ZipFile(docx_path, "r") as z:
        names = z.namelist()
        if "word/comments.xml" not in names:
            return False
        doc_xml = z.read("word/document.xml")
        comments_xml = z.read("word/comments.xml")

    doc_root = etree.fromstring(doc_xml)
    comments_root = etree.fromstring(comments_xml)
    ordered_ids = _comment_ids_in_document_order(doc_root)

    comments_by_id = {
        comment.get(f"{WQ}id"): comment
        for comment in comments_root.findall(f"{WQ}comment")
    }

    ordered_comment_elements = []
    seen_ids = set()
    query_number = 1
    for comment_id in ordered_ids:
        comment = comments_by_id.get(comment_id)
        if comment is None:
            continue
        _renumber_comment_element_author_query(comment, query_number)
        query_number += 1
        ordered_comment_elements.append(comment)
        seen_ids.add(comment_id)

    for comment in comments_root.findall(f"{WQ}comment"):
        comment_id = comment.get(f"{WQ}id")
        if comment_id in seen_ids:
            continue
        _renumber_comment_element_author_query(comment, query_number)
        query_number += 1
        ordered_comment_elements.append(comment)

    if not ordered_comment_elements:
        return False

    for child in list(comments_root):
        comments_root.remove(child)
    for comment in ordered_comment_elements:
        comments_root.append(comment)

    new_comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/comments.xml":
                zout.writestr(item, new_comments_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    os.replace(tmp_path, docx_path)
    return True


#build standard comment text for author tracked-change operations
def _build_authors_tracked_change_comment(plan):
    reason = plan.get("reason", "")
    if reason is None:
        reason = ""
    reason = reason.strip()

    if reason == "":
        return "Authors paragraph was changed to match the required naming pattern."

    return "Authors paragraph was changed because: " + reason


#apply title plan actions to document (tracked change + comment)
def _apply_title_plan(input_path, output_path, title_plan):
    action = title_plan.get("action", "none")
    if action == "none":
        return False

    anchor = title_plan.get("anchor_pos", 0)
    message = title_plan.get("message", "")
    needs_title_case_fix = title_plan.get("needs_title_case_fix", False)
    needs_style_fix = title_plan.get("style_ok") is False and title_plan.get("title_found") is not False

    wrote_output = False

    styles_loaded = _ensure_template_styles_available(
        input_path,
        output_path,
        [TITLE_REQUIRED_STYLE],
    )
    if styles_loaded is True:
        wrote_output = True

    if needs_title_case_fix is True or needs_style_fix is True:
        source_for_edit = input_path
        if wrote_output is True:
            source_for_edit = output_path

        doc_root, body_paras = _load_body_paras_raw(source_for_edit)
        if anchor < len(body_paras):
            para_el = body_paras[anchor]

            if needs_title_case_fix is True:
                new_text = title_plan.get("corrected_title", title_plan.get("title_text", ""))
                _remove_paragraph_content_keep_style(para_el)
                _append_plain_text_run(para_el, new_text)

            if needs_style_fix is True:
                old_style = title_plan.get("title_style", "") or "Normal"
                _apply_tracked_style_change(para_el, TITLE_REQUIRED_STYLE_ID, old_style, change_id=10)

            new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
            _save_doc_xml_to_zip(source_for_edit, output_path, new_doc_xml)
            wrote_output = True

    if message.strip() != "":
        if wrote_output is True:
            _write_single_comment_docx(
                input_path=output_path,
                output_path=output_path,
                paragraph_index=anchor,
                message=message,
            )
        else:
            _write_single_comment_docx(
                input_path=input_path,
                output_path=output_path,
                paragraph_index=anchor,
                message=message,
            )
        wrote_output = True

    if wrote_output is False:
        shutil.copy2(input_path, output_path)
        return False

    return True


#apply author plan actions to document
def _apply_author_plan(input_path, output_path, plan):
    action = plan.get("action", "none")
    authors_style_ok = plan.get("authors_style_ok", True)
    current_input_path = input_path
    applied_any = False

    if action == "none" and authors_style_ok is not False:
        return False

    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        [
            AUTHORS_REQUIRED_STYLE,
            AFFILIATIONS_REQUIRED_STYLE,
        ],
    )
    if styles_loaded is True:
        current_input_path = output_path
        applied_any = True

    anchor = plan.get("anchor_pos", 0)
    old_author_style = plan.get("authors_current_style", "") or "Normal"

    def _apply_style_if_needed(para_el):
        if authors_style_ok is False:
            _apply_tracked_style_change(para_el, AUTHORS_REQUIRED_STYLE_ID, old_author_style, change_id=20)

    def _style_fix_pass(path_in, path_out):
        """Standalone pass: apply tracked style change only."""
        doc_root2, body_paras2 = _load_body_paras_raw(path_in)
        if anchor < len(body_paras2):
            _apply_style_if_needed(body_paras2[anchor])
            new_xml = etree.tostring(doc_root2, xml_declaration=True, encoding="UTF-8", standalone=True)
            _save_doc_xml_to_zip(path_in, path_out, new_xml)

    if action == "add_missing_authors_comment":
        # Insert placeholder Authors + Author Affiliations paragraphs (tracked
        # insertions) immediately after the title, then attach the comment to the
        # Authors line so the editor sees it on the right line.
        doc_root, body_paras = _load_body_paras_raw(current_input_path)
        if anchor < len(body_paras):
            anchor_el = body_paras[anchor]
            # addnext inserts directly after anchor_el.  By adding affiliations
            # first then authors, the final order is: title → authors → affiliations.
            affiliations_placeholder = _make_inserted_styled_paragraph(
                "^aInstitution 1, Country, ^bInstitution 2, Country",
                AFFILIATIONS_REQUIRED_STYLE_ID, change_id=2
            )
            authors_placeholder = _make_inserted_authors_paragraph(
                "First I. Last^a, First I. Last^b, and First I. Last^c", change_id=1
            )
            anchor_el.addnext(affiliations_placeholder)
            anchor_el.addnext(authors_placeholder)
        new_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        _save_doc_xml_to_zip(current_input_path, output_path, new_xml)
        _write_single_comment_docx(
            input_path=output_path,
            output_path=output_path,
            paragraph_index=anchor + 1,
            message=plan["message"],
        )
        return True

    if action == "add_authors_naming_comment":
        _write_single_comment_docx(
            input_path=current_input_path,
            output_path=output_path,
            paragraph_index=anchor,
            message=plan["message"],
        )
        if authors_style_ok is False:
            _style_fix_pass(output_path, output_path)
        return True

    if action == "add_affiliation_mismatch_comment":
        _write_single_comment_docx(
            input_path=current_input_path,
            output_path=output_path,
            paragraph_index=anchor,
            message=plan["message"],
        )
        if authors_style_ok is False:
            _style_fix_pass(output_path, output_path)
        return True

    # Style-only fix (naming/position already correct, style is wrong).
    if action == "fix_author_style":
        _style_fix_pass(current_input_path, output_path)
        _write_single_comment_docx(
            input_path=output_path,
            output_path=output_path,
            paragraph_index=anchor,
            message=plan["message"],
        )
        return True

    # Also handle action == "none" with style wrong (fix style silently).
    if action == "none" and authors_style_ok is False:
        _style_fix_pass(current_input_path, output_path)
        return True

    # Keep python-docx for its API (add_paragraph, insert helpers) but use
    # _save_doc_xml_to_zip instead of doc.save() so the input zip's media files
    # and relationships are always copied through intact.
    doc = DocxDocument(current_input_path)
    change_id = 1

    def _flush(src_path):
        new_xml = etree.tostring(doc.element, xml_declaration=True, encoding="UTF-8", standalone=True)
        _save_doc_xml_to_zip(src_path, output_path, new_xml)

    if action == "insert_default_authors":
        insert_at = min(plan.get("insert_pos", anchor + 1), len(doc.paragraphs))
        _remove_blank_paragraph_at_index(doc, insert_at)
        _insert_paragraph_at_index(
            doc,
            insert_at,
            _make_plain_styled_paragraph(
                plan["corrected_authors_line"],
                AUTHORS_REQUIRED_STYLE_ID,
            ),
        )
        append_pos = plan.get("append_affiliation_pos")
        if append_pos is not None:
            if append_pos >= insert_at:
                append_pos += 1
            if append_pos < len(doc.paragraphs):
                paragraph = doc.paragraphs[append_pos]
                new_text = _append_affiliation_before_notes(
                    paragraph.text,
                    plan.get("append_affiliation_text", ""),
                )
                _remove_paragraph_content_keep_style(paragraph._element)
                _append_text_with_line_breaks(paragraph._element, new_text)

        affiliations_insert_at = plan.get("affiliations_insert_pos")
        if affiliations_insert_at is None:
            affiliations_insert_at = insert_at + 1
        elif affiliations_insert_at > insert_at:
            affiliations_insert_at += 1
        for i, line in enumerate(plan.get("default_affiliations", []), start=1):
            _insert_paragraph_at_index(
                doc,
                affiliations_insert_at + i - 1,
                _make_plain_styled_paragraph(line, AFFILIATIONS_REQUIRED_STYLE_ID),
            )

        _flush(current_input_path)
        _write_single_comment_docx(
            input_path=output_path,
            output_path=output_path,
            paragraph_index=insert_at,
            message="Replace these placeholder author names with the correct author list.",
        )
        if len(plan.get("default_affiliations", [])) > 0:
            _write_single_comment_docx(
                input_path=output_path,
                output_path=output_path,
                paragraph_index=affiliations_insert_at,
                message="Replace these placeholder affiliations with the authors' correct institutional affiliations.",
            )
        elif append_pos is not None:
            _write_single_comment_docx(
                input_path=output_path,
                output_path=output_path,
                paragraph_index=append_pos,
                message="Replace these placeholder affiliations with the authors' correct institutional affiliations.",
            )
        return True

    if action == "replace_in_place":
        if anchor < len(doc.paragraphs):
            paragraph = doc.paragraphs[anchor]
            old_text = paragraph.text
            new_text = plan["corrected_authors_line"]

            if plan.get("replace_first_line_only") is True:
                old_authors_line = _extract_authors_line_only(old_text)
                affiliations_lines = _extract_authors_affiliations_section_lines(old_text)

                _remove_paragraph_content_keep_style(paragraph._element)
                _append_tracked_replace(paragraph._element, old_authors_line, new_text, change_id)
                _apply_style_if_needed(paragraph._element)

                if len(affiliations_lines) > 0:
                    affiliations_paragraph = doc.add_paragraph()
                    for i, line in enumerate(affiliations_lines):
                        run = affiliations_paragraph.add_run(line)
                        if i < len(affiliations_lines) - 1:
                            run.add_break()
                    _apply_tracked_style_change(
                        affiliations_paragraph._element,
                        AFFILIATIONS_REQUIRED_STYLE_ID,
                        old_style_name="Normal",
                        change_id=30,
                    )
                    paragraph._element.addnext(affiliations_paragraph._element)

                _flush(current_input_path)
                _write_single_comment_docx(
                    input_path=output_path, output_path=output_path,
                    paragraph_index=anchor, message=_build_authors_tracked_change_comment(plan),
                )
                return True

            _remove_paragraph_content_keep_style(paragraph._element)
            _append_tracked_replace(paragraph._element, old_text, new_text, change_id)
            _apply_style_if_needed(paragraph._element)

        _flush(current_input_path)
        _write_single_comment_docx(
            input_path=output_path, output_path=output_path,
            paragraph_index=anchor, message=_build_authors_tracked_change_comment(plan),
        )
        return True

    if action == "move_only":
        target = plan["target_pos"]
        if anchor < len(doc.paragraphs):
            paragraph = doc.paragraphs[anchor]
            old_text = paragraph.text
            _remove_paragraph_content_keep_style(paragraph._element)
            _append_tracked_replace(paragraph._element, old_text, "", change_id)
            _apply_style_if_needed(paragraph._element)
            insert_paragraph = _make_inserted_authors_paragraph(old_text, change_id + 2)
            _insert_paragraph_at_index(doc, target, insert_paragraph)

        _flush(current_input_path)
        _write_single_comment_docx(
            input_path=output_path, output_path=output_path,
            paragraph_index=anchor, message=_build_authors_tracked_change_comment(plan),
        )
        return True

    if action == "move_and_replace":
        target = plan["target_pos"]
        if anchor < len(doc.paragraphs):
            paragraph = doc.paragraphs[anchor]
            old_text = paragraph.text
            new_text = plan["corrected_authors_line"]
            _remove_paragraph_content_keep_style(paragraph._element)
            _append_tracked_replace(paragraph._element, old_text, "", change_id)
            _apply_style_if_needed(paragraph._element)
            insert_paragraph = _make_inserted_authors_paragraph(new_text, change_id + 2)
            _insert_paragraph_at_index(doc, target, insert_paragraph)

        _flush(current_input_path)
        _write_single_comment_docx(
            input_path=output_path, output_path=output_path,
            paragraph_index=anchor, message=_build_authors_tracked_change_comment(plan),
        )
        return True

    return applied_any


#physically merge abstract paragraphs into one, appending later paragraph runs to the first
def _merge_abstract_paragraph_breaks(input_path, output_path, paragraph_positions):
    """Physically merge abstract paragraphs into one by moving runs and removing extra paragraphs."""
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    INLINE_TAGS = {"r", "hyperlink", "bookmarkStart", "bookmarkEnd", "proofErr",
                   "ins", "del", "sdt", "smartTag", "customXml"}

    doc_root, paras = _load_body_paras_raw(input_path)

    para_els = [paras[pos] for pos in paragraph_positions if pos < len(paras)]
    if len(para_els) < 2:
        return

    first_el = para_els[0]
    for later_el in para_els[1:]:
        space_run = etree.SubElement(first_el, f"{WQ}r")
        space_t   = etree.SubElement(space_run, f"{WQ}t")
        space_t.text = " "
        space_t.set(XML_SPACE, "preserve")
        for child in list(later_el):
            if child.tag.split("}")[-1] in INLINE_TAGS:
                first_el.append(child)
        parent = later_el.getparent()
        if parent is not None:
            parent.remove(later_el)

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)


def _apply_abstract_plan(input_path, output_path, abstract_plan):
    action = abstract_plan.get("action", "none")
    style_fixes = abstract_plan.get("style_fixes", [])
    current_input_path = input_path
    applied_any = False

    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        [
            ABSTRACT_REQUIRED_HEADING_STYLE,
            ABSTRACT_REQUIRED_PARAGRAPH_STYLE,
            PRACTITIONER_REQUIRED_HEADING_STYLE,
            PRACTITIONER_REQUIRED_NOTE_STYLE,
            AFFILIATIONS_REQUIRED_STYLE,
        ],
    )
    if styles_loaded is True:
        current_input_path = output_path
        applied_any = True

    if len(style_fixes) > 0:
        style_applied = _apply_tracked_style_fixes(
            current_input_path,
            output_path,
            style_fixes,
            change_id_start=7000,
        )
        if style_applied is True:
            current_input_path = output_path
            applied_any = True

    if abstract_plan.get("insert_default_affiliations") is True:
        affiliation_text = str(abstract_plan.get("default_affiliations_text", "")).strip()
        if affiliation_text != "":
            insert_at = abstract_plan.get("default_affiliations_insert_pos", 0)
            _insert_raw_body_paragraph(
                current_input_path,
                output_path,
                insert_at,
                _make_plain_styled_paragraph(affiliation_text, AFFILIATIONS_REQUIRED_STYLE_ID),
            )
            current_input_path = output_path
            applied_any = True
            _shift_abstract_plan_positions_after_insert(abstract_plan, insert_at)
            _write_single_comment_docx(
                input_path=current_input_path,
                output_path=output_path,
                paragraph_index=insert_at,
                message="Replace these placeholder affiliations with the authors' correct institutional affiliations.",
            )
            current_input_path = output_path

    if action == "none":
        return applied_any

    if action == "merge_abstract_paragraphs":
        positions = abstract_plan.get("merge_paragraph_positions", [])
        if len(positions) > 1:
            _merge_abstract_paragraph_breaks(current_input_path, output_path, positions)
            current_input_path = output_path
            applied_any = True

        for item in abstract_plan.get("comments", []):
            msg = item.get("message", "").strip()
            if msg:
                _write_single_comment_docx(
                    input_path=current_input_path,
                    output_path=output_path,
                    paragraph_index=item.get("anchor_pos", 0),
                    message=msg,
                )
                current_input_path = output_path
                applied_any = True
        return applied_any

    if action == "add_abstract_comments":
        comment_items = abstract_plan.get("comments", [])
        if len(comment_items) == 0:
            return applied_any

        for item in comment_items:
            item_anchor = item.get("anchor_pos", 0)
            item_message = item.get("message", "").strip()
            if item_message != "":
                _write_single_comment_docx(
                    input_path=current_input_path,
                    output_path=output_path,
                    paragraph_index=item_anchor,
                    message=item_message,
                )
                current_input_path = output_path
                applied_any = True
        return applied_any

    if action == "add_abstract_comment":
        item_message = abstract_plan.get("message", "").strip()
        if item_message == "":
            return applied_any

        _write_single_comment_docx(
            input_path=current_input_path,
            output_path=output_path,
            paragraph_index=abstract_plan["anchor_pos"],
            message=item_message,
        )
        return True

    return applied_any


def _apply_missing_practitioner_stub(input_path, output_path, abstract_plan):
    if abstract_plan.get("ensure_practitioner_section") is not True and abstract_plan.get("insert_practitioner_stub") is not True:
        return False

    practitioner_comment_message = str(
        abstract_plan.get("practitioner_missing_comment", "")
    ).strip()

    current_input_path = input_path
    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        [
            ABSTRACT_REQUIRED_PARAGRAPH_STYLE,
            PRACTITIONER_REQUIRED_HEADING_STYLE,
            PRACTITIONER_REQUIRED_NOTE_STYLE,
        ],
    )
    if styles_loaded is True:
        current_input_path = output_path

    heading_style_id = _resolve_style_id_from_doc(
        current_input_path,
        PRACTITIONER_REQUIRED_HEADING_STYLE,
        PRACTITIONER_REQUIRED_HEADING_STYLE_ID,
    )
    note_style_id = _resolve_style_id_from_doc(
        current_input_path,
        PRACTITIONER_REQUIRED_NOTE_STYLE,
        PRACTITIONER_REQUIRED_NOTE_STYLE_ID,
    )
    def _raw_style_id(para_element):
        pPr = para_element.find(f"{WQ}pPr")
        if pPr is None:
            return ""
        pStyle = pPr.find(f"{WQ}pStyle")
        if pStyle is None:
            return ""
        return pStyle.get(f"{WQ}val", "")

    def _raw_text(para_element):
        return "".join(t.text or "" for t in para_element.iter(f"{WQ}t"))

    def _is_practitioner_heading_text(text):
        low = text.strip().lower()
        if low == "practioner notes" or low == "practitioner notes":
            return True
        return False

    def _is_practitioner_section_end(text):
        low = text.strip().lower()
        if low == "":
            return False
        if low == "keywords" or low.startswith("keywords:"):
            return True
        if low == "citation" or low.startswith("citation:"):
            return True
        if low == "introduction":
            return True
        return False

    def _has_visible_text(text):
        cleaned = text
        cleaned = cleaned.replace("\u200b", "")
        cleaned = cleaned.replace("\ufeff", "")
        cleaned = cleaned.replace("\u2060", "")
        if cleaned.strip() == "":
            return False
        return True

    def _is_abstract_heading_candidate(paragraph):
        low = paragraph.text.strip().lower()
        if low == "abstract":
            return True
        if paragraph.style == ABSTRACT_REQUIRED_HEADING_STYLE and "abstract" in low:
            return True
        return False

    insert_after_pos = abstract_plan.get("practitioner_insert_after_pos", -1)
    if insert_after_pos is None:
        insert_after_pos = -1
    if insert_after_pos < 0:
        insert_after_pos = abstract_plan.get("anchor_pos", 0)

    # Recompute from the current (post-edit) document to avoid stale indexes
    # from earlier planning stages after title/author/abstract rewrites.
    current_paragraphs = load_paragraphs(current_input_path)
    current_front = get_front_page(current_paragraphs)

    current_abstract_heading = None
    for paragraph in current_front:
        if _is_abstract_heading_candidate(paragraph) is True:
            current_abstract_heading = paragraph
            break

    if current_abstract_heading is not None:
        current_abstract_paragraph_pos = None
        for paragraph in current_front:
            if paragraph.index <= current_abstract_heading.index:
                continue
            if (
                paragraph.style == ABSTRACT_REQUIRED_HEADING_STYLE
                and _is_abstract_heading_candidate(paragraph) is False
                and _has_visible_text(paragraph.text) is True
            ):
                break
            if _is_abstract_body_end_marker(paragraph) is True:
                break
            if _has_visible_text(paragraph.text) is True:
                current_abstract_paragraph_pos = paragraph.index
                break

        if current_abstract_paragraph_pos is not None:
            insert_after_pos = current_abstract_paragraph_pos
        else:
            insert_after_pos = current_abstract_heading.index
    else:
        current_abstract_state = abstractFound(current_input_path)
        if current_abstract_state.get("affiliations_found") is True:
            insert_after_pos = current_abstract_state.get("affiliations_end_pos", insert_after_pos)

    doc_root, body_paras = _load_body_paras_raw(current_input_path)
    body = doc_root.find(f"{WQ}body")
    if body is None:
        return False

    heading_index = -1
    for i, para_element in enumerate(body_paras):
        if _is_practitioner_heading_text(_raw_text(para_element)):
            heading_index = i
            break

    changed = False

    if heading_index >= 0:
        heading_element = body_paras[heading_index]
        current_style = _raw_style_id(heading_element)
        if current_style != heading_style_id:
            _apply_tracked_style_change(
                heading_element,
                heading_style_id,
                current_style,
                7300,
            )
            changed = True

        has_note = False
        i = heading_index + 1
        while i < len(body_paras):
            text = _raw_text(body_paras[i])
            if _is_practitioner_section_end(text):
                break
            if _raw_style_id(body_paras[i]) == note_style_id:
                has_note = True
                break
            i += 1

        if has_note is False:
            note_para = _make_inserted_styled_paragraph(" ", note_style_id, 7302)
            if heading_index + 1 < len(body_paras):
                body_paras[heading_index + 1].addprevious(note_para)
            else:
                sect_pr = body.find(f"{WQ}sectPr")
                if sect_pr is not None:
                    sect_pr.addprevious(note_para)
                else:
                    body.append(note_para)
            changed = True

        if changed is False and practitioner_comment_message == "":
            return False

        if changed is True:
            new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
            _save_doc_xml_to_zip(current_input_path, output_path, new_doc_xml)

        if practitioner_comment_message != "":
            comment_input_path = current_input_path
            if changed is True:
                comment_input_path = output_path
            _write_single_comment_docx(
                input_path=comment_input_path,
                output_path=output_path,
                paragraph_index=heading_index,
                message=practitioner_comment_message,
            )
        return True

    insert_at = insert_after_pos + 1
    if insert_at < 0:
        insert_at = 0

    heading_para = _make_inserted_styled_paragraph(
        "Practitioner Notes",
        heading_style_id,
        7302,
        PRACTITIONER_INSERTED_HEADING_SIZE_HALF_POINTS,
    )
    note_para = _make_inserted_styled_paragraph(" ", note_style_id, 7304)

    if insert_at < len(body_paras):
        anchor = body_paras[insert_at]
        anchor.addprevious(heading_para)
        anchor.addprevious(note_para)
    else:
        sect_pr = body.find(f"{WQ}sectPr")
        if sect_pr is not None:
            sect_pr.addprevious(heading_para)
            sect_pr.addprevious(note_para)
        else:
            body.append(heading_para)
            body.append(note_para)

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(current_input_path, output_path, new_doc_xml)

    if practitioner_comment_message != "":
        _write_single_comment_docx(
            input_path=output_path,
            output_path=output_path,
            paragraph_index=insert_at,
            message=practitioner_comment_message,
        )
    return True


def _apply_missing_keywords_stub(input_path, output_path, keywords_plan):
    if keywords_plan.get("ensure_keywords_section") is not True:
        return False

    generated_keywords = keywords_plan.get("generated_keywords") or []

    current_input_path = input_path
    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        # When we have an AI-generated keyword list to insert, also pull in
        # the body text style so the new body paragraph is template-clean.
        [KEYWORDS_REQUIRED_HEADING_STYLE]
        + ([KEYWORDS_REQUIRED_TEXT_STYLE] if generated_keywords else []),
    )
    if styles_loaded is True:
        current_input_path = output_path

    heading_style_id = _resolve_style_id_from_doc(
        current_input_path,
        KEYWORDS_REQUIRED_HEADING_STYLE,
        KEYWORDS_REQUIRED_HEADING_STYLE_ID,
    )
    note_style_id = _resolve_style_id_from_doc(
        current_input_path,
        PRACTITIONER_REQUIRED_NOTE_STYLE,
        PRACTITIONER_REQUIRED_NOTE_STYLE_ID,
    )
    text_style_id = _resolve_style_id_from_doc(
        current_input_path,
        KEYWORDS_REQUIRED_TEXT_STYLE,
        KEYWORDS_REQUIRED_TEXT_STYLE_ID,
    )

    doc_root, body_paras = _load_body_paras_raw(current_input_path)
    body = doc_root.find(f"{WQ}body")
    if body is None:
        return False

    def _raw_style_id(para_element):
        pPr = para_element.find(f"{WQ}pPr")
        if pPr is None:
            return ""
        pStyle = pPr.find(f"{WQ}pStyle")
        if pStyle is None:
            return ""
        return pStyle.get(f"{WQ}val", "")

    def _raw_text(para_element):
        return "".join(t.text or "" for t in para_element.iter(f"{WQ}t"))

    def _is_keywords_heading_text(text):
        low = text.strip().lower()
        if low == "keywords":
            return True
        if low.startswith("keywords:"):
            return True
        return False

    def _is_practitioner_heading_text(text):
        low = text.strip().lower()
        if low == "practioner notes" or low == "practitioner notes":
            return True
        return False

    def _is_practitioner_section_end(text):
        low = text.strip().lower()
        if low == "":
            return False
        if low == "keywords" or low.startswith("keywords:"):
            return True
        if low == "citation" or low.startswith("citation:"):
            return True
        if low == "introduction" or low.startswith("introduction:"):
            return True
        return False

    for para_element in body_paras:
        if _is_keywords_heading_text(_raw_text(para_element)) is True:
            return False

    practitioner_heading_index = -1
    for i, para_element in enumerate(body_paras):
        if _is_practitioner_heading_text(_raw_text(para_element)) is True:
            practitioner_heading_index = i
            break

    insert_at = 0
    remove_positions = []
    if practitioner_heading_index >= 0:
        insert_after = practitioner_heading_index
        i = practitioner_heading_index + 1
        while i < len(body_paras):
            text = _raw_text(body_paras[i])
            if _is_practitioner_section_end(text) is True:
                break
            if _raw_style_id(body_paras[i]) == note_style_id:
                insert_after = i
            elif text.strip() == "":
                remove_positions.append(i)
            else:
                insert_after = i
            i += 1
        insert_at = insert_after + 1
    else:
        intro_index = -1
        for i, para_element in enumerate(body_paras):
            low = _raw_text(para_element).strip().lower()
            if low == "introduction" or low.startswith("introduction:"):
                intro_index = i
                break
        if intro_index >= 0:
            insert_at = intro_index
        else:
            insert_at = keywords_plan.get("anchor_pos", 0)

    if insert_at < 0:
        insert_at = 0

    for pos in sorted(remove_positions, reverse=True):
        if pos < 0 or pos >= len(body_paras):
            continue
        parent = body_paras[pos].getparent()
        if parent is None:
            continue
        parent.remove(body_paras[pos])
        if pos < insert_at:
            insert_at -= 1
    if remove_positions:
        body_paras = [el for el in body if el.tag == f"{WQ}p"]

    heading_para = _make_inserted_styled_paragraph(
        "Keywords",
        heading_style_id,
        7400,
        PRACTITIONER_INSERTED_HEADING_SIZE_HALF_POINTS,
    )

    body_para = None
    if generated_keywords:
        # JUTLP template: single line, comma-separated. Mirrors the
        # screenshot example "i.e. artificial intelligence, curriculum
        # design, leadership, belonging".
        body_text = ", ".join(generated_keywords)
        body_para = _make_inserted_styled_paragraph(
            body_text,
            text_style_id,
            7401,
        )

    if insert_at < len(body_paras):
        anchor = body_paras[insert_at]
        anchor.addprevious(heading_para)
        if body_para is not None:
            anchor.addprevious(body_para)
    else:
        sect_pr = body.find(f"{WQ}sectPr")
        if sect_pr is not None:
            sect_pr.addprevious(heading_para)
            if body_para is not None:
                sect_pr.addprevious(body_para)
        else:
            body.append(heading_para)
            if body_para is not None:
                body.append(body_para)

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(current_input_path, output_path, new_doc_xml)

    comment_message = str(
        keywords_plan.get(
            "missing_keywords_comment",
            "Keywords missing, please add up to 5 keywords pertaining to the manuscript.",
        )
    ).strip()
    if comment_message != "":
        _write_single_comment_docx(
            input_path=output_path,
            output_path=output_path,
            paragraph_index=insert_at,
            message=comment_message,
        )
    return True


def _apply_inline_keywords_split(input_path, output_path, keywords_plan):
    if keywords_plan.get("split_inline_keywords") is not True:
        return False

    inline_keywords_text = str(keywords_plan.get("inline_keywords_text", "")).strip()
    if inline_keywords_text == "":
        return False

    paragraph_index = keywords_plan.get("keywords_heading_pos")
    if paragraph_index is None:
        return False

    inline_heading_text = str(keywords_plan.get("keywords_heading_text", "")).strip()

    heading_style_id = _resolve_style_id_from_doc(
        input_path,
        KEYWORDS_REQUIRED_HEADING_STYLE,
        KEYWORDS_REQUIRED_HEADING_STYLE_ID,
    )
    text_style_id = _resolve_style_id_from_doc(
        input_path,
        KEYWORDS_REQUIRED_TEXT_STYLE,
        KEYWORDS_REQUIRED_TEXT_STYLE_ID,
    )

    doc_root, body_paras = _load_body_paras_raw(input_path)

    def _raw_text(para_element):
        return "".join(t.text or "" for t in para_element.iter(f"{WQ}t")).strip()

    para_el = None
    old_text = ""
    if paragraph_index >= 0 and paragraph_index < len(body_paras):
        possible_para = body_paras[paragraph_index]
        possible_text = _raw_text(possible_para)
        if _extract_inline_keywords_text(possible_text) != "":
            para_el = possible_para
            old_text = possible_text

    if para_el is None:
        for possible_para in body_paras:
            possible_text = _raw_text(possible_para)
            possible_inline = _extract_inline_keywords_text(possible_text)
            if possible_inline == inline_keywords_text:
                para_el = possible_para
                old_text = possible_text
                break
            if inline_heading_text != "" and possible_text == inline_heading_text:
                para_el = possible_para
                old_text = possible_text
                break

    if para_el is None:
        return False

    old_style = _read_raw_paragraph_style_name(para_el)
    if old_style != heading_style_id:
        _apply_tracked_style_change(para_el, heading_style_id, old_style, 7140)

    _remove_paragraph_content_keep_style(para_el)
    _append_tracked_replace(para_el, old_text, "Keywords", 7142)

    keyword_para = _make_inserted_styled_paragraph(
        inline_keywords_text,
        text_style_id,
        7144,
    )
    para_el.addnext(keyword_para)

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


#apply keywords plan action(s)
def _apply_keywords_plan(input_path, output_path, keywords_plan):
    action = keywords_plan.get("action", "none")
    style_fixes = keywords_plan.get("style_fixes", [])
    current_input_path = input_path
    applied_any = False

    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        [
            KEYWORDS_REQUIRED_HEADING_STYLE,
            KEYWORDS_REQUIRED_TEXT_STYLE,
        ],
    )
    if styles_loaded is True:
        current_input_path = output_path
        applied_any = True

    split_applied = _apply_inline_keywords_split(
        current_input_path,
        output_path,
        keywords_plan,
    )
    if split_applied is True:
        current_input_path = output_path
        applied_any = True

    style_fixes_to_apply = style_fixes
    if keywords_plan.get("split_inline_keywords") is True:
        style_fixes_to_apply = []
        for fix in style_fixes:
            if fix.get("paragraph_index") != keywords_plan.get("keywords_heading_pos"):
                style_fixes_to_apply.append(fix)

    if len(style_fixes_to_apply) > 0:
        style_applied = _apply_tracked_style_fixes(
            current_input_path,
            output_path,
            style_fixes_to_apply,
            change_id_start=7100,
        )
        if style_applied is True:
            current_input_path = output_path
            applied_any = True

    if action == "none":
        return applied_any

    if action == "add_keywords_comment":
        item_message = keywords_plan.get("message", "").strip()
        if item_message == "":
            return applied_any

        _write_single_comment_docx(
            input_path=current_input_path,
            output_path=output_path,
            paragraph_index=keywords_plan["anchor_pos"],
            message=item_message,
        )
        return True

    return applied_any


def _raw_paragraph_text(para_element):
    return "".join(t.text or "" for t in para_element.iter(f"{WQ}t")).strip()


def _normalise_docx_target(target):
    clean = str(target).replace("\\", "/")
    if clean.startswith("/"):
        clean = clean[1:]
    if clean.startswith("word/"):
        return clean
    return "word/" + clean


def _next_relationship_id(rels_root):
    highest = 0
    for rel in rels_root.findall(f"{PACKAGE_REL_Q}Relationship"):
        rel_id = rel.get("Id", "")
        if rel_id.startswith("rId") is False:
            continue
        number_text = rel_id[3:]
        if number_text.isdigit() is False:
            continue
        number = int(number_text)
        if number > highest:
            highest = number
    return "rId" + str(highest + 1)


def _next_footer_name(names):
    index = 1
    while True:
        name = "word/footer" + str(index) + ".xml"
        if name not in names:
            return name
        index += 1


def _find_relationship_target(rels_root, rel_id):
    for rel in rels_root.findall(f"{PACKAGE_REL_Q}Relationship"):
        if rel.get("Id") == rel_id:
            return rel.get("Target", "")
    return ""


def _find_first_footer_ref(sect_pr):
    first_ref = None

    for footer_ref in sect_pr.findall(f"{WQ}footerReference"):
        ref_type = footer_ref.get(f"{WQ}type", "")
        if ref_type == "first":
            first_ref = footer_ref

    if first_ref is not None:
        return first_ref
    return None


#find or create the first page footer part for the citation text
def _ensure_first_footer(doc_root, rels_root, names):
    body = doc_root.find(f"{WQ}body")
    if body is None:
        return "", False

    sect_pr = body.find(f"{WQ}sectPr")
    if sect_pr is None:
        sect_pr = etree.SubElement(body, f"{WQ}sectPr")

    footer_ref = _find_first_footer_ref(sect_pr)
    if footer_ref is not None:
        rel_id = footer_ref.get(f"{RQ}id", "")
        target = _find_relationship_target(rels_root, rel_id)
        if target != "":
            title_pg = sect_pr.find(f"{WQ}titlePg")
            if title_pg is None:
                etree.SubElement(sect_pr, f"{WQ}titlePg")
            return _normalise_docx_target(target), False

    footer_name = _next_footer_name(names)
    rel_id = _next_relationship_id(rels_root)

    new_rel = etree.SubElement(rels_root, f"{PACKAGE_REL_Q}Relationship")
    new_rel.set("Id", rel_id)
    new_rel.set("Type", FOOTER_REL_TYPE)
    new_rel.set("Target", footer_name.replace("word/", ""))

    new_ref = etree.Element(f"{WQ}footerReference")
    new_ref.set(f"{WQ}type", "first")
    new_ref.set(f"{RQ}id", rel_id)
    sect_pr.insert(0, new_ref)

    title_pg = sect_pr.find(f"{WQ}titlePg")
    if title_pg is None:
        etree.SubElement(sect_pr, f"{WQ}titlePg")

    return footer_name, True


def _patch_footer_content_type(ct_xml, footer_name):
    tree = etree.fromstring(ct_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    part_name = "/" + footer_name

    for override in tree.findall(f"{{{ns}}}Override"):
        if override.get("PartName") == part_name:
            return ct_xml

    new_override = etree.SubElement(tree, f"{{{ns}}}Override")
    new_override.set("PartName", part_name)
    new_override.set("ContentType", FOOTER_CONTENT_TYPE)

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


#make one footer paragraph using the Citation style from the template
def _make_footer_citation_paragraph(text, style_id):
    paragraph = etree.Element(f"{WQ}p", nsmap={"w": W})
    p_pr = etree.SubElement(paragraph, f"{WQ}pPr")
    p_style = etree.SubElement(p_pr, f"{WQ}pStyle")
    p_style.set(f"{WQ}val", style_id)

    run = etree.SubElement(paragraph, f"{WQ}r")
    text_element = etree.SubElement(run, f"{WQ}t")
    text_element.text = text
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return paragraph


#replace any old footer citation and write the new citation text
def _write_citation_to_footer(footer_xml, citation_text):
    if footer_xml is None:
        footer_root = etree.Element(f"{WQ}ftr", nsmap={"w": W})
    else:
        footer_root = etree.fromstring(footer_xml)

    for paragraph in list(footer_root.findall(f"{WQ}p")):
        p_style = paragraph.find(f"{WQ}pPr/{WQ}pStyle")
        if p_style is None:
            continue
        if p_style.get(f"{WQ}val", "") == CITATION_FOOTER_STYLE_ID:
            footer_root.remove(paragraph)

    insert_before = None
    for paragraph in footer_root.findall(f"{WQ}p"):
        p_style = paragraph.find(f"{WQ}pPr/{WQ}pStyle")
        if p_style is None:
            continue
        if p_style.get(f"{WQ}val", "") == "Footer":
            insert_before = paragraph
            break

    citation_heading = _make_footer_citation_paragraph("Citation:", CITATION_FOOTER_STYLE_ID)
    citation_para = _make_footer_citation_paragraph(citation_text, CITATION_FOOTER_STYLE_ID)

    if insert_before is None:
        footer_root.append(citation_heading)
        footer_root.append(citation_para)
    else:
        insert_before.addprevious(citation_heading)
        insert_before.addprevious(citation_para)

    return etree.tostring(footer_root, xml_declaration=True, encoding="UTF-8", standalone=True)


#copy the body citation into the first page footer, then remove the old body citation
def _move_body_citation_to_footer(input_path, output_path, citation_text, remove_positions):
    if citation_text.strip() == "":
        return False

    with zipfile.ZipFile(input_path, "r") as z:
        names = z.namelist()
        doc_xml = z.read("word/document.xml")
        rels_xml = z.read("word/_rels/document.xml.rels")
        ct_xml = z.read("[Content_Types].xml")

    doc_root = etree.fromstring(doc_xml)
    rels_root = etree.fromstring(rels_xml)
    footer_name, footer_created = _ensure_first_footer(doc_root, rels_root, names)
    if footer_name == "":
        return False

    footer_xml = None
    if footer_name in names:
        with zipfile.ZipFile(input_path, "r") as z:
            footer_xml = z.read(footer_name)
    footer_needs_write = footer_created
    if footer_name not in names:
        footer_needs_write = True

    body = doc_root.find(f"{WQ}body")
    body_paras = []
    if body is not None:
        body_paras = [el for el in body if el.tag == f"{WQ}p"]

    sorted_positions = sorted(set(remove_positions), reverse=True)
    for paragraph_index in sorted_positions:
        if paragraph_index < 0 or paragraph_index >= len(body_paras):
            continue
        body.remove(body_paras[paragraph_index])

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    new_rels_xml = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    new_ct_xml = _patch_footer_content_type(ct_xml, footer_name)
    new_footer_xml = _write_citation_to_footer(footer_xml, citation_text)

    tmp_path = output_path + ".tmp"
    with zipfile.ZipFile(input_path, "r") as zin, \
         zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, new_doc_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, new_rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, new_ct_xml)
            elif item.filename == footer_name:
                zout.writestr(item, new_footer_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
        if footer_needs_write is True:
            zout.writestr(footer_name, new_footer_xml)

    os.replace(tmp_path, output_path)
    return True


#make Introduction start on a new page
def _apply_intro_page_break(input_path, output_path):
    doc_root, body_paras = _load_body_paras_raw(input_path)

    for para_element in body_paras:
        # Strip any leading section number ("1. Introduction") so the page break
        # is inserted regardless of whether the author numbered their headings.
        text = _normalise_heading_text(
            strip_leading_section_number(_raw_paragraph_text(para_element))
        )
        if text != "introduction":
            continue

        p_pr = para_element.find(f"{WQ}pPr")
        if p_pr is None:
            p_pr = etree.Element(f"{WQ}pPr")
            para_element.insert(0, p_pr)

        page_break = p_pr.find(f"{WQ}pageBreakBefore")
        if page_break is not None:
            return False

        etree.SubElement(p_pr, f"{WQ}pageBreakBefore")
        new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
        return True

    return False


#apply citation plan action(s)
def _apply_citation_plan(input_path, output_path, citation_plan):
    action = citation_plan.get("action", "none")
    style_fixes = citation_plan.get("style_fixes", [])
    current_input_path = input_path
    applied_any = False
    citation_text = str(citation_plan.get("citation_text", "")).strip()

    if action == "move_citation_to_footer" and citation_text != "":
        styles_loaded = _ensure_template_styles_available(
            current_input_path,
            output_path,
            [
                CITATION_FOOTER_STYLE,
            ],
        )
        if styles_loaded is True:
            current_input_path = output_path
            applied_any = True

    if len(style_fixes) > 0:
        style_applied = _apply_tracked_style_fixes(
            current_input_path,
            output_path,
            style_fixes,
            change_id_start=7200,
        )
        if style_applied is True:
            current_input_path = output_path
            applied_any = True

    if action == "none":
        return applied_any

    if action == "move_citation_to_footer":
        remove_positions = citation_plan.get("citation_remove_positions", [])
        moved = _move_body_citation_to_footer(
            current_input_path,
            output_path,
            citation_text,
            remove_positions,
        )
        if moved is True:
            return True
        return applied_any

    if action == "add_citation_comment":
        item_message = citation_plan.get("message", "").strip()
        if item_message == "":
            return applied_any

        _write_single_comment_docx(
            input_path=current_input_path,
            output_path=output_path,
            paragraph_index=citation_plan["anchor_pos"],
            message=item_message,
        )
        return True

    return applied_any


#helper functions for tracked body paragraph formatting changes
def _remove_tracked_property_change(property_element, change_tag):
    existing = property_element.find(change_tag)
    if existing is not None:
        property_element.remove(existing)


def _copy_properties_without_change(property_element, property_tag, change_tag):
    if property_element is None:
        old_properties = etree.Element(property_tag)
    else:
        old_properties = copy.deepcopy(property_element)

    _remove_tracked_property_change(old_properties, change_tag)
    return old_properties


def _set_tracked_change_details(change_element, change_id):
    change_element.set(f"{WQ}id", str(change_id))
    change_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
    change_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)


def _find_or_make_child(parent_element, child_tag):
    child = parent_element.find(child_tag)
    if child is None:
        child = etree.SubElement(parent_element, child_tag)
    return child


def _apply_tracked_body_paragraph_format(para_element, fix, change_id):
    p_pr = para_element.find(f"{WQ}pPr")
    old_p_pr = _copy_properties_without_change(p_pr, f"{WQ}pPr", f"{WQ}pPrChange")

    if p_pr is None:
        p_pr = etree.Element(f"{WQ}pPr")
        para_element.insert(0, p_pr)

    _remove_tracked_property_change(p_pr, f"{WQ}pPrChange")

    if fix.get("alignment") is True:
        jc = _find_or_make_child(p_pr, f"{WQ}jc")
        jc.set(f"{WQ}val", "both")

    if fix.get("line_spacing") is True or fix.get("spacing") is True:
        spacing = _find_or_make_child(p_pr, f"{WQ}spacing")
        spacing.set(f"{WQ}line", BODY_REQUIRED_LINE_SPACING_XML)
        spacing.set(f"{WQ}lineRule", "auto")
        spacing.set(f"{WQ}after", BODY_REQUIRED_SPACE_AFTER_XML)
        spacing.set(f"{WQ}before", "0")

    if fix.get("indent") is True:
        ind = _find_or_make_child(p_pr, f"{WQ}ind")
        ind.set(f"{WQ}left", "0")
        ind.set(f"{WQ}start", "0")
        ind.set(f"{WQ}firstLine", "0")
        if f"{WQ}hanging" in ind.attrib:
            del ind.attrib[f"{WQ}hanging"]
        if f"{WQ}hangingChars" in ind.attrib:
            del ind.attrib[f"{WQ}hangingChars"]

    p_pr_change = etree.SubElement(p_pr, f"{WQ}pPrChange")
    _set_tracked_change_details(p_pr_change, change_id)
    p_pr_change.append(old_p_pr)


def _run_has_body_text(run_element):
    text_elements = run_element.findall(f"{WQ}t")
    for text_element in text_elements:
        if text_element.text is None:
            continue
        if text_element.text.strip() != "":
            return True
    return False


def _apply_tracked_body_run_format(run_element, fix, change_id):
    r_pr = run_element.find(f"{WQ}rPr")
    old_r_pr = _copy_properties_without_change(r_pr, f"{WQ}rPr", f"{WQ}rPrChange")

    if r_pr is None:
        r_pr = etree.Element(f"{WQ}rPr")
        run_element.insert(0, r_pr)

    _remove_tracked_property_change(r_pr, f"{WQ}rPrChange")

    if fix.get("font_name") is True:
        r_fonts = _find_or_make_child(r_pr, f"{WQ}rFonts")
        r_fonts.set(f"{WQ}ascii", BODY_REQUIRED_FONT_NAME)
        r_fonts.set(f"{WQ}hAnsi", BODY_REQUIRED_FONT_NAME)
        r_fonts.set(f"{WQ}cs", BODY_REQUIRED_FONT_NAME)
        r_fonts.set(f"{WQ}eastAsia", BODY_REQUIRED_FONT_NAME)

    if fix.get("font_size") is True:
        sz = _find_or_make_child(r_pr, f"{WQ}sz")
        sz.set(f"{WQ}val", BODY_REQUIRED_FONT_SIZE_XML)
        sz_cs = _find_or_make_child(r_pr, f"{WQ}szCs")
        sz_cs.set(f"{WQ}val", BODY_REQUIRED_FONT_SIZE_XML)

    r_pr_change = etree.SubElement(r_pr, f"{WQ}rPrChange")
    _set_tracked_change_details(r_pr_change, change_id)
    r_pr_change.append(old_r_pr)


def _apply_body_run_format(run_element, fix):
    r_pr = run_element.find(f"{WQ}rPr")
    if r_pr is None:
        r_pr = etree.Element(f"{WQ}rPr")
        run_element.insert(0, r_pr)

    _remove_tracked_property_change(r_pr, f"{WQ}rPrChange")

    if fix.get("font_name") is True:
        r_fonts = _find_or_make_child(r_pr, f"{WQ}rFonts")
        r_fonts.set(f"{WQ}ascii", BODY_REQUIRED_FONT_NAME)
        r_fonts.set(f"{WQ}hAnsi", BODY_REQUIRED_FONT_NAME)
        r_fonts.set(f"{WQ}cs", BODY_REQUIRED_FONT_NAME)
        r_fonts.set(f"{WQ}eastAsia", BODY_REQUIRED_FONT_NAME)

    if fix.get("font_size") is True:
        sz = _find_or_make_child(r_pr, f"{WQ}sz")
        sz.set(f"{WQ}val", BODY_REQUIRED_FONT_SIZE_XML)
        sz_cs = _find_or_make_child(r_pr, f"{WQ}szCs")
        sz_cs.set(f"{WQ}val", BODY_REQUIRED_FONT_SIZE_XML)


def _apply_tracked_body_format_fixes(input_path, output_path, format_fixes, change_id_start):
    if format_fixes is None:
        return False
    if len(format_fixes) == 0:
        return False

    doc_root, body_paras = _load_body_paras_raw(input_path)
    changed = False
    change_id = change_id_start
    seen_positions = set()

    for fix in format_fixes:
        paragraph_index = fix.get("paragraph_index")
        if paragraph_index is None:
            continue
        if paragraph_index in seen_positions:
            continue
        if paragraph_index < 0 or paragraph_index >= len(body_paras):
            continue

        para_element = body_paras[paragraph_index]

        paragraph_change_needed = False
        if fix.get("alignment") is True:
            paragraph_change_needed = True
        if fix.get("line_spacing") is True:
            paragraph_change_needed = True
        if fix.get("spacing") is True:
            paragraph_change_needed = True
        if fix.get("indent") is True:
            paragraph_change_needed = True

        if paragraph_change_needed is True:
            _apply_tracked_body_paragraph_format(para_element, fix, change_id)
            change_id += 1
            changed = True

        run_change_needed = False
        if fix.get("font_name") is True:
            run_change_needed = True
        if fix.get("font_size") is True:
            run_change_needed = True

        if run_change_needed is True:
            direct_runs = para_element.findall(f"{WQ}r")
            for run_element in direct_runs:
                if _run_has_body_text(run_element) is False:
                    continue
                _apply_body_run_format(run_element, fix)
                changed = True

        seen_positions.add(paragraph_index)

    if changed is False:
        return False

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


def _apply_document_body_plan(input_path, output_path, body_plan):
    action = body_plan.get("action", "none")
    applied_any = False
    current_input_path = input_path

    format_fixes = body_plan.get("format_fixes", [])
    format_applied = _apply_tracked_body_format_fixes(
        current_input_path,
        output_path,
        format_fixes,
        change_id_start=7300,
    )
    if format_applied is True:
        current_input_path = output_path
        applied_any = True

    if action == "none":
        return applied_any

    if action == "apply_document_body_formatting":
        return applied_any

    if action == "add_document_body_comments":
        comment_items = body_plan.get("comments", [])
        if len(comment_items) == 0:
            return applied_any

        i = 0
        while i < len(comment_items):
            item = comment_items[i]
            item_anchor = item.get("anchor_pos", 0)
            item_message = item.get("message", "").strip()
            if item_message != "":
                if i == 0:
                    _write_single_comment_docx(
                        input_path=current_input_path,
                        output_path=output_path,
                        paragraph_index=item_anchor,
                        message=item_message,
                    )
                else:
                    _write_single_comment_docx(
                        input_path=output_path,
                        output_path=output_path,
                        paragraph_index=item_anchor,
                        message=item_message,
                    )
            i += 1
        return True

    if action == "add_document_body_comment":
        _write_single_comment_docx(
            input_path=current_input_path,
            output_path=output_path,
            paragraph_index=body_plan["anchor_pos"],
            message=body_plan["message"],
        )
        return True

    return applied_any


def _heading_should_be_centered(text):
    clean = _normalise_heading_text(text)
    clean = clean.rstrip(":.")

    for heading_text in CENTER_ALIGNED_HEADING_TEXTS:
        if clean == heading_text:
            return True

    return False


def _heading_should_be_left_aligned(text):
    clean = _normalise_heading_text(text)
    clean = clean.rstrip(":.")

    for heading_text in LEFT_ALIGNED_HEADING_TEXTS:
        if clean == heading_text:
            return True

    return False


def _read_direct_paragraph_alignment(para_element):
    p_pr = para_element.find(f"{WQ}pPr")
    if p_pr is None:
        return ""

    jc = p_pr.find(f"{WQ}jc")
    if jc is None:
        return ""

    value = jc.get(f"{WQ}val", "")
    if value is None:
        return ""

    return value


def _style_already_centres_heading(style_name):
    if style_name == HEADING_1_STYLE_ID:
        return True
    if style_name == "Heading1":
        return True
    if style_name == "ReferencesHeading":
        return True
    return False


#align headings that have a journal-specific alignment rule
def _apply_heading_center_alignment(input_path, output_path):
    doc_root, body_paras = _load_body_paras_raw(input_path)
    changed = False

    for para_element in body_paras:
        text = _raw_paragraph_text(para_element)
        style_name = _read_raw_paragraph_style_name(para_element)
        if _heading_should_be_left_aligned(text) is True or style_name in TITLE_LEFT_ALIGNED_STYLE_IDS:
            direct_alignment = _read_direct_paragraph_alignment(para_element)
            if direct_alignment == "left":
                continue

            p_pr = para_element.find(f"{WQ}pPr")
            if p_pr is None:
                p_pr = etree.Element(f"{WQ}pPr")
                para_element.insert(0, p_pr)

            jc = _find_or_make_child(p_pr, f"{WQ}jc")
            jc.set(f"{WQ}val", "left")
            changed = True
            continue

        if _heading_should_be_centered(text) is False:
            continue

        direct_alignment = _read_direct_paragraph_alignment(para_element)
        if direct_alignment == "center":
            continue

        if direct_alignment == "" and _style_already_centres_heading(style_name) is True:
            continue

        p_pr = para_element.find(f"{WQ}pPr")
        if p_pr is None:
            p_pr = etree.Element(f"{WQ}pPr")
            para_element.insert(0, p_pr)

        jc = _find_or_make_child(p_pr, f"{WQ}jc")
        jc.set(f"{WQ}val", "center")
        changed = True

    if changed is False:
        return False

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


def _ensure_ppr_child(p_pr, tag):
    child = p_pr.find(f"{WQ}{tag}")
    if child is not None:
        val = child.get(f"{WQ}val", "")
        if val.lower() in ("0", "false", "off"):
            child.attrib.pop(f"{WQ}val", None)
            return True
        return False
    child = etree.Element(f"{WQ}{tag}")
    change = p_pr.find(f"{WQ}pPrChange")
    p_pr.insert(list(p_pr).index(change), child) if change is not None else p_pr.append(child)
    return True


def _ensure_keep_next(para_element, keep_lines=False):
    p_pr = para_element.find(f"{WQ}pPr")
    if p_pr is None:
        p_pr = etree.Element(f"{WQ}pPr")
        para_element.insert(0, p_pr)
    changed = _ensure_ppr_child(p_pr, "keepNext")
    if keep_lines is True:
        changed = _ensure_ppr_child(p_pr, "keepLines") or changed
    return changed


def _apply_heading_keep_next(input_path, output_path):
    doc_root, body_paras = _load_body_paras_raw(input_path)
    changed = False

    for i, para_element in enumerate(body_paras):
        if _read_raw_paragraph_style_name(para_element) not in HEADING_KEEP_NEXT_STYLE_IDS:
            continue
        changed = _ensure_keep_next(para_element, True) or changed
        j = i + 1
        while j < len(body_paras) and _raw_paragraph_text(body_paras[j]) == "":
            changed = _ensure_keep_next(body_paras[j]) or changed
            j += 1

    if changed is False:
        return False

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


# Styles that belong to non-body roles — do not "fix" these to Normal.
_NON_BODY_STYLES = frozenset({
    "Normal",
    "Heading 1", "Heading 2", "Heading 3",
    "Heading1", "Heading2", "Heading3",
    TITLE_REQUIRED_STYLE, AUTHORS_REQUIRED_STYLE, AFFILIATIONS_REQUIRED_STYLE,
    TITLE_REQUIRED_STYLE_ID, AUTHORS_REQUIRED_STYLE_ID, AFFILIATIONS_REQUIRED_STYLE_ID,
    ABSTRACT_REQUIRED_HEADING_STYLE, ABSTRACT_REQUIRED_PARAGRAPH_STYLE,
    ABSTRACT_REQUIRED_HEADING_STYLE_ID, ABSTRACT_REQUIRED_PARAGRAPH_STYLE_ID,
    PRACTITIONER_REQUIRED_HEADING_STYLE, PRACTITIONER_REQUIRED_NOTE_STYLE,
    PRACTITIONER_REQUIRED_HEADING_STYLE_ID, PRACTITIONER_REQUIRED_NOTE_STYLE_ID,
    KEYWORDS_REQUIRED_HEADING_STYLE, KEYWORDS_REQUIRED_TEXT_STYLE,
    KEYWORDS_REQUIRED_HEADING_STYLE_ID, KEYWORDS_REQUIRED_TEXT_STYLE_ID,
    CITATION_REQUIRED_HEADING_STYLE, CITATION_REQUIRED_TEXT_STYLE,
    CITATION_REQUIRED_HEADING_STYLE_ID, CITATION_REQUIRED_TEXT_STYLE_ID,
    "Figure Number", "Figure Title",
    "FigureNumber", "FigureTitle",
    "Figure/Table Number", "Figure/Table Title",
    "FigureTableNumber", "FigureTableTitle",
    FIGURE_NUMBER_STYLE, FIGURE_TITLE_STYLE,
    FIGURE_NUMBER_STYLE_ID, FIGURE_TITLE_STYLE_ID,
    "Table Number", "Table Title", "Table Text", "Table Emphasis", "Table Note",
    "TableNumber", "TableTitle", "TableText", "TableEmphasis", "TableNote",
    REFERENCE_ENTRY_REQUIRED_STYLE,
    REFERENCE_ENTRY_REQUIRED_STYLE_ID,
    "Quote", QUOTE_STYLE_ID,
    "CommentText",
    "Caption",
})

_MAX_BODY_STYLE_FIXES = 20
_MAX_PSEUDO_HEADING_CHARS = 120


def _run_has_direct_bold(run_element):
    r_pr = run_element.find(f"{WQ}rPr")
    if r_pr is None:
        return False
    bold = r_pr.find(f"{WQ}b")
    if bold is None:
        return False
    val = bold.get(f"{WQ}val")
    return val is None or val.lower() not in ("0", "false", "off")


def _raw_paragraph_is_pseudo_heading(para_element, text, style_name):
    if style_name not in ("", "Normal") and (
        style_name in _NON_BODY_STYLES or style_name.startswith("Heading")
    ):
        return False
    if text == "" or len(text) > _MAX_PSEUDO_HEADING_CHARS:
        return False
    if text.endswith((".", "!", "?")):
        return False
    if (
        DOT_POINT_PATTERN.match(text)
        or _looks_like_table_number(text)
        or _looks_like_figure_number(text)
    ):
        return False

    runs = []
    for run in para_element.iter(f"{WQ}r"):
        run_text = "".join(t.text or "" for t in run.findall(f"{WQ}t"))
        if run_text.strip() != "":
            runs.append(run)

    return bool(runs) and all(_run_has_direct_bold(run) for run in runs)


def _required_body_heading_style(text):
    normalised = _normalise_heading_text(text)
    normalised = normalised.rstrip(":.")
    heading_names = [
        "introduction",
        "literature",
        "literature review",
        "method",
        "methods",
        "results",
        "discussion",
        "acknowledgements",
        "acknowledgments",
        "conclusion",
        "conclusions",
        "concusion",
        "concusions",
    ]

    for heading_name in heading_names:
        if normalised == heading_name:
            return HEADING_1_STYLE_ID

    for canonical in METHOD_REQUIRED_SUBHEADINGS + DISCUSSION_REQUIRED_SUBHEADINGS:
        if normalised in _all_alias_norms(canonical):
            return HEADING_2_STYLE_ID

    return ""


_TABLE_LABEL_RE = r"(?:\d+(?:\.\d+)?[a-z]?|[a-z]\d+[a-z]?)"


def _looks_like_table_number(text):
    normalised = _normalise_heading_text(text)
    match = re.match(rf"^table\s+{_TABLE_LABEL_RE}[.:]?$", normalised)
    if match is not None:
        return True
    return False


def _looks_like_figure_number(text):
    normalised = _normalise_heading_text(text)
    match = re.match(rf"^figure\s+{_TABLE_LABEL_RE}[.:]?$", normalised)
    if match is not None:
        return True
    return False


def _has_long_quote(text):
    for segment in _extract_quoted_segments(text):
        if _count_title_words(segment) > INLINE_QUOTE_MAX_WORDS:
            return True
    return False


def _quote_tail_ok(text):
    tail = text.strip().strip(" .,:;")
    return tail == "" or re.fullmatch(r"\([^)]+\)", tail) is not None


def _is_standalone_long_quote(text):
    clean = text.strip()
    for match in QUOTE_SEGMENT_PATTERN.finditer(clean):
        if _count_title_words(match.group(1)) <= INLINE_QUOTE_MAX_WORDS:
            continue
        if clean[:match.start()].strip() == "" and _quote_tail_ok(clean[match.end():]):
            return True
    return False


def _needs_quote_style(text, style_name):
    return style_name not in ("Quote", QUOTE_STYLE_ID) and _is_standalone_long_quote(text)


def _remove_bold_italic_from_normal_body_paragraphs(input_path, output_path):
    doc_root, body_paras = _load_body_paras_raw(input_path)
    intro_idx = None
    refs_idx = None

    for i, para in enumerate(body_paras):
        style = _read_raw_paragraph_style_name(para)
        text = _normalise_heading_text(_raw_paragraph_text(para)).rstrip(":.")
        if text in {"references", "reference list"} and refs_idx is None:
            refs_idx = i
        if style in {"Heading1", "Heading 1"}:
            if text.startswith("introduction") and intro_idx is None:
                intro_idx = i

    if intro_idx is None:
        return False

    changed = False
    for i, para in enumerate(body_paras):
        if i <= intro_idx:
            continue
        if refs_idx is not None and i >= refs_idx:
            continue
        if _read_raw_paragraph_style_name(para) not in {"", "Normal"}:
            continue

        for run in para.iter(f"{WQ}r"):
            if run.find(f"{WQ}t") is None and run.find(f"{WQ}delText") is None:
                continue
            r_pr = run.find(f"{WQ}rPr")
            if r_pr is None:
                continue
            for tag in ("b", "bCs", "i", "iCs"):
                el = r_pr.find(f"{WQ}{tag}")
                if el is not None:
                    r_pr.remove(el)
                    changed = True

    if changed is False:
        return False

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(input_path, output_path, new_doc_xml)
    return True


def _set_border(border_element, val, size="4"):
    border_element.set(f"{WQ}val", val)
    if val == "single":
        border_element.set(f"{WQ}sz", size)
    elif val == "none":
        border_element.set(f"{WQ}sz", "0")
    elif f"{WQ}sz" in border_element.attrib:
        del border_element.attrib[f"{WQ}sz"]
    border_element.set(f"{WQ}space", "0")
    border_element.set(f"{WQ}color", "auto")


def _remove_child(parent, child_tag):
    child = parent.find(child_tag)
    if child is not None:
        parent.remove(child)


def _remove_table_cell_shading(cell_element):
    changed = False
    tc_pr = cell_element.find(f"{WQ}tcPr")
    if tc_pr is not None and tc_pr.find(f"{WQ}shd") is not None:
        _remove_child(tc_pr, f"{WQ}shd")
        changed = True
    for paragraph in cell_element.findall(f"{WQ}p"):
        p_pr = paragraph.find(f"{WQ}pPr")
        if p_pr is not None and p_pr.find(f"{WQ}shd") is not None:
            _remove_child(p_pr, f"{WQ}shd")
            changed = True
    return changed


def _run_has_table_text(run_element):
    return run_element.find(f"{WQ}t") is not None or run_element.find(f"{WQ}delText") is not None


def _table_can_take_apa_borders(table_element, rows):
    has_merges = (
        table_element.find(f".//{WQ}gridSpan") is not None
        or table_element.find(f".//{WQ}vMerge") is not None
    )
    return has_merges is False


def _apply_tracked_table_properties(table_element, change_id):
    tbl_pr = table_element.find(f"{WQ}tblPr")

    if tbl_pr is None:
        tbl_pr = etree.Element(f"{WQ}tblPr")
        table_element.insert(0, tbl_pr)

    _remove_tracked_property_change(tbl_pr, f"{WQ}tblPrChange")

    tbl_style = tbl_pr.find(f"{WQ}tblStyle")
    if tbl_style is None:
        tbl_style = etree.Element(f"{WQ}tblStyle")
        tbl_pr.insert(0, tbl_style)
    tbl_style.set(f"{WQ}val", TABLE_GRID_STYLE_ID)

    borders = _find_or_make_child(tbl_pr, f"{WQ}tblBorders")
    for tag in ("top", "bottom"):
        _set_border(_find_or_make_child(borders, f"{WQ}{tag}"), "single")
    for tag in ("left", "right", "insideH", "insideV"):
        _set_border(_find_or_make_child(borders, f"{WQ}{tag}"), "none", "0")

    tbl_look = _find_or_make_child(tbl_pr, f"{WQ}tblLook")
    tbl_look.set(f"{WQ}firstRow", "1")
    tbl_look.set(f"{WQ}firstColumn", "0")
    tbl_look.set(f"{WQ}noHBand", "1")
    tbl_look.set(f"{WQ}noVBand", "1")


def _apply_tracked_table_cell_borders(cell_element, row_index, last_row_index, change_id):
    tc_pr = cell_element.find(f"{WQ}tcPr")

    if tc_pr is None:
        tc_pr = etree.Element(f"{WQ}tcPr")
        cell_element.insert(0, tc_pr)

    _remove_tracked_property_change(tc_pr, f"{WQ}tcPrChange")

    borders = _find_or_make_child(tc_pr, f"{WQ}tcBorders")
    for tag in ("top", "bottom", "left", "right", "insideH", "insideV"):
        _remove_child(borders, f"{WQ}{tag}")
    if row_index == 0:
        _set_border(_find_or_make_child(borders, f"{WQ}bottom"), "single")
    elif row_index == last_row_index:
        _set_border(_find_or_make_child(borders, f"{WQ}top"), "nil")
    elif row_index == 1:
        _set_border(_find_or_make_child(borders, f"{WQ}bottom"), "nil")
    else:
        _set_border(_find_or_make_child(borders, f"{WQ}top"), "nil")
        _set_border(_find_or_make_child(borders, f"{WQ}bottom"), "nil")


def _run_bool_value(run_properties, tag):
    el = run_properties.find(f"{WQ}{tag}")
    if el is None:
        return None
    val = el.get(f"{WQ}val")
    if val is None:
        return "1"
    return val


def _table_run_text(run_element):
    parts = []
    for tag in (f"{WQ}t", f"{WQ}delText"):
        for text_element in run_element.findall(tag):
            parts.append(text_element.text or "")
    return "".join(parts)


def _table_run_should_keep_fonts(run_element, run_properties):
    r_fonts = run_properties.find(f"{WQ}rFonts")
    if r_fonts is None:
        return False

    font_values = [
        r_fonts.get(f"{WQ}{name}", "")
        for name in ("ascii", "hAnsi", "cs", "eastAsia")
    ]
    has_symbol_font = any(
        "symbol" in font.lower() or "wingdings" in font.lower() or "webdings" in font.lower()
        for font in font_values
    )
    if has_symbol_font is False:
        return False

    text = _table_run_text(run_element).strip()
    return text != "" and not any(char.isalnum() for char in text)


def _apply_tracked_table_run_format(run_element, change_id):
    if _run_has_table_text(run_element) is False:
        return False

    run_properties = run_element.find(f"{WQ}rPr")
    if run_properties is None:
        return False

    remove_tags = ["b", "bCs", "i", "iCs", "sz", "szCs"]
    if _table_run_should_keep_fonts(run_element, run_properties) is False:
        remove_tags.append("rFonts")

    present_tags = [tag for tag in remove_tags if run_properties.find(f"{WQ}{tag}") is not None]
    had_change = run_properties.find(f"{WQ}rPrChange") is not None
    _remove_tracked_property_change(run_properties, f"{WQ}rPrChange")
    if not present_tags:
        return had_change

    for tag in present_tags:
        _remove_child(run_properties, f"{WQ}{tag}")

    return True


def _apply_table_paragraph_style(paragraph, paragraph_style):
    p_pr = paragraph.find(f"{WQ}pPr")
    if p_pr is None:
        p_pr = etree.Element(f"{WQ}pPr")
        paragraph.insert(0, p_pr)
    _remove_tracked_property_change(p_pr, f"{WQ}pPrChange")
    p_style = p_pr.find(f"{WQ}pStyle")
    if p_style is None:
        p_style = etree.Element(f"{WQ}pStyle")
        p_pr.insert(0, p_style)
    p_style.set(f"{WQ}val", paragraph_style)
    _remove_child(p_pr, f"{WQ}jc")
    _remove_child(p_pr, f"{WQ}spacing")


def _target_table_paragraph_style(emphasis):
    return TABLE_EMPHASIS_STYLE_ID if emphasis else TABLE_TEXT_STYLE_ID


def _table_cell_text(cell_element):
    return " ".join(
        _raw_paragraph_text(paragraph)
        for paragraph in cell_element.findall(f"{WQ}p")
        if _raw_paragraph_text(paragraph) != ""
    ).strip()


def _table_row_is_total(row_element):
    cells = row_element.findall(f"{WQ}tc")
    if not cells:
        return False
    return _table_cell_text(cells[0]).lower().startswith("total")


def _paragraph_is_inside_table(paragraph):
    return any(parent.tag == f"{WQ}tbl" for parent in paragraph.iterancestors())


def _set_paragraph_text_direct(paragraph, text):
    text_nodes = list(paragraph.iter(f"{WQ}t"))
    if not text_nodes:
        run = etree.SubElement(paragraph, f"{WQ}r")
        text_nodes = [etree.SubElement(run, f"{WQ}t")]
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ""


_COMBINED_TABLE_CAPTION_RE = re.compile(
    rf"^table\s+({_TABLE_LABEL_RE})([.:])\s+(.+)$",
    re.IGNORECASE,
)
_COMBINED_FIGURE_CAPTION_RE = re.compile(
    rf"^figure\s+({_TABLE_LABEL_RE})([.:]?)\s+(.+)$",
    re.IGNORECASE,
)
_TABLE_CAPTION_PROSE_START_RE = re.compile(
    r"^(?:presents?|shows?|reports?|lists?|summari[sz]es?|describes?|provides?|contains?|displays?)\b",
    re.IGNORECASE,
)


def _caption_title_looks_like_prose(text):
    stripped = text.strip()
    return (
        stripped != ""
        and stripped[0].islower()
        and _TABLE_CAPTION_PROSE_START_RE.match(stripped) is not None
    )


def _split_combined_table_caption(paragraph, text):
    match = _COMBINED_TABLE_CAPTION_RE.match(text)
    if match is None:
        return False

    parent = paragraph.getparent()
    if parent is None:
        return False
    paragraph_index = list(parent).index(paragraph)
    for sibling in list(parent)[paragraph_index + 1:]:
        if sibling.tag == f"{WQ}tbl":
            break
        if sibling.tag == f"{WQ}p" and _raw_paragraph_text(sibling) == "":
            continue
        return False

    raw_title = match.group(3).strip()
    if _caption_title_looks_like_prose(raw_title):
        return False

    table_number = "Table " + match.group(1) + match.group(2)
    table_title = _to_title_case_title(raw_title)
    _apply_table_paragraph_style(paragraph, TABLE_NUMBER_STYLE_ID)
    _remove_paragraph_content_keep_style(paragraph)
    _append_plain_text_run(paragraph, table_number)

    title_paragraph = _make_plain_styled_paragraph(table_title, TABLE_TITLE_STYLE_ID)
    parent.insert(paragraph_index + 1, title_paragraph)
    return True


def _paragraph_has_drawing(paragraph):
    return (
        paragraph.find(f".//{WQ}drawing") is not None
        or paragraph.find(f".//{WQ}pict") is not None
    )


def _next_nonblank_sibling_has_drawing(paragraph):
    parent = paragraph.getparent()
    if parent is None:
        return False
    paragraph_index = list(parent).index(paragraph)
    for sibling in list(parent)[paragraph_index + 1:]:
        if sibling.tag != f"{WQ}p":
            return False
        if _raw_paragraph_text(sibling) == "" and not _paragraph_has_drawing(sibling):
            continue
        return _paragraph_has_drawing(sibling)
    return False


def _split_combined_figure_caption(paragraph, text):
    match = _COMBINED_FIGURE_CAPTION_RE.match(text)
    if match is None or not _next_nonblank_sibling_has_drawing(paragraph):
        return False

    raw_title = match.group(3).strip()
    if _caption_title_looks_like_prose(raw_title):
        return False

    figure_number = "Figure " + match.group(1)
    figure_title = _to_title_case_title(raw_title)
    parent = paragraph.getparent()
    paragraph_index = list(parent).index(paragraph)
    _apply_table_paragraph_style(paragraph, FIGURE_NUMBER_STYLE_ID)
    _remove_paragraph_content_keep_style(paragraph)
    _append_plain_text_run(paragraph, figure_number)

    title_paragraph = _make_plain_styled_paragraph(figure_title, FIGURE_TITLE_STYLE_ID)
    parent.insert(paragraph_index + 1, title_paragraph)
    return True


def _apply_table_caption_formatting(body_paras_raw):
    changed = False
    next_table_title = False
    next_table_note = False

    for paragraph in body_paras_raw:
        if _paragraph_is_inside_table(paragraph) is True:
            continue
        text = _raw_paragraph_text(paragraph)
        if text == "":
            continue

        required_style = ""
        replacement = None
        if _split_combined_table_caption(paragraph, text) is True:
            next_table_title = False
            next_table_note = True
            changed = True
            continue
        elif _looks_like_table_number(text) is True:
            required_style = TABLE_NUMBER_STYLE_ID
            next_table_title = True
            next_table_note = False
        elif next_table_title is True and text.lower().startswith("note."):
            required_style = TABLE_NOTE_STYLE_ID
            next_table_title = False
            next_table_note = False
        elif next_table_title is True:
            required_style = TABLE_TITLE_STYLE_ID
            replacement = _to_title_case_title(text)
            next_table_title = False
            next_table_note = True
        elif next_table_note is True and text.lower().startswith("note."):
            required_style = TABLE_NOTE_STYLE_ID
            next_table_note = False

        if required_style == "":
            continue

        p_pr = paragraph.find(f"{WQ}pPr")
        old_style = _read_raw_paragraph_style_name(paragraph)
        had_change = p_pr is not None and p_pr.find(f"{WQ}pPrChange") is not None
        had_jc = p_pr is not None and p_pr.find(f"{WQ}jc") is not None
        if old_style != required_style or had_change is True or had_jc is True:
            _apply_table_paragraph_style(paragraph, required_style)
            changed = True
        if replacement is not None and replacement != text:
            _set_paragraph_text_direct(paragraph, replacement)
            changed = True

    return changed


def _apply_tracked_table_formatting(input_path, output_path):
    doc_root, _ = _load_body_paras_raw(input_path)
    if not list(doc_root.iter(f"{WQ}tbl")):
        return False

    current_input_path = input_path
    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        TABLE_TEMPLATE_STYLE_NAMES,
        replace_existing=True,
    )
    if styles_loaded is True:
        current_input_path = output_path

    doc_root, body_paras_raw = _load_body_paras_raw(current_input_path)
    tables = list(doc_root.iter(f"{WQ}tbl"))
    if not tables:
        return styles_loaded

    change_id = 8000
    changed = _apply_table_caption_formatting(body_paras_raw)

    for table in tables:
        rows = table.findall(f"{WQ}tr")
        if not rows:
            continue

        can_restyle_borders = _table_can_take_apa_borders(table, rows)
        if can_restyle_borders is True:
            _apply_tracked_table_properties(table, change_id)
            change_id += 1
            changed = True

        last_row_index = len(rows) - 1
        for row_index, row in enumerate(rows):
            cells = row.findall(f"{WQ}tc")
            row_emphasis = row_index == 0 or _table_row_is_total(row) is True
            for cell in cells:
                if _remove_table_cell_shading(cell) is True:
                    changed = True
                if can_restyle_borders is True:
                    _apply_tracked_table_cell_borders(cell, row_index, last_row_index, change_id)
                    change_id += 1
                    changed = True

                for paragraph in cell.findall(f"{WQ}p"):
                    if _raw_paragraph_text(paragraph) != "":
                        old_style = _read_raw_paragraph_style_name(paragraph)
                        paragraph_style = _target_table_paragraph_style(row_emphasis)
                        p_pr = paragraph.find(f"{WQ}pPr")
                        had_change = p_pr is not None and p_pr.find(f"{WQ}pPrChange") is not None
                        had_jc = p_pr is not None and p_pr.find(f"{WQ}jc") is not None
                        had_spacing = p_pr is not None and p_pr.find(f"{WQ}spacing") is not None
                        if old_style != paragraph_style or had_change or had_jc or had_spacing:
                            _apply_table_paragraph_style(paragraph, paragraph_style)
                            change_id += 1
                            changed = True
                    for run in paragraph.iter(f"{WQ}r"):
                        if _apply_tracked_table_run_format(run, change_id) is True:
                            change_id += 1
                            changed = True

    if changed is False:
        return styles_loaded

    new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _save_doc_xml_to_zip(current_input_path, output_path, new_doc_xml)
    return True


def _apply_body_and_reference_style_fixes(input_path, output_path):
    """Apply tracked style changes for body and reference paragraphs:
    - Required body headings, table captions, and figure captions -> tracked change to journal styles
    - Body paragraphs not using Normal → tracked change to Normal
    - Reference paragraphs not using APA 7 Reference List Entry → tracked change to that style
    IDs start at 5000/6000 to avoid clashing with other tracked changes.
    """
    current_input_path = input_path
    styles_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        [
            "Heading 1",
            "Heading 2",
            "Table Number",
            "Table Title",
            FIGURE_NUMBER_STYLE,
            FIGURE_TITLE_STYLE,
            "Quote",
            "Quote Char",
        ],
        replace_existing=True,
    )
    if styles_loaded is True:
        current_input_path = output_path

    # The APA 7 reference style must use justified alignment. Replace any
    # author-supplied definition (which commonly sets w:jc=left) with the
    # canonical template style so every reference entry is justified — the
    # same replace_existing approach already used for Normal and table styles.
    # A plain add (replace_existing=False) would leave a pre-existing
    # left-aligned definition untouched.
    reference_style_loaded = _ensure_template_styles_available(
        current_input_path,
        output_path,
        [REFERENCE_ENTRY_REQUIRED_STYLE],
        replace_existing=True,
    )
    if reference_style_loaded is True:
        current_input_path = output_path
        styles_loaded = True

    # Use python-docx only for metadata (style names, text); raw lxml for XML writes.
    doc = DocxDocument(current_input_path)
    paras = doc.paragraphs
    doc_root, body_paras_raw = _load_body_paras_raw(current_input_path)

    def _raw_style(el):
        pPr = el.find(f"{WQ}pPr")
        if pPr is None:
            return ""
        pStyle = pPr.find(f"{WQ}pStyle")
        return pStyle.get(f"{WQ}val", "") if pStyle is not None else ""

    def _raw_text(el):
        return "".join(t.text or "" for t in el.iter(f"{WQ}t")).strip()

    intro_idx = None
    refs_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip().lower()
        if text.startswith("introduction") and intro_idx is None:
            intro_idx = i
        if text == "references" and refs_idx is None:
            refs_idx = i

    body_fixes = 0
    heading_fixes = 0
    reference_fixes = 0
    changed = False
    quote_style_needed = False
    next_table_title = False
    next_figure_title = False

    for i, raw_el in enumerate(body_paras_raw):
        text = _raw_text(raw_el)
        if not text:
            continue
        sname = _raw_style(raw_el)

        if refs_idx is not None and i > refs_idx:
            if sname == REFERENCE_ENTRY_REQUIRED_STYLE_ID:
                continue
            _apply_tracked_style_change(
                raw_el,
                REFERENCE_ENTRY_REQUIRED_STYLE_ID,
                sname,
                change_id=5000 + reference_fixes,
            )
            reference_fixes += 1
            changed = True
            continue

        required_style = ""
        if refs_idx is None or i < refs_idx:
            heading_style = _required_body_heading_style(text)
            is_body = intro_idx is not None and i > intro_idx
            if _split_combined_figure_caption(raw_el, text) is True:
                next_table_title = False
                next_figure_title = False
                changed = True
                continue
            elif _looks_like_table_number(text) is True:
                required_style = TABLE_NUMBER_STYLE_ID
                next_table_title = True
                next_figure_title = False
            elif _looks_like_figure_number(text) is True:
                required_style = FIGURE_NUMBER_STYLE_ID
                next_figure_title = True
                next_table_title = False
            elif is_body is True and _needs_quote_style(text, sname) is True:
                quote_style_needed = True
                required_style = QUOTE_STYLE_ID
                next_table_title = False
                next_figure_title = False
            elif heading_style != "":
                required_style = heading_style
                next_table_title = False
                next_figure_title = False
            elif next_table_title is True:
                required_style = TABLE_TITLE_STYLE_ID
                next_table_title = False
            elif next_figure_title is True:
                required_style = FIGURE_TITLE_STYLE_ID
                next_figure_title = False
            elif is_body is True and _raw_paragraph_is_pseudo_heading(raw_el, text, sname) is True:
                required_style = HEADING_2_STYLE_ID
                next_table_title = False
                next_figure_title = False
            else:
                required_style = ""

        if required_style != "":
            if sname != required_style:
                _apply_tracked_style_change(
                    raw_el,
                    required_style,
                    sname,
                    change_id=6100 + heading_fixes,
                )
                heading_fixes += 1
                changed = True
            continue

        if intro_idx is not None and i > intro_idx:
            if sname in _NON_BODY_STYLES or sname.startswith("Heading"):
                continue
            if body_fixes >= _MAX_BODY_STYLE_FIXES:
                continue
            _apply_tracked_style_change(raw_el, "Normal", sname, change_id=6000 + body_fixes)
            body_fixes += 1
            changed = True

    if quote_style_needed is True:
        quote_style_loaded = _ensure_template_styles_available(
            current_input_path,
            output_path,
            ["Quote", "Quote Char"],
            replace_existing=True,
        )
        if quote_style_loaded is True:
            current_input_path = output_path
            changed = True

    if changed:
        new_doc_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        _save_doc_xml_to_zip(current_input_path, output_path, new_doc_xml)
        return True

    if styles_loaded is True:
        return True

    if os.path.abspath(current_input_path) != os.path.abspath(output_path):
        shutil.copy2(current_input_path, output_path)
    return False


def _tracked_replace_boundary_ok(text, start, end):
    if start < 0:
        return False
    before = text[start - 1] if start > 0 else ""
    after = text[end] if end < len(text) else ""
    starts_with_word = bool(text[start:end]) and re.match(r"\w", text[start]) is not None
    ends_with_word = bool(text[start:end]) and re.match(r"\w", text[end - 1]) is not None
    if starts_with_word and before and re.match(r"\w", before) is not None:
        return False
    if ends_with_word and after and re.match(r"\w", after) is not None:
        return False
    return True


def _find_tracked_replace_match(text, find_text, start_at=0):
    start = text.find(find_text, start_at)
    while start >= 0:
        end = start + len(find_text)
        if _tracked_replace_boundary_ok(text, start, end):
            return start
        start = text.find(find_text, start + 1)
    return -1


def _apply_intra_paragraph_tracked_replace(paragraph_element, find_text, replace_text, change_id):
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    nsmap = {"w": W}

    direct_runs = paragraph_element.findall(f"{WQ}r")
    if not direct_runs:
        return False

    text_parts = []
    paragraph_text_parts = []
    offset = 0
    for run_element in direct_runs:
        for text_element in run_element.findall(f"{WQ}t"):
            run_text = text_element.text or ""
            text_parts.append((run_element, text_element, run_text, offset))
            paragraph_text_parts.append(run_text)
            offset += len(run_text)
    quote_spans = find_quote_spans("".join(paragraph_text_parts))

    for run_element, text_element, run_text, run_offset in text_parts:
        start = _find_tracked_replace_match(run_text, find_text, 0)
        while start >= 0:
            match_start = run_offset + start
            match_end = match_start + len(find_text)
            overlaps_quote = any(
                match_start < quote_end and quote_start < match_end
                for quote_start, quote_end in quote_spans
            )
            if not overlaps_quote:
                break
            start = _find_tracked_replace_match(run_text, find_text, start + 1)
        if start < 0:
            continue

        end = start + len(find_text)
        before_text = run_text[:start]
        after_text = run_text[end:]
        run_properties = run_element.find(f"{WQ}rPr")

        if before_text:
            text_element.text = before_text
            if before_text[0].isspace() or before_text[-1].isspace():
                text_element.set(xml_space, "preserve")
        else:
            text_element.text = ""

        delete_element = etree.Element(f"{WQ}del", nsmap=nsmap)
        delete_element.set(f"{WQ}id", str(change_id))
        delete_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
        delete_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)
        delete_run = etree.SubElement(delete_element, f"{WQ}r")
        if run_properties is not None:
            delete_run.append(copy.deepcopy(run_properties))
        delete_text_element = etree.SubElement(delete_run, f"{WQ}delText")
        delete_text_element.text = find_text
        if find_text[0].isspace() or find_text[-1].isspace():
            delete_text_element.set(xml_space, "preserve")

        insert_element = etree.Element(f"{WQ}ins", nsmap=nsmap)
        insert_element.set(f"{WQ}id", str(change_id + 1))
        insert_element.set(f"{WQ}author", TRACKED_CHANGE_AUTHOR)
        insert_element.set(f"{WQ}date", TRACKED_CHANGE_DATE)
        insert_run = etree.SubElement(insert_element, f"{WQ}r")
        if run_properties is not None:
            insert_run.append(copy.deepcopy(run_properties))
        insert_text_element = etree.SubElement(insert_run, f"{WQ}t")
        insert_text_element.text = replace_text
        if replace_text and (replace_text[0].isspace() or replace_text[-1].isspace()):
            insert_text_element.set(xml_space, "preserve")

        after_run_element = None
        if after_text:
            after_run_element = etree.Element(f"{WQ}r", nsmap=nsmap)
            if run_properties is not None:
                after_run_element.append(copy.deepcopy(run_properties))
            after_text_element = etree.SubElement(after_run_element, f"{WQ}t")
            after_text_element.text = after_text
            if after_text[0].isspace() or after_text[-1].isspace():
                after_text_element.set(xml_space, "preserve")

        parent = paragraph_element
        run_position = list(parent).index(run_element)
        parent.insert(run_position + 1, delete_element)
        parent.insert(run_position + 2, insert_element)
        if after_run_element is not None:
            parent.insert(run_position + 3, after_run_element)

        return True

    return False


def _apply_body_edit_plan(input_path, output_path, body_edit_plan):
    if body_edit_plan.get("action") != "apply_body_edits":
        return False
    edits = body_edit_plan.get("edits") or []
    if not edits:
        return False
    doc = DocxDocument(input_path)
    paragraphs = list(doc.paragraphs)
    edits_by_paragraph = {}
    for edit in edits:
        paragraph_index = edit.get("paragraph_index")
        if paragraph_index is None or paragraph_index < 0 or paragraph_index >= len(paragraphs):
            continue
        edits_by_paragraph.setdefault(paragraph_index, []).append(edit)
    change_id = 1000
    applied_any = False
    for paragraph_index, paragraph_edits in edits_by_paragraph.items():
        paragraph_element = paragraphs[paragraph_index]._element
        for edit in paragraph_edits:
            while True:
                success = _apply_intra_paragraph_tracked_replace(
                    paragraph_element, edit["find"], edit["replace"], change_id,
                )
                if not success:
                    break
                applied_any = True
                change_id += 2
    if not applied_any:
        return False
    doc.save(output_path)
    return True


def _apply_editorial_review_comment_plan(input_path, output_path, plan):
    if plan.get("action") != "add_editorial_review_comments":
        return False
    comment_items = plan.get("comments", [])
    if not comment_items:
        return False
    for i, item in enumerate(comment_items):
        item_anchor = item.get("anchor_pos", 0)
        item_message = (item.get("message") or "").strip()
        if item_message:
            src = input_path if i == 0 else output_path
            _write_single_comment_docx(
                input_path=src, output_path=output_path,
                paragraph_index=item_anchor, message=item_message,
            )
    return True


def _template_heading_1_name(section_name):
    """Return the exact Heading 1 name used by the JUTLP 2026 template model."""
    for heading_name in CANONICAL_STRUCTURE["main_sections"]:
        if _normalise_heading_text(heading_name) == _normalise_heading_text(section_name):
            return heading_name
    return section_name


# Built from the shared SECTION_RENAME_MAP (single source of truth, see
# canonical_jultp_template) so this silent Heading 1 rename covers exactly the
# same wordings as the tracked heading-corrections pass. Keys are normalised
# aliases; values are the canonical JUTLP heading label.
HEADING_1_TEXT_NORMALIZATIONS = {
    alias: _template_heading_1_name(canonical)
    for alias, canonical in SECTION_RENAME_MAP.items()
}


def _canonical_heading_1_text(text):
    clean = _normalise_heading_text(text).rstrip(":.")
    return HEADING_1_TEXT_NORMALIZATIONS.get(clean)


def _apply_heading_1_text_normalization(input_path, output_path):
    """Apply direct Heading 1 text fixes without comments or tracked changes."""
    doc = DocxDocument(input_path)
    changed = False

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name != "Heading 1":
            continue

        replacement = _canonical_heading_1_text(paragraph.text)
        if replacement is None:
            continue
        if paragraph.text.strip() == replacement:
            continue

        paragraph.text = replacement
        changed = True

    if changed is False:
        return False

    doc.save(output_path)
    return True


def _apply_heading_2_title_case(input_path, output_path):
    doc = DocxDocument(input_path)
    changed = False

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name != "Heading 2":
            continue
        replacement = _to_title_case_title(paragraph.text)
        if paragraph.text == replacement:
            continue
        paragraph.text = replacement
        changed = True

    if changed is False:
        return False

    doc.save(output_path)
    return True


#run all section planners and apply all generated edits/comments in sequence
def build_edited_document(input_path, output_path=None):
    auto_output_path = output_path is None
    if output_path is None:
        parent_dir = os.path.dirname(input_path)
        temp_number = 1
        output_path = os.path.join(parent_dir, ".copybot_working.docx")
        while os.path.exists(output_path):
            temp_number += 1
            output_path = os.path.join(parent_dir, f".copybot_working_{temp_number}.docx")

    title_state = titleFound(input_path)
    title_plan = titleFormatCheck(input_path, title_state)

    current_input_path = input_path
    applied_any = False

    title_applied = _apply_title_plan(current_input_path, output_path, title_plan)
    if title_applied is True:
        current_input_path = output_path
        applied_any = True

    abstract_state = abstractFound(current_input_path)
    abstract_plan = abstractFormatCheck(current_input_path, abstract_state)
    abstract_applied = _apply_abstract_plan(current_input_path, output_path, abstract_plan)
    if abstract_applied is True:
        current_input_path = output_path
        applied_any = True

    keywords_state = keywordsFound(current_input_path)
    keywords_plan = keywordsFormatCheck(current_input_path, keywords_state)
    keywords_applied = _apply_keywords_plan(current_input_path, output_path, keywords_plan)
    if keywords_applied is True:
        current_input_path = output_path
        applied_any = True

    citation_state = citationFound(current_input_path)
    citation_plan = citationFormatCheck(current_input_path, citation_state)
    citation_applied = _apply_citation_plan(current_input_path, output_path, citation_plan)
    if citation_applied is True:
        current_input_path = output_path
        applied_any = True

    intro_page_break_applied = _apply_intro_page_break(current_input_path, output_path)
    if intro_page_break_applied is True:
        current_input_path = output_path
        applied_any = True

    heading_1_text_normalization_applied = _apply_heading_1_text_normalization(
        current_input_path, output_path
    )
    if heading_1_text_normalization_applied is True:
        current_input_path = output_path
        applied_any = True

    normal_style_applied = _apply_normal_style_fix(current_input_path, output_path)
    if normal_style_applied is True:
        current_input_path = output_path
        applied_any = True

    body_state = documentBodyFound(current_input_path)
    body_plan = documentBodyFormatCheck(current_input_path, body_state)
    body_applied = _apply_document_body_plan(current_input_path, output_path, body_plan)
    if body_applied is True:
        current_input_path = output_path
        applied_any = True

    author_state = authorFound(current_input_path)
    author_plan = authorFormatCheck(current_input_path, author_state)
    author_applied = _apply_author_plan(current_input_path, output_path, author_plan)
    if author_applied is True:
        current_input_path = output_path
        applied_any = True

    body_edit_plan = build_body_edit_plan(current_input_path)
    body_edit_applied = _apply_body_edit_plan(current_input_path, output_path, body_edit_plan)
    if body_edit_applied is True:
        current_input_path = output_path
        applied_any = True

    narrative_citation_corrections = []
    narrative_citation_applied = False
    _, narrative_citation_corrections = apply_narrative_citation_corrections(
        current_input_path, output_path, 1200
    )
    if narrative_citation_corrections:
        current_input_path = output_path
        narrative_citation_applied = True
        applied_any = True

    editorial_review_plan = build_editorial_review_comment_plan(current_input_path)
    editorial_review_applied = _apply_editorial_review_comment_plan(
        current_input_path, output_path, editorial_review_plan
    )
    if editorial_review_applied is True:
        current_input_path = output_path
        applied_any = True

    style_fixes_applied = _apply_body_and_reference_style_fixes(current_input_path, output_path)
    if style_fixes_applied is True:
        current_input_path = output_path
        applied_any = True

    heading_2_title_case_applied = _apply_heading_2_title_case(
        current_input_path, output_path
    )
    if heading_2_title_case_applied is True:
        current_input_path = output_path
        applied_any = True

    _, heading_corrections = apply_heading_corrections(current_input_path, output_path, 9500)
    if heading_corrections:
        current_input_path = output_path
        applied_any = True

    table_formatting_applied = _apply_tracked_table_formatting(current_input_path, output_path)
    if table_formatting_applied is True:
        current_input_path = output_path
        applied_any = True

    practitioner_stub_applied = _apply_missing_practitioner_stub(
        current_input_path, output_path, abstract_plan
    )
    if practitioner_stub_applied is True:
        current_input_path = output_path
        applied_any = True

    keywords_stub_applied = _apply_missing_keywords_stub(
        current_input_path, output_path, keywords_plan
    )
    if keywords_stub_applied is True:
        current_input_path = output_path
        applied_any = True

    heading_alignment_applied = _apply_heading_center_alignment(current_input_path, output_path)
    if heading_alignment_applied is True:
        current_input_path = output_path
        applied_any = True

    heading_keep_next_applied = _apply_heading_keep_next(current_input_path, output_path)
    if heading_keep_next_applied is True:
        current_input_path = output_path
        applied_any = True

    bold_italic_removed = _remove_bold_italic_from_normal_body_paragraphs(
        current_input_path,
        output_path,
    )
    if bold_italic_removed is True:
        current_input_path = output_path
        applied_any = True

    front_page_asset_plan = build_front_page_asset_check_plan(current_input_path)
    front_page_asset_applied = _apply_front_page_asset_plan(
        current_input_path,
        output_path,
        front_page_asset_plan,
    )
    if front_page_asset_applied is True:
        current_input_path = output_path
        applied_any = True

    front_page_textbox_normalised = _normalise_front_page_textbox_layout(
        current_input_path,
        output_path,
    )
    if front_page_textbox_normalised is True:
        current_input_path = output_path
        applied_any = True

    _, run_font_actions = apply_run_font_corrections(current_input_path, output_path, 9000)
    if run_font_actions:
        current_input_path = output_path
        applied_any = True

    if applied_any is False:
        shutil.copy2(input_path, output_path)

    _renumber_author_queries_by_anchor_order(output_path)
    _remove_line_numbers(output_path)

    if auto_output_path is True:
        parent_dir = os.path.dirname(output_path)
        output_name = build_output_filename_from_author_line(
            input_path,
            parent_dir,
            author_plan.get("corrected_authors_line", ""),
            ignore_path=output_path,
        )
        final_output_path = os.path.join(parent_dir, output_name)
        os.replace(output_path, final_output_path)
        output_path = final_output_path

    return {
        "output_path": output_path,
        "plan": {
            "title": title_plan,
            "author": author_plan,
            "abstract": abstract_plan,
            "keywords": keywords_plan,
            "citation": citation_plan,
            "document_body": body_plan,
            "body_edit": body_edit_plan,
            "narrative_citation": {
                "action": (
                    "apply_narrative_citation_corrections"
                    if narrative_citation_corrections
                    else "none"
                ),
                "corrections": narrative_citation_corrections,
            },
            "editorial_review": editorial_review_plan,
            "front_page_assets": front_page_asset_plan,
            "heading_corrections": heading_corrections,
        },
        "title_plan": title_plan,
        "author_plan": author_plan,
        "abstract_plan": abstract_plan,
        "keywords_plan": keywords_plan,
        "citation_plan": citation_plan,
        "document_body_plan": body_plan,
        "body_edit_plan": body_edit_plan,
        "narrative_citation_corrections": narrative_citation_corrections,
        "editorial_review_plan": editorial_review_plan,
        "front_page_asset_plan": front_page_asset_plan,
        "intro_page_break_applied": intro_page_break_applied,
        "heading_1_text_normalization_applied": heading_1_text_normalization_applied,
        "normal_style_applied": normal_style_applied,
        "heading_alignment_applied": heading_alignment_applied,
        "heading_keep_next_applied": heading_keep_next_applied,
        "heading_corrections": heading_corrections,
        "table_formatting_applied": table_formatting_applied,
        "bold_italic_removed": bold_italic_removed,
        "title_applied": title_applied,
        "author_applied": author_applied,
        "abstract_applied": abstract_applied,
        "keywords_applied": keywords_applied,
        "citation_applied": citation_applied,
        "document_body_applied": body_applied,
        "body_edit_applied": body_edit_applied,
        "narrative_citation_applied": narrative_citation_applied,
        "editorial_review_applied": editorial_review_applied,
        "practitioner_stub_applied": practitioner_stub_applied,
        "keywords_stub_applied": keywords_stub_applied,
        "front_page_asset_applied": front_page_asset_applied,
        "front_page_textbox_normalised": front_page_textbox_normalised,
        "applied": applied_any,
    }


if __name__ == "__main__":
    import sys

    input_path = "tests/sample_docs_JUTLP/Polly_2025_JUTLP.docx"
    output_path = "tests/generated_outputs/reviewed_Polly_2025_JUTLP_author_tracking.docx"

    if len(sys.argv) > 1:
        input_path = sys.argv[1]
    if len(sys.argv) > 2:
        output_path = sys.argv[2]

    result = build_edited_document(input_path, output_path)
    print(result["plan"])
    print(result["output_path"])
