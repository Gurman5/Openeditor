"""Build a reviewed .docx file with inline Word comments.

This module takes validation findings and writes them back into a copy of the
original document as Word comments.

High-level flow:
1. Collect findings that should become comments.
2. Pick where each comment should attach.
3. Insert Word XML markers (`commentRangeStart`, `commentRangeEnd`,
   `commentReference`) into `word/document.xml`.
4. Write/update `word/comments.xml` plus required relationship/content-type
   entries.
"""

import os
import re
import zipfile
from copy import deepcopy

from docx import Document as DocxDocument
from lxml import etree

from app.services.document_analysis_services import (
    _is_main_section_heading,
    _text_matches_main_section,
    load_paragraphs,
)
from app.services.jutlp_validator import _matches_section, validate
from app.services.timestamps import now_sydney_iso

_SECTION_ALIASES = {
    "methods": "Method",
    "methodology": "Method",
    "literature review": "Literature",
    "findings": "Results",
    "conclusions": "Conclusion",
    "acknowledgments": "Acknowledgements",
}

# WordprocessingML namespace used for manual XML edits inside .docx files.
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"

COMMENT_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
COMMENT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)

AUTHOR = "CopyEditor AI"
INITIALS = "AI"
DATE = now_sydney_iso()

PASS_MARK = "☑"
ISSUE_MARK = "⚠"


REFERENCE_ISSUE_SUMMARY_PREFIX = "Reference list issue summary:"


def _reference_entry_number(rule_id: str, prefixes: tuple[str, ...]) -> int | None:
    prefix_group = "|".join(prefixes)
    match = re.match(rf"(?:{prefix_group})(\d+)", rule_id)
    if match is None:
        return None
    return int(match.group(1))


def _reference_issue_summary_comment(ref_results: list[dict]) -> str | None:
    citation_warning_count = 0
    unverified_entries: set[int] = set()
    reference_format_count = 0
    other_reference_issue_count = 0

    for result in ref_results:
        status = str(result.get("status", "")).lower()
        if status not in {"fail", "warn"}:
            continue

        rule_id = str(result.get("rule_id", ""))
        if rule_id.startswith("CONS"):
            citation_warning_count += 1
            continue

        if status == "fail" and rule_id.startswith(("CREF", "HREF", "DOIT")):
            entry_num = _reference_entry_number(rule_id, ("CREF", "HREF", "DOIT"))
            if entry_num is not None:
                unverified_entries.add(entry_num)
            continue

        if rule_id.startswith("REFE"):
            reference_format_count += 1
            continue

        if rule_id.startswith("REF"):
            other_reference_issue_count += 1

    issue_parts: list[str] = []
    if citation_warning_count:
        issue_parts.append(
            f"{citation_warning_count} citation-consistency warning"
            f"{'s' if citation_warning_count != 1 else ''}"
        )
    if unverified_entries:
        count = len(unverified_entries)
        issue_parts.append(
            f"{count} reference entr"
            f"{'ies' if count != 1 else 'y'} "
            "could not be automatically verified"
        )
    if reference_format_count:
        issue_parts.append(
            f"{reference_format_count} reference-format issue"
            f"{'s' if reference_format_count != 1 else ''}"
        )
    if other_reference_issue_count:
        issue_parts.append(
            f"{other_reference_issue_count} other reference-list issue"
            f"{'s' if other_reference_issue_count != 1 else ''}"
        )
    if not issue_parts:
        return None

    return (
        REFERENCE_ISSUE_SUMMARY_PREFIX
        + " "
        + "; ".join(issue_parts)
        + ". Please review the reference list before publication."
    )


def _format_summary_line(mark: str, result: dict) -> str:
    rule_id = str(result.get("rule_id", "")).strip() or "UNKNOWN"
    message = str(result.get("message", "")).strip()
    if message:
        return f"{mark} {rule_id}: {message}"
    return f"{mark} {rule_id}"


def _append_validation_summary(docx_doc: DocxDocument, all_results: list[dict]) -> None:
    docx_doc.add_page_break()
    docx_doc.add_heading("CopyEditor AI Validation Summary", level=1)
    passed = [r for r in all_results if str(r.get("status", "")).lower() == "pass"]
    flagged = [r for r in all_results if str(r.get("status", "")).lower() in {"fail", "warn"}]

    docx_doc.add_paragraph(f"{PASS_MARK} Passed checks:")
    if passed:
        for result in passed:
            docx_doc.add_paragraph(_format_summary_line(PASS_MARK, result))
    else:
        docx_doc.add_paragraph(f"{PASS_MARK} None")

    docx_doc.add_paragraph(f"{ISSUE_MARK} Flagged checks:")
    if flagged:
        for result in flagged:
            docx_doc.add_paragraph(_format_summary_line(ISSUE_MARK, result))
    else:
        docx_doc.add_paragraph(f"{ISSUE_MARK} None")


def _format_author_query_text(comment_id: int, text: str) -> str:
    prefix, query_text = _author_query_parts(comment_id, text)
    return prefix + query_text


def _author_query_parts(comment_id: int, text: str) -> tuple[str, str]:
    clean_text = (text or "").strip()
    match = re.match(r"^(Author Query\s+\d+\.\s*)(.*)$", clean_text, flags=re.IGNORECASE)
    if match is not None:
        return match.group(1), match.group(2)
    return f"Author Query {comment_id}. ", clean_text


def _append_author_query_runs(parent: etree._Element, comment_id: int, text: str) -> None:
    prefix, query_text = _author_query_parts(comment_id, text)

    prefix_run = etree.SubElement(parent, f"{WQ}r")
    prefix_rPr = etree.SubElement(prefix_run, f"{WQ}rPr")
    etree.SubElement(prefix_rPr, f"{WQ}b")
    prefix_t = etree.SubElement(prefix_run, f"{WQ}t")
    prefix_t.text = prefix
    prefix_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    text_run = etree.SubElement(parent, f"{WQ}r")
    t = etree.SubElement(text_run, f"{WQ}t")
    t.text = query_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _make_comment_element(comment_id: int, text: str) -> etree._Element:
    """Create one `<w:comment>` node for `word/comments.xml`."""
    nsmap = {"w": W}
    comment = etree.Element(f"{WQ}comment", nsmap=nsmap)
    comment.set(f"{WQ}id", str(comment_id))
    comment.set(f"{WQ}author", AUTHOR)
    comment.set(f"{WQ}date", DATE)
    comment.set(f"{WQ}initials", INITIALS)

    p = etree.SubElement(comment, f"{WQ}p")
    pPr = etree.SubElement(p, f"{WQ}pPr")
    pStyle = etree.SubElement(pPr, f"{WQ}pStyle")
    pStyle.set(f"{WQ}val", "CommentText")

    ref_run = etree.SubElement(p, f"{WQ}r")
    ref_rPr = etree.SubElement(ref_run, f"{WQ}rPr")
    ref_rStyle = etree.SubElement(ref_rPr, f"{WQ}rStyle")
    ref_rStyle.set(f"{WQ}val", "CommentReference")
    etree.SubElement(ref_run, f"{WQ}annotationRef")

    _append_author_query_runs(p, comment_id, text)

    return comment


def _inject_comment_markers(para_element: etree._Element, comment_id: int) -> None:
    """Fallback comment anchor: wrap the whole paragraph.

    We only use this when phrase-level anchoring is not possible.
    """
    nsmap = {"w": W}

    start = etree.Element(f"{WQ}commentRangeStart", nsmap=nsmap)
    start.set(f"{WQ}id", str(comment_id))

    end = etree.Element(f"{WQ}commentRangeEnd", nsmap=nsmap)
    end.set(f"{WQ}id", str(comment_id))

    ref_run = etree.Element(f"{WQ}r", nsmap=nsmap)
    rPr = etree.SubElement(ref_run, f"{WQ}rPr")
    rStyle = etree.SubElement(rPr, f"{WQ}rStyle")
    rStyle.set(f"{WQ}val", "CommentReference")
    comment_ref = etree.SubElement(ref_run, f"{WQ}commentReference")
    comment_ref.set(f"{WQ}id", str(comment_id))

    para_element.insert(0, start)
    para_element.append(end)
    para_element.append(ref_run)


def _build_comment_ref_run(comment_id: int) -> etree._Element:
    """Build `<w:r><w:commentReference .../></w:r>` for the inline comment mark."""
    nsmap = {"w": W}
    ref_run = etree.Element(f"{WQ}r", nsmap=nsmap)
    r_pr = etree.SubElement(ref_run, f"{WQ}rPr")
    r_style = etree.SubElement(r_pr, f"{WQ}rStyle")
    r_style.set(f"{WQ}val", "CommentReference")
    comment_ref = etree.SubElement(ref_run, f"{WQ}commentReference")
    comment_ref.set(f"{WQ}id", str(comment_id))
    return ref_run


def _run_text(run_element: etree._Element) -> str:
    """Return the full visible text for a single Word run (`w:r`)."""
    return "".join((t.text or "") for t in run_element.findall(f"{WQ}t"))


def _set_run_text(run_element: etree._Element, text: str) -> None:
    """Replace all text nodes in a run with one new text node."""
    for child in list(run_element):
        if child.tag == f"{WQ}t":
            run_element.remove(child)
    t = etree.SubElement(run_element, f"{WQ}t")
    t.text = text
    if text and (text[0].isspace() or text[-1].isspace()):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _split_run(run_element: etree._Element, offset: int) -> etree._Element:
    """Split one run at `offset` and return the new right-side run.

    Example:
    - input run text: "Discussion"
    - offset: 4
    - result: current run = "Disc", new run = "ussion"
    """
    text = _run_text(run_element)
    if offset <= 0 or offset >= len(text):
        return run_element

    left = text[:offset]
    right = text[offset:]
    _set_run_text(run_element, left)

    new_run = etree.Element(f"{WQ}r", nsmap=run_element.nsmap)
    r_pr = run_element.find(f"{WQ}rPr")
    if r_pr is not None:
        new_run.append(deepcopy(r_pr))
    _set_run_text(new_run, right)

    parent = run_element.getparent()
    insert_at = list(parent).index(run_element) + 1
    parent.insert(insert_at, new_run)
    return new_run


def _find_span(full_text: str, target_phrase: str | None) -> tuple[int, int] | None:
    """Find `[start, end)` character offsets to anchor a comment in one paragraph.

    If a target phrase is provided, we anchor to that phrase.
    Otherwise, we anchor to the first visible token as a safe fallback.
    """
    if not full_text:
        return None

    if target_phrase:
        start = full_text.lower().find(target_phrase.lower())
        if start >= 0:
            return start, start + len(target_phrase)

    # Fallback: comment just the first token instead of the full paragraph.
    match = re.search(r"\S+", full_text)
    if not match:
        return None
    return match.start(), match.end()


def _infer_target_phrase(result: dict, para_records: list, para_idx: int) -> str | None:
    """Guess which phrase inside a paragraph should receive the comment.

    This gives phrase-level anchors better precision than paragraph-level
    anchoring.
    """
    if result.get("source") == "llm":
        # Use the LLM's verbatim quote first — most precise anchor available.
        quote = (result.get("quote") or "").strip()
        if quote:
            return quote
        # No quote: return None so _find_span falls back to first token of
        # the body paragraph (acceptable — better than anchoring the heading).
        return None

    rule_id = result["rule_id"]
    if rule_id in _SEC_RULE_TO_SECTION and rule_id != "SEC009":
        return _SEC_RULE_TO_SECTION[rule_id]
    if rule_id.startswith("MET"):
        match = re.match(r"MET(\d+)", rule_id)
        if match:
            pos = int(match.group(1)) - 1
            if 0 <= pos < len(_METHOD_SUBSECTIONS):
                return _METHOD_SUBSECTIONS[pos]
        return "Method"
    if rule_id.startswith("DIS"):
        match = re.match(r"DIS(\d+)", rule_id)
        if match:
            pos = int(match.group(1)) - 1
            if 0 <= pos < len(_DISCUSSION_SUBSECTIONS):
                return _DISCUSSION_SUBSECTIONS[pos]
        return "Discussion"
    if rule_id == "FP001":
        return "Title"
    if rule_id == "FP002":
        return "Authors"
    if rule_id == "FP003":
        return "Affiliations"
    if rule_id == "FP004":
        return "Abstract"
    if rule_id in ("FP007", "FP008"):
        return "Abstract"
    if rule_id in {"FP005", "FP009"}:
        return "Practitioner Notes"
    if rule_id in {"FP006", "FP010"}:
        return "Keywords"
    if rule_id == "STY001":
        return "Guidance Notes"
    if rule_id.startswith("ANDOR") or rule_id.startswith("ETC") or rule_id.startswith("DOTPT"):
        phrase = result.get("anchor_phrase")
        if phrase:
            return phrase
        return None
    if rule_id == "AFF001":
        phrase = result.get("anchor_phrase")
        if phrase:
            return phrase
        return None
    if (
        rule_id.startswith("REF")
        or rule_id.startswith("CREF")
        or rule_id.startswith("HREF")
        or rule_id.startswith("DOIT")
    ):
        if 0 <= para_idx < len(para_records):
            return para_records[para_idx].text
    return None


def _inject_comment_markers_for_phrase(
    para_element: etree._Element,
    comment_id: int,
    target_phrase: str | None,
) -> None:
    """Anchor a comment to a specific phrase inside one paragraph.

    How it works:
    1. Gather all text runs in the paragraph.
    2. Build one combined string so we can locate phrase offsets.
    3. Map offsets back to concrete run elements.
    4. Split runs if phrase boundaries are inside a run.
    5. Insert start/end/reference markers around the matched run range.
    """
    runs = []
    for node in list(para_element):
        if node.tag != f"{WQ}r":
            continue
        txt = _run_text(node)
        if txt:
            runs.append((node, txt))

    if not runs:
        # No text runs -> cannot anchor by phrase.
        _inject_comment_markers(para_element, comment_id)
        return

    full_text = "".join(txt for _, txt in runs)
    span = _find_span(full_text, target_phrase)
    if span is None:
        # No phrase match -> fallback to paragraph-level anchor.
        _inject_comment_markers(para_element, comment_id)
        return
    start_char, end_char = span

    cursor = 0
    start_run = None
    start_off = 0
    end_run = None
    end_off = 0
    for run_el, txt in runs:
        # Move across the paragraph text and find run boundaries for span start/end.
        nxt = cursor + len(txt)
        if start_run is None and start_char < nxt:
            start_run = run_el
            start_off = start_char - cursor
        if end_run is None and end_char <= nxt:
            end_run = run_el
            end_off = end_char - cursor
            break
        cursor = nxt

    if start_run is None or end_run is None:
        # Defensive fallback for unexpected offset mapping issues.
        _inject_comment_markers(para_element, comment_id)
        return

    if start_run is end_run:
        # Span is inside one run. Split at end first, then at start.
        run_len = len(_run_text(start_run))
        if end_off < run_len:
            _split_run(start_run, end_off)
        match_start_run = start_run if start_off <= 0 else _split_run(start_run, start_off)
        match_end_run = match_start_run
    else:
        # Span crosses multiple runs. Split start/end boundary runs as needed.
        end_len = len(_run_text(end_run))
        if end_off < end_len:
            _split_run(end_run, end_off)
        match_start_run = start_run if start_off <= 0 else _split_run(start_run, start_off)
        match_end_run = end_run

    nsmap = {"w": W}
    start = etree.Element(f"{WQ}commentRangeStart", nsmap=nsmap)
    start.set(f"{WQ}id", str(comment_id))
    end = etree.Element(f"{WQ}commentRangeEnd", nsmap=nsmap)
    end.set(f"{WQ}id", str(comment_id))
    ref_run = _build_comment_ref_run(comment_id)

    start_at = list(para_element).index(match_start_run)
    para_element.insert(start_at, start)
    end_at = list(para_element).index(match_end_run)
    para_element.insert(end_at + 1, end)
    ref_at = list(para_element).index(end) + 1
    para_element.insert(ref_at, ref_run)


_CANONICAL_ORDER = [
    "Introduction", "Literature", "Method", "Results",
    "Discussion", "Conclusion", "Acknowledgements", "References",
]

_SEC_RULE_TO_SECTION = {f"SEC{i + 1:03d}": s for i, s in enumerate(_CANONICAL_ORDER)}

_METHOD_SUBSECTIONS = [
    "Research Design",
    "Participants",
    "Measures",
    "Procedure",
    "Analysis",
]

_DISCUSSION_SUBSECTIONS = [
    "Practical Implications",
    "Theoretical Implications",
    "Limitations and Future Research",
]


def _find_heading(para_records: list, text: str) -> int | None:
    """Find a section heading paragraph index by title text.

    A paragraph counts as a heading when it carries the Heading-1 style OR is a
    recognised section heading by text (an author who marked the heading only
    visually, before the restyle pass runs). This mirrors the validator's
    text-based section recognition so anchoring agrees with it.
    """
    for i, p in enumerate(para_records):
        if p.text == text and (p.style == "Heading 1" or _is_main_section_heading(p)):
            return i
    for i, p in enumerate(para_records):
        if text.lower() in p.text.lower() and (p.style == "Heading 1" or _is_main_section_heading(p)):
            return i
    return None


def _find_first_heading(para_records: list) -> int | None:
    """Return the first non-empty Heading 1 paragraph index."""
    for i, p in enumerate(para_records):
        if p.style == "Heading 1" and not p.is_empty:
            return i
    return None


def _nth_heading_after(para_records: list, start_idx: int | None, number: int) -> int | None:
    if start_idx is None:
        return None
    headings = [
        i for i, p in enumerate(para_records)
        if i > start_idx and p.style == "Heading 1" and not p.is_empty
    ]
    if number <= 0 or len(headings) < number:
        return None
    return headings[number - 1]


def _find_first_non_empty_para(para_records: list) -> int:
    """Return the first non-empty paragraph index (fallback anchor)."""
    for i, p in enumerate(para_records):
        if not p.is_empty:
            return i
    return 0


def _find_front_matter_anchor(para_records: list) -> int:
    """Pick a stable anchor in the front matter area.

    Preference:
    1) Article title paragraph.
    2) First non-empty paragraph before Introduction.
    3) First heading in document.
    4) First non-empty paragraph in document.
    """
    intro_idx = _find_heading(para_records, "Introduction")
    upper_bound = intro_idx if intro_idx is not None else len(para_records)

    for i, p in enumerate(para_records[:upper_bound]):
        if p.style == "Article Title" and not p.is_empty:
            return i

    for i, p in enumerate(para_records[:upper_bound]):
        if not p.is_empty:
            return i

    first_heading = _find_first_heading(para_records)
    if first_heading is not None:
        return first_heading

    return _find_first_non_empty_para(para_records)


def _front_matter_end(para_records: list) -> int:
    """Return the paragraph index where front matter ends."""
    intro_idx = _find_heading(para_records, "Introduction")
    return intro_idx if intro_idx is not None else len(para_records)


def _find_front_style(para_records: list, styles: tuple[str, ...]) -> int | None:
    """Find first paragraph in front matter with one of the given styles."""
    end = _front_matter_end(para_records)
    for i, p in enumerate(para_records[:end]):
        if p.style in styles and not p.is_empty:
            return i
    return None


def _find_front_text(para_records: list, text: str) -> int | None:
    """Find front-matter paragraph by exact text, case-insensitive."""
    end = _front_matter_end(para_records)
    needle = text.strip().lower()
    for i, p in enumerate(para_records[:end]):
        if p.text.strip().lower() == needle:
            return i
    return None


def _find_after(para_records: list, start_idx: int, limit: int | None = None) -> int | None:
    """Find first non-empty paragraph after `start_idx` and before `limit`."""
    upper = len(para_records) if limit is None else min(limit, len(para_records))
    for i in range(max(start_idx + 1, 0), upper):
        if not para_records[i].is_empty:
            return i
    return None


def _find_last_front_non_empty(para_records: list) -> int | None:
    """Find the last non-empty front-matter paragraph."""
    end = _front_matter_end(para_records)
    for i in range(end - 1, -1, -1):
        if not para_records[i].is_empty:
            return i
    return None


def _find_front_matter_rule_anchor(rule_id: str, para_records: list) -> int:
    """Resolve a specific front-matter anchor for FP/STY rules."""
    intro_idx = _front_matter_end(para_records)
    title_idx = _find_front_style(para_records, ("Article Title",))
    authors_idx = _find_front_style(para_records, ("Authors",))
    affiliations_idx = _find_front_style(para_records, ("Author Affiliations",))
    abstract_idx = _find_front_text(para_records, "Abstract")
    prac_idx = _find_front_text(para_records, "Practitioner Notes")
    keywords_idx = _find_front_text(para_records, "Keywords")

    if rule_id == "FP001":
        if title_idx is not None:
            return title_idx
        return _find_front_matter_anchor(para_records)

    if rule_id == "FP002":
        if authors_idx is not None:
            return authors_idx
        if title_idx is not None:
            next_idx = _find_after(para_records, title_idx, intro_idx)
            if next_idx is not None:
                return next_idx
            return title_idx
        return _find_front_matter_anchor(para_records)

    if rule_id == "FP003":
        if affiliations_idx is not None:
            return affiliations_idx
        if authors_idx is not None:
            next_idx = _find_after(para_records, authors_idx, intro_idx)
            if next_idx is not None:
                return next_idx
            return authors_idx
        if title_idx is not None:
            first_after = _find_after(para_records, title_idx, intro_idx)
            if first_after is not None:
                second_after = _find_after(para_records, first_after, intro_idx)
                if second_after is not None:
                    return second_after
                return first_after
            return title_idx
        return _find_front_matter_anchor(para_records)

    if rule_id in {"FP004", "FP007", "FP008"}:
        if abstract_idx is not None:
            return abstract_idx
        if title_idx is not None:
            return title_idx
        return _find_front_matter_anchor(para_records)

    if rule_id in {"FP005", "FP009"}:
        if prac_idx is not None:
            return prac_idx
        if keywords_idx is not None:
            return keywords_idx
        if abstract_idx is not None:
            after_abstract = _find_after(para_records, abstract_idx, intro_idx)
            if after_abstract is not None:
                return after_abstract
            return abstract_idx
        return _find_front_matter_anchor(para_records)

    if rule_id in {"FP006", "FP010"}:
        if keywords_idx is not None:
            return keywords_idx
        if prac_idx is not None:
            after_prac = _find_after(para_records, prac_idx, intro_idx)
            if after_prac is not None:
                return after_prac
            return prac_idx
        if abstract_idx is not None:
            last_front = _find_last_front_non_empty(para_records)
            if last_front is not None:
                return last_front
            return abstract_idx
        return _find_front_matter_anchor(para_records)

    if rule_id == "STY001":
        for i, p in enumerate(para_records):
            if p.style == "Guidance Notes":
                return i
        return _find_front_matter_anchor(para_records)

    return _find_front_matter_anchor(para_records)


def _find_missing_section_anchor(section_name: str, para_records: list) -> int:
    """Anchor missing canonical sections near their expected slot in the structure."""
    if section_name not in _CANONICAL_ORDER:
        return _find_preceding_section(section_name, para_records)

    found = _find_heading(para_records, section_name)
    if found is not None:
        return found

    intro_idx = _find_heading(para_records, "Introduction")
    if section_name == "Literature":
        anchor = _nth_heading_after(para_records, intro_idx, 1)
        if anchor is not None:
            return anchor
    if section_name == "Method":
        anchor = _nth_heading_after(para_records, intro_idx, 2)
        if anchor is not None:
            return anchor
    if section_name == "Results":
        method_idx = _find_heading(para_records, "Method")
        anchor = _nth_heading_after(para_records, method_idx, 1)
        if anchor is not None:
            return anchor

    target_idx = _CANONICAL_ORDER.index(section_name)

    prev_canon_idx = None
    prev_para_idx = None
    for i in range(target_idx - 1, -1, -1):
        para_idx = _find_heading(para_records, _CANONICAL_ORDER[i])
        if para_idx is not None:
            prev_canon_idx = i
            prev_para_idx = para_idx
            break

    next_canon_idx = None
    next_para_idx = None
    for i in range(target_idx + 1, len(_CANONICAL_ORDER)):
        para_idx = _find_heading(para_records, _CANONICAL_ORDER[i])
        if para_idx is not None:
            next_canon_idx = i
            next_para_idx = para_idx
            break

    lower = prev_para_idx if prev_para_idx is not None else -1
    upper = next_para_idx if next_para_idx is not None else len(para_records)

    in_between_headings = [
        i
        for i, p in enumerate(para_records)
        if p.style == "Heading 1" and not p.is_empty and lower < i < upper
    ]

    if in_between_headings:
        start_c = prev_canon_idx + 1 if prev_canon_idx is not None else 0
        end_c = next_canon_idx if next_canon_idx is not None else len(_CANONICAL_ORDER)
        missing_sections = [
            s for s in _CANONICAL_ORDER[start_c:end_c]
            if _find_heading(para_records, s) is None
        ]
        if section_name in missing_sections:
            slot = missing_sections.index(section_name)
            return in_between_headings[min(slot, len(in_between_headings) - 1)]
        return in_between_headings[0]

    # If no unlabeled heading candidate exists, prefer whichever neighboring
    # canonical section is closer to the target slot.
    if prev_canon_idx is not None and next_canon_idx is not None:
        dist_prev = target_idx - prev_canon_idx
        dist_next = next_canon_idx - target_idx
        if dist_next < dist_prev:
            return next_para_idx
        return prev_para_idx

    if next_para_idx is not None:
        return next_para_idx
    if prev_para_idx is not None:
        return prev_para_idx

    first_heading = _find_first_heading(para_records)
    if first_heading is not None:
        return first_heading

    return _find_first_non_empty_para(para_records)


def _find_preceding_section(section_name: str, para_records: list) -> int:
    """Fallback anchor when a target section heading is missing.

    We walk backward through canonical section order and use the nearest earlier
    section that exists in the document.
    """
    idx = _CANONICAL_ORDER.index(section_name) if section_name in _CANONICAL_ORDER else -1
    if idx > 0:
        for prev in reversed(_CANONICAL_ORDER[:idx]):
            found = _find_heading(para_records, prev)
            if found is not None:
                return found

    first_heading = _find_first_heading(para_records)
    if first_heading is not None:
        return first_heading

    return _find_first_non_empty_para(para_records)


# ── Missing-section stub insertion ───────────────────────────────────────────
# When a required canonical section is absent (or only present as a combined
# heading like "Results and Discussion"), the validator's SEC*/DIS*/MET* fail
# results would otherwise anchor to the nearest neighbouring heading. Instead
# we insert a tracked Heading-1 stub for the missing section at its correct
# canonical slot and anchor those comments to it.

# Rule families whose owning section, when absent, should get a stub heading.
_SUBSECTION_RULE_SECTION = {"DIS": "Discussion", "MET": "Method"}


def _section_for_rule(rule_id: str) -> str | None:
    """Return the canonical section a fail-rule belongs to, or None.

    SEC001-008 map via _SEC_RULE_TO_SECTION (SEC009 is ordering, skipped).
    DIS* → Discussion, MET* → Method (subsection-missing families).
    """
    if rule_id == "SEC009":
        return None
    if rule_id in _SEC_RULE_TO_SECTION:
        return _SEC_RULE_TO_SECTION[rule_id]
    for prefix, section in _SUBSECTION_RULE_SECTION.items():
        if rule_id.startswith(prefix):
            return section
    return None


def _section_heading_present(section: str, para_records: list) -> bool:
    """True if a Heading-1 paragraph is the section itself (exact / 'name:'
    prefix), NOT merely a substring.

    Uses the validator's alias-aware exact matcher so a combined
    "Results and Discussion" heading counts as Results but NOT as Discussion
    — which is exactly when we want to insert a separate Discussion stub.
    """
    from app.domain.canonical_jultp_template import CANONICAL_STRUCTURE
    aliases = CANONICAL_STRUCTURE["section_aliases"].get(section, [])
    # Discussion has no aliases, so "Results and Discussion" won't match it.
    # Accept a recognised unstyled section heading (not just Heading-1 style):
    # the restyle pass applies Heading 1 later, so before it runs a present
    # section would otherwise look absent and get a spurious duplicate stub.
    for p in para_records:
        if p.is_empty:
            continue
        if not _matches_section(p.text, section, aliases):
            continue
        if p.style == "Heading 1" or _is_main_section_heading(p):
            return True
    return False


def _resolve_h1_style_id(all_paras: list) -> str:
    """Return the pStyle val used by existing Heading-1 paragraphs.

    Documents vary: "Heading 1" vs "Heading1". Read it from the first H1 in
    the body so the inserted stub matches; fall back to "Heading 1".
    """
    for p in all_paras:
        pPr = p.find(f"{WQ}pPr")
        if pPr is None:
            continue
        pStyle = pPr.find(f"{WQ}pStyle")
        if pStyle is None:
            continue
        val = pStyle.get(f"{WQ}val", "")
        if val.replace(" ", "").lower() == "heading1":
            return val
    return "Heading 1"


def _make_inserted_heading_paragraph(section: str, style_id: str, change_id: int) -> etree._Element:
    """Build a tracked-insertion Heading-1 stub paragraph for ``section``."""
    nsmap = {"w": W}
    p = etree.Element(f"{WQ}p", nsmap=nsmap)
    pPr = etree.SubElement(p, f"{WQ}pPr")
    pStyle = etree.SubElement(pPr, f"{WQ}pStyle")
    pStyle.set(f"{WQ}val", style_id)

    ins = etree.SubElement(p, f"{WQ}ins")
    ins.set(f"{WQ}id", str(change_id))
    ins.set(f"{WQ}author", AUTHOR)
    ins.set(f"{WQ}date", DATE)
    run = etree.SubElement(ins, f"{WQ}r")
    t = etree.SubElement(run, f"{WQ}t")
    t.text = section
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p


def _para_style_val(para_el: etree._Element) -> str:
    pPr = para_el.find(f"{WQ}pPr")
    if pPr is None:
        return ""
    pStyle = pPr.find(f"{WQ}pStyle")
    return pStyle.get(f"{WQ}val", "") if pStyle is not None else ""


def _para_text(para_el: etree._Element) -> str:
    return "".join(t.text or "" for t in para_el.iter(f"{WQ}t"))


def _heading_para_index(all_paras: list, section: str, aliases: list) -> int | None:
    """Index within ``all_paras`` of the Heading-1 paragraph that IS ``section``
    (exact/alias/prefix). Operates on live lxml paragraphs."""
    h1 = _resolve_h1_style_id(all_paras)
    for i, p in enumerate(all_paras):
        text = _para_text(p)
        # Heading-1 styled OR a recognised unstyled section heading by text.
        if _para_style_val(p) != h1 and not _text_matches_main_section(text):
            continue
        if _matches_section(text, section, aliases):
            return i
    return None


def _canonical_insert_index(section: str, all_paras: list) -> int:
    """Return the index within ``all_paras`` to insert the stub for ``section``.

    Place it just after the block of the nearest PRESENT canonical
    predecessor (i.e. before that predecessor's following Heading-1), else
    just before the nearest PRESENT successor, else at the end.
    """
    from app.domain.canonical_jultp_template import CANONICAL_STRUCTURE

    def aliases_for(s):
        return CANONICAL_STRUCTURE["section_aliases"].get(s, [])

    target = _CANONICAL_ORDER.index(section)
    h1 = _resolve_h1_style_id(all_paras)

    def _next_h1_after(idx: int) -> int:
        for j in range(idx + 1, len(all_paras)):
            if _para_style_val(all_paras[j]) == h1:
                return j
        return len(all_paras)

    # Nearest present predecessor → insert at the end of its block.
    for i in range(target - 1, -1, -1):
        pred = _heading_para_index(all_paras, _CANONICAL_ORDER[i], aliases_for(_CANONICAL_ORDER[i]))
        if pred is not None:
            return _next_h1_after(pred)

    # No present predecessor → insert before nearest present successor.
    for i in range(target + 1, len(_CANONICAL_ORDER)):
        succ = _heading_para_index(all_paras, _CANONICAL_ORDER[i], aliases_for(_CANONICAL_ORDER[i]))
        if succ is not None:
            return succ

    return len(all_paras)


_REF_ENTRY_STYLE = "APA 7 Reference List Entry"
_REF_YEAR_RE = re.compile(r'\(\d{4}[a-z]?\)')


def _find_reference_entry_indexes(para_records: list) -> list[int]:
    """Return para_records list-positions for reference entries.

    Uses the same two-tier logic as reference_checker._select_reference_entries
    so "Entry N" numbering in comment text matches comment anchor position:

    Tier 1: paragraphs styled as the reference entry style (preferred).
    Tier 2 (only when no styled entries exist): year + length heuristic.
    """
    refs_idx = _find_heading(para_records, "References")
    if refs_idx is None:
        # No References heading: doc-wide styled scan only (mirrors extract_references)
        return [
            i for i, p in enumerate(para_records)
            if p.style == _REF_ENTRY_STYLE and not p.is_empty
        ]

    # Collect window: paragraphs inside the References section
    window: list[tuple[int, object]] = []
    for i in range(refs_idx + 1, len(para_records)):
        p = para_records[i]
        if p.style == "Heading 1" and not p.is_empty:
            break
        if not p.is_empty:
            window.append((i, p))

    # Two-tier: prefer styled entries; fall back to year+length heuristic
    styled = [(i, p) for i, p in window if p.style == _REF_ENTRY_STYLE]
    if styled:
        return [i for i, p in styled]
    return [
        i for i, p in window
        if _REF_YEAR_RE.search(p.text) and len(p.text) > 20
    ]


def _find_target_para_index(result: dict, para_records: list) -> int:
    """Choose the best paragraph index for a finding.

    This is a coarse anchor selection. Phrase-level anchoring then chooses
    a specific substring inside that paragraph.
    """

    if result.get("source") == "llm":
        return _find_para_by_section(result.get("section"), para_records)

    # Rule-to-location mapping: each family points to a section or heading.
    rule_id = result["rule_id"]

    if rule_id in _SEC_RULE_TO_SECTION and rule_id != "SEC009":
        section = _SEC_RULE_TO_SECTION[rule_id]
        return _find_missing_section_anchor(section, para_records)

    if rule_id == "SEC009":
        found = _find_heading(para_records, "Introduction")
        if found is not None:
            return found
        first_heading = _find_first_heading(para_records)
        if first_heading is not None:
            return first_heading
        return _find_first_non_empty_para(para_records)

    if rule_id.startswith("MET"):
        found = _find_heading(para_records, "Method")
        if found is not None:
            return found
        return _find_missing_section_anchor("Method", para_records)

    if rule_id.startswith("DIS"):
        found = _find_heading(para_records, "Discussion")
        if found is not None:
            return found
        return _find_missing_section_anchor("Discussion", para_records)

    if rule_id.startswith("SPE"):
        # Anchor "combined Results+Discussion" at the combined heading itself
        found = _find_heading(para_records, "Results")
        if found is not None:
            return found
        found = _find_heading(para_records, "Discussion")
        if found is not None:
            return found
        return _find_first_non_empty_para(para_records)

    if rule_id.startswith("FP") or rule_id == "STY001":
        return _find_front_matter_rule_anchor(rule_id, para_records)

    elif rule_id.startswith("CONS"):
        found = _find_heading(para_records, "References")
        if found is not None:
            return found
        return _find_missing_section_anchor("References", para_records)

    elif rule_id.startswith("REF") and not rule_id.startswith("REFE"):
        found = _find_heading(para_records, "References")
        if found is not None:
            return found
        return _find_missing_section_anchor("References", para_records)

    elif rule_id.startswith("CREF") and rule_id != "CREF000":
        match = re.match(r"CREF(\d+)", rule_id)
        if match:
            entry_num = int(match.group(1))
            ref_paras = _find_reference_entry_indexes(para_records)
            if entry_num <= len(ref_paras):
                return ref_paras[entry_num - 1]
        found = _find_heading(para_records, "References")
        if found is not None:
            return found
        return _find_missing_section_anchor("References", para_records)

    elif rule_id.startswith("REFE"):
        match = re.match(r"REFE(\d+)_", rule_id)
        if match:
            entry_num = int(match.group(1))
            ref_paras = _find_reference_entry_indexes(para_records)
            if entry_num <= len(ref_paras):
                return ref_paras[entry_num - 1]
        found = _find_heading(para_records, "References")
        if found is not None:
            return found
        return _find_missing_section_anchor("References", para_records)

    elif rule_id.startswith("HREF") or rule_id.startswith("DOIT"):
        # Both families anchor at the per-reference entry whose number is
        # encoded in the rule id: HREF005, DOIT005, DOIT005_2 all → entry 5.
        # The optional `_N` suffix on DOIT is used when a single reference
        # contains multiple DOIs — they share an anchor paragraph.
        match = re.match(r"(?:HREF|DOIT)(\d+)", rule_id)
        if match:
            entry_num = int(match.group(1))
            ref_paras = _find_reference_entry_indexes(para_records)
            if entry_num <= len(ref_paras):
                return ref_paras[entry_num - 1]
        found = _find_heading(para_records, "References")
        if found is not None:
            return found
        return _find_missing_section_anchor("References", para_records)

    if rule_id.startswith("ANDOR") or rule_id.startswith("ETC") or rule_id.startswith("DOTPT"):
        stored_idx = result.get("para_idx")
        if stored_idx is not None:
            return min(int(stored_idx), len(para_records) - 1)
        return _find_first_non_empty_para(para_records)

    if rule_id == "AFF001":
        stored_idx = result.get("para_idx")
        if stored_idx is not None:
            return min(int(stored_idx), len(para_records) - 1)
        # Fall back to the Author Affiliations paragraph.
        affiliations_idx = _find_front_style(para_records, ("Author Affiliations",))
        if affiliations_idx is not None:
            return affiliations_idx
        return _find_front_matter_anchor(para_records)

    return _find_first_non_empty_para(para_records)

def _find_para_by_section(section_name: str | None, para_records: list) -> int:
    """Find paragraph index from an LLM-provided section name.

    Returns the first non-empty body paragraph AFTER the section heading,
    not the heading itself, so LLM comments land on content not the label.
    """
    if not section_name:
        return _find_first_non_empty_para(para_records)

    raw = section_name.strip()
    if not raw:
        return _find_first_non_empty_para(para_records)

    canonical = _SECTION_ALIASES.get(raw.lower(), raw)

    heading_idx = _find_heading(para_records, canonical)
    if heading_idx is not None:
        # Walk forward to the first non-empty, non-heading paragraph
        for i in range(heading_idx + 1, len(para_records)):
            p = para_records[i]
            if not p.is_empty and p.style not in ("Heading 1", "Heading 2"):
                return i
        return heading_idx  # fallback: use heading if section is empty

    low = canonical.lower()
    if low in {"title", "front matter", "frontmatter"}:
        return _find_front_matter_anchor(para_records)
    if low in {"front page", "frontpage"}:
        return _find_front_matter_rule_anchor("FP004", para_records)
    if low in {"abstract", "summary"}:
        return _find_front_matter_rule_anchor("FP004", para_records)
    if low in {"keywords", "keyword"}:
        return _find_front_matter_rule_anchor("FP010", para_records)
    if low in {"practitioner notes", "practitioner note"}:
        return _find_front_matter_rule_anchor("FP009", para_records)
    if low in {"authors", "author"}:
        return _find_front_matter_rule_anchor("FP002", para_records)
    if low in {"affiliations", "author affiliations"}:
        return _find_front_matter_rule_anchor("FP003", para_records)

    if low in {"throughout manuscript", "manuscript", "overall", "across manuscript"}:
        intro = _find_heading(para_records, "Introduction")
        if intro is not None:
            return intro
        first_heading = _find_first_heading(para_records)
        if first_heading is not None:
            return first_heading

    if canonical in _CANONICAL_ORDER:
        return _find_missing_section_anchor(canonical, para_records)

    return _find_first_non_empty_para(para_records)

def _patch_rels(rels_xml: bytes) -> bytes:
    """Ensure `document.xml.rels` contains the comments relationship."""
    tree = etree.fromstring(rels_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    for rel in tree:
        if rel.get("Type") == COMMENT_REL_TYPE:
            return rels_xml

    new_rel = etree.SubElement(tree, f"{{{ns}}}Relationship")
    new_rel.set("Id", "rIdComments")
    new_rel.set("Type", COMMENT_REL_TYPE)
    new_rel.set("Target", "comments.xml")

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _patch_content_types(ct_xml: bytes) -> bytes:
    """Ensure `[Content_Types].xml` declares `word/comments.xml`."""
    tree = etree.fromstring(ct_xml)
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    part_name = "/word/comments.xml"

    for override in tree.findall(f"{{{ns}}}Override"):
        if override.get("PartName") == part_name:
            return ct_xml

    new_override = etree.SubElement(tree, f"{{{ns}}}Override")
    new_override.set("PartName", part_name)
    new_override.set("ContentType", COMMENT_CONTENT_TYPE)

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


##helper function foruse with _llm_notes_to_results
def _extract_field(obj, name: str, default=None):
    """Read a field from dict/object safely with a default."""
    if isinstance(obj, dict):
        return obj.get(name, default)
    return getattr(obj, name, default)


def _llm_notes_to_results(llm_results) -> list[dict]:
    """Normalize LLM editorial notes into validator-like result dictionaries."""
    if not llm_results:
        return []

    notes = _extract_field(llm_results, "notes", []) or []
    mapped = []

    for i, note in enumerate(notes, start=1):
        severity = str(_extract_field(note, "severity", "low")).lower()

        # Keep comments focused on actionable items.
        if severity == "low":
            continue

        status = "fail" if severity == "high" else "warn"
        message = str(_extract_field(note, "message", "")).strip()
        suggestion = str(_extract_field(note, "suggestion", "")).strip()
        section = str(_extract_field(note, "section", "")).strip()

        combined_message = message
        if suggestion:
            combined_message = f"{message} Suggested fix: {suggestion}"

        quote = str(_extract_field(note, "quote", "")).strip()

        mapped.append(
            {
                "rule_id": f"LLM{i:03d}",
                "status": status,
                "message": combined_message,
                "section": section,
                "quote": quote,
                "source": "llm",
            }
        )

    return mapped


def _dedupe_pending_comments(pending_comments: list[dict]) -> list[dict]:
    """Collapse exact duplicate comments for the same document anchor."""
    deduped = []
    seen = set()
    for item in pending_comments:
        message_key = " ".join(str(item.get("message", "")).split()).casefold()
        phrase_key = " ".join(str(item.get("target_phrase") or "").split()).casefold()
        key = (item.get("para_idx"), phrase_key, message_key)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def _comment_group_key(item: dict) -> str | None:
    """Return a family key for high-volume, same-kind comment groups."""
    rule_id = str(item.get("rule_id", ""))
    for prefix in ("ANDOR", "ETC", "DOTPT"):
        if rule_id.startswith(prefix):
            return prefix

    if item.get("source") == "llm":
        message = str(item.get("message", "")).casefold()
        if "acronym" in message:
            return "LLM_ACRONYM"
        if "abbreviation" in message:
            return "LLM_ABBREVIATION"
    return None


def _grouped_comment_message(group_key: str, items: list[dict]) -> str:
    count = len(items)
    if group_key == "ANDOR":
        return (
            f"Avoid 'and/or'. {count} occurrence{'s' if count != 1 else ''} "
            "were found in the manuscript. Please revise these constructions "
            "for clearer academic wording."
        )
    if group_key == "ETC":
        return (
            f"Avoid informal 'etc.' or 'etcetera' wording. {count} occurrence"
            f"{'s' if count != 1 else ''} were found. Please list the remaining "
            "items explicitly, or use a precise alternative where appropriate."
        )
    if group_key == "DOTPT":
        return (
            f"Convert bullet/dot-point lists to continuous prose. {count} list "
            f"issue{'s' if count != 1 else ''} were found in the manuscript."
        )
    if group_key == "LLM_ACRONYM":
        return (
            f"Please check acronym use across the manuscript. {count} similar "
            "acronym notes were grouped here to keep the comment pane manageable. "
            "Expand acronyms on first use and use the acronym form thereafter."
        )
    if group_key == "LLM_ABBREVIATION":
        return (
            f"Please check abbreviation use across the manuscript. {count} "
            "similar abbreviation notes were grouped here to keep the comment "
            "pane manageable."
        )
    return items[0]["message"]


def _group_pending_comments(pending_comments: list[dict]) -> list[dict]:
    """Reduce high-volume same-family findings to one Word comment each."""
    grouped: dict[str, list[dict]] = {}
    output: list[dict] = []

    for item in pending_comments:
        key = _comment_group_key(item)
        if key is None:
            output.append(item)
            continue
        grouped.setdefault(key, []).append(item)

    for key, items in grouped.items():
        first = min(items, key=lambda item: item["para_idx"])
        merged = dict(first)
        merged["message"] = _grouped_comment_message(key, items)
        # The summary belongs at the first relevant paragraph; don't try to
        # highlight one repeated phrase for a document-wide grouped comment.
        merged["target_phrase"] = first.get("target_phrase")
        output.append(merged)

    return output


def _patch_hyperlink_color_styles(styles_xml: bytes) -> bytes:
    """Set the Hyperlink character style to black via raw lxml (no python-docx)."""
    root = etree.fromstring(styles_xml)
    for style in root.findall(f"{WQ}style"):
        name_el = style.find(f"{WQ}name")
        if name_el is not None and name_el.get(f"{WQ}val") == "Hyperlink":
            rPr = style.find(f"{WQ}rPr")
            if rPr is None:
                rPr = etree.SubElement(style, f"{WQ}rPr")
            color = rPr.find(f"{WQ}color")
            if color is None:
                color = etree.SubElement(rPr, f"{WQ}color")
            color.set(f"{WQ}val", "000000")
            color.attrib.pop(f"{WQ}themeColor", None)
            color.attrib.pop(f"{WQ}themeShade", None)
            break
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def generate_commented_docx(
    input_path: str,
    output_path: str,
    report: dict,
    llm_results=None,
    ref_results: list | None = None,
) -> None:
    """Generate a reviewed `.docx` with inline comments and a validation summary page."""
    if ref_results is None:
        ref_results = []

    _REF_SECTION_NAMES = frozenset(("references", "reference list", "reference"))
    llm_mapped_results = [
        r for r in _llm_notes_to_results(llm_results)
        if str(r.get("section", "")).strip().lower() not in _REF_SECTION_NAMES
    ]
    all_results = report["results"] + ref_results + llm_mapped_results

    fails = [
        r for r in all_results
        if r["status"] == "fail"
        or (r["status"] == "warn" and r["rule_id"].startswith("REFE") and not r["rule_id"].endswith("_DOI"))
        # Advisory subsection-deviation notes (Method/Discussion carry
        # subheadings not in the template's expected set) are warn-level but
        # should still surface as a comment, anchored at their section heading.
        or (r["status"] == "warn" and r["rule_id"] in ("MET900", "DIS900"))
    ]

    ref_summary_message = _reference_issue_summary_comment(ref_results)
    para_records_for_summary = None
    if ref_summary_message is not None:
        para_records_for_summary = load_paragraphs(input_path)
        if _find_heading(para_records_for_summary, "References") is not None:
            fails.insert(0, {
                "rule_id": "REFSUMMARY",
                "status": "fail",
                "message": ref_summary_message,
            })

    if not fails:
        import shutil
        shutil.copy2(input_path, output_path)
        return

    with zipfile.ZipFile(input_path, "r") as z:
        names        = z.namelist()
        doc_xml      = z.read("word/document.xml")
        rels_xml     = z.read("word/_rels/document.xml.rels")
        ct_xml       = z.read("[Content_Types].xml")
        has_styles   = "word/styles.xml" in names
        styles_xml   = z.read("word/styles.xml") if has_styles else None
        has_comments = "word/comments.xml" in names
        comments_xml = z.read("word/comments.xml") if has_comments else None

    new_styles_xml = _patch_hyperlink_color_styles(styles_xml) if styles_xml else None

    # Parse document.xml with lxml directly — all XML preserved including images.
    doc_root  = etree.fromstring(doc_xml)
    body      = doc_root.find(f"{WQ}body")
    all_paras = [el for el in (body if body is not None else []) if el.tag == f"{WQ}p"]

    # Anchor lookup uses original input (paragraph indices match original structure).
    para_records = para_records_for_summary or load_paragraphs(input_path)

    if comments_xml:
        comments_root = etree.fromstring(comments_xml)
        existing_ids  = [int(el.get(f"{WQ}id", 0)) for el in comments_root.findall(f"{WQ}comment")]
        next_id = max(existing_ids, default=0) + 1
    else:
        comments_root = etree.Element(f"{WQ}comments", nsmap={"w": W})
        next_id = 1

    # ── Insert tracked Heading-1 stubs for missing canonical sections ────────
    # A missing section's SEC*/DIS*/MET* comments would otherwise anchor to a
    # neighbouring heading. Instead, insert a tracked stub heading at the
    # correct canonical slot and anchor those comments to it. The stub is
    # tracked via <w:ins> so the editor can accept (real heading) or reject.
    inserted_headings: dict[str, etree._Element] = {}
    # Anchor ELEMENT (and phrase) for each non-stub comment, captured BEFORE any
    # stub is inserted. Inserting a stub for a missing section shifts the
    # indices in ``all_paras``, so an index computed against the original
    # paragraph order would point at the wrong paragraph afterwards (e.g. a
    # Discussion comment landing on Results once Literature/Method stubs push
    # everything down). lxml element references survive insertions, so we anchor
    # by element instead. Keyed by ``id(result)``.
    non_stub_anchor: dict[int, tuple[etree._Element, str | None]] = {}
    if body is not None:
        missing_sections = []
        for result in fails:
            section = _section_for_rule(result.get("rule_id", ""))
            if section is None or section in missing_sections:
                continue
            if not _section_heading_present(section, para_records):
                missing_sections.append(section)
        # Capture non-stub anchor elements now, while all_paras still matches
        # para_records one-to-one.
        for result in fails:
            section = _section_for_rule(result.get("rule_id", ""))
            if section in missing_sections:
                continue  # anchored to its stub instead
            idx = _find_target_para_index(result, para_records)
            idx = min(idx, len(all_paras) - 1)
            if idx < 0:
                continue
            non_stub_anchor[id(result)] = (
                all_paras[idx],
                _infer_target_phrase(result, para_records, idx),
            )
        # Insert in canonical order so a freshly-inserted predecessor stub is
        # seen as "present" when positioning the next one.
        missing_sections.sort(key=_CANONICAL_ORDER.index)
        h1_style = _resolve_h1_style_id(all_paras)
        # Stub change-ids start high to avoid colliding with downstream passes.
        stub_change_id = 9000
        for section in missing_sections:
            insert_at = _canonical_insert_index(section, all_paras)
            stub = _make_inserted_heading_paragraph(section, h1_style, stub_change_id)
            stub_change_id += 1
            if insert_at < len(all_paras):
                all_paras[insert_at].addprevious(stub)
            else:
                sect_pr = body.find(f"{WQ}sectPr")
                if sect_pr is not None:
                    sect_pr.addprevious(stub)
                else:
                    body.append(stub)
            all_paras.insert(insert_at, stub)
            inserted_headings[section] = stub

    pending_comments = []
    for result in fails:
        rule_id = result.get("rule_id", "")
        section = _section_for_rule(rule_id)
        stub_element = inserted_headings.get(section) if section else None
        if stub_element is not None:
            anchor_el = stub_element
            target_phrase = None
        else:
            captured = non_stub_anchor.get(id(result))
            if captured is not None:
                anchor_el, target_phrase = captured
            else:
                # Fallback (e.g. body is None): resolve against live paragraphs.
                idx = _find_target_para_index(result, para_records)
                idx = min(idx, len(all_paras) - 1)
                if idx < 0 or not all_paras:
                    continue
                anchor_el = all_paras[idx]
                target_phrase = _infer_target_phrase(result, para_records, idx)
        # Element-based position keeps dedupe/sort stable across insertions.
        para_idx = all_paras.index(anchor_el)
        pending_comments.append({
            "para_idx": para_idx,
            "target_phrase": target_phrase,
            "message": result["message"],
            "rule_id": rule_id,
            "source": result.get("source", ""),
            "stub_element": stub_element,
            "anchor_el": anchor_el,
        })

    pending_comments = _dedupe_pending_comments(pending_comments)
    pending_comments = _group_pending_comments(pending_comments)
    pending_comments.sort(key=lambda item: item["para_idx"])

    for i, item in enumerate(pending_comments):
        comment_id = next_id + i
        target_phrase = item["target_phrase"]

        comments_root.append(_make_comment_element(comment_id, item["message"]))
        anchor_el = item.get("anchor_el")
        if anchor_el is None:
            anchor_el = item.get("stub_element")
        if anchor_el is None:
            anchor_el = all_paras[item["para_idx"]]
        _inject_comment_markers_for_phrase(anchor_el, comment_id, target_phrase)

    new_doc_xml      = etree.tostring(doc_root,       xml_declaration=True, encoding="UTF-8", standalone=True)
    new_comments_xml = etree.tostring(comments_root,  xml_declaration=True, encoding="UTF-8", standalone=True)
    new_rels_xml     = _patch_rels(rels_xml)
    new_ct_xml       = _patch_content_types(ct_xml)

    tmp_path = output_path + ".tmp"
    with zipfile.ZipFile(input_path, "r") as zin, \
         zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, new_doc_xml)
            elif item.filename == "word/styles.xml" and new_styles_xml is not None:
                zout.writestr(item, new_styles_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, new_rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, new_ct_xml)
            elif item.filename == "word/comments.xml":
                zout.writestr(item, new_comments_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
        if not has_comments:
            zout.writestr("word/comments.xml", new_comments_xml)

    os.replace(tmp_path, output_path)


def add_document_summary_comment(docx_path: str, text: str) -> None:
    """Attach a single document-level Word comment to the first body paragraph.

    Used by the input-normalisation pass to record a one-line summary of what
    was stripped from the author's submission. Self-contained: creates
    `word/comments.xml`, the rels entry, and the content-type override if the
    document doesn't already carry comments.
    """
    with zipfile.ZipFile(docx_path, "r") as z:
        names = z.namelist()
        doc_xml = z.read("word/document.xml")
        rels_xml = z.read("word/_rels/document.xml.rels")
        ct_xml = z.read("[Content_Types].xml")
        has_comments = "word/comments.xml" in names
        comments_xml = z.read("word/comments.xml") if has_comments else None

    doc_root = etree.fromstring(doc_xml)
    body = doc_root.find(f"{WQ}body")
    if body is None:
        return
    first_para = next(
        (el for el in body if el.tag == f"{WQ}p"),
        None,
    )
    if first_para is None:
        return

    if comments_xml:
        comments_root = etree.fromstring(comments_xml)
        existing_ids = [
            int(el.get(f"{WQ}id", 0)) for el in comments_root.findall(f"{WQ}comment")
        ]
        next_id = max(existing_ids, default=0) + 1
    else:
        comments_root = etree.Element(f"{WQ}comments", nsmap={"w": W})
        next_id = 1

    comments_root.append(_make_comment_element(next_id, text))
    _inject_comment_markers(first_para, next_id)

    new_doc_xml = etree.tostring(
        doc_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    new_comments_xml = etree.tostring(
        comments_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    new_rels_xml = _patch_rels(rels_xml)
    new_ct_xml = _patch_content_types(ct_xml)

    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, new_doc_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, new_rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, new_ct_xml)
            elif item.filename == "word/comments.xml":
                zout.writestr(item, new_comments_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
        if not has_comments:
            zout.writestr("word/comments.xml", new_comments_xml)

    os.replace(tmp_path, docx_path)


def add_paragraph_comment(docx_path: str, para_index: int, message: str) -> None:
    """Attach a Word comment to a specific body paragraph (zero-based index).

    If `para_index` is out of range it is clamped to the nearest valid index.
    Modifies `docx_path` in-place via an atomic tmp-file replace.
    """
    with zipfile.ZipFile(docx_path, "r") as z:
        names        = z.namelist()
        doc_xml      = z.read("word/document.xml")
        rels_xml     = z.read("word/_rels/document.xml.rels")
        ct_xml       = z.read("[Content_Types].xml")
        has_comments = "word/comments.xml" in names
        comments_xml = z.read("word/comments.xml") if has_comments else None

    doc_root  = etree.fromstring(doc_xml)
    body      = doc_root.find(f"{WQ}body")
    if body is None:
        return
    all_paras = [el for el in body if el.tag == f"{WQ}p"]
    if not all_paras:
        return
    target = max(0, min(para_index, len(all_paras) - 1))

    if comments_xml:
        comments_root = etree.fromstring(comments_xml)
        existing_ids  = [int(el.get(f"{WQ}id", 0)) for el in comments_root.findall(f"{WQ}comment")]
        next_id = max(existing_ids, default=0) + 1
    else:
        comments_root = etree.Element(f"{WQ}comments", nsmap={"w": W})
        next_id = 1

    comments_root.append(_make_comment_element(next_id, message))
    _inject_comment_markers(all_paras[target], next_id)

    new_doc_xml      = etree.tostring(doc_root,      xml_declaration=True, encoding="UTF-8", standalone=True)
    new_comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    new_rels_xml     = _patch_rels(rels_xml)
    new_ct_xml       = _patch_content_types(ct_xml)

    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, new_doc_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, new_rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, new_ct_xml)
            elif item.filename == "word/comments.xml":
                zout.writestr(item, new_comments_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
        if not has_comments:
            zout.writestr("word/comments.xml", new_comments_xml)
    os.replace(tmp_path, docx_path)


if __name__ == "__main__":
    import sys
    from pprint import pprint

    input_path = (
        sys.argv[1]
        if len(sys.argv) > 1
        else "tests/jutlp_sample_docx_test_pack/04_structure_and_endmatter_issues.docx"
    )
    output_path = sys.argv[2] if len(sys.argv) > 2 else "output_reviewed.docx"

    report = validate(input_path)

    pprint({"fail": report["fail"], "pass": report["pass"]})
    generate_commented_docx(input_path, output_path, report)
    print(f"\nCommented document saved to: {output_path}")
