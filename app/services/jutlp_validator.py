import re

from docx import Document as DocxDocument
from docx.oxml.ns import qn as _qn

from app.domain.canonical_jultp_template import (
    CANONICAL_STRUCTURE,
    strip_leading_section_number,
    subsection_alias_match,
)
from app.services.document_analysis_services import (
    extract_practitioner_notes,
    get_front_page,
    get_section_bounds,
    load_paragraphs,
    parse_docx_structure,
)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WQ = f"{{{_W}}}"


def _result(rule_id: str, status: str, message: str) -> dict:
    return {"rule_id": rule_id, "status": status, "message": message}


def _matches_section(
    heading_text: str,
    canonical: str,
    aliases: list[str] | None = None,
) -> bool:
    """True if heading_text matches canonical or any alias, exactly or as a
    prefix-with-colon (e.g. 'Introduction: Background' matches 'Introduction').

    Handles:
    - Exact match (case-insensitive): "Introduction" == "Introduction"
    - Alias match: "Methods" in aliases ["Methods", "Methodology"]
    - Titled heading: "Introduction: Background" startswith "introduction:"
    """
    h = strip_leading_section_number(heading_text).strip().lower()
    for name in [canonical] + (aliases or []):
        n = name.strip().lower()
        if h == n or h.startswith(n + ":"):
            return True
    return False


def _section_present(found_headings: list[str], canonical: str) -> bool:
    """True if any heading in found_headings fuzzy-matches canonical or its aliases."""
    aliases = CANONICAL_STRUCTURE["section_aliases"].get(canonical, [])
    return any(_matches_section(h, canonical, aliases) for h in found_headings)


def _required_section_location(section: str) -> str | None:
    if section == "Literature":
        return "after 'Introduction'"
    if section == "Method":
        return "after 'Literature'"
    return None


def check_required_sections(parsed: dict) -> list[dict]:
    results = []
    found = parsed["main_sections"]
    aliases = CANONICAL_STRUCTURE["section_aliases"]

    for i, section in enumerate(CANONICAL_STRUCTURE["main_sections"]):
        rule_id = f"SEC{i + 1:03d}"
        section_aliases = aliases.get(section, [])

        if any(_matches_section(h, section, section_aliases) for h in found):
            results.append(_result(rule_id, "pass", f"'{section}' found"))
            continue

        combined = next(
            (h for h in found if section.lower() in h.lower()),
            None,
        )
        if combined:
            results.append(_result(
                rule_id, "fail",
                f"'{section}' found combined with another section ('{combined}') — "
                f"JUTLP requires it as a separate top-level section",
            ))
        else:
            message = CANONICAL_STRUCTURE.get("missing_section_queries", {}).get(section)
            if message is None:
                message = f"Required section '{section}' not found"
                location = _required_section_location(section)
                if location is not None:
                    message += f"; it should appear {location}"
            results.append(_result(rule_id, "fail", message))

    return results


def check_section_order(parsed: dict) -> list[dict]:
    found = parsed["main_sections"]
    canonical = CANONICAL_STRUCTURE["main_sections"]
    aliases = CANONICAL_STRUCTURE["section_aliases"]

    def canonical_index(name):
        for i, s in enumerate(canonical):
            if _matches_section(name, s, aliases.get(s, [])):
                return i
        return None

    indices = [canonical_index(s) for s in found if canonical_index(s) is not None]

    if indices == sorted(indices):
        return [_result("SEC009", "pass", "Sections are in the correct order")]

    found_positions = {}
    for position, heading in enumerate(found):
        for section in canonical:
            if section in found_positions:
                continue
            if _matches_section(heading, section, aliases.get(section, [])):
                found_positions[section] = position
                break

    pair_messages = []
    ordered_pairs = [
        ("Introduction", "Literature"),
        ("Literature", "Method"),
    ]
    for previous, section in ordered_pairs:
        if previous not in found_positions or section not in found_positions:
            continue
        if found_positions[section] < found_positions[previous]:
            pair_messages.append(f"'{section}' should appear after '{previous}'")

    if pair_messages:
        return [_result(
            "SEC009",
            "fail",
            "Sections are out of the expected order: " + "; ".join(pair_messages),
        )]
    return [_result("SEC009", "fail", "Sections are out of the expected order")]


def _looks_like_subheading(text: str) -> bool:
    """True if ``text`` is plausibly a subheading, not a mis-styled body line.

    Mirrors the guards in ``_is_pseudo_heading2``: short and not ending in
    sentence-terminal punctuation. ``extract_subsections`` can over-capture a
    long Heading-2-styled sentence as a "subsection"; this keeps such body lines
    from being reported as non-template subheadings.
    """
    t = (text or "").strip()
    return bool(t) and len(t) <= 80 and t[-1] not in ".!?"


def _extra_subheadings(found_subs: list[str], required: list[str]) -> list[str]:
    """Return heading-like found subheadings that match no required subsection
    (or alias).

    Reuses ``subsection_alias_match`` per-found-heading so an aliased heading
    (e.g. 'Setting of the Research' → 'Research Design') is treated as expected.
    """
    return [
        f for f in found_subs
        if _looks_like_subheading(f)
        and not any(subsection_alias_match(req, [f]) for req in required)
    ]


def _extra_subheading_result(rule_id: str, section: str, extras: list[str],
                             required: list[str]) -> dict:
    expected = ", ".join(required)
    listed = ", ".join(f"'{e}'" for e in extras)
    return _result(
        rule_id, "warn",
        f"{section} contains subheading(s) not in the JUTLP expected set "
        f"({expected}): {listed}. Consider aligning with the expected "
        f"subheadings or confirm they are necessary.",
    )


def check_method_subsections(parsed: dict) -> list[dict]:
    # Skip entirely when Method section is absent — absence is already reported
    # by check_required_sections; running here would cascade-fail every MET rule.
    if not _section_present(parsed["main_sections"], "Method"):
        return []

    results = []
    found_subs = parsed["subsections"].get("Method", [])
    required = CANONICAL_STRUCTURE["required_subsections"]["Method"]

    for i, sub in enumerate(required):
        rule_id = f"MET{i + 1:03d}"
        if subsection_alias_match(sub, found_subs):
            results.append(_result(rule_id, "pass", f"Method subsection '{sub}' found"))
        else:
            results.append(_result(rule_id, "fail", f"Required Method subsection '{sub}' missing"))

    extras = _extra_subheadings(found_subs, required)
    if extras:
        results.append(_extra_subheading_result("MET900", "Method", extras, required))

    return results


def check_discussion_subsections(parsed: dict) -> list[dict]:
    found = parsed["main_sections"]
    discussion_present = _section_present(found, "Discussion")

    required = CANONICAL_STRUCTURE["required_subsections"]["Discussion"]

    if not discussion_present:
        # Discussion absent or only combined — all subsections fail.
        return [
            _result(f"DIS{i + 1:03d}", "fail", f"Required Discussion subsection '{sub}' missing")
            for i, sub in enumerate(required)
        ]

    results = []
    found_subs = parsed["subsections"].get("Discussion", [])

    for i, sub in enumerate(required):
        rule_id = f"DIS{i + 1:03d}"
        if subsection_alias_match(sub, found_subs):
            results.append(_result(rule_id, "pass", f"Discussion subsection '{sub}' found"))
        else:
            results.append(_result(rule_id, "fail", f"Required Discussion subsection '{sub}' missing"))

    extras = _extra_subheadings(found_subs, required)
    if extras:
        results.append(_extra_subheading_result("DIS900", "Discussion", extras, required))

    return results


def check_front_page(parsed: dict) -> list[dict]:
    fp = parsed["front_page"]

    checks = [
        ("FP001", fp["title_found"],                      "Article title found",             "Article title (style 'Article Title') not found"),
        ("FP002", fp["authors_found"],                    "Authors block found",              "Authors block (style 'Authors') not found"),
        ("FP003", fp["affiliations_found"],               "Affiliations block found",         "Affiliations (style 'Author Affiliations') not found"),
        ("FP004", fp["abstract_heading_found"],           "Abstract heading found",           "Abstract heading not found"),
        ("FP005", fp["practitioner_notes_heading_found"], "Practitioner Notes heading found", "Practitioner Notes heading not found"),
        ("FP006", fp["keywords_heading_found"],           "Keywords heading found",           "Keywords heading not found"),
    ]

    return [
        _result(rule_id, "pass" if ok else "fail", pass_msg if ok else fail_msg)
        for rule_id, ok, pass_msg, fail_msg in checks
    ]


def check_abstract(parsed: dict) -> list[dict]:
    front = CANONICAL_STRUCTURE["front_page"]
    abstract = parsed["front_page"]["abstract"]
    max_words = front["abstract_max_words"]
    max_lines = front["abstract_max_lines"]
    region = front["abstract_line_region"]

    if not abstract["found"]:
        return [
            _result("FP007", "fail", "Abstract body not found"),
            _result("FP008", "fail", "Abstract length cannot be checked — abstract not found"),
        ]

    # JUTLP rule: the abstract must fit lines 7–23 of the front page (≤ max_lines
    # lines) AND be ≤ max_words words — both are length limits on one element.
    wc = abstract["word_count"]
    lc = abstract.get("line_count", 0)
    over_words = wc > max_words
    over_lines = lc > max_lines

    if over_words or over_lines:
        parts = []
        if over_lines:
            parts.append(
                f"~{lc} lines (must fit lines {region[0]}–{region[1]}, i.e. max {max_lines})"
            )
        if over_words:
            parts.append(f"{wc} words (max {max_words})")
        status = "fail"
        msg = "Abstract is too long: " + " and ".join(parts) + "."
    else:
        status = "pass"
        msg = (
            f"Abstract length OK ({wc} words, ~{lc} lines; "
            f"limits: {max_words} words / {max_lines} lines)"
        )

    return [
        _result("FP007", "pass", "Abstract body found"),
        _result("FP008", status, msg),
    ]


def check_practitioner_notes(parsed: dict) -> list[dict]:
    count = parsed["front_page"]["practitioner_notes_count"]
    required = CANONICAL_STRUCTURE["front_page"]["practitioner_notes_count"]

    if count == required:
        return [_result("FP009", "pass", f"Practitioner notes: {count} (required {required})")]
    return [_result("FP009", "fail", f"Practitioner notes: {count} found, {required} required")]


def check_keywords(parsed: dict) -> list[dict]:
    count = len(parsed["front_page"]["keywords"])
    max_count = CANONICAL_STRUCTURE["front_page"]["keywords_max_count"]

    if count == 0:
        return [_result("FP010", "fail", "No keywords found")]
    if count <= max_count:
        return [_result("FP010", "pass", f"{count} keyword(s) found (max {max_count})")]
    return [_result("FP010", "fail", f"{count} keywords found — exceeds maximum of {max_count}")]


def check_forbidden_styles(parsed: dict) -> list[dict]:
    style_counts = parsed["style_counts"]
    forbidden = CANONICAL_STRUCTURE["style_rules"]["forbidden_styles"]

    return [
        _result(
            f"STY{i + 1:03d}",
            "fail" if style in style_counts else "pass",
            f"Forbidden style '{style}' found — remove before submission"
            if style in style_counts
            else f"Forbidden style '{style}' not present",
        )
        for i, style in enumerate(forbidden)
    ]


def check_combined_results_discussion(parsed: dict) -> list[dict]:
    for s in parsed["main_sections"]:
        clean = s.lower()
        has_results_label = "result" in clean or "finding" in clean
        if has_results_label and "discussion" in clean:
            return [_result("SPE001", "warn",
                f"Combined '{s}' section found. "
                + CANONICAL_STRUCTURE["combined_results_discussion_query"],
            )]
    return [_result("SPE001", "pass", "Results and Discussion are in separate sections")]


def check_heading3_usage(parsed: dict) -> list[dict]:
    count = parsed["style_counts"].get("Heading 3", 0)
    if count:
        return [_result("STY002", "warn",
            f"{count} 'Heading 3' paragraph(s) found — level 3 headings are rarely necessary and should be avoided")]
    return [_result("STY002", "pass", "No Heading 3 paragraphs found")]


def check_conclusion_length(docx_path: str) -> list[dict]:
    paragraphs = load_paragraphs(docx_path)
    bounds = get_section_bounds(paragraphs, "Conclusion")
    if bounds is None:
        bounds = get_section_bounds(paragraphs, "Conclusions")
    if bounds is None:
        return []  # missing conclusion already caught by check_required_sections
    start, end = bounds
    body_paras = [
        p for p in paragraphs[start + 1 : end]
        if not p.is_empty and p.style not in ("Heading 1", "Heading 2", "Heading 3")
    ]
    if len(body_paras) <= 3:
        return [_result("CON001", "pass", f"Conclusion is {len(body_paras)} paragraph(s) (within 2–3 paragraph guideline)")]
    return [_result("CON001", "warn",
        f"Conclusion is {len(body_paras)} paragraphs — guidelines recommend no more than 2–3")]


def _para_has_page_break_before(para) -> bool:
    pPr = para._element.find(f"{_WQ}pPr")
    if pPr is None:
        return False
    pb = pPr.find(f"{_WQ}pageBreakBefore")
    if pb is None:
        return False
    val = pb.get(f"{_WQ}val", "true")
    return val.lower() not in ("false", "0", "off")


def _preceding_para_ends_with_page_break(all_paras: list, heading_para) -> bool:
    for i, p in enumerate(all_paras):
        if p is heading_para and i > 0:
            prev = all_paras[i - 1]
            for br in prev._element.findall(f".//{_WQ}br"):
                if br.get(f"{_WQ}type") == "page":
                    return True
    return False


def _section_has_page_break(docx_path: str, canonical_name: str) -> bool:
    """Return True if the heading matching canonical_name has a page break before it."""
    doc = DocxDocument(docx_path)
    paras = doc.paragraphs
    aliases = CANONICAL_STRUCTURE["section_aliases"].get(canonical_name, [])
    heading = next(
        (p for p in paras
         if p.style.name.startswith("Heading")
         and _matches_section(p.text.strip(), canonical_name, aliases)),
        None,
    )
    if heading is None:
        return False
    return _para_has_page_break_before(heading) or _preceding_para_ends_with_page_break(paras, heading)


def check_page_break_before_introduction(docx_path: str, parsed: dict) -> list[dict]:
    if not _section_present(parsed["main_sections"], "Introduction"):
        return []
    if _section_has_page_break(docx_path, "Introduction"):
        return [_result("SPE002", "pass", "Page break before Introduction found")]
    return [_result("SPE002", "warn",
        "No page break found before Introduction — the JUTLP template requires one")]


def check_page_break_before_references(docx_path: str, parsed: dict) -> list[dict]:
    if not _section_present(parsed["main_sections"], "References"):
        return []
    if _section_has_page_break(docx_path, "References"):
        return [_result("SPE003", "pass", "Page break before References found")]
    return [_result("SPE003", "warn",
        "No page break found before References — the JUTLP template requires one")]


# Styles excluded from informal-language and dot-point scanning.
_INFORMAL_SKIP_STYLES = frozenset({"APA 7 Reference List Entry"})

_ANDOR_RE = re.compile(r"\band/or\b", re.IGNORECASE)
_ETC_RE = re.compile(r"\betc\.|\betc\b|etcetera\b|et\s+cetera\b", re.IGNORECASE)


def check_informal_language(docx_path: str) -> list[dict]:
    """Flag 'and/or' and 'etc.'/'etcetera' — ambiguous or vague in academic writing."""
    paragraphs = load_paragraphs(docx_path)
    andor_results: list[dict] = []
    etc_results: list[dict] = []

    for para in paragraphs:
        if not para.text or para.style in _INFORMAL_SKIP_STYLES:
            continue

        for match in _ANDOR_RE.finditer(para.text):
            n = len(andor_results) + 1
            andor_results.append({
                "rule_id": f"ANDOR{n:03d}",
                "status": "fail",
                "message": "Avoid 'and/or'.",
                "para_idx": para.index,
                "anchor_phrase": match.group(),
            })

        for match in _ETC_RE.finditer(para.text):
            n = len(etc_results) + 1
            etc_results.append({
                "rule_id": f"ETC{n:03d}",
                "status": "fail",
                "message": (
                    f"Avoid '{match.group()}' — vague and informal in academic writing. "
                    "Suggested fix: list the remaining items explicitly, or use "
                    "'among others' / 'and so on' if an exhaustive list is not feasible."
                ),
                "para_idx": para.index,
                "anchor_phrase": match.group(),
            })

    if not andor_results:
        andor_results = [_result("ANDOR000", "pass", "No 'and/or' constructions found")]
    if not etc_results:
        etc_results = [_result("ETC000", "pass", "No 'etc.' or 'etcetera' usage found")]

    return andor_results + etc_results


def check_dot_points(docx_path: str) -> list[dict]:
    """Flag bullet/numbered list paragraphs — continuous prose is expected in JUTLP."""
    doc = DocxDocument(docx_path)

    results: list[dict] = []
    in_practitioner_notes = False
    group_start_idx = None
    group_anchor = ""
    group_count = 0

    def _flush_group() -> None:
        nonlocal group_start_idx, group_anchor, group_count
        if group_start_idx is None:
            return
        n = len(results) + 1
        message = "Avoid bullet/dot-point lists — continuous prose is expected in academic journal manuscripts."
        if group_count > 1:
            message = (
                "This list contains "
                + str(group_count)
                + " bullet/dot-point items. Convert the list to continuous prose for an academic journal manuscript."
            )
        results.append({
            "rule_id": f"DOTPT{n:03d}",
            "status": "fail",
            "message": message,
            "para_idx": group_start_idx,
            "anchor_phrase": group_anchor,
        })
        group_start_idx = None
        group_anchor = ""
        group_count = 0

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue
        low = text.lower()
        if low in {"practitioner notes", "practioner notes"}:
            _flush_group()
            in_practitioner_notes = True
            continue
        if low == "keywords":
            _flush_group()
            in_practitioner_notes = False
        if in_practitioner_notes:
            continue
        style_name = p.style.name if p.style else ""
        if style_name in _INFORMAL_SKIP_STYLES:
            _flush_group()
            continue

        is_list = "list" in style_name.lower()
        if not is_list:
            pPr = p._p.pPr
            if pPr is not None and pPr.find(_qn("w:numPr")) is not None:
                is_list = True

        if not is_list:
            _flush_group()
            continue

        words = text.split()
        anchor = " ".join(words[:4]) if len(words) >= 4 else text
        if group_start_idx is None:
            group_start_idx = i
            group_anchor = anchor
        group_count += 1

    _flush_group()

    if not results:
        return [_result("DOTPT000", "pass", "No bullet or dot-point lists found")]

    return results


# ── New constants for additional checks ──────────────────────────────────────

# Styles that belong to non-body roles and are exempt from the Normal-style check.
_FIGURE_NUMBER_STYLES = frozenset({
    "Figure Number", "FigureNumber", "Figure/Table Number", "FigureTableNumber",
})
_FIGURE_TITLE_STYLES = frozenset({
    "Figure Title", "FigureTitle", "Figure/Table Title", "FigureTableTitle",
})

_BODY_SKIP_STYLES = frozenset({
    "Article Title", "Authors", "Author Affiliations",
    "Heading Front Page", "Front Page Text", "Practitioner Notes",
    "Heading 1", "Heading 2", "Heading 3",
    *_FIGURE_NUMBER_STYLES,
    *_FIGURE_TITLE_STYLES,
    "Table Number", "Table Title", "Table Text", "Table Emphasis",
    "APA 7 Reference List Entry",
    "Quote",
    "Guidance Notes",
    "CommentText",
    "Default Paragraph Style",
    "Caption",
})

_QUOTE_CHARS = frozenset({'"', '\u201c', '\u201d'})
_FIG_NUM_RE = re.compile(r"Figure\s+(\d+)", re.IGNORECASE)
_TAB_NUM_RE = re.compile(r"Table\s+(\d+)", re.IGNORECASE)
_ABBREV_RE  = re.compile(r"\b[A-Z]{2,5}\b")
_EQUAL_CONTRIB_RE = re.compile(r"equal.*contribut|contribut.*equal", re.IGNORECASE)
_PN_MAX_WORDS = 30


def check_abstract_single_paragraph(parsed: dict) -> list[dict]:
    """FP011: JUTLP requires an unstructured single-paragraph abstract."""
    abstract = parsed["front_page"]["abstract"]
    if not abstract["found"]:
        return []
    count = abstract.get("paragraph_count", 1)
    if count <= 1:
        return [_result("FP011", "pass", "Abstract is a single paragraph")]
    return [_result("FP011", "fail",
        f"Abstract spans {count} paragraphs — JUTLP requires an unstructured single-paragraph abstract")]


def check_practitioner_note_length(docx_path: str) -> list[dict]:
    """FP012: Each practitioner note should be no more than 2 lines (~30 words)."""
    paragraphs = load_paragraphs(docx_path)
    notes = extract_practitioner_notes(paragraphs)
    if not notes:
        return []
    flags = []
    for i, note in enumerate(notes, start=1):
        wc = len(note.text.split())
        if wc > _PN_MAX_WORDS:
            flags.append(_result("FP012", "warn",
                f"Practitioner note {i} is {wc} words — each note should be no more than 2 lines (~{_PN_MAX_WORDS} words)"))
    if not flags:
        return [_result("FP012", "pass", "All practitioner notes are within the 2-line limit")]
    return flags


def check_keyword_abbreviations(parsed: dict) -> list[dict]:
    """FP013: Keywords should use full words, no abbreviations."""
    keywords = parsed["front_page"]["keywords"]
    if not keywords:
        return []
    flagged = [kw for kw in keywords if _ABBREV_RE.search(kw)]
    if not flagged:
        return [_result("FP013", "pass", "No abbreviations found in keywords")]
    return [_result("FP013", "warn",
        f"Keyword(s) may contain abbreviations: {', '.join(flagged)} — "
        "use full words (e.g. 'artificial intelligence' not 'AI')")]


def check_body_paragraph_styles(docx_path: str) -> list[dict]:
    """STY003: Body paragraphs should use the Normal style."""
    paragraphs = load_paragraphs(docx_path)
    wrong = [
        p for p in paragraphs
        if not p.is_empty
        and p.style not in _BODY_SKIP_STYLES
        and not p.style.startswith("Heading")
        and p.style != "Normal"
    ]
    if not wrong:
        return [_result("STY003", "pass", "All body paragraphs use the Normal style")]
    unique_styles = ", ".join(sorted({p.style for p in wrong}))
    return [_result("STY003", "warn",
        f"{len(wrong)} body paragraph(s) use non-Normal styles ({unique_styles}) — "
        "body text should use the Normal style (11pt Arial, left-justified, 1.15 line-spacing)")]


def check_reference_paragraph_styles(docx_path: str) -> list[dict]:
    """STY004: Reference paragraphs should use 'APA 7 Reference List Entry' style."""
    paragraphs = load_paragraphs(docx_path)
    bounds = get_section_bounds(paragraphs, "References")
    if bounds is None:
        return []
    start, end = bounds
    ref_paras = [p for p in paragraphs[start + 1:end] if not p.is_empty]
    if not ref_paras:
        return []
    wrong = [p for p in ref_paras if p.style != "APA 7 Reference List Entry"]
    if not wrong:
        return [_result("STY004", "pass", "All reference paragraphs use 'APA 7 Reference List Entry' style")]
    return [_result("STY004", "warn",
        f"{len(wrong)} reference paragraph(s) do not use 'APA 7 Reference List Entry' style")]


def check_block_quote_style(docx_path: str) -> list[dict]:
    """SPE004: Quotes ≥40 words should use the Quote style, not Normal."""
    paragraphs = load_paragraphs(docx_path)
    flags = [
        p for p in paragraphs
        if not p.is_empty
        and p.style == "Normal"
        and len(p.text) > 0
        and (p.text[0] in _QUOTE_CHARS or p.text[-1] in _QUOTE_CHARS)
        and len(p.text.split()) >= 40
    ]
    if not flags:
        return [_result("SPE004", "pass", "No long inline quotes found that should use the Quote style")]
    return [_result("SPE004", "warn",
        f"{len(flags)} paragraph(s) appear to be block quotes (≥40 words with quotation marks) "
        "but use Normal style — quotes longer than 40 words should use the Quote style")]


def check_figure_table_structure(docx_path: str) -> list[dict]:
    """FIG001/FIG002/TAB001/TAB002: Number→title order and sequential numbering."""
    paragraphs = load_paragraphs(docx_path)
    results = []
    fig_numbers: list[int] = []
    tab_numbers: list[int] = []
    fig_fails = 0
    tab_fails = 0

    for i, p in enumerate(paragraphs):
        if p.style in _FIGURE_NUMBER_STYLES:
            nxt = next((pp for pp in paragraphs[i + 1:i + 4] if not pp.is_empty), None)
            if nxt is None or nxt.style not in _FIGURE_TITLE_STYLES:
                fig_fails += 1
                results.append(_result("FIG001", "fail",
                    f"'{p.text}' is not immediately followed by a Figure Title paragraph"))
            m = _FIG_NUM_RE.search(p.text)
            if m:
                fig_numbers.append(int(m.group(1)))

        elif p.style == "Table Number":
            nxt = next((pp for pp in paragraphs[i + 1:i + 4] if not pp.is_empty), None)
            if nxt is None or nxt.style != "Table Title":
                tab_fails += 1
                results.append(_result("TAB001", "fail",
                    f"'{p.text}' is not immediately followed by a Table Title paragraph"))
            m = _TAB_NUM_RE.search(p.text)
            if m:
                tab_numbers.append(int(m.group(1)))

    if fig_numbers:
        if fig_fails == 0:
            results.append(_result("FIG001", "pass",
                f"All figure number paragraphs are followed by Figure Title ({len(fig_numbers)} figure(s))"))
        expected = list(range(1, len(fig_numbers) + 1))
        if fig_numbers == expected:
            results.append(_result("FIG002", "pass",
                f"Figure numbers are sequential ({len(fig_numbers)} figure(s))"))
        else:
            results.append(_result("FIG002", "fail",
                f"Figure numbers are not sequential: found {fig_numbers}, expected {expected}"))
    else:
        results.append(_result("FIG002", "pass", "No figures found"))

    if tab_numbers:
        if tab_fails == 0:
            results.append(_result("TAB001", "pass",
                f"All table number paragraphs are followed by Table Title ({len(tab_numbers)} table(s))"))
        expected = list(range(1, len(tab_numbers) + 1))
        if tab_numbers == expected:
            results.append(_result("TAB002", "pass",
                f"Table numbers are sequential ({len(tab_numbers)} table(s))"))
        else:
            results.append(_result("TAB002", "fail",
                f"Table numbers are not sequential: found {tab_numbers}, expected {expected}"))
    else:
        results.append(_result("TAB002", "pass", "No tables found"))

    return results


def check_no_equal_contributing(docx_path: str) -> list[dict]:
    """AFF001: JUTLP does not accept equal contributing author statements."""
    paragraphs = load_paragraphs(docx_path)
    front = get_front_page(paragraphs)
    for p in front:
        m = _EQUAL_CONTRIB_RE.search(p.text)
        if m:
            words = p.text.split()
            anchor = " ".join(words[:6]) if len(words) >= 6 else p.text
            return [{
                "rule_id": "AFF001",
                "status": "fail",
                "message": (
                    "Equal contributing authors note found - "
                    "JUTLP does not accept equal contributing author statements. "
                    "Please remove this note from the affiliations block."
                ),
                "para_idx": p.index,
                "anchor_phrase": anchor,
            }]
    return [_result("AFF001", "pass", "No equal contributing authors note found")]


def build_report(results: list[dict]) -> dict:
    return {
        "total": len(results),
        "pass": sum(1 for r in results if r["status"] == "pass"),
        "warn": sum(1 for r in results if r["status"] == "warn"),
        "fail": sum(1 for r in results if r["status"] == "fail"),
        "results": results,
    }


def validate(docx_path: str) -> dict:
    parsed = parse_docx_structure(docx_path)

    results = []
    results += check_required_sections(parsed)
    results += check_section_order(parsed)
    results += check_method_subsections(parsed)
    results += check_discussion_subsections(parsed)
    results += check_front_page(parsed)
    results += check_abstract(parsed)
    results += check_abstract_single_paragraph(parsed)
    results += check_practitioner_notes(parsed)
    results += check_practitioner_note_length(docx_path)
    results += check_keywords(parsed)
    results += check_keyword_abbreviations(parsed)
    results += check_forbidden_styles(parsed)
    results += check_combined_results_discussion(parsed)
    results += check_heading3_usage(parsed)
    results += check_conclusion_length(docx_path)
    results += check_page_break_before_introduction(docx_path, parsed)
    results += check_page_break_before_references(docx_path, parsed)
    results += check_body_paragraph_styles(docx_path)
    results += check_reference_paragraph_styles(docx_path)
    results += check_block_quote_style(docx_path)
    results += check_figure_table_structure(docx_path)
    results += check_no_equal_contributing(docx_path)
    results += check_informal_language(docx_path)
    results += check_dot_points(docx_path)

    return build_report(results)


if __name__ == "__main__":
    import sys
    from pprint import pprint

    path = sys.argv[1] if len(sys.argv) > 1 else "tests/jutlp_sample_docx_test_pack/01_valid_identified.docx"
    pprint(validate(path))
