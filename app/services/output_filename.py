import os
import re
from datetime import datetime

from docx import Document

from app.services.document_analysis_services import get_front_page, load_paragraphs

JOURNAL_NAME = "JUTLP"


def _remove_superscript_runs_from_author_line(docx_path: str, paragraph_index: int | None, fallback: str):
    if paragraph_index is None:
        return fallback, False

    doc = Document(docx_path)
    if paragraph_index >= len(doc.paragraphs):
        return fallback, False

    text = ""
    removed = False
    for run in doc.paragraphs[paragraph_index].runs:
        if run.font.superscript is True:
            removed = True
        else:
            text += run.text or ""

    text = "\n".join(re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines() if line.strip())
    return text or fallback, removed


def _get_author_line_and_markers(docx_path: str):
    paragraphs = load_paragraphs(docx_path)
    front = get_front_page(paragraphs)

    author_line = ""
    author_index = None
    author_style_found = False
    markers = []

    for p in front:
        if p.style == "Authors" and p.is_empty is False:
            author_line = p.text
            author_index = p.index
            author_style_found = True
            break

    if author_line == "":
        non_empty = []
        for p in front:
            if p.is_empty is False:
                non_empty.append(p)
        if len(non_empty) > 1:
            author_line = non_empty[1].text
            author_index = non_empty[1].index

    for p in front:
        if p.is_empty is True:
            continue
        if author_index is not None and p.index <= author_index:
            continue

        lower_text = p.text.strip().lower()
        if lower_text in ["abstract", "keywords", "citation", "introduction"]:
            break

        found = re.findall(r"(?:^|[;,\n])\s*([a-zA-Z])\s+", p.text)
        prefix_match = re.match(r"^\s*([a-zA-Z])(?=[A-ZÀ-Ö])", p.text)
        if prefix_match is not None:
            found.append(prefix_match.group(1))

        for marker in found:
            clean_marker = marker.lower()
            if clean_marker not in markers:
                markers.append(clean_marker)

    author_line, removed_superscript = _remove_superscript_runs_from_author_line(docx_path, author_index, author_line)
    return author_line, markers, author_style_found, removed_superscript


def _first_author_last_name(author_line: str, markers: list[str]) -> str:
    lines = [line.strip() for line in author_line.splitlines() if line.strip()]
    text = (lines[0] if lines else author_line).strip()

    if "," in text:
        first_author = text.split(",", 1)[0].strip()
    elif " and " in text:
        first_author = text.split(" and ", 1)[0].strip()
    else:
        first_author = text

    first_author = re.sub(r"\([^)]*\)", " ", first_author)
    first_author = re.sub(r"\^[A-Za-z0-9]+", " ", first_author)
    first_author = first_author.replace("*", " ")
    first_author = first_author.replace("#", " ")
    first_author = re.sub(r"\d+", " ", first_author)

    words = re.findall(r"[A-Za-z][A-Za-z'-]*", first_author)
    title_words = ["dr", "prof", "professor", "mr", "mrs", "ms", "miss"]

    name_words = []
    for word in words:
        if word.lower() not in title_words:
            name_words.append(word)

    if len(name_words) == 0:
        return "UnknownAuthor"

    last_name = name_words[-1]

    for marker in markers:
        if marker == "":
            continue
        if last_name.lower().endswith(marker):
            if len(last_name) > len(marker) + 2:
                last_name = last_name[:-len(marker)]
                break

    last_name = re.sub(r"[^A-Za-z0-9]", "", last_name)
    if last_name == "":
        return "UnknownAuthor"

    return last_name[0].upper() + last_name[1:]


def _get_document_year(docx_path: str) -> int:
    paragraphs = load_paragraphs(docx_path)
    front = get_front_page(paragraphs)

    citation_found = False
    for p in front:
        if p.text.strip().lower() == "citation":
            citation_found = True
            continue
        if citation_found is True and p.is_empty is False:
            match = re.search(r"\((20\d{2})[a-z]?\)", p.text)
            if match is not None:
                return int(match.group(1))
            break

    for p in front:
        match = re.search(r"\((20\d{2})[a-z]?\)", p.text)
        if match is not None:
            return int(match.group(1))

    return datetime.now().year


def _unique_output_filename(last_name: str, year: int, output_dir: str, ignore_path: str | None = None) -> str:
    copy_edit_number = 1
    while True:
        filename = f"{last_name}_{JOURNAL_NAME}_{year}_CopyEdit{copy_edit_number}.docx"
        output_path = os.path.join(output_dir, filename)

        if os.path.exists(output_path) is False or os.path.abspath(output_path) == os.path.abspath(ignore_path or ""):
            return filename

        copy_edit_number += 1


def build_output_filename(docx_path: str, output_dir: str, ignore_path: str | None = None) -> str:
    author_line, markers, _, _ = _get_author_line_and_markers(docx_path)
    last_name = _first_author_last_name(author_line, markers)
    year = _get_document_year(docx_path)
    return _unique_output_filename(last_name, year, output_dir, ignore_path)


def build_output_filename_from_author_line(
    docx_path: str,
    output_dir: str,
    author_line: str,
    ignore_path: str | None = None,
) -> str:
    if author_line.strip() == "":
        return build_output_filename(docx_path, output_dir, ignore_path)
    doc_author_line, markers, author_style_found, removed_superscript = _get_author_line_and_markers(docx_path)
    if author_style_found is True and removed_superscript is True:
        author_line = doc_author_line
    last_name = _first_author_last_name(author_line, markers)
    year = _get_document_year(docx_path)
    return _unique_output_filename(last_name, year, output_dir, ignore_path)
