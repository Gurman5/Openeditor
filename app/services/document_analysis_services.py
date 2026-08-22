import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domain.canonical_jultp_template import (
    CANONICAL_STRUCTURE,
    strip_leading_section_number,
)
from app.domain.models import ParagraphRecord

_QN_R = qn("w:r")
_QN_T = qn("w:t")
_QN_TAB = qn("w:tab")
_QN_BR = qn("w:br")
_QN_CR = qn("w:cr")
_QN_DEL = qn("w:del")
_QN_P = qn("w:p")


def _run_skipped(run_el, para_p) -> bool:
    """True if ``run_el`` should NOT contribute to ``para_p``'s text.

    Walks up to the run's nearest ancestors:
    * inside a tracked deletion (``<w:del>``) → skip (accepted-changes view);
    * inside a *nested* paragraph (the first ``<w:p>`` ancestor is not
      ``para_p``) → skip. This excludes runs that live in a floating text box /
      drawing (``w:txbxContent``) anchored on this paragraph, which would
      otherwise merge the box's text into the host paragraph (e.g. the JUTLP
      editorial "Editors [LEAVE BLANK]…" box bleeding into the Abstract
      heading and breaking the exact-text section match).
    """
    anc = run_el.getparent()
    while anc is not None:
        tag = anc.tag
        if tag == _QN_DEL:
            return True
        if tag == _QN_P:
            return anc is not para_p
        anc = anc.getparent()
    return True


def _accepted_paragraph_text(paragraph) -> str:
    """Return a paragraph's *own* text as if all tracked changes were accepted.

    python-docx ``paragraph.text`` reads only the runs that are *direct*
    children of the paragraph, so it DROPS text inside ``<w:ins>`` (tracked
    insertions) and never includes deletions — a broken hybrid for documents
    that carry tracked changes. This walks the paragraph's runs (including
    those nested in ``<w:ins>``), skips runs inside ``<w:del>`` and runs inside
    nested paragraphs (floating text boxes / drawings), and mirrors
    python-docx's run-text rendering (``w:t`` → text, ``w:tab`` → tab,
    ``w:br``/``w:cr`` → newline) so line breaks and tabs are preserved.
    """
    para_p = paragraph._p
    parts: list[str] = []
    for run in para_p.iter(_QN_R):
        if _run_skipped(run, para_p):
            continue
        for child in run:
            tag = child.tag
            if tag == _QN_T:
                parts.append(child.text or "")
            elif tag == _QN_TAB:
                parts.append("\t")
            elif tag == _QN_BR or tag == _QN_CR:
                parts.append("\n")
    return "".join(parts)


def _text_matches_main_section(text: str) -> bool:
    """True if ``text`` is exactly a JUTLP main-section name or alias.

    Matching is case-insensitive and tolerates a single trailing ``:`` or ``.``
    (e.g. ``"Discussion:"``). It requires the *whole* paragraph to equal a
    section name — a body sentence that merely starts with "Discussion" never
    matches.

    Why this exists: the Sam-fix assembly pass restyles a paragraph to
    ``Heading 1`` whenever its text equals one of these names, regardless of the
    style the author left it in (see ``output_generation_samfix.
    _required_body_heading_style``). That pass runs in Phase 3, but the
    structure validator runs in Phase 1 and would otherwise miss a section the
    author marked only visually (bold/Normal) — reporting e.g. "Discussion
    missing" for a section that is plainly present and will be restyled.
    Recognising the heading by text here closes that ordering gap without
    reordering the pipeline, mirroring ``_is_pseudo_heading2`` for subheadings.
    """
    if not text:
        return False
    normalised = strip_leading_section_number(text).strip().lower().rstrip(":.").strip()
    if not normalised:
        return False
    aliases = CANONICAL_STRUCTURE["section_aliases"]
    for section in CANONICAL_STRUCTURE["main_sections"]:
        for name in [section] + aliases.get(section, []):
            if normalised == name.strip().lower():
                return True
    return False


def _is_main_section_heading(p: ParagraphRecord) -> bool:
    """True for a real ``Heading 1`` paragraph OR a non-empty paragraph whose
    text equals a canonical JUTLP section name (an as-yet-unstyled heading)."""
    if p.is_empty:
        return False
    return p.style == "Heading 1" or _text_matches_main_section(p.text)


##read a docx file from a given path
def read_docx(path: Path) -> Document:
    return Document(path)


##helper function for parsing docx into seperate ParagraphRecord objects
def load_paragraphs(docx_path: str) -> list[ParagraphRecord]:
    doc = Document(docx_path)
    records = []

    for i, p in enumerate(doc.paragraphs):
        # Use the accepted-changes view so tracked insertions the bot makes
        # (e.g. an inserted Keywords heading) are visible when a pass re-reads
        # the output, and tracked deletions don't leave phantom gaps.
        text = _accepted_paragraph_text(p).strip()
        style = p.style.name if p.style else "NO_STYLE"

        # Detect pseudo-headings: paragraphs where every run is bold. Some
        # authors style headings visually (bold + larger font) without using
        # Heading 1/Heading 2 styles. We use this signal in the prompt
        # builder to treat these as headings instead of body paragraphs.
        # Guard against empty runs (true for empty paragraphs) and make sure
        # at least one run is actually bold — python-docx returns None for
        # "inherit from style" and False for explicit not-bold.
        runs = [r for r in p.runs if (r.text or "").strip()]
        is_all_bold = bool(runs) and all(r.bold for r in runs)

        records.append(
            ParagraphRecord(
                index=i,
                text=text,
                style=style,
                is_empty=(text == ""),
                is_all_bold=is_all_bold,
            )
        )

    return records


##Helper Functions
##find the first paragraph matching text and optional style
def find_first_paragraph(paragraphs: list[ParagraphRecord], text: str, style: str | None = None):
    for p in paragraphs:
        if p.text == text and (style is None or p.style == style):
            return p
    return None


##count occurrences of each paragraph style
def count_styles(paragraphs: list[ParagraphRecord]) -> dict:
    counts = {}
    for p in paragraphs:
        counts[p.style] = counts.get(p.style, 0) + 1
    return counts


##get all paragraphs for a specific style
def find_all_by_style(paragraphs: list[ParagraphRecord], style: str):
    return [p for p in paragraphs if p.style == style]


##get all heading 1 paragraphs
def find_all_heading1(paragraphs: list[ParagraphRecord]):
    return [p for p in paragraphs if p.style == "Heading 1"]


def extract_heading_level_1_sections(paragraphs: list[ParagraphRecord]) -> list[dict]:
    """Return all non-empty Heading 1 sections in document order.

    `position` is the paragraph index in the Word document, which gives callers
    a stable anchor for comments, section bounds, and ordering checks.
    """
    return [
        {"text": p.text, "position": p.index}
        for p in paragraphs
        if _is_main_section_heading(p)
    ]


def detect_heading_level_1(docx_path: str) -> list[dict]:
    """Scan a .docx file and return its Heading 1 section text and positions."""
    return extract_heading_level_1_sections(load_paragraphs(docx_path))


##get all heading 2 paragraphs
def find_all_heading2(paragraphs: list[ParagraphRecord]):
    return [p for p in paragraphs if p.style == "Heading 2"]


def _heading_matches(heading_text: str, query: str) -> bool:
    """True if heading_text exactly matches query, or starts with 'query:'.

    Handles titled headings like 'Introduction: Background and Context'
    matching the query 'Introduction', and strips a leading section number
    ("3. Methodology" → "Methodology") so numbered headings match before the
    heading-correction pass removes the number. Does NOT do alias expansion —
    callers that need alias support should iterate over aliases themselves.
    """
    h = strip_leading_section_number(heading_text).strip().lower()
    q = query.strip().lower()
    return h == q or h.startswith(q + ":")


##extract non-empty main section titles from heading 1 (or unstyled equivalents)
def extract_main_sections(paragraphs: list[ParagraphRecord]) -> list[str]:
    # Strip leading heading numbers ("2. Literature Review" -> "Literature
    # Review") so every consumer of parsed["main_sections"] matches against
    # clean, unnumbered names.
    return [
        strip_leading_section_number(p.text)
        for p in paragraphs
        if _is_main_section_heading(p)
    ]


##Define secion boundaries - used for inspecting each section
def get_heading1_positions(paragraphs: list[ParagraphRecord]) -> list[int]:
    return [p.index for p in paragraphs if p.style == "Heading 1" and not p.is_empty]


#returns the start and end line of an inout section
def get_section_bounds(paragraphs: list[ParagraphRecord], section_name: str):
    # Match the section by its canonical name OR any alias ("3. Methodology"
    # → Method), so subsection extraction locates the section even when the
    # author used a numbered/alternative heading the rename pass fixes later.
    aliases = CANONICAL_STRUCTURE.get("section_aliases", {}).get(section_name, [])
    candidates = [section_name, *aliases]
    heading_positions = [
        p.index for p in paragraphs
        if _is_main_section_heading(p)
        and any(_heading_matches(p.text, c) for c in candidates)
    ]

    if not heading_positions:
        return None

    start = heading_positions[0]
    next_h1 = [
        p.index for p in paragraphs
        if _is_main_section_heading(p) and p.index > start
    ]

    end = next_h1[0] if next_h1 else len(paragraphs)
    return start, end


#match a paragraph that looks like a Heading 2 caption (e.g. "Table 1.", "Figure 3:").
_CAPTION_LIKE = re.compile(r"^(?:table|figure)\s+\d+[.:]?\s*$", re.IGNORECASE)


def _is_pseudo_heading2(p: ParagraphRecord) -> bool:
    """Return True if a Normal-styled paragraph looks structurally like an H2.

    Some manuscripts mark subheadings only visually (bold, short, no
    terminal punctuation) without applying the Heading 2 style. The bot's
    Sam-fix pass restyles these later, but the validator runs first and
    would otherwise miss them entirely. Accepting pseudo-H2s here closes
    that gap without reordering the pipeline.

    A paragraph qualifies when ALL of these hold:

    * has visible text and every run is bold (``is_all_bold``),
    * is short (≤80 characters — real subheadings rarely exceed this),
    * does not end in terminal punctuation ``.!?`` — distinguishes
      headings from short sentences that just happen to be bold-emphasised,
    * is not a Table/Figure caption like ``Table 1.``.
    """
    if p.is_empty or not p.is_all_bold:
        return False
    text = p.text.strip()
    if not text or len(text) > 80:
        return False
    if text[-1] in ".!?":
        return False
    if _CAPTION_LIKE.match(text):
        return False
    return True


#get bounds of subsections
def extract_subsections(paragraphs: list[ParagraphRecord], section_name: str) -> list[str]:
    bounds = get_section_bounds(paragraphs, section_name)
    if bounds is None:
        return []

    start, end = bounds
    # Accept paragraphs styled "Heading 2" OR structurally-equivalent
    # pseudo-headings (bold + short + no terminal punctuation). The pseudo
    # branch catches manuscripts that visually mark subheadings without
    # applying the template style; Sam-fix restyles them later in the
    # pipeline but the validator runs first.
    return [
        p.text for p in paragraphs[start:end]
        if not p.is_empty
        and (p.style == "Heading 2" or _is_pseudo_heading2(p))
    ]


#extract the front page (contains title, authors, affiliations, abstract (<250 words), practitioner notes, and keywords.)
def get_front_page(paragraphs: list[ParagraphRecord]) -> list[ParagraphRecord]:
    # Use fuzzy match so "Introduction: Background" also works as the boundary
    intro = next(
        (p for p in paragraphs
         if _is_main_section_heading(p)
         and _heading_matches(p.text, "Introduction")
         and not p.is_empty),
        None,
    )
    if intro is None:
        return paragraphs
    return [p for p in paragraphs if p.index < intro.index]


##summarise required front page elements
def front_page_summary(paragraphs: list[ParagraphRecord]) -> dict:
    front = get_front_page(paragraphs)

    return {
        "title_found": any(p.style == "Article Title" for p in front),
        "authors_found": any(p.style == "Authors" for p in front),
        "affiliations_found": any(p.style == "Author Affiliations" for p in front),
        "abstract_heading_found": any(p.text == "Abstract" for p in front),
        "practitioner_notes_heading_found": any(p.text == "Practitioner Notes" for p in front),
        "keywords_heading_found": any(p.text == "Keywords" for p in front),
    }


#Abstract extractor/word count
def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+(?:[-']\w+)*\b", text))


def estimate_line_count(text: str, words_per_line: int) -> int:
    """Estimate how many rendered lines ``text`` occupies.

    Takes the larger of the explicit line count (hard ``\\n`` breaks) and a
    words-per-line estimate, so an abstract padded with manual breaks is not
    undercounted. Used to enforce the front-page abstract line limit."""
    if not text.strip():
        return 0
    explicit_lines = text.count("\n") + 1
    words = word_count(text)
    if words_per_line <= 0 or words == 0:
        return explicit_lines
    estimated = (words + words_per_line - 1) // words_per_line
    return max(explicit_lines, estimated)


# Styles / texts that mark the end of the abstract body (start of the next
# front-page element or the body). Detected so the abstract length is measured
# correctly even when the manuscript does NOT use the template "Front Page Text"
# style (authors frequently leave the abstract as Normal).
_ABSTRACT_BOUNDARY_STYLES = frozenset({
    "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading Front Page",
})
_ABSTRACT_BOUNDARY_TEXTS = frozenset({
    "practitioner notes", "practioner notes", "keywords", "keyword", "citation",
})


def _is_abstract_boundary(p: ParagraphRecord) -> bool:
    if p.style in _ABSTRACT_BOUNDARY_STYLES:
        return True
    low = p.text.strip().lower()
    if low in _ABSTRACT_BOUNDARY_TEXTS:
        return True
    return low.startswith(("keywords:", "keyword:", "citation:"))


##extract abstract paragraphs and return counts
def extract_abstract(paragraphs: list[ParagraphRecord]) -> dict:
    abstract_heading = find_first_paragraph(paragraphs, "Abstract")
    if abstract_heading is None:
        return {
            "found": False,
            "paragraphs": [],
            "paragraph_count": 0,
            "word_count": 0,
            "line_count": 0,
        }

    # Collect every non-empty paragraph between the Abstract heading and the
    # next front-page boundary (Practitioner Notes / Keywords / Citation) or the
    # first body heading (e.g. Introduction). Style-agnostic so a Normal-styled
    # abstract is still measured — the previous "Front Page Text"-only filter
    # silently counted such abstracts as 0 words and never flagged length.
    abstract_paras = []
    for p in paragraphs:
        if p.index <= abstract_heading.index:
            continue
        if p.is_empty:
            continue
        if _is_abstract_boundary(p):
            break
        abstract_paras.append(p)

    combined_text = "\n".join(p.text for p in abstract_paras)
    wpl = CANONICAL_STRUCTURE["front_page"]["abstract_words_per_line_estimate"]

    return {
        "found": True,
        "paragraphs": abstract_paras,
        "paragraph_count": len(abstract_paras),
        "word_count": word_count(combined_text),
        "line_count": estimate_line_count(combined_text, wpl),
    }


#exctract practitioner words
def extract_practitioner_notes(paragraphs: list[ParagraphRecord]) -> list[ParagraphRecord]:
    heading = find_first_paragraph(paragraphs, "Practitioner Notes")
    if heading is None:
        return []

    notes = []
    for p in paragraphs:
        if p.index <= heading.index:
            continue
        if p.text == "Keywords":
            break
        if p.style == "Practitioner Notes" and not p.is_empty:
            notes.append(p)

    return notes


##extract and clean keyword list
def extract_keywords(paragraphs: list[ParagraphRecord]) -> list[str]:
    heading = find_first_paragraph(paragraphs, "Keywords")
    if heading is None:
        return []

    keyword_lines = []
    for p in paragraphs:
        if p.index <= heading.index:
            continue
        if p.style == "Heading 1":
            break
        if p.text == "Introduction":
            break
        if not p.is_empty:
            keyword_lines.append(p.text)

    if not keyword_lines:
        return []

    text = " ".join(keyword_lines)
    parts = [x.strip() for x in text.split(",")]
    return [p for p in parts if p]


#combine everything and create doc structure
def parse_docx_structure(docx_path: str) -> dict:
    paragraphs = load_paragraphs(docx_path)

    # Try canonical name then common aliases so "Methods"/"Methodology" still
    # populates the "Method" subsections bucket used by the validator.
    method_subs = (
        extract_subsections(paragraphs, "Method")
        or extract_subsections(paragraphs, "Methods")
        or extract_subsections(paragraphs, "Methodology")
    )

    return {
        "front_page": {
            **front_page_summary(paragraphs),
            "abstract": extract_abstract(paragraphs),
            "practitioner_notes_count": len(extract_practitioner_notes(paragraphs)),
            "keywords": extract_keywords(paragraphs),
        },
        "main_sections": extract_main_sections(paragraphs),
        "heading_level_1_sections": extract_heading_level_1_sections(paragraphs),
        "subsections": {
            "Method": method_subs,
            "Discussion": extract_subsections(paragraphs, "Discussion"),
        },
        "style_counts": count_styles(paragraphs),
    }


if __name__ == "__main__":
    parsed = parse_docx_structure("tests/jutlp_sample_docx_test_pack/01_valid_identified.docx")
    from pprint import pprint
    pprint(parsed)
